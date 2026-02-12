"""Report generator for heat treatment simulation results.

Generates Word documents (.docx) and PDF reports with simulation configuration,
results summary, and embedded plots.
"""
from io import BytesIO
from datetime import datetime
from typing import Optional, List, Tuple
import tempfile
import os

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from fpdf import FPDF


class SimulationReportGenerator:
    """Generate Word reports for completed simulations."""

    def __init__(self, simulation):
        """Initialize report generator.

        Parameters
        ----------
        simulation : Simulation
            The simulation model instance
        """
        self.sim = simulation
        self.doc = Document()

    def generate_report(self) -> bytes:
        """Generate complete simulation report.

        Returns
        -------
        bytes
            The Word document as bytes
        """
        self._add_header()
        self._add_geometry_section()
        self._add_heat_treatment_section()
        self._add_solver_section()
        self._add_results_summary()
        self._add_temperature_plots()
        self._add_phase_fractions()
        self._add_hardness_prediction()

        # Save to bytes
        output = BytesIO()
        self.doc.save(output)
        output.seek(0)
        return output.getvalue()

    def _add_header(self) -> None:
        """Add report header with simulation info."""
        # Title
        title = self.doc.add_heading(f'Heat Treatment Simulation Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Simulation name and description
        self.doc.add_heading(self.sim.name, level=1)
        if self.sim.description:
            self.doc.add_paragraph(self.sim.description)

        # Info table
        table = self.doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.LEFT

        cells = [
            ('Steel Grade', self.sim.steel_grade.designation if self.sim.steel_grade else 'N/A'),
            ('Data Source', self.sim.steel_grade.data_source if self.sim.steel_grade else 'N/A'),
            ('Status', self.sim.status.title()),
            ('Generated', datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')),
        ]

        for i, (label, value) in enumerate(cells):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = str(value)
            # Bold the label
            for run in table.rows[i].cells[0].paragraphs[0].runs:
                run.bold = True

        self.doc.add_paragraph()  # Spacing

    def _add_geometry_section(self) -> None:
        """Add geometry configuration section."""
        self.doc.add_heading('Geometry Configuration', level=2)

        geom = self.sim.geometry_dict
        geom_type = self.sim.geometry_type

        # Geometry type
        p = self.doc.add_paragraph()
        p.add_run('Type: ').bold = True
        p.add_run(self.sim.geometry_label)

        # Dimensions table
        table = self.doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.rows[0].cells[0].text = 'Parameter'
        table.rows[0].cells[1].text = 'Value'
        for cell in table.rows[0].cells:
            for run in cell.paragraphs[0].runs:
                run.bold = True

        if geom_type == 'cylinder':
            self._add_table_row(table, 'Radius', f"{geom.get('radius', 0) * 1000:.1f} mm")
            self._add_table_row(table, 'Length', f"{geom.get('length', 0) * 1000:.1f} mm")
        elif geom_type == 'plate':
            self._add_table_row(table, 'Thickness', f"{geom.get('thickness', 0) * 1000:.1f} mm")
            self._add_table_row(table, 'Width', f"{geom.get('width', 0) * 1000:.1f} mm")
            self._add_table_row(table, 'Length', f"{geom.get('length', 0) * 1000:.1f} mm")
        elif geom_type == 'hollow_cylinder':
            self._add_table_row(table, 'Outer Diameter', f"{geom.get('outer_diameter', 0) * 1000:.1f} mm")
            self._add_table_row(table, 'Inner Diameter', f"{geom.get('inner_diameter', 0) * 1000:.1f} mm")
            self._add_table_row(table, 'Length', f"{geom.get('length', 0) * 1000:.1f} mm")
        elif geom_type == 'ring':
            self._add_table_row(table, 'Inner Radius', f"{geom.get('inner_radius', 0) * 1000:.1f} mm")
            self._add_table_row(table, 'Outer Radius', f"{geom.get('outer_radius', 0) * 1000:.1f} mm")
            self._add_table_row(table, 'Length', f"{geom.get('length', 0) * 1000:.1f} mm")
        elif geom_type == 'cad':
            cad = self.sim.cad_analysis_dict
            if cad:
                self._add_table_row(table, 'Source File', self.sim.cad_filename or 'Unknown')
                self._add_table_row(table, 'Volume', f"{cad.get('volume', 0) * 1e6:.2f} cm³")
                self._add_table_row(table, 'Surface Area', f"{cad.get('surface_area', 0) * 1e4:.1f} cm²")
                self._add_table_row(table, 'Characteristic Length', f"{cad.get('characteristic_length', 0) * 1000:.2f} mm")
                self._add_table_row(table, 'Equivalent Type', self.sim.cad_equivalent_type or 'auto')

        self.doc.add_paragraph()  # Spacing

    def _add_heat_treatment_section(self) -> None:
        """Add heat treatment configuration section."""
        self.doc.add_heading('Heat Treatment Configuration', level=2)

        ht = self.sim.ht_config
        if not ht:
            self.doc.add_paragraph('Heat treatment not configured.')
            return

        # Heating phase
        heating = ht.get('heating', {})
        if heating.get('enabled', False):
            self.doc.add_heading('Heating Phase', level=3)
            table = self._create_param_table()
            self._add_table_row(table, 'Initial Temperature', f"{heating.get('initial_temperature', 25):.0f} °C")
            self._add_table_row(table, 'Target Temperature', f"{heating.get('target_temperature', 850):.0f} °C")
            self._add_table_row(table, 'Hold Time', f"{heating.get('hold_time', 60):.0f} min")
            self._add_table_row(table, 'Furnace Atmosphere', heating.get('furnace_atmosphere', 'air').title())
            self._add_table_row(table, 'Furnace HTC', f"{heating.get('furnace_htc', 25):.0f} W/m²K")
            if heating.get('use_radiation', True):
                self._add_table_row(table, 'Emissivity', f"{heating.get('furnace_emissivity', 0.85):.2f}")
            if heating.get('cold_furnace', False):
                self._add_table_row(table, 'Cold Furnace Start', f"{heating.get('furnace_start_temperature', 25):.0f} °C")
                self._add_table_row(table, 'Ramp Rate', f"{heating.get('furnace_ramp_rate', 5):.1f} °C/min")
            self.doc.add_paragraph()

        # Transfer phase
        transfer = ht.get('transfer', {})
        if transfer.get('enabled', False):
            self.doc.add_heading('Transfer Phase', level=3)
            table = self._create_param_table()
            self._add_table_row(table, 'Duration', f"{transfer.get('duration', 10):.0f} s")
            self._add_table_row(table, 'Ambient Temperature', f"{transfer.get('ambient_temperature', 25):.0f} °C")
            self._add_table_row(table, 'HTC', f"{transfer.get('htc', 10):.0f} W/m²K")
            if transfer.get('use_radiation', True):
                self._add_table_row(table, 'Emissivity', f"{transfer.get('emissivity', 0.85):.2f}")
            self.doc.add_paragraph()

        # Quenching phase
        quenching = ht.get('quenching', {})
        if quenching:
            self.doc.add_heading('Quenching Phase', level=3)
            table = self._create_param_table()
            media = quenching.get('media', 'water')
            from app.models.simulation import QUENCH_MEDIA_LABELS, AGITATION_LABELS, calculate_quench_htc
            self._add_table_row(table, 'Media', QUENCH_MEDIA_LABELS.get(media, media.title()))
            self._add_table_row(table, 'Media Temperature', f"{quenching.get('media_temperature', 25):.0f} °C")
            agitation = quenching.get('agitation', 'moderate')
            self._add_table_row(table, 'Agitation', AGITATION_LABELS.get(agitation, agitation.title()))
            # Calculate effective HTC
            htc = quenching.get('htc_override')
            if htc is None:
                htc = calculate_quench_htc(media, agitation, quenching.get('media_temperature', 25))
            self._add_table_row(table, 'Effective HTC', f"{htc:.0f} W/m²K")
            self._add_table_row(table, 'Duration', f"{quenching.get('duration', 300):.0f} s")
            self.doc.add_paragraph()

        # Tempering phase
        tempering = ht.get('tempering', {})
        if tempering.get('enabled', False):
            self.doc.add_heading('Tempering Phase', level=3)
            table = self._create_param_table()
            self._add_table_row(table, 'Temperature', f"{tempering.get('temperature', 550):.0f} °C")
            self._add_table_row(table, 'Hold Time', f"{tempering.get('hold_time', 120):.0f} min")
            self._add_table_row(table, 'Cooling Method', tempering.get('cooling_method', 'air').title())
            self._add_table_row(table, 'HTC', f"{tempering.get('htc', 25):.0f} W/m²K")
            self.doc.add_paragraph()

    def _add_solver_section(self) -> None:
        """Add solver settings section."""
        self.doc.add_heading('Solver Settings', level=2)

        solver = self.sim.solver_dict
        table = self._create_param_table()
        self._add_table_row(table, 'Spatial Nodes', str(solver.get('n_nodes', 51)))
        self._add_table_row(table, 'Time Step', f"{solver.get('dt', 0.1):.3f} s")
        self._add_table_row(table, 'Max Simulation Time', f"{solver.get('max_time', 1800):.0f} s")
        if solver.get('auto_dt', False):
            self._add_table_row(table, 'Auto Time Step', 'Enabled')

        self.doc.add_paragraph()

    def _add_results_summary(self) -> None:
        """Add results summary section."""
        self.doc.add_heading('Results Summary', level=2)

        # Status and timing
        table = self._create_param_table()
        self._add_table_row(table, 'Status', self.sim.status.title())
        if self.sim.started_at:
            self._add_table_row(table, 'Started', self.sim.started_at.strftime('%Y-%m-%d %H:%M:%S'))
        if self.sim.completed_at:
            self._add_table_row(table, 'Completed', self.sim.completed_at.strftime('%Y-%m-%d %H:%M:%S'))
        if self.sim.duration_seconds:
            self._add_table_row(table, 'Computation Time', f"{self.sim.duration_seconds:.1f} s")

        # t8/5 from results
        results = self.sim.results.all()
        full_cycle = next((r for r in results if r.result_type == 'full_cycle'), None)
        if full_cycle and full_cycle.t_800_500:
            self._add_table_row(table, 't₈/₅ Cooling Time', f"{full_cycle.t_800_500:.2f} s")

        self.doc.add_paragraph()

    def _add_temperature_plots(self) -> None:
        """Add temperature plots section."""
        self.doc.add_heading('Temperature Results', level=2)

        results = self.sim.results.all()

        # Full cycle plot
        full_cycle = next((r for r in results if r.result_type == 'full_cycle'), None)
        if full_cycle and full_cycle.plot_image:
            self.doc.add_heading('Full Heat Treatment Cycle', level=3)
            self._add_plot(full_cycle.plot_image)

        # Temperature profile
        profile = next((r for r in results if r.result_type == 'temperature_profile'), None)
        if profile and profile.plot_image:
            self.doc.add_heading('Temperature Profile', level=3)
            self._add_plot(profile.plot_image)

        # Cooling rate
        rate = next((r for r in results if r.result_type == 'cooling_rate'), None)
        if rate and rate.plot_image:
            self.doc.add_heading('Cooling Rate', level=3)
            self._add_plot(rate.plot_image)

    def _add_phase_fractions(self) -> None:
        """Add phase fraction results section."""
        results = self.sim.results.all()
        phase_result = next((r for r in results if r.result_type == 'phase_fraction'), None)

        if not phase_result:
            return

        self.doc.add_heading('Phase Transformation', level=2)

        # Plot
        if phase_result.plot_image:
            self._add_plot(phase_result.plot_image)

        # Phase fractions table
        phases = phase_result.phases_dict
        if phases:
            table = self.doc.add_table(rows=1, cols=2)
            table.style = 'Table Grid'
            table.rows[0].cells[0].text = 'Phase'
            table.rows[0].cells[1].text = 'Fraction (%)'
            for cell in table.rows[0].cells:
                for run in cell.paragraphs[0].runs:
                    run.bold = True

            for phase, fraction in phases.items():
                if fraction > 0.01:  # Only show phases > 1%
                    row = table.add_row()
                    row.cells[0].text = phase.replace('_', ' ').title()
                    row.cells[1].text = f"{fraction * 100:.1f}"

        self.doc.add_paragraph()

    def _add_hardness_prediction(self) -> None:
        """Add hardness prediction results section."""
        results = self.sim.results.all()
        hardness_result = next((r for r in results if r.result_type == 'hardness_prediction'), None)

        if not hardness_result:
            return

        self.doc.add_heading('Hardness Prediction', level=2)

        # Plot
        if hardness_result.plot_image:
            self._add_plot(hardness_result.plot_image)

        # Hardness values table
        data = hardness_result.data_dict
        if data and data.get('hardness_hv'):
            hv = data.get('hardness_hv', {})
            hrc = data.get('hardness_hrc', {})

            table = self.doc.add_table(rows=1, cols=3)
            table.style = 'Table Grid'
            table.rows[0].cells[0].text = 'Position'
            table.rows[0].cells[1].text = 'HV'
            table.rows[0].cells[2].text = 'HRC'
            for cell in table.rows[0].cells:
                for run in cell.paragraphs[0].runs:
                    run.bold = True

            positions = [
                ('Center', 'center'),
                ('1/3 R', 'one_third'),
                ('2/3 R', 'two_thirds'),
                ('Surface', 'surface'),
            ]

            for label, key in positions:
                row = table.add_row()
                row.cells[0].text = label
                row.cells[1].text = f"{hv.get(key, 0):.0f}" if hv.get(key) else '-'
                row.cells[2].text = f"{hrc.get(key, 0):.0f}" if hrc.get(key) else '-'

            self.doc.add_paragraph()

            # Carbon equivalent and DI
            if data.get('carbon_equivalent'):
                p = self.doc.add_paragraph()
                p.add_run('Carbon Equivalent CE(IIW): ').bold = True
                p.add_run(f"{data['carbon_equivalent']:.3f}")

            if data.get('ideal_diameter'):
                p = self.doc.add_paragraph()
                p.add_run('Ideal Diameter DI: ').bold = True
                p.add_run(f"{data['ideal_diameter']:.2f} in")

            # Mechanical properties
            if data.get('uts_mpa'):
                self.doc.add_heading('Estimated Mechanical Properties', level=3)
                table = self.doc.add_table(rows=1, cols=5)
                table.style = 'Table Grid'
                headers = ['Position', 'UTS (MPa)', 'YS (MPa)', 'El. (%)', 'Toughness']
                for i, h in enumerate(headers):
                    table.rows[0].cells[i].text = h
                    for run in table.rows[0].cells[i].paragraphs[0].runs:
                        run.bold = True

                for label, key in positions:
                    row = table.add_row()
                    row.cells[0].text = label
                    row.cells[1].text = f"{data['uts_mpa'].get(key, 0):.0f}" if data['uts_mpa'].get(key) else '-'
                    row.cells[2].text = f"{data['ys_mpa'].get(key, 0):.0f}" if data['ys_mpa'].get(key) else '-'
                    row.cells[3].text = f"{data.get('elongation_pct', {}).get(key, 0):.1f}" if data.get('elongation_pct', {}).get(key) else '-'
                    row.cells[4].text = data.get('toughness_rating', {}).get(key, '-').title()

                self.doc.add_paragraph()

    def _create_param_table(self):
        """Create a parameter table with header row."""
        table = self.doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        table.rows[0].cells[0].text = 'Parameter'
        table.rows[0].cells[1].text = 'Value'
        for cell in table.rows[0].cells:
            for run in cell.paragraphs[0].runs:
                run.bold = True
        return table

    def _add_table_row(self, table, label: str, value: str) -> None:
        """Add a row to a table."""
        row = table.add_row()
        row.cells[0].text = label
        row.cells[1].text = value

    def _add_plot(self, plot_bytes: bytes, width: float = 6.0) -> None:
        """Add a plot image to the document."""
        self.doc.add_picture(BytesIO(plot_bytes), width=Inches(width))
        # Center the image
        last_paragraph = self.doc.paragraphs[-1]
        last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph()  # Spacing


class SimulationPDFReportGenerator:
    """Generate PDF reports for completed simulations using fpdf2."""

    def __init__(self, simulation):
        """Initialize PDF report generator.

        Parameters
        ----------
        simulation : Simulation
            The simulation model instance
        """
        self.sim = simulation
        self.pdf = FPDF()
        self.pdf.set_auto_page_break(auto=True, margin=15)
        self.temp_files = []  # Track temp files for cleanup

    def generate_report(self) -> bytes:
        """Generate complete PDF simulation report.

        Returns
        -------
        bytes
            The PDF document as bytes
        """
        try:
            self.pdf.add_page()
            self._add_header()
            self._add_geometry_section()
            self._add_heat_treatment_section()
            self._add_solver_section()
            self._add_results_summary()
            self._add_temperature_plots()
            self._add_phase_fractions()
            self._add_hardness_prediction()

            # Get PDF as bytes
            return bytes(self.pdf.output())
        finally:
            # Clean up temp files
            for f in self.temp_files:
                try:
                    os.unlink(f)
                except OSError:
                    pass

    def _add_header(self) -> None:
        """Add report header."""
        # Title
        self.pdf.set_font('Helvetica', 'B', 24)
        self.pdf.set_text_color(26, 82, 118)  # Dark blue
        self.pdf.cell(0, 15, 'Heat Treatment Simulation Report', align='C', new_x='LMARGIN', new_y='NEXT')

        # Line under title
        self.pdf.set_draw_color(26, 82, 118)
        self.pdf.set_line_width(0.5)
        self.pdf.line(10, self.pdf.get_y(), 200, self.pdf.get_y())
        self.pdf.ln(5)

        # Simulation name
        self.pdf.set_font('Helvetica', 'B', 16)
        self.pdf.set_text_color(40, 116, 166)
        self.pdf.cell(0, 10, self.sim.name, align='C', new_x='LMARGIN', new_y='NEXT')

        if self.sim.description:
            self.pdf.set_font('Helvetica', 'I', 11)
            self.pdf.set_text_color(100, 100, 100)
            self.pdf.multi_cell(0, 6, self.sim.description, align='C')

        self.pdf.ln(5)

        # Info table
        self._add_info_table([
            ('Steel Grade', self.sim.steel_grade.designation if self.sim.steel_grade else 'N/A'),
            ('Data Source', self.sim.steel_grade.data_source if self.sim.steel_grade else 'N/A'),
            ('Status', self.sim.status.title()),
            ('Generated', datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')),
        ])
        self.pdf.ln(5)

    def _add_section_heading(self, title: str) -> None:
        """Add a section heading."""
        self.pdf.ln(3)
        self.pdf.set_font('Helvetica', 'B', 14)
        self.pdf.set_text_color(40, 116, 166)
        self.pdf.cell(0, 8, title, new_x='LMARGIN', new_y='NEXT')
        # Line under heading
        self.pdf.set_draw_color(40, 116, 166)
        self.pdf.set_line_width(0.3)
        self.pdf.line(10, self.pdf.get_y(), 200, self.pdf.get_y())
        self.pdf.ln(3)

    def _add_subsection_heading(self, title: str) -> None:
        """Add a subsection heading."""
        self.pdf.set_font('Helvetica', 'B', 12)
        self.pdf.set_text_color(52, 152, 219)
        self.pdf.cell(0, 7, title, new_x='LMARGIN', new_y='NEXT')
        self.pdf.ln(1)

    def _add_info_table(self, rows: List[Tuple[str, str]]) -> None:
        """Add a simple two-column info table."""
        self.pdf.set_font('Helvetica', '', 10)
        self.pdf.set_text_color(0, 0, 0)

        col_width = 45
        for label, value in rows:
            self.pdf.set_font('Helvetica', 'B', 10)
            self.pdf.cell(col_width, 7, label + ':', border=1)
            self.pdf.set_font('Helvetica', '', 10)
            self.pdf.cell(0, 7, str(value), border=1, new_x='LMARGIN', new_y='NEXT')

    def _add_param_table(self, rows: List[Tuple[str, str]]) -> None:
        """Add a parameter/value table."""
        self.pdf.set_font('Helvetica', '', 10)
        self.pdf.set_text_color(0, 0, 0)

        # Header
        self.pdf.set_fill_color(236, 240, 241)
        self.pdf.set_font('Helvetica', 'B', 10)
        self.pdf.cell(80, 7, 'Parameter', border=1, fill=True)
        self.pdf.cell(0, 7, 'Value', border=1, fill=True, new_x='LMARGIN', new_y='NEXT')

        # Rows
        self.pdf.set_font('Helvetica', '', 10)
        for label, value in rows:
            self.pdf.cell(80, 7, label, border=1)
            self.pdf.cell(0, 7, str(value), border=1, new_x='LMARGIN', new_y='NEXT')

    def _add_geometry_section(self) -> None:
        """Add geometry configuration section."""
        self._add_section_heading('Geometry Configuration')

        geom = self.sim.geometry_dict
        geom_type = self.sim.geometry_type

        self.pdf.set_font('Helvetica', '', 11)
        self.pdf.set_text_color(0, 0, 0)
        self.pdf.cell(0, 7, f'Type: {self.sim.geometry_label}', new_x='LMARGIN', new_y='NEXT')
        self.pdf.ln(2)

        rows = []
        if geom_type == 'cylinder':
            rows = [
                ('Radius', f"{geom.get('radius', 0) * 1000:.1f} mm"),
                ('Length', f"{geom.get('length', 0) * 1000:.1f} mm"),
            ]
        elif geom_type == 'plate':
            rows = [
                ('Thickness', f"{geom.get('thickness', 0) * 1000:.1f} mm"),
                ('Width', f"{geom.get('width', 0) * 1000:.1f} mm"),
                ('Length', f"{geom.get('length', 0) * 1000:.1f} mm"),
            ]
        elif geom_type == 'hollow_cylinder':
            rows = [
                ('Outer Diameter', f"{geom.get('outer_diameter', 0) * 1000:.1f} mm"),
                ('Inner Diameter', f"{geom.get('inner_diameter', 0) * 1000:.1f} mm"),
                ('Length', f"{geom.get('length', 0) * 1000:.1f} mm"),
            ]
        elif geom_type == 'ring':
            rows = [
                ('Inner Radius', f"{geom.get('inner_radius', 0) * 1000:.1f} mm"),
                ('Outer Radius', f"{geom.get('outer_radius', 0) * 1000:.1f} mm"),
                ('Length', f"{geom.get('length', 0) * 1000:.1f} mm"),
            ]
        elif geom_type == 'cad':
            cad = self.sim.cad_analysis_dict
            if cad:
                rows = [
                    ('Source File', self.sim.cad_filename or 'Unknown'),
                    ('Volume', f"{cad.get('volume', 0) * 1e6:.2f} cm3"),
                    ('Surface Area', f"{cad.get('surface_area', 0) * 1e4:.1f} cm2"),
                    ('Characteristic Length', f"{cad.get('characteristic_length', 0) * 1000:.2f} mm"),
                    ('Equivalent Type', self.sim.cad_equivalent_type or 'auto'),
                ]

        if rows:
            self._add_param_table(rows)

    def _add_heat_treatment_section(self) -> None:
        """Add heat treatment configuration section."""
        self._add_section_heading('Heat Treatment Configuration')

        ht = self.sim.ht_config
        if not ht:
            self.pdf.set_font('Helvetica', '', 11)
            self.pdf.cell(0, 7, 'Heat treatment not configured.', new_x='LMARGIN', new_y='NEXT')
            return

        # Heating phase
        heating = ht.get('heating', {})
        if heating.get('enabled', False):
            self._add_subsection_heading('Heating Phase')
            rows = [
                ('Initial Temperature', f"{heating.get('initial_temperature', 25):.0f} C"),
                ('Target Temperature', f"{heating.get('target_temperature', 850):.0f} C"),
                ('Hold Time', f"{heating.get('hold_time', 60):.0f} min"),
                ('Furnace Atmosphere', heating.get('furnace_atmosphere', 'air').title()),
                ('Furnace HTC', f"{heating.get('furnace_htc', 25):.0f} W/m2K"),
            ]
            if heating.get('use_radiation', True):
                rows.append(('Emissivity', f"{heating.get('furnace_emissivity', 0.85):.2f}"))
            self._add_param_table(rows)
            self.pdf.ln(3)

        # Transfer phase
        transfer = ht.get('transfer', {})
        if transfer.get('enabled', False):
            self._add_subsection_heading('Transfer Phase')
            rows = [
                ('Duration', f"{transfer.get('duration', 10):.0f} s"),
                ('Ambient Temperature', f"{transfer.get('ambient_temperature', 25):.0f} C"),
                ('HTC', f"{transfer.get('htc', 10):.0f} W/m2K"),
            ]
            self._add_param_table(rows)
            self.pdf.ln(3)

        # Quenching phase
        quenching = ht.get('quenching', {})
        if quenching:
            self._add_subsection_heading('Quenching Phase')
            media = quenching.get('media', 'water')
            from app.models.simulation import QUENCH_MEDIA_LABELS, AGITATION_LABELS, calculate_quench_htc
            agitation = quenching.get('agitation', 'moderate')
            htc = quenching.get('htc_override')
            if htc is None:
                htc = calculate_quench_htc(media, agitation, quenching.get('media_temperature', 25))
            rows = [
                ('Media', QUENCH_MEDIA_LABELS.get(media, media.title())),
                ('Media Temperature', f"{quenching.get('media_temperature', 25):.0f} C"),
                ('Agitation', AGITATION_LABELS.get(agitation, agitation.title())),
                ('Effective HTC', f"{htc:.0f} W/m2K"),
                ('Duration', f"{quenching.get('duration', 300):.0f} s"),
            ]
            self._add_param_table(rows)
            self.pdf.ln(3)

        # Tempering phase
        tempering = ht.get('tempering', {})
        if tempering.get('enabled', False):
            self._add_subsection_heading('Tempering Phase')
            rows = [
                ('Temperature', f"{tempering.get('temperature', 550):.0f} C"),
                ('Hold Time', f"{tempering.get('hold_time', 120):.0f} min"),
                ('Cooling Method', tempering.get('cooling_method', 'air').title()),
                ('HTC', f"{tempering.get('htc', 25):.0f} W/m2K"),
            ]
            self._add_param_table(rows)

    def _add_solver_section(self) -> None:
        """Add solver settings section."""
        self._add_section_heading('Solver Settings')

        solver = self.sim.solver_dict
        rows = [
            ('Spatial Nodes', str(solver.get('n_nodes', 51))),
            ('Time Step', f"{solver.get('dt', 0.1):.3f} s"),
            ('Max Simulation Time', f"{solver.get('max_time', 1800):.0f} s"),
        ]
        if solver.get('auto_dt', False):
            rows.append(('Auto Time Step', 'Enabled'))
        self._add_param_table(rows)

    def _add_results_summary(self) -> None:
        """Add results summary section."""
        self._add_section_heading('Results Summary')

        rows = [('Status', self.sim.status.title())]
        if self.sim.started_at:
            rows.append(('Started', self.sim.started_at.strftime('%Y-%m-%d %H:%M:%S')))
        if self.sim.completed_at:
            rows.append(('Completed', self.sim.completed_at.strftime('%Y-%m-%d %H:%M:%S')))
        if self.sim.duration_seconds:
            rows.append(('Computation Time', f"{self.sim.duration_seconds:.1f} s"))

        # t8/5 from results
        results = self.sim.results.all()
        full_cycle = next((r for r in results if r.result_type == 'full_cycle'), None)
        if full_cycle and full_cycle.t_800_500:
            rows.append(('t8/5 Cooling Time', f"{full_cycle.t_800_500:.2f} s"))

        self._add_param_table(rows)

    def _add_plot(self, plot_bytes: bytes, title: str = None) -> None:
        """Add a plot image to the PDF."""
        if title:
            self._add_subsection_heading(title)

        # Save to temp file (fpdf2 needs file path for images)
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as f:
            f.write(plot_bytes)
            temp_path = f.name
            self.temp_files.append(temp_path)

        # Calculate image dimensions to fit page width
        page_width = self.pdf.w - 20  # margins
        self.pdf.image(temp_path, x=10, w=page_width)
        self.pdf.ln(5)

    def _add_temperature_plots(self) -> None:
        """Add temperature plots section."""
        self._add_section_heading('Temperature Results')

        results = self.sim.results.all()

        # Full cycle plot
        full_cycle = next((r for r in results if r.result_type == 'full_cycle'), None)
        if full_cycle and full_cycle.plot_image:
            self._add_plot(full_cycle.plot_image, 'Full Heat Treatment Cycle')

        # Temperature profile - new page for readability
        profile = next((r for r in results if r.result_type == 'temperature_profile'), None)
        if profile and profile.plot_image:
            self.pdf.add_page()
            self._add_plot(profile.plot_image, 'Temperature Profile')

        # Cooling rate
        rate = next((r for r in results if r.result_type == 'cooling_rate'), None)
        if rate and rate.plot_image:
            self._add_plot(rate.plot_image, 'Cooling Rate')

    def _add_phase_fractions(self) -> None:
        """Add phase fraction results section."""
        results = self.sim.results.all()
        phase_result = next((r for r in results if r.result_type == 'phase_fraction'), None)

        if not phase_result:
            return

        self.pdf.add_page()
        self._add_section_heading('Phase Transformation')

        # Plot
        if phase_result.plot_image:
            self._add_plot(phase_result.plot_image)

        # Phase fractions table
        phases = phase_result.phases_dict
        if phases:
            self.pdf.set_font('Helvetica', 'B', 10)
            self.pdf.set_fill_color(236, 240, 241)
            self.pdf.cell(80, 7, 'Phase', border=1, fill=True)
            self.pdf.cell(0, 7, 'Fraction (%)', border=1, fill=True, new_x='LMARGIN', new_y='NEXT')

            self.pdf.set_font('Helvetica', '', 10)
            for phase, fraction in phases.items():
                if fraction > 0.01:
                    phase_name = phase.replace('_', ' ').title()
                    self.pdf.cell(80, 7, phase_name, border=1)
                    self.pdf.cell(0, 7, f"{fraction * 100:.1f}", border=1, new_x='LMARGIN', new_y='NEXT')

    def _add_hardness_prediction(self) -> None:
        """Add hardness prediction results section."""
        results = self.sim.results.all()
        hardness_result = next((r for r in results if r.result_type == 'hardness_prediction'), None)

        if not hardness_result:
            return

        self.pdf.add_page()
        self._add_section_heading('Hardness Prediction')

        # Plot
        if hardness_result.plot_image:
            self._add_plot(hardness_result.plot_image)

        # Hardness values table
        data = hardness_result.data_dict
        if data and data.get('hardness_hv'):
            hv = data.get('hardness_hv', {})
            hrc = data.get('hardness_hrc', {})

            self.pdf.set_font('Helvetica', 'B', 10)
            self.pdf.set_fill_color(236, 240, 241)
            self.pdf.cell(60, 7, 'Position', border=1, fill=True)
            self.pdf.cell(40, 7, 'HV', border=1, fill=True)
            self.pdf.cell(0, 7, 'HRC', border=1, fill=True, new_x='LMARGIN', new_y='NEXT')

            self.pdf.set_font('Helvetica', '', 10)
            positions = [
                ('Center', 'center'),
                ('1/3 R', 'one_third'),
                ('2/3 R', 'two_thirds'),
                ('Surface', 'surface'),
            ]
            for label, key in positions:
                hv_val = f"{hv.get(key, 0):.0f}" if hv.get(key) else '-'
                hrc_val = f"{hrc.get(key, 0):.0f}" if hrc.get(key) else '-'
                self.pdf.cell(60, 7, label, border=1)
                self.pdf.cell(40, 7, hv_val, border=1)
                self.pdf.cell(0, 7, hrc_val, border=1, new_x='LMARGIN', new_y='NEXT')

            self.pdf.ln(5)

            # Carbon equivalent and DI
            self.pdf.set_font('Helvetica', '', 11)
            if data.get('carbon_equivalent'):
                self.pdf.cell(0, 7, f"Carbon Equivalent CE(IIW): {data['carbon_equivalent']:.3f}", new_x='LMARGIN', new_y='NEXT')
            if data.get('ideal_diameter'):
                self.pdf.cell(0, 7, f"Ideal Diameter DI: {data['ideal_diameter']:.2f} in", new_x='LMARGIN', new_y='NEXT')

            # Mechanical properties
            if data.get('uts_mpa'):
                self.pdf.ln(5)
                self.pdf.set_font('Helvetica', 'B', 12)
                self.pdf.cell(0, 8, 'Estimated Mechanical Properties', new_x='LMARGIN', new_y='NEXT')

                self.pdf.set_font('Helvetica', 'B', 10)
                self.pdf.set_fill_color(236, 240, 241)
                self.pdf.cell(35, 7, 'Position', border=1, fill=True)
                self.pdf.cell(30, 7, 'UTS (MPa)', border=1, fill=True)
                self.pdf.cell(30, 7, 'YS (MPa)', border=1, fill=True)
                self.pdf.cell(25, 7, 'El. (%)', border=1, fill=True)
                self.pdf.cell(0, 7, 'Toughness', border=1, fill=True, new_x='LMARGIN', new_y='NEXT')

                self.pdf.set_font('Helvetica', '', 10)
                positions = [
                    ('Center', 'center'),
                    ('1/3 R', 'one_third'),
                    ('2/3 R', 'two_thirds'),
                    ('Surface', 'surface'),
                ]
                for label, key in positions:
                    uts_val = f"{data['uts_mpa'].get(key, 0):.0f}" if data['uts_mpa'].get(key) else '-'
                    ys_val = f"{data['ys_mpa'].get(key, 0):.0f}" if data['ys_mpa'].get(key) else '-'
                    el_val = f"{data.get('elongation_pct', {}).get(key, 0):.1f}" if data.get('elongation_pct', {}).get(key) else '-'
                    tough_val = data.get('toughness_rating', {}).get(key, '-').title()
                    self.pdf.cell(35, 7, label, border=1)
                    self.pdf.cell(30, 7, uts_val, border=1)
                    self.pdf.cell(30, 7, ys_val, border=1)
                    self.pdf.cell(25, 7, el_val, border=1)
                    self.pdf.cell(0, 7, tough_val, border=1, new_x='LMARGIN', new_y='NEXT')


def generate_simulation_report(simulation) -> bytes:
    """Generate a Word report for a simulation.

    Parameters
    ----------
    simulation : Simulation
        The simulation model instance

    Returns
    -------
    bytes
        The Word document as bytes
    """
    generator = SimulationReportGenerator(simulation)
    return generator.generate_report()


def generate_simulation_pdf_report(simulation) -> bytes:
    """Generate a PDF report for a simulation.

    Parameters
    ----------
    simulation : Simulation
        The simulation model instance

    Returns
    -------
    bytes
        The PDF document as bytes
    """
    generator = SimulationPDFReportGenerator(simulation)
    return generator.generate_report()
