"""Generate compliance Word documents from simulation snapshots."""
import io
import json
from datetime import datetime

from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


class ComplianceReportGenerator:
    """Generate Word compliance report from a simulation snapshot."""

    def __init__(self, snapshot):
        self.snapshot = snapshot
        self.sim = snapshot.simulation
        self.doc = Document()

    def generate(self):
        """Generate the full compliance report.

        Returns
        -------
        bytes
            Word document as bytes
        """
        self._set_styles()
        self._add_cover_page()
        self._add_material_section()
        self._add_config_section()
        self._add_results_section()
        self._add_lineage_section()
        self._add_change_log_section()

        buf = io.BytesIO()
        self.doc.save(buf)
        buf.seek(0)
        return buf.getvalue()

    def _set_styles(self):
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Calibri'
        font.size = Pt(10)

    def _add_cover_page(self):
        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('\n\n\n')

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run('Simulation Compliance Report')
        run.font.size = Pt(24)
        run.bold = True

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'\n{self.sim.name}')
        run.font.size = Pt(16)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'\nVersion {self.snapshot.version}')
        run.font.size = Pt(14)

        p = self.doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f'\nGenerated: {datetime.utcnow().strftime("%Y-%m-%d %H:%M UTC")}')
        run.font.size = Pt(10)

        self.doc.add_page_break()

    def _add_material_section(self):
        self.doc.add_heading('1. Material Data', level=1)
        self.doc.add_heading('1.1 Steel Grade', level=2)

        self.doc.add_paragraph(
            f'Designation: {self.snapshot.steel_grade_designation}'
        )
        self.doc.add_paragraph(
            f'Data Source: {self.snapshot.steel_grade_data_source or "N/A"}'
        )

        # Composition
        comp = self.snapshot.composition_dict
        if comp:
            self.doc.add_heading('1.2 Chemical Composition (wt%)', level=2)
            table = self.doc.add_table(rows=1, cols=2)
            table.style = 'Light Grid Accent 1'
            table.alignment = WD_TABLE_ALIGNMENT.LEFT
            hdr = table.rows[0].cells
            hdr[0].text = 'Element'
            hdr[1].text = 'Content (%)'

            for elem in ['carbon', 'manganese', 'silicon', 'chromium', 'nickel',
                         'molybdenum', 'vanadium', 'tungsten', 'copper',
                         'phosphorus', 'sulfur', 'nitrogen', 'boron']:
                val = comp.get(elem)
                if val and val > 0:
                    row = table.add_row().cells
                    row[0].text = elem.title()
                    row[1].text = f'{val:.4f}'

        # Phase Diagram
        pd = self.snapshot.phase_diagram_dict
        if pd and pd.get('transformation_temps'):
            self.doc.add_heading('1.3 Transformation Temperatures', level=2)
            table = self.doc.add_table(rows=1, cols=2)
            table.style = 'Light Grid Accent 1'
            hdr = table.rows[0].cells
            hdr[0].text = 'Parameter'
            hdr[1].text = 'Temperature (°C)'

            for key, val in pd['transformation_temps'].items():
                row = table.add_row().cells
                row[0].text = key
                row[1].text = str(val)

        # Properties
        props = self.snapshot.material_props_dict
        if props:
            self.doc.add_heading('1.4 Material Properties', level=2)
            table = self.doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'
            hdr = table.rows[0].cells
            hdr[0].text = 'Property'
            hdr[1].text = 'Type'
            hdr[2].text = 'Units'

            for p in props:
                row = table.add_row().cells
                row[0].text = p.get('name', '')
                row[1].text = p.get('type', '')
                row[2].text = p.get('units', '') or '-'

    def _add_config_section(self):
        self.doc.add_heading('2. Simulation Configuration', level=1)

        # Geometry
        self.doc.add_heading('2.1 Geometry', level=2)
        self.doc.add_paragraph(f'Type: {self.snapshot.geometry_type}')
        geo = self.snapshot.geometry_dict
        if geo:
            table = self.doc.add_table(rows=1, cols=2)
            table.style = 'Light Grid Accent 1'
            hdr = table.rows[0].cells
            hdr[0].text = 'Parameter'
            hdr[1].text = 'Value'
            for key, val in geo.items():
                row = table.add_row().cells
                row[0].text = key
                row[1].text = f'{val:.4f} m' if isinstance(val, (int, float)) else str(val)

        if self.snapshot.cad_filename:
            self.doc.add_paragraph(f'CAD File: {self.snapshot.cad_filename}')

        # Heat Treatment
        ht = self.snapshot.ht_config
        if ht:
            self.doc.add_heading('2.2 Heat Treatment', level=2)
            for phase_name in ['heating', 'transfer', 'quenching', 'tempering']:
                phase = ht.get(phase_name, {})
                if phase.get('enabled', False):
                    self.doc.add_heading(f'{phase_name.title()}', level=3)
                    table = self.doc.add_table(rows=1, cols=2)
                    table.style = 'Light Grid Accent 1'
                    hdr = table.rows[0].cells
                    hdr[0].text = 'Parameter'
                    hdr[1].text = 'Value'
                    for key, val in phase.items():
                        if key != 'enabled':
                            row = table.add_row().cells
                            row[0].text = key.replace('_', ' ').title()
                            row[1].text = str(val)

        # Solver
        solver = self.snapshot.solver_dict
        if solver:
            self.doc.add_heading('2.3 Solver Settings', level=2)
            table = self.doc.add_table(rows=1, cols=2)
            table.style = 'Light Grid Accent 1'
            hdr = table.rows[0].cells
            hdr[0].text = 'Setting'
            hdr[1].text = 'Value'
            for key, val in solver.items():
                row = table.add_row().cells
                row[0].text = key.replace('_', ' ').title()
                row[1].text = str(val)

    def _add_results_section(self):
        self.doc.add_heading('3. Results Summary', level=1)

        table = self.doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        hdr = table.rows[0].cells
        hdr[0].text = 'Metric'
        hdr[1].text = 'Value'

        metrics = [
            ('Status', self.snapshot.status or 'Unknown'),
            ('Duration', f'{self.snapshot.duration_seconds:.1f} s' if self.snapshot.duration_seconds else 'N/A'),
            ('t8/5', f'{self.snapshot.t_800_500:.2f} s' if self.snapshot.t_800_500 else 'N/A'),
            ('Surface Hardness', f'{self.snapshot.predicted_hardness_surface:.0f} HV' if self.snapshot.predicted_hardness_surface else 'N/A'),
            ('Center Hardness', f'{self.snapshot.predicted_hardness_center:.0f} HV' if self.snapshot.predicted_hardness_center else 'N/A'),
        ]

        for label, val in metrics:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = val

    def _add_lineage_section(self):
        self.doc.add_heading('4. Lineage Metadata', level=1)

        table = self.doc.add_table(rows=1, cols=2)
        table.style = 'Light Grid Accent 1'
        hdr = table.rows[0].cells
        hdr[0].text = 'Field'
        hdr[1].text = 'Value'

        entries = [
            ('Simulation ID', str(self.sim.id)),
            ('Simulation Name', self.sim.name),
            ('Snapshot Version', str(self.snapshot.version)),
            ('Started At', self.snapshot.started_at.strftime('%Y-%m-%d %H:%M:%S') if self.snapshot.started_at else 'N/A'),
            ('Completed At', self.snapshot.completed_at.strftime('%Y-%m-%d %H:%M:%S') if self.snapshot.completed_at else 'N/A'),
        ]

        for label, val in entries:
            row = table.add_row().cells
            row[0].text = label
            row[1].text = val

        # Drift check
        from .lineage_service import LineageService
        drifts = LineageService.check_drift(self.snapshot)
        if drifts:
            self.doc.add_heading('4.1 Material Drift Detected', level=2)
            self.doc.add_paragraph(
                'The following material data has changed since this simulation was run:'
            )
            table = self.doc.add_table(rows=1, cols=3)
            table.style = 'Light Grid Accent 1'
            hdr = table.rows[0].cells
            hdr[0].text = 'Field'
            hdr[1].text = 'At Run Time'
            hdr[2].text = 'Current Value'

            for d in drifts:
                row = table.add_row().cells
                row[0].text = d['field']
                row[1].text = str(d['snapshot_value']) if d['snapshot_value'] is not None else '-'
                row[2].text = str(d['current_value']) if d['current_value'] is not None else '-'
        else:
            self.doc.add_paragraph('No material drift detected. Frozen data matches current database.')

    def _add_change_log_section(self):
        self.doc.add_heading('5. Material Change Log', level=1)

        from app.models import MaterialChangeLog

        # Get changes for this steel grade since the previous snapshot
        grade_id = self.sim.steel_grade_id
        changes = MaterialChangeLog.query.filter(
            MaterialChangeLog.steel_grade_id == grade_id,
        ).order_by(MaterialChangeLog.changed_at.desc()).limit(50).all()

        if changes:
            table = self.doc.add_table(rows=1, cols=5)
            table.style = 'Light Grid Accent 1'
            hdr = table.rows[0].cells
            hdr[0].text = 'Date'
            hdr[1].text = 'User'
            hdr[2].text = 'Action'
            hdr[3].text = 'Entity'
            hdr[4].text = 'Field'

            for c in changes:
                row = table.add_row().cells
                row[0].text = c.changed_at.strftime('%Y-%m-%d %H:%M') if c.changed_at else ''
                row[1].text = c.changed_by_username or 'system'
                row[2].text = c.action
                row[3].text = c.entity_type
                row[4].text = c.field_name or '-'
        else:
            self.doc.add_paragraph('No material changes recorded for this steel grade.')
