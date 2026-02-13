"""Validation report generator — compares simulation vs measured data.

Generates Word (.docx) and PDF reports documenting simulation accuracy.
"""
from io import BytesIO
from datetime import datetime
from typing import Optional
import tempfile
import os

import numpy as np
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from fpdf import FPDF

from app.services.comparison_service import ComparisonService
from app.services import visualization


class ValidationReportGenerator:
    """Generate Word validation reports for simulation vs measured data."""

    def __init__(self, simulation):
        self.sim = simulation
        self.doc = Document()

    def generate_report(self) -> bytes:
        self._add_header()
        self._add_simulation_summary()
        self._add_measured_data_summary()
        self._add_comparison_metrics()
        self._add_overlay_plots()
        self._add_interpretation()
        self._add_conclusions()

        output = BytesIO()
        self.doc.save(output)
        output.seek(0)
        return output.getvalue()

    def _add_header(self):
        title = self.doc.add_heading('Simulation Validation Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        self.doc.add_heading(self.sim.name, level=1)
        if self.sim.description:
            self.doc.add_paragraph(self.sim.description)

        table = self.doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'
        cells = [
            ('Steel Grade', self.sim.steel_grade.designation if self.sim.steel_grade else 'N/A'),
            ('Geometry', self.sim.geometry_label),
            ('Status', self.sim.status.title()),
            ('Report Date', datetime.utcnow().strftime('%Y-%m-%d %H:%M UTC')),
        ]
        for i, (label, value) in enumerate(cells):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = str(value)
            for run in table.rows[i].cells[0].paragraphs[0].runs:
                run.bold = True
        self.doc.add_paragraph()

    def _add_simulation_summary(self):
        self.doc.add_heading('Simulation Configuration', level=2)
        ht = self.sim.ht_config

        rows = []
        if ht.get('heating', {}).get('enabled'):
            h = ht['heating']
            rows.append(('Austenitizing', f"{h.get('target_temperature', 850):.0f} °C, hold {h.get('hold_time', 0):.0f} min"))
        if ht.get('transfer', {}).get('enabled'):
            rows.append(('Transfer', f"{ht['transfer'].get('duration', 0):.0f} s"))
        q = ht.get('quenching', {})
        if q:
            rows.append(('Quenching', f"{q.get('media', 'water').title()}, {q.get('agitation', 'none')}, {q.get('media_temperature', 25):.0f} °C"))
        if ht.get('tempering', {}).get('enabled'):
            t = ht['tempering']
            rows.append(('Tempering', f"{t.get('temperature', 550):.0f} °C, hold {t.get('hold_time', 60):.0f} min"))

        # Geometry
        geom = self.sim.geometry_dict
        rows.append(('Geometry Type', self.sim.geometry_label))

        if rows:
            table = self.doc.add_table(rows=len(rows), cols=2)
            table.style = 'Table Grid'
            for i, (label, value) in enumerate(rows):
                table.rows[i].cells[0].text = label
                table.rows[i].cells[1].text = value
                for run in table.rows[i].cells[0].paragraphs[0].runs:
                    run.bold = True

        # Key results
        cycle = self.sim.results.filter_by(result_type='full_cycle').first()
        if cycle and cycle.t_800_500:
            self.doc.add_paragraph(f't₈₋₅ = {cycle.t_800_500:.1f} s')
        self.doc.add_paragraph()

    def _add_measured_data_summary(self):
        self.doc.add_heading('Measured Data', level=2)

        measured_list = self.sim.measured_data.all()
        if not measured_list:
            self.doc.add_paragraph('No measured data uploaded.')
            return

        table = self.doc.add_table(rows=len(measured_list) + 1, cols=5)
        table.style = 'Table Grid'
        headers = ['Name', 'Process Step', 'Channels', 'Points', 'Duration (s)']
        for j, h in enumerate(headers):
            table.rows[0].cells[j].text = h
            for run in table.rows[0].cells[j].paragraphs[0].runs:
                run.bold = True

        for i, md in enumerate(measured_list, 1):
            table.rows[i].cells[0].text = md.name or ''
            table.rows[i].cells[1].text = (md.process_step or 'full').title()
            table.rows[i].cells[2].text = str(md.num_channels)
            table.rows[i].cells[3].text = str(md.num_points)
            table.rows[i].cells[4].text = f'{md.duration_seconds:.0f}' if md.duration_seconds else '-'
        self.doc.add_paragraph()

    def _add_comparison_metrics(self):
        self.doc.add_heading('Comparison Metrics', level=2)

        metrics = ComparisonService.compare_simulation(self.sim)
        if not metrics:
            self.doc.add_paragraph('Comparison could not be computed.')
            return

        # Overall metrics table
        table = self.doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'
        overall = [
            ('RMS Error (°C)', f'{metrics.rms_error:.1f}'),
            ('R²', f'{metrics.r_squared:.4f}'),
            ('Peak Temp Diff (°C)', f'{metrics.peak_temp_diff:.1f}'),
            ('Max |Error| (°C)', f'{metrics.max_abs_error:.1f}'),
            ('Time Offset (s)', f'{metrics.time_offset:.1f}'),
            ('Overall Rating', metrics.rating.title()),
        ]
        for i, (label, value) in enumerate(overall):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value
            for run in table.rows[i].cells[0].paragraphs[0].runs:
                run.bold = True
        self.doc.add_paragraph()

        # Phase-by-phase metrics
        if metrics.phase_metrics:
            self.doc.add_heading('Per-Phase Metrics', level=3)
            pm_table = self.doc.add_table(rows=len(metrics.phase_metrics) + 1, cols=4)
            pm_table.style = 'Table Grid'
            pm_headers = ['Phase', 'RMS (°C)', 'R²', 'Rating']
            for j, h in enumerate(pm_headers):
                pm_table.rows[0].cells[j].text = h
                for run in pm_table.rows[0].cells[j].paragraphs[0].runs:
                    run.bold = True

            for i, (phase, pm) in enumerate(metrics.phase_metrics.items(), 1):
                pm_table.rows[i].cells[0].text = phase.title()
                pm_table.rows[i].cells[1].text = f'{pm.get("rms", 0):.1f}'
                pm_table.rows[i].cells[2].text = f'{pm.get("r_squared", 0):.4f}'
                pm_table.rows[i].cells[3].text = pm.get('rating', 'unknown').title()
            self.doc.add_paragraph()

    def _add_overlay_plots(self):
        self.doc.add_heading('Overlay Plots', level=2)

        cycle = self.sim.results.filter_by(result_type='full_cycle').first()
        if not cycle:
            return

        sim_times = np.array(cycle.time_array)
        sim_temps = np.array(cycle.data_dict.get('center', cycle.value_array))

        measured_list = self.sim.measured_data.all()
        for md in measured_list:
            for ch in md.available_channels:
                meas_times = np.array(md.get_channel_times(ch))
                meas_temps = np.array(md.get_channel_data(ch))
                if len(meas_times) < 2:
                    continue

                plot_bytes = visualization.create_comparison_plot(
                    sim_times=sim_times,
                    sim_temps=sim_temps,
                    measured_data=[{
                        'name': f'{md.name} - {ch}',
                        'times': meas_times,
                        'temps': meas_temps,
                    }],
                    title=f'Sim vs {ch} ({md.process_step or "full"})'
                )
                self._add_plot(plot_bytes)

    def _add_interpretation(self):
        self.doc.add_heading('Engineering Interpretation', level=2)

        metrics = ComparisonService.compare_simulation(self.sim)
        if not metrics:
            self.doc.add_paragraph('No metrics available for interpretation.')
            return

        if metrics.rating == 'good':
            text = (
                'The simulation shows good agreement with measured data. '
                f'The RMS error of {metrics.rms_error:.1f}°C and R² of {metrics.r_squared:.4f} '
                'indicate that the thermal model captures the heat treatment process accurately. '
                'The simulation can be used with confidence for process optimization.'
            )
        elif metrics.rating == 'acceptable':
            text = (
                'The simulation shows acceptable agreement with measured data. '
                f'The RMS error of {metrics.rms_error:.1f}°C suggests moderate deviations. '
                'Consider reviewing boundary conditions, material properties, or thermocouple placement. '
                'The simulation may be used for trend analysis but quantitative predictions '
                'should be treated with caution.'
            )
        else:
            text = (
                'The simulation shows poor agreement with measured data. '
                f'The RMS error of {metrics.rms_error:.1f}°C indicates significant deviations. '
                'Recommendations: (1) verify material thermal properties, '
                '(2) check heat transfer coefficients and quench media conditions, '
                '(3) review thermocouple placement and contact quality, '
                '(4) consider geometry simplification effects.'
            )
        self.doc.add_paragraph(text)

    def _add_conclusions(self):
        self.doc.add_heading('Conclusions', level=2)

        metrics = ComparisonService.compare_simulation(self.sim)
        if not metrics:
            return

        bullets = [
            f'Overall validation rating: {metrics.rating.upper()}',
            f'RMS temperature error: {metrics.rms_error:.1f}°C',
            f'Correlation coefficient R² = {metrics.r_squared:.4f}',
        ]
        if metrics.time_offset:
            bullets.append(f'Detected time offset: {metrics.time_offset:.1f}s (auto-corrected in analysis)')
        if metrics.phase_metrics:
            best_phase = min(metrics.phase_metrics.items(), key=lambda x: x[1].get('rms', 999))
            worst_phase = max(metrics.phase_metrics.items(), key=lambda x: x[1].get('rms', 0))
            bullets.append(f'Best phase match: {best_phase[0].title()} (RMS {best_phase[1].get("rms", 0):.1f}°C)')
            if best_phase[0] != worst_phase[0]:
                bullets.append(f'Worst phase match: {worst_phase[0].title()} (RMS {worst_phase[1].get("rms", 0):.1f}°C)')

        for b in bullets:
            self.doc.add_paragraph(b, style='List Bullet')

    def _add_plot(self, plot_bytes: bytes, width: float = 6.0):
        self.doc.add_picture(BytesIO(plot_bytes), width=Inches(width))
        last_p = self.doc.paragraphs[-1]
        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        self.doc.add_paragraph()


class ValidationPDFReportGenerator:
    """Generate PDF validation reports using fpdf2."""

    def __init__(self, simulation):
        self.sim = simulation
        self.pdf = FPDF()
        self.pdf.set_auto_page_break(auto=True, margin=15)
        self.temp_files = []

    def generate_report(self) -> bytes:
        try:
            self.pdf.add_page()
            self._add_header()
            self._add_simulation_summary()
            self._add_measured_summary()
            self._add_comparison_metrics()
            self._add_overlay_plots()
            self._add_interpretation()
            self._add_conclusions()
            return bytes(self.pdf.output())
        finally:
            for f in self.temp_files:
                try:
                    os.unlink(f)
                except OSError:
                    pass

    def _add_header(self):
        self.pdf.set_font('Helvetica', 'B', 22)
        self.pdf.set_text_color(26, 82, 118)
        self.pdf.cell(0, 15, 'Simulation Validation Report', align='C', new_x='LMARGIN', new_y='NEXT')
        self.pdf.set_draw_color(26, 82, 118)
        self.pdf.set_line_width(0.5)
        self.pdf.line(10, self.pdf.get_y(), 200, self.pdf.get_y())
        self.pdf.ln(5)

        self.pdf.set_font('Helvetica', 'B', 16)
        self.pdf.set_text_color(0, 0, 0)
        self.pdf.cell(0, 10, self.sim.name, new_x='LMARGIN', new_y='NEXT')
        self.pdf.set_font('Helvetica', '', 10)
        self.pdf.set_text_color(100, 100, 100)
        grade = self.sim.steel_grade.designation if self.sim.steel_grade else 'N/A'
        self.pdf.cell(0, 6, f'Steel: {grade}  |  Geometry: {self.sim.geometry_label}  |  Date: {datetime.utcnow().strftime("%Y-%m-%d")}',
                      new_x='LMARGIN', new_y='NEXT')
        self.pdf.ln(5)

    def _add_simulation_summary(self):
        self.pdf.set_font('Helvetica', 'B', 14)
        self.pdf.set_text_color(26, 82, 118)
        self.pdf.cell(0, 8, 'Simulation Configuration', new_x='LMARGIN', new_y='NEXT')
        self.pdf.set_text_color(0, 0, 0)
        self.pdf.set_font('Helvetica', '', 10)

        ht = self.sim.ht_config
        lines = []
        if ht.get('heating', {}).get('enabled'):
            h = ht['heating']
            lines.append(f"Austenitizing: {h.get('target_temperature', 850):.0f} C, hold {h.get('hold_time', 0):.0f} min")
        q = ht.get('quenching', {})
        if q:
            lines.append(f"Quench: {q.get('media', 'water').title()}, {q.get('agitation', 'none')}")
        if ht.get('tempering', {}).get('enabled'):
            t = ht['tempering']
            lines.append(f"Tempering: {t.get('temperature', 550):.0f} C, hold {t.get('hold_time', 60):.0f} min")
        for line in lines:
            self.pdf.cell(0, 5, line, new_x='LMARGIN', new_y='NEXT')
        self.pdf.ln(3)

    def _add_measured_summary(self):
        self.pdf.set_font('Helvetica', 'B', 14)
        self.pdf.set_text_color(26, 82, 118)
        self.pdf.cell(0, 8, 'Measured Data', new_x='LMARGIN', new_y='NEXT')
        self.pdf.set_text_color(0, 0, 0)
        self.pdf.set_font('Helvetica', '', 10)

        measured_list = self.sim.measured_data.all()
        if not measured_list:
            self.pdf.cell(0, 5, 'No measured data uploaded.', new_x='LMARGIN', new_y='NEXT')
            return

        for md in measured_list:
            dur = f'{md.duration_seconds:.0f}s' if md.duration_seconds else '-'
            self.pdf.cell(0, 5, f'{md.name} | {(md.process_step or "full").title()} | {md.num_channels} ch, {md.num_points} pts, {dur}',
                          new_x='LMARGIN', new_y='NEXT')
        self.pdf.ln(3)

    def _add_comparison_metrics(self):
        self.pdf.set_font('Helvetica', 'B', 14)
        self.pdf.set_text_color(26, 82, 118)
        self.pdf.cell(0, 8, 'Comparison Metrics', new_x='LMARGIN', new_y='NEXT')
        self.pdf.set_text_color(0, 0, 0)

        metrics = ComparisonService.compare_simulation(self.sim)
        if not metrics:
            self.pdf.set_font('Helvetica', '', 10)
            self.pdf.cell(0, 5, 'Comparison not available.', new_x='LMARGIN', new_y='NEXT')
            return

        # Metrics table
        self.pdf.set_font('Helvetica', 'B', 10)
        col_w = 47.5
        self.pdf.cell(col_w, 6, 'RMS (C)', border=1, align='C')
        self.pdf.cell(col_w, 6, 'R2', border=1, align='C')
        self.pdf.cell(col_w, 6, 'Peak Diff (C)', border=1, align='C')
        self.pdf.cell(col_w, 6, 'Rating', border=1, align='C', new_x='LMARGIN', new_y='NEXT')

        self.pdf.set_font('Helvetica', '', 10)
        self.pdf.cell(col_w, 6, f'{metrics.rms_error:.1f}', border=1, align='C')
        self.pdf.cell(col_w, 6, f'{metrics.r_squared:.4f}', border=1, align='C')
        self.pdf.cell(col_w, 6, f'{metrics.peak_temp_diff:.1f}', border=1, align='C')
        self.pdf.cell(col_w, 6, metrics.rating.title(), border=1, align='C', new_x='LMARGIN', new_y='NEXT')
        self.pdf.ln(5)

    def _add_overlay_plots(self):
        self.pdf.set_font('Helvetica', 'B', 14)
        self.pdf.set_text_color(26, 82, 118)
        self.pdf.cell(0, 8, 'Overlay Plots', new_x='LMARGIN', new_y='NEXT')
        self.pdf.set_text_color(0, 0, 0)

        cycle = self.sim.results.filter_by(result_type='full_cycle').first()
        if not cycle:
            return

        sim_times = np.array(cycle.time_array)
        sim_temps = np.array(cycle.data_dict.get('center', cycle.value_array))

        measured_list = self.sim.measured_data.all()
        for md in measured_list:
            for ch in md.available_channels:
                meas_times = np.array(md.get_channel_times(ch))
                meas_temps = np.array(md.get_channel_data(ch))
                if len(meas_times) < 2:
                    continue

                plot_bytes = visualization.create_comparison_plot(
                    sim_times=sim_times, sim_temps=sim_temps,
                    measured_data=[{'name': f'{md.name} - {ch}', 'times': meas_times, 'temps': meas_temps}],
                    title=f'Sim vs {ch} ({md.process_step or "full"})'
                )
                self._add_plot(plot_bytes, f'{ch} comparison')

    def _add_interpretation(self):
        self.pdf.set_font('Helvetica', 'B', 14)
        self.pdf.set_text_color(26, 82, 118)
        self.pdf.cell(0, 8, 'Engineering Interpretation', new_x='LMARGIN', new_y='NEXT')
        self.pdf.set_text_color(0, 0, 0)
        self.pdf.set_font('Helvetica', '', 10)

        metrics = ComparisonService.compare_simulation(self.sim)
        if not metrics:
            return

        if metrics.rating == 'good':
            text = (f'The simulation shows good agreement (RMS {metrics.rms_error:.1f}C, '
                    f'R2 {metrics.r_squared:.4f}). The model can be used with confidence.')
        elif metrics.rating == 'acceptable':
            text = (f'Acceptable agreement (RMS {metrics.rms_error:.1f}C). '
                    'Review boundary conditions and material properties for improvement.')
        else:
            text = (f'Poor agreement (RMS {metrics.rms_error:.1f}C). '
                    'Verify material properties, HTC values, and thermocouple placement.')
        self.pdf.multi_cell(0, 5, text)
        self.pdf.ln(3)

    def _add_conclusions(self):
        self.pdf.set_font('Helvetica', 'B', 14)
        self.pdf.set_text_color(26, 82, 118)
        self.pdf.cell(0, 8, 'Conclusions', new_x='LMARGIN', new_y='NEXT')
        self.pdf.set_text_color(0, 0, 0)
        self.pdf.set_font('Helvetica', '', 10)

        metrics = ComparisonService.compare_simulation(self.sim)
        if not metrics:
            return

        bullets = [
            f'Overall rating: {metrics.rating.upper()}',
            f'RMS error: {metrics.rms_error:.1f}C',
            f'R2 = {metrics.r_squared:.4f}',
        ]
        for b in bullets:
            self.pdf.cell(5, 5, '-')
            self.pdf.cell(0, 5, b, new_x='LMARGIN', new_y='NEXT')

    def _add_plot(self, plot_bytes: bytes, title: str = None):
        with tempfile.NamedTemporaryFile(suffix='.png', delete=False) as f:
            f.write(plot_bytes)
            temp_path = f.name
        self.temp_files.append(temp_path)

        if self.pdf.get_y() > 180:
            self.pdf.add_page()
        if title:
            self.pdf.set_font('Helvetica', 'I', 9)
            self.pdf.cell(0, 5, title, new_x='LMARGIN', new_y='NEXT')
        self.pdf.image(temp_path, w=180)
        self.pdf.ln(5)


def generate_validation_report(simulation) -> bytes:
    """Convenience function to generate Word validation report."""
    return ValidationReportGenerator(simulation).generate_report()


def generate_validation_pdf_report(simulation) -> bytes:
    """Convenience function to generate PDF validation report."""
    return ValidationPDFReportGenerator(simulation).generate_report()
