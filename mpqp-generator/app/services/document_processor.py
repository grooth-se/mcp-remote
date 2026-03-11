"""Document text extraction for PDF, Word, and Excel files."""
import os
import logging

logger = logging.getLogger(__name__)

SUPPORTED_EXTENSIONS = {'.pdf', '.docx', '.doc', '.xlsx', '.xls'}


def extract_text(file_path):
    """Extract text from a document file.

    Returns dict with 'text', 'page_count', 'format', and 'metadata'.
    """
    ext = os.path.splitext(file_path)[1].lower()

    if ext == '.pdf':
        return _extract_pdf(file_path)
    elif ext in ('.docx', '.doc'):
        return _extract_word(file_path)
    elif ext in ('.xlsx', '.xls'):
        return _extract_excel(file_path)
    else:
        return {'text': '', 'page_count': 0, 'format': ext, 'metadata': {}, 'error': f'Unsupported format: {ext}'}


def _extract_pdf_worker(file_path):
    """Worker function for PDF extraction — runs in a subprocess."""
    import fitz  # PyMuPDF

    doc = fitz.open(file_path)
    pages = [page.get_text() for page in doc]
    text = '\n\n'.join(pages)
    metadata = {
        'title': doc.metadata.get('title', ''),
        'author': doc.metadata.get('author', ''),
        'subject': doc.metadata.get('subject', ''),
    }
    page_count = len(doc)
    doc.close()
    return {
        'text': text,
        'page_count': page_count,
        'format': 'PDF',
        'metadata': metadata,
    }


def _pdf_subprocess_target(file_path, result_queue):
    """Subprocess target for PDF extraction — must be module-level for pickling."""
    try:
        result = _extract_pdf_worker(file_path)
        result_queue.put(result)
    except Exception as e:
        result_queue.put({'text': '', 'page_count': 0, 'format': 'PDF', 'metadata': {}, 'error': str(e)})


def _extract_pdf(file_path):
    """Extract text from PDF using PyMuPDF.

    Uses a subprocess with timeout to protect against corrupted PDFs that
    can hang or segfault PyMuPDF (a C extension). signal.SIGALRM does not
    work from background threads, so we use multiprocessing instead.
    """
    from multiprocessing import Process, Queue

    result_queue = Queue()
    proc = Process(target=_pdf_subprocess_target, args=(file_path, result_queue))
    proc.start()
    proc.join(timeout=30)

    if proc.is_alive():
        logger.warning(f'PDF extraction timed out for {file_path}, killing subprocess')
        proc.kill()
        proc.join()
        return {'text': '', 'page_count': 0, 'format': 'PDF', 'metadata': {}, 'error': 'extraction timed out'}

    if proc.exitcode != 0:
        logger.error(f'PDF extraction crashed for {file_path} (exit code {proc.exitcode})')
        return {'text': '', 'page_count': 0, 'format': 'PDF', 'metadata': {}, 'error': f'subprocess crashed (exit code {proc.exitcode})'}

    try:
        return result_queue.get_nowait()
    except Exception:
        return {'text': '', 'page_count': 0, 'format': 'PDF', 'metadata': {}, 'error': 'no result from subprocess'}


def _extract_word(file_path):
    """Extract text from Word documents using python-docx."""
    try:
        from docx import Document

        doc = Document(file_path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]

        # Also extract text from tables
        for table in doc.tables:
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    paragraphs.append(row_text)

        text = '\n'.join(paragraphs)
        metadata = {
            'title': doc.core_properties.title or '',
            'author': doc.core_properties.author or '',
        }

        return {
            'text': text,
            'page_count': 0,  # python-docx doesn't expose page count
            'format': 'DOCX',
            'metadata': metadata,
        }
    except ImportError:
        logger.error('python-docx not installed')
        return {'text': '', 'page_count': 0, 'format': 'DOCX', 'metadata': {}, 'error': 'python-docx not installed'}
    except Exception as e:
        logger.error(f'Word extraction failed for {file_path}: {e}')
        return {'text': '', 'page_count': 0, 'format': 'DOCX', 'metadata': {}, 'error': str(e)}


def _extract_excel(file_path):
    """Extract text from Excel files using openpyxl."""
    try:
        from openpyxl import load_workbook

        wb = load_workbook(file_path, data_only=True)
        sheets = []
        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            rows = []
            for row in ws.iter_rows(values_only=True):
                cells = [str(c) if c is not None else '' for c in row]
                row_text = ' | '.join(c for c in cells if c)
                if row_text:
                    rows.append(row_text)
            if rows:
                sheets.append(f'[Sheet: {sheet_name}]\n' + '\n'.join(rows))

        text = '\n\n'.join(sheets)
        wb.close()

        return {
            'text': text,
            'page_count': len(wb.sheetnames),
            'format': 'XLSX',
            'metadata': {'sheets': wb.sheetnames},
        }
    except ImportError:
        logger.error('openpyxl not installed')
        return {'text': '', 'page_count': 0, 'format': 'XLSX', 'metadata': {}, 'error': 'openpyxl not installed'}
    except Exception as e:
        logger.error(f'Excel extraction failed for {file_path}: {e}')
        return {'text': '', 'page_count': 0, 'format': 'XLSX', 'metadata': {}, 'error': str(e)}


def get_file_info(file_path):
    """Get basic file information."""
    ext = os.path.splitext(file_path)[1].lower()
    size = os.path.getsize(file_path) if os.path.exists(file_path) else 0
    return {
        'name': os.path.basename(file_path),
        'extension': ext,
        'size': size,
        'size_human': _human_size(size),
        'supported': ext in SUPPORTED_EXTENSIONS,
    }


def _human_size(size_bytes):
    for unit in ('B', 'KB', 'MB', 'GB'):
        if size_bytes < 1024:
            return f'{size_bytes:.1f} {unit}'
        size_bytes /= 1024
    return f'{size_bytes:.1f} TB'
