"""Template filler for generating Word (.docx) documents from LLM output.

Takes the generated markdown content and produces a formatted Word document
using python-docx — either from a blank template or filling a user-provided one.
"""
import os
import re
import logging
from datetime import datetime

from flask import current_app

logger = logging.getLogger(__name__)


def create_word_document(content, metadata, output_path=None):
    """Create a Word document from generated markdown content.

    Args:
        content: Generated text (markdown-formatted with ## headers)
        metadata: Dict with project_name, customer, product_type, doc_type
        output_path: Optional output path (auto-generated if None)

    Returns:
        Path to the created .docx file, or None on failure.
    """
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError:
        logger.error('python-docx not installed')
        return None

    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Title page
    doc.add_paragraph('')  # Spacer
    title = doc.add_heading(metadata.get('doc_type', 'MPQP'), level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = doc.add_heading(metadata.get('project_name', 'Document'), level=1)
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph('')
    info_lines = [
        f"Customer: {metadata.get('customer', 'N/A')}",
        f"Product Type: {metadata.get('product_type', 'N/A')}",
        f"Date: {datetime.utcnow().strftime('%Y-%m-%d')}",
        f"Status: DRAFT — For Review",
    ]
    for line in info_lines:
        p = doc.add_paragraph(line)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_page_break()

    # Parse markdown content into sections
    sections = _parse_markdown_sections(content)

    for section in sections:
        level = section.get('level', 2)
        heading_level = min(level, 4)  # docx supports up to heading 9
        doc.add_heading(section['title'], level=heading_level)

        # Add body paragraphs
        for para_text in section['paragraphs']:
            if para_text.startswith('- ') or para_text.startswith('* '):
                # Bullet point
                doc.add_paragraph(para_text[2:], style='List Bullet')
            elif re.match(r'^\d+\.\s', para_text):
                # Numbered list
                doc.add_paragraph(re.sub(r'^\d+\.\s', '', para_text), style='List Number')
            elif para_text.startswith('|'):
                # Table row — skip for now (complex to parse)
                doc.add_paragraph(para_text, style='Normal')
            else:
                doc.add_paragraph(para_text)

    # Save
    if not output_path:
        generated_dir = current_app.config['GENERATED_FOLDER']
        os.makedirs(generated_dir, exist_ok=True)
        safe_name = ''.join(c if c.isalnum() or c in '.-_' else '_'
                            for c in metadata.get('project_name', 'document'))
        output_path = os.path.join(
            generated_dir,
            f"{metadata.get('doc_type', 'MPQP')}_{safe_name}.docx"
        )

    doc.save(output_path)
    logger.info(f'Word document created: {output_path}')
    return output_path


def fill_template(template_path, content, metadata, output_path=None):
    """Fill an existing Word template with generated content.

    Looks for placeholder patterns like {{SECTION_NAME}} in the template
    and replaces them with generated content.

    Args:
        template_path: Path to the .docx template
        content: Generated text (markdown)
        metadata: Dict with project info
        output_path: Output path for the filled document

    Returns:
        Path to the output file, or None on failure.
    """
    try:
        from docx import Document
    except ImportError:
        logger.error('python-docx not installed')
        return None

    if not os.path.exists(template_path):
        logger.error(f'Template not found: {template_path}')
        return None

    doc = Document(template_path)

    # Replace metadata placeholders
    replacements = {
        '{{PROJECT_NAME}}': metadata.get('project_name', ''),
        '{{CUSTOMER}}': metadata.get('customer', ''),
        '{{PRODUCT_TYPE}}': metadata.get('product_type', ''),
        '{{DATE}}': datetime.utcnow().strftime('%Y-%m-%d'),
        '{{DOC_TYPE}}': metadata.get('doc_type', 'MPQP'),
    }

    # Parse sections from generated content
    sections = _parse_markdown_sections(content)
    for section in sections:
        key = '{{' + section['title'].upper().replace(' ', '_') + '}}'
        replacements[key] = '\n'.join(section['paragraphs'])

    # Apply replacements to all paragraphs
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            if key in paragraph.text:
                paragraph.text = paragraph.text.replace(key, value)

    # Also check tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for key, value in replacements.items():
                    if key in cell.text:
                        cell.text = cell.text.replace(key, value)

    # Save
    if not output_path:
        generated_dir = current_app.config['GENERATED_FOLDER']
        os.makedirs(generated_dir, exist_ok=True)
        safe_name = ''.join(c if c.isalnum() or c in '.-_' else '_'
                            for c in metadata.get('project_name', 'document'))
        output_path = os.path.join(
            generated_dir,
            f"{metadata.get('doc_type', 'MPQP')}_{safe_name}_filled.docx"
        )

    doc.save(output_path)
    logger.info(f'Template filled: {output_path}')
    return output_path


def _parse_markdown_sections(content):
    """Parse markdown content into sections with headers and paragraphs."""
    sections = []
    current_section = None

    for line in content.split('\n'):
        stripped = line.strip()

        # Check for markdown headers
        header_match = re.match(r'^(#{1,4})\s+(.+)$', stripped)
        if header_match:
            if current_section:
                sections.append(current_section)
            level = len(header_match.group(1))
            current_section = {
                'title': header_match.group(2).strip(),
                'level': level,
                'paragraphs': [],
            }
        elif current_section is not None:
            if stripped:
                current_section['paragraphs'].append(stripped)
        elif stripped:
            # Content before first header
            if not sections and not current_section:
                current_section = {'title': 'Introduction', 'level': 2, 'paragraphs': []}
            if current_section:
                current_section['paragraphs'].append(stripped)

    if current_section:
        sections.append(current_section)

    return sections
