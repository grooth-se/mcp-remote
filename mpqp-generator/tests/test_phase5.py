"""Phase 5 tests: document generation, template filling, Word export."""
import os
import tempfile
import re


# --- Template filler / markdown parser tests ---

def test_parse_markdown_sections():
    from app.services.template_filler import _parse_markdown_sections

    md = """## Scope and Purpose
This section covers the scope.
- Item one
- Item two

## Material Specifications
Inconel 625 per ASTM B443.
1. First requirement
2. Second requirement

### Sub-section
Details here.
"""
    sections = _parse_markdown_sections(md)
    assert len(sections) == 3
    assert sections[0]['title'] == 'Scope and Purpose'
    assert sections[0]['level'] == 2
    assert '- Item one' in sections[0]['paragraphs']
    assert sections[1]['title'] == 'Material Specifications'
    assert sections[2]['title'] == 'Sub-section'
    assert sections[2]['level'] == 3


def test_parse_markdown_no_header():
    from app.services.template_filler import _parse_markdown_sections

    md = "Some text before any header.\nMore text."
    sections = _parse_markdown_sections(md)
    assert len(sections) == 1
    assert sections[0]['title'] == 'Introduction'


def test_create_word_document(app):
    from app.services.template_filler import create_word_document

    content = """## Scope
This document covers manufacturing.

## Materials
- Inconel 625
- F22 steel
"""
    metadata = {
        'project_name': 'Test Project',
        'customer': 'Aker Solutions',
        'product_type': 'TTR',
        'doc_type': 'MPQP',
    }

    with tempfile.TemporaryDirectory() as tmpdir:
        output = os.path.join(tmpdir, 'test_output.docx')
        result = create_word_document(content, metadata, output_path=output)
        assert result is not None
        assert os.path.exists(result)
        assert result.endswith('.docx')
        assert os.path.getsize(result) > 0


def test_create_word_document_auto_path(app):
    from app.services.template_filler import create_word_document

    content = "## Section One\nContent here."
    metadata = {'project_name': 'AutoPath', 'doc_type': 'ITP'}

    result = create_word_document(content, metadata)
    assert result is not None
    assert os.path.exists(result)
    assert 'ITP' in os.path.basename(result)
    # Clean up
    os.remove(result)


def test_fill_template(app):
    """Test filling a Word template with placeholders."""
    try:
        from docx import Document
    except ImportError:
        import pytest
        pytest.skip('python-docx not installed')

    from app.services.template_filler import fill_template

    # Create a simple template with placeholders
    with tempfile.TemporaryDirectory() as tmpdir:
        template_path = os.path.join(tmpdir, 'template.docx')
        doc = Document()
        doc.add_paragraph('Project: {{PROJECT_NAME}}')
        doc.add_paragraph('Customer: {{CUSTOMER}}')
        doc.add_paragraph('Date: {{DATE}}')
        doc.save(template_path)

        content = "## Scope\nTest scope content."
        metadata = {'project_name': 'Filled Project', 'customer': 'TechnipFMC'}
        output = os.path.join(tmpdir, 'filled.docx')

        result = fill_template(template_path, content, metadata, output_path=output)
        assert result is not None
        assert os.path.exists(result)

        # Verify content was replaced
        filled_doc = Document(result)
        texts = [p.text for p in filled_doc.paragraphs]
        assert any('Filled Project' in t for t in texts)
        assert any('TechnipFMC' in t for t in texts)


# --- Document generator tests ---

def test_format_requirements(app, db):
    from app.services.document_generator import _format_requirements
    from app.models.generation import GenerationJob

    job = GenerationJob(
        new_project_name='Test',
        extracted_requirements={
            'customer_name': 'Aker',
            'materials': ['Inconel 625', 'F22'],
            'standards': ['API 6A', 'NORSOK M-630'],
            'testing_requirements': ['UT', 'MPI'],
        },
    )
    db.session.add(job)
    db.session.commit()

    result = _format_requirements(job)
    assert 'Aker' in result
    assert 'Inconel 625' in result
    assert 'API 6A' in result
    assert 'UT' in result


def test_format_requirements_empty(app, db):
    from app.services.document_generator import _format_requirements
    from app.models.generation import GenerationJob

    job = GenerationJob(new_project_name='Empty')
    db.session.add(job)
    db.session.commit()

    result = _format_requirements(job)
    assert 'No specific requirements' in result


def test_get_template_structure_default(app, db):
    from app.services.document_generator import _get_template_structure
    from app.models.generation import GenerationJob

    job = GenerationJob(new_project_name='Test')
    db.session.add(job)
    db.session.commit()

    structure = _get_template_structure(job)
    assert 'Scope' in structure
    assert 'Material' in structure


def test_get_doc_type_default(app, db):
    from app.services.document_generator import _get_doc_type
    from app.models.generation import GenerationJob

    job = GenerationJob(new_project_name='Test')
    db.session.add(job)
    db.session.commit()

    assert _get_doc_type(job) == 'MPQP'


def test_save_generated_content(app, db):
    from app.services.document_generator import _save_generated_content
    from app.models.generation import GenerationJob, DocumentVersion

    job = GenerationJob(new_project_name='SaveTest', current_version=1)
    db.session.add(job)
    db.session.commit()

    result = _save_generated_content(job, '## Scope\nTest content here.', 'MPQP')
    assert result['filepath'] is not None
    assert os.path.exists(result['filepath'])
    assert result['version'] == 1
    assert result['content_length'] > 0

    # Check version record
    versions = DocumentVersion.query.filter_by(generation_job_id=job.id).all()
    assert len(versions) == 1
    assert versions[0].version_number == 1

    # Clean up
    os.remove(result['filepath'])


def test_gather_reference_text_empty(app, db):
    from app.services.document_generator import _gather_reference_text
    from app.models.generation import GenerationJob

    job = GenerationJob(new_project_name='NoRefs', selected_references=[])
    db.session.add(job)
    db.session.commit()

    text = _gather_reference_text(job)
    assert 'No reference' in text


def test_gather_reference_text_with_docs(app, db):
    from app.services.document_generator import _gather_reference_text
    from app.models.generation import GenerationJob
    from app.models.project import Project
    from app.models.document import Document

    p = Project(project_number='P-REF', folder_path='/tmp')
    db.session.add(p)
    db.session.flush()

    doc = Document(
        project_id=p.id,
        file_name='MPQP_test.pdf',
        file_path='/tmp/MPQP_test.pdf',
        document_type='MPQP',
        extracted_text='This is the reference MPQP content with scope and materials.',
    )
    db.session.add(doc)
    db.session.flush()

    job = GenerationJob(new_project_name='WithRefs', selected_references=[p.id])
    db.session.add(job)
    db.session.commit()

    text = _gather_reference_text(job)
    assert 'reference MPQP content' in text


# --- Route tests ---

def test_generate_doc_route_no_job(app, db, logged_in_client):
    from app.models.generation import GenerationJob
    job = GenerationJob(new_project_name='Will redirect', status='pending')
    db.session.add(job)
    db.session.commit()

    csrf = _get_csrf(logged_in_client, f'/generate/status/{job.id}')
    resp = logged_in_client.post('/generate/generate-document/9999',
                                 data={'csrf_token': csrf},
                                 follow_redirects=True)
    assert resp.status_code == 200


def test_download_no_document(app, db, logged_in_client):
    from app.models.generation import GenerationJob
    job = GenerationJob(new_project_name='No Doc', status='pending')
    db.session.add(job)
    db.session.commit()

    resp = logged_in_client.get(f'/generate/download/{job.id}', follow_redirects=True)
    assert resp.status_code == 200
    assert b'No generated document' in resp.data


def test_download_with_document(app, db, logged_in_client):
    from app.models.generation import GenerationJob

    with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False) as f:
        f.write('# Test Document\n\nContent here.')
        filepath = f.name

    try:
        job = GenerationJob(
            new_project_name='Download Test',
            status='completed',
            generated_document_path=filepath,
        )
        db.session.add(job)
        db.session.commit()

        resp = logged_in_client.get(f'/generate/download/{job.id}')
        assert resp.status_code == 200
        assert b'Test Document' in resp.data
    finally:
        os.unlink(filepath)


def test_export_word_no_document(app, db, logged_in_client):
    from app.models.generation import GenerationJob
    job = GenerationJob(new_project_name='No Doc', status='pending')
    db.session.add(job)
    db.session.commit()

    resp = logged_in_client.post(
        f'/generate/export-word/{job.id}',
        data={'csrf_token': _get_csrf(logged_in_client, f'/generate/status/{job.id}')},
        follow_redirects=True,
    )
    assert resp.status_code == 200
    assert b'No generated document' in resp.data


def test_export_word_success(app, db, logged_in_client):
    from app.models.generation import GenerationJob

    with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False) as f:
        f.write('# MPQP - Test\n\n---\n\n## Scope\nScope content.\n\n## Materials\n- Inconel 625\n')
        filepath = f.name

    try:
        job = GenerationJob(
            new_project_name='Word Export Test',
            status='completed',
            generated_document_path=filepath,
        )
        db.session.add(job)
        db.session.commit()

        resp = logged_in_client.post(
            f'/generate/export-word/{job.id}',
            data={'csrf_token': _get_csrf(logged_in_client, f'/generate/status/{job.id}')},
        )
        assert resp.status_code == 200
        # Should return a .docx file
        assert 'application/' in resp.content_type or 'octet-stream' in resp.content_type
    finally:
        os.unlink(filepath)


def test_status_page_shows_generate_button(app, db, logged_in_client):
    from app.models.generation import GenerationJob
    job = GenerationJob(
        new_project_name='Gen Button Test',
        status='review',
        selected_references=[1],
    )
    db.session.add(job)
    db.session.commit()

    resp = logged_in_client.get(f'/generate/status/{job.id}')
    assert resp.status_code == 200
    assert b'Generate Document' in resp.data
    assert b'generate-document' in resp.data


def test_status_page_completed_shows_download(app, db, logged_in_client):
    from app.models.generation import GenerationJob

    with tempfile.NamedTemporaryFile(mode='w', suffix='.md', delete=False) as f:
        f.write('# Test\nContent')
        filepath = f.name

    try:
        job = GenerationJob(
            new_project_name='Completed Test',
            status='completed',
            generated_document_path=filepath,
        )
        db.session.add(job)
        db.session.commit()

        resp = logged_in_client.get(f'/generate/status/{job.id}')
        assert resp.status_code == 200
        assert b'Download Markdown' in resp.data
        assert b'Export Word' in resp.data
    finally:
        os.unlink(filepath)


# --- Helper ---

def _get_csrf(client, url='/generate/jobs'):
    resp = client.get(url)
    data = resp.data.decode()
    match = re.search(r'name="csrf_token" value="([^"]+)"', data)
    return match.group(1) if match else ''
