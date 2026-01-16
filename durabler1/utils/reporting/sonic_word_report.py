"""
Word report generator for Sonic Resonance (E1875) test results.

Populates a Word template with test data and results.
"""

import re
from pathlib import Path
from typing import Dict, Any, Optional
from docx import Document
from docx.shared import Inches, Cm


class SonicReportGenerator:
    """
    Generate Sonic Resonance test reports from Word template.

    Uses placeholder syntax {{placeholder_name}} in template.
    """

    def __init__(self, template_path: Path):
        """
        Initialize report generator.

        Parameters
        ----------
        template_path : Path
            Path to Word template file
        """
        self.template_path = template_path
        if not template_path.exists():
            raise FileNotFoundError(f"Template not found: {template_path}")

    def generate_report(
        self,
        output_path: Path,
        data: Dict[str, Any],
        chart_path: Optional[Path] = None,
        logo_path: Optional[Path] = None
    ) -> Path:
        """
        Generate report by populating template with data.

        Parameters
        ----------
        output_path : Path
            Path for output Word document
        data : Dict[str, Any]
            Dictionary of placeholder values
        chart_path : Path, optional
            Path to velocity chart image
        logo_path : Path, optional
            Path to logo image

        Returns
        -------
        Path
            Path to generated report
        """
        doc = Document(self.template_path)

        # Replace placeholders in paragraphs
        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph, data, chart_path, logo_path)

        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, data, chart_path, logo_path)

        doc.save(output_path)
        return output_path

    def _replace_in_paragraph(
        self,
        paragraph,
        data: Dict[str, Any],
        chart_path: Optional[Path],
        logo_path: Optional[Path]
    ):
        """Replace placeholders in a paragraph."""
        full_text = paragraph.text

        # Handle logo placeholder
        if '{{logo}}' in full_text and logo_path and logo_path.exists():
            paragraph.clear()
            run = paragraph.add_run()
            run.add_picture(str(logo_path), height=Cm(1.5))
            return

        # Handle chart placeholder
        if '{{chart}}' in full_text and chart_path and chart_path.exists():
            paragraph.clear()
            run = paragraph.add_run()
            run.add_picture(str(chart_path), width=Inches(5.5))
            return

        # Find all placeholders
        placeholders = re.findall(r'\{\{([^}]+)\}\}', full_text)

        if not placeholders:
            return

        # Replace each placeholder
        for placeholder in placeholders:
            key = placeholder
            value = data.get(key, '')

            if value is None:
                value = ''

            # Format numeric values
            if isinstance(value, float):
                if abs(value) < 0.001:
                    value = f"{value:.4g}"
                elif abs(value) < 1:
                    value = f"{value:.4f}"
                else:
                    value = f"{value:.2f}"
            elif isinstance(value, int):
                value = str(value)

            full_text = full_text.replace(f'{{{{{placeholder}}}}}', str(value))

        # Update paragraph text
        if paragraph.runs:
            first_run = paragraph.runs[0]
            font_name = first_run.font.name
            font_size = first_run.font.size
            bold = first_run.bold
            italic = first_run.italic

            paragraph.clear()
            new_run = paragraph.add_run(full_text)
            new_run.font.name = font_name
            new_run.font.size = font_size
            new_run.bold = bold
            new_run.italic = italic
        else:
            paragraph.text = full_text

    @staticmethod
    def prepare_report_data(
        test_info: Dict[str, str],
        specimen_data: Dict[str, Any],
        velocity_data: Dict[str, Any],
        results: Any
    ) -> Dict[str, Any]:
        """
        Prepare data dictionary for report generation.

        Parameters
        ----------
        test_info : Dict[str, str]
            Test information (project, customer, etc.)
        specimen_data : Dict[str, Any]
            Specimen geometry data
        velocity_data : Dict[str, Any]
            Velocity measurements
        results : SonicResults
            Analysis results

        Returns
        -------
        Dict[str, Any]
            Flattened dictionary ready for template
        """
        data = {}

        # Test information
        data['test_project'] = test_info.get('test_project', '')
        data['customer'] = test_info.get('customer', '')
        data['customer_order'] = test_info.get('customer_order', '')
        data['product_sn'] = test_info.get('product_sn', '')
        data['specimen_id'] = test_info.get('specimen_id', '')
        data['location_orientation'] = test_info.get('location_orientation', '')
        data['material'] = test_info.get('material', '')
        data['certificate_number'] = test_info.get('certificate_number', '')
        data['test_date'] = test_info.get('test_date', '')
        data['temperature'] = test_info.get('temperature', '23')
        data['test_standard'] = 'Modified ASTM E1875'
        data['test_equipment'] = 'Ultrasonic Tester'

        # Specimen geometry
        data['specimen_type'] = specimen_data.get('specimen_type', '')
        data['diameter'] = specimen_data.get('diameter', '-')
        data['side_length'] = specimen_data.get('side_length', '-')
        data['length'] = specimen_data.get('length', '')
        data['mass'] = specimen_data.get('mass', '')

        # Velocity measurements
        data['vl1'] = velocity_data.get('vl1', '-')
        data['vl2'] = velocity_data.get('vl2', '-')
        data['vl3'] = velocity_data.get('vl3', '-')
        data['vs1'] = velocity_data.get('vs1', '-')
        data['vs2'] = velocity_data.get('vs2', '-')
        data['vs3'] = velocity_data.get('vs3', '-')

        # Results
        if results:
            data['density'] = f"{results.density.value:.1f}"
            data['vl_avg'] = f"{results.longitudinal_velocity.value:.1f}"
            data['vs_avg'] = f"{results.shear_velocity.value:.1f}"

            data['poissons_ratio'] = f"{results.poissons_ratio.value:.4f}"
            data['poissons_ratio_unc'] = f"±{results.poissons_ratio.uncertainty:.4f}"

            data['shear_modulus'] = f"{results.shear_modulus.value:.2f}"
            data['shear_modulus_unc'] = f"±{results.shear_modulus.uncertainty:.2f}"

            data['youngs_modulus'] = f"{results.youngs_modulus.value:.2f}"
            data['youngs_modulus_unc'] = f"±{results.youngs_modulus.uncertainty:.2f}"

            data['flexural_frequency'] = f"{results.flexural_frequency.value:.1f}"
            data['flexural_frequency_unc'] = f"±{results.flexural_frequency.uncertainty:.1f}"

            data['torsional_frequency'] = f"{results.torsional_frequency.value:.1f}"
            data['torsional_frequency_unc'] = f"±{results.torsional_frequency.uncertainty:.1f}"

            data['validity_status'] = 'VALID' if results.is_valid else 'CHECK'
            data['validity_notes'] = results.validity_notes
        else:
            data['density'] = '-'
            data['vl_avg'] = '-'
            data['vs_avg'] = '-'
            data['poissons_ratio'] = '-'
            data['poissons_ratio_unc'] = '-'
            data['shear_modulus'] = '-'
            data['shear_modulus_unc'] = '-'
            data['youngs_modulus'] = '-'
            data['youngs_modulus_unc'] = '-'
            data['flexural_frequency'] = '-'
            data['flexural_frequency_unc'] = '-'
            data['torsional_frequency'] = '-'
            data['torsional_frequency_unc'] = '-'
            data['validity_status'] = '-'
            data['validity_notes'] = ''

        # Signatures
        data['tested_by'] = ''
        data['tested_date'] = ''
        data['reviewed_by'] = ''
        data['reviewed_date'] = ''
        data['approved_by'] = ''
        data['approved_date'] = ''

        return data
