"""
KIC Fracture Toughness Report Generator.

Generates Word document reports for KIC (ASTM E399) fracture toughness tests.
Reports include test information, specimen dimensions, material properties,
results with uncertainties, validity assessment, and force-displacement plot.
"""

from dataclasses import dataclass
from pathlib import Path
from typing import Optional, Dict, Any
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. Report generation will be limited.")


class KICReportGenerator:
    """
    Generator for KIC fracture toughness test reports.

    Creates Word documents with all test information, results,
    and validity assessment per ASTM E399.

    Parameters
    ----------
    template_path : Path, optional
        Path to Word template file. If None, creates report from scratch.
    """

    def __init__(self, template_path: Optional[Path] = None):
        self.template_path = template_path

        if not DOCX_AVAILABLE:
            print("Warning: python-docx not available. Install with: pip install python-docx")

    def generate_report(self,
                        output_path: Path,
                        test_info: Dict[str, str],
                        dimensions: Dict[str, str],
                        material_props: Dict[str, str],
                        results: Any,
                        chart_path: Optional[Path] = None,
                        logo_path: Optional[Path] = None) -> Path:
        """
        Generate KIC test report.

        Parameters
        ----------
        output_path : Path
            Output file path for the Word document
        test_info : Dict[str, str]
            Test information (certificate, project, customer, etc.)
        dimensions : Dict[str, str]
            Specimen dimensions (W, B, a_0, S, etc.)
        material_props : Dict[str, str]
            Material properties (yield strength, E, nu)
        results : KICResult
            Analysis results from KICAnalyzer
        chart_path : Path, optional
            Path to chart image to include
        logo_path : Path, optional
            Path to company logo

        Returns
        -------
        Path
            Path to generated report file
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for report generation")

        # Create document from template or new
        if self.template_path and self.template_path.exists():
            doc = Document(self.template_path)
            self._fill_template(doc, test_info, dimensions, material_props, results,
                               chart_path, logo_path)
        else:
            doc = self._create_report_from_scratch(
                test_info, dimensions, material_props, results, chart_path, logo_path
            )

        # Save document
        doc.save(output_path)
        return output_path

    def _create_report_from_scratch(self,
                                     test_info: Dict[str, str],
                                     dimensions: Dict[str, str],
                                     material_props: Dict[str, str],
                                     results: Any,
                                     chart_path: Optional[Path],
                                     logo_path: Optional[Path]) -> Document:
        """Create report without template."""
        doc = Document()

        # Set up styles
        style = doc.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)

        # Header with logo
        if logo_path and logo_path.exists():
            doc.add_picture(str(logo_path), width=Inches(2))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Title
        title = doc.add_heading('KIC Fracture Toughness Test Report', level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Subtitle
        subtitle = doc.add_paragraph('ASTM E399 - Plane-Strain Fracture Toughness')
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Test Information table
        doc.add_heading('Test Information', level=1)
        table = doc.add_table(rows=7, cols=2)
        table.style = 'Table Grid'

        info_data = [
            ('Certificate Number:', test_info.get('certificate_number', '')),
            ('Test Project:', test_info.get('test_project', '')),
            ('Customer:', test_info.get('customer', '')),
            ('Specimen ID:', test_info.get('specimen_id', '')),
            ('Material:', test_info.get('material', '')),
            ('Test Date:', test_info.get('test_date', '')),
            ('Temperature:', f"{test_info.get('temperature', '23')} C"),
        ]

        for i, (label, value) in enumerate(info_data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value
            # Bold the label
            table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

        doc.add_paragraph()

        # Specimen Dimensions table
        doc.add_heading('Specimen Dimensions', level=1)
        table = doc.add_table(rows=6, cols=3)
        table.style = 'Table Grid'

        # Header row
        headers = ['Parameter', 'Value', 'Unit']
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

        dim_data = [
            ('Specimen Type', dimensions.get('specimen_type', ''), '-'),
            ('Width W', dimensions.get('W', ''), 'mm'),
            ('Thickness B', dimensions.get('B', ''), 'mm'),
            ('Crack length a_0', dimensions.get('a_0', ''), 'mm'),
            ('Span S', dimensions.get('S', '-'), 'mm'),
        ]

        for i, (param, value, unit) in enumerate(dim_data):
            table.rows[i+1].cells[0].text = param
            table.rows[i+1].cells[1].text = str(value)
            table.rows[i+1].cells[2].text = unit

        doc.add_paragraph()

        # Material Properties table
        doc.add_heading('Material Properties', level=1)
        table = doc.add_table(rows=4, cols=3)
        table.style = 'Table Grid'

        # Header row
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

        mat_data = [
            ('Yield strength sigma_ys', material_props.get('yield_strength', ''), 'MPa'),
            ("Young's modulus E", material_props.get('youngs_modulus', ''), 'GPa'),
            ("Poisson's ratio nu", material_props.get('poissons_ratio', ''), '-'),
        ]

        for i, (param, value, unit) in enumerate(mat_data):
            table.rows[i+1].cells[0].text = param
            table.rows[i+1].cells[1].text = str(value)
            table.rows[i+1].cells[2].text = unit

        doc.add_paragraph()

        # Results table
        doc.add_heading('Test Results', level=1)
        table = doc.add_table(rows=7, cols=4)
        table.style = 'Table Grid'

        # Header row
        result_headers = ['Parameter', 'Value', 'U (k=2)', 'Unit']
        for i, header in enumerate(result_headers):
            table.rows[0].cells[i].text = header
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

        if results:
            results_data = [
                ('P_max', f'{results.P_max.value:.2f}', f'+/-{results.P_max.uncertainty:.2f}', 'kN'),
                ('P_Q (5% secant)', f'{results.P_Q.value:.2f}', f'+/-{results.P_Q.uncertainty:.2f}', 'kN'),
                ('P_max/P_Q ratio', f'{results.P_ratio:.3f}', '', '-'),
                ('K_Q (conditional)', f'{results.K_Q.value:.2f}', f'+/-{results.K_Q.uncertainty:.2f}', 'MPa*sqrt(m)'),
                ('K_IC', f'{results.K_IC.value:.2f}' if results.K_IC else 'CONDITIONAL',
                 f'+/-{results.K_IC.uncertainty:.2f}' if results.K_IC else '', 'MPa*sqrt(m)'),
                ('Compliance', f'{results.compliance:.4f}', '', 'mm/kN'),
            ]

            for i, (param, value, unc, unit) in enumerate(results_data):
                table.rows[i+1].cells[0].text = param
                table.rows[i+1].cells[1].text = value
                table.rows[i+1].cells[2].text = unc
                table.rows[i+1].cells[3].text = unit

        doc.add_paragraph()

        # Force-Displacement Plot
        if chart_path and chart_path.exists():
            doc.add_heading('Force vs Displacement', level=1)
            doc.add_picture(str(chart_path), width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()

        # Validity Assessment
        doc.add_heading('Validity Assessment per ASTM E399', level=1)

        validity_status = "VALID" if results and results.is_valid else "CONDITIONAL"
        status_para = doc.add_paragraph()
        status_run = status_para.add_run(f"Overall Status: {validity_status}")
        status_run.bold = True
        if results and results.is_valid:
            status_run.font.color.rgb = None  # Use default (black)
        else:
            # Can't easily set color without more complex code, so just make it bold
            pass

        doc.add_paragraph()

        if results and results.validity_notes:
            for note in results.validity_notes:
                p = doc.add_paragraph(note, style='List Bullet')

        doc.add_paragraph()

        # Signatures
        doc.add_heading('Approval', level=1)
        sig_table = doc.add_table(rows=3, cols=3)
        sig_table.style = 'Table Grid'

        sig_headers = ['Role', 'Name', 'Date']
        for i, header in enumerate(sig_headers):
            sig_table.rows[0].cells[i].text = header
            sig_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

        sig_roles = ['Tested by:', 'Reviewed by:', 'Approved by:']
        for i, role in enumerate(sig_roles):
            sig_table.rows[i+1].cells[0].text = role if i < 2 else ''

        # Can only have 3 rows, adjust
        sig_table.rows[1].cells[0].text = 'Tested by:'
        sig_table.rows[2].cells[0].text = 'Reviewed by:'

        doc.add_paragraph()

        # Footer
        footer_para = doc.add_paragraph()
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        footer_run = footer_para.add_run(
            f"Report generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
        )
        footer_run.font.size = Pt(9)
        footer_run.font.color.rgb = None

        return doc

    def _fill_template(self,
                       doc: Document,
                       test_info: Dict[str, str],
                       dimensions: Dict[str, str],
                       material_props: Dict[str, str],
                       results: Any,
                       chart_path: Optional[Path],
                       logo_path: Optional[Path]):
        """Fill in template placeholders."""
        # Build replacement dictionary
        replacements = {
            '{{certificate_number}}': test_info.get('certificate_number', ''),
            '{{test_project}}': test_info.get('test_project', ''),
            '{{customer}}': test_info.get('customer', ''),
            '{{specimen_id}}': test_info.get('specimen_id', ''),
            '{{material}}': test_info.get('material', ''),
            '{{test_date}}': test_info.get('test_date', ''),
            '{{temperature}}': test_info.get('temperature', '23'),
            '{{specimen_type}}': dimensions.get('specimen_type', ''),
            '{{W}}': dimensions.get('W', ''),
            '{{B}}': dimensions.get('B', ''),
            '{{a_0}}': dimensions.get('a_0', ''),
            '{{S}}': dimensions.get('S', '-'),
            '{{B_n}}': dimensions.get('B_n', '-'),
            '{{yield_strength}}': material_props.get('yield_strength', ''),
            '{{youngs_modulus}}': material_props.get('youngs_modulus', ''),
            '{{poissons_ratio}}': material_props.get('poissons_ratio', ''),
        }

        if results:
            replacements.update({
                '{{P_max}}': f'{results.P_max.value:.2f}',
                '{{P_max_uncertainty}}': f'{results.P_max.uncertainty:.2f}',
                '{{P_Q}}': f'{results.P_Q.value:.2f}',
                '{{P_Q_uncertainty}}': f'{results.P_Q.uncertainty:.2f}',
                '{{P_ratio}}': f'{results.P_ratio:.3f}',
                '{{K_Q}}': f'{results.K_Q.value:.2f}',
                '{{K_Q_uncertainty}}': f'{results.K_Q.uncertainty:.2f}',
                '{{K_IC}}': f'{results.K_IC.value:.2f}' if results.K_IC else 'CONDITIONAL',
                '{{K_IC_uncertainty}}': f'{results.K_IC.uncertainty:.2f}' if results.K_IC else '-',
                '{{compliance}}': f'{results.compliance:.4f}',
                '{{is_valid}}': 'VALID' if results.is_valid else 'CONDITIONAL',
                '{{validity_notes}}': '\n'.join(results.validity_notes) if results.validity_notes else '',
            })

        # Replace in paragraphs
        for para in doc.paragraphs:
            for key, value in replacements.items():
                if key in para.text:
                    para.text = para.text.replace(key, str(value))

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for key, value in replacements.items():
                            if key in para.text:
                                para.text = para.text.replace(key, str(value))

        # Insert chart if placeholder exists
        if chart_path and chart_path.exists():
            for para in doc.paragraphs:
                if '{{chart}}' in para.text:
                    para.text = ''
                    run = para.add_run()
                    run.add_picture(str(chart_path), width=Inches(5.5))
                    break

        # Insert logo if placeholder exists
        if logo_path and logo_path.exists():
            for para in doc.paragraphs:
                if '{{logo}}' in para.text:
                    para.text = ''
                    run = para.add_run()
                    run.add_picture(str(logo_path), width=Inches(2))
                    break
