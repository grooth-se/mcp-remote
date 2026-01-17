"""
Word report generator for FCGR E647 test results.

Populates a Word template with test data, Paris law results, plots, and photos.
"""

import re
from pathlib import Path
from typing import Dict, Any, Optional, List
from docx import Document
from docx.shared import Inches, Cm


class FCGRReportGenerator:
    """
    Generate FCGR test reports from Word template.

    Uses placeholder syntax {{placeholder_name}} in template.
    """

    def __init__(self, template_path: Path):
        """
        Initialize report generator.

        Parameters
        ----------
        template_path : Path
            Path to Word template file
        """
        self.template_path = template_path
        if not template_path.exists():
            raise FileNotFoundError(f"Template not found: {template_path}")

    def generate_report(
        self,
        output_path: Path,
        data: Dict[str, Any],
        plot1_path: Optional[Path] = None,
        plot2_path: Optional[Path] = None,
        logo_path: Optional[Path] = None,
        photo_paths: Optional[List[Path]] = None
    ) -> Path:
        """
        Generate report by populating template with data.

        Parameters
        ----------
        output_path : Path
            Path for output Word document
        data : Dict[str, Any]
            Dictionary of placeholder values
        plot1_path : Path, optional
            Path to Crack Length vs Cycles plot image
        plot2_path : Path, optional
            Path to da/dN vs Delta-K (Paris law) plot image
        logo_path : Path, optional
            Path to logo image to insert
        photo_paths : List[Path], optional
            List of paths to crack surface photos

        Returns
        -------
        Path
            Path to generated report
        """
        doc = Document(self.template_path)

        # Replace placeholders in paragraphs
        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph, data, plot1_path, plot2_path, logo_path, photo_paths)

        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, data, plot1_path, plot2_path, logo_path, photo_paths)

        doc.save(output_path)
        return output_path

    def _replace_in_paragraph(
        self,
        paragraph,
        data: Dict[str, Any],
        plot1_path: Optional[Path],
        plot2_path: Optional[Path],
        logo_path: Optional[Path],
        photo_paths: Optional[List[Path]]
    ):
        """Replace placeholders in a paragraph."""
        full_text = paragraph.text

        # Handle logo placeholder
        if '{{logo}}' in full_text and logo_path and logo_path.exists():
            paragraph.clear()
            run = paragraph.add_run()
            run.add_picture(str(logo_path), height=Cm(1.5))
            return

        # Handle plot1 placeholder (Crack Length vs Cycles)
        if '{{plot1}}' in full_text and plot1_path and plot1_path.exists():
            paragraph.clear()
            run = paragraph.add_run()
            run.add_picture(str(plot1_path), width=Inches(3.2))
            return

        # Handle plot2 placeholder (da/dN vs Delta-K)
        if '{{plot2}}' in full_text and plot2_path and plot2_path.exists():
            paragraph.clear()
            run = paragraph.add_run()
            run.add_picture(str(plot2_path), width=Inches(3.2))
            return

        # Handle combined plots placeholder (both plots side by side)
        if '{{plots}}' in full_text:
            paragraph.clear()
            if plot1_path and plot1_path.exists():
                run = paragraph.add_run()
                run.add_picture(str(plot1_path), width=Inches(3.0))
            if plot2_path and plot2_path.exists():
                run = paragraph.add_run("  ")  # Space between plots
                run = paragraph.add_run()
                run.add_picture(str(plot2_path), width=Inches(3.0))
            return

        # Handle photos placeholder
        if '{{photos}}' in full_text:
            paragraph.clear()
            if photo_paths:
                for i, photo_path in enumerate(photo_paths):
                    if photo_path and photo_path.exists():
                        run = paragraph.add_run()
                        run.add_picture(str(photo_path), width=Inches(2.5))
                        if i < len(photo_paths) - 1:
                            paragraph.add_run("  ")  # Space between photos
            else:
                paragraph.add_run("No crack surface photos attached.")
            return

        # Find all placeholders
        placeholders = re.findall(r'\{\{([^}]+)\}\}', full_text)

        if not placeholders:
            return

        # Replace each placeholder
        for placeholder in placeholders:
            key = placeholder
            value = data.get(key, '')

            # Handle None values
            if value is None:
                value = ''

            # Format numeric values
            if isinstance(value, float):
                if abs(value) < 1e-10:
                    value = "0"
                elif abs(value) < 0.0001 or abs(value) > 10000:
                    value = f"{value:.4e}"
                elif abs(value) < 1:
                    value = f"{value:.4f}"
                else:
                    value = f"{value:.2f}"
            elif isinstance(value, int):
                value = str(value)

            full_text = full_text.replace(f'{{{{{placeholder}}}}}', str(value))

        # Update paragraph text while preserving formatting
        if paragraph.runs:
            first_run = paragraph.runs[0]
            font_name = first_run.font.name
            font_size = first_run.font.size
            bold = first_run.bold
            italic = first_run.italic

            paragraph.clear()
            new_run = paragraph.add_run(full_text)
            new_run.font.name = font_name
            new_run.font.size = font_size
            new_run.bold = bold
            new_run.italic = italic
        else:
            paragraph.text = full_text

    @staticmethod
    def prepare_report_data(
        test_info: Dict[str, str],
        specimen_data: Dict[str, Any],
        material_data: Dict[str, Any],
        test_params: Dict[str, Any],
        results: Any,
        validity_notes: Optional[List[str]] = None
    ) -> Dict[str, Any]:
        """
        Prepare data dictionary for report generation.

        Parameters
        ----------
        test_info : Dict[str, str]
            Test information (project, customer, certificate, etc.)
        specimen_data : Dict[str, Any]
            Specimen geometry data
        material_data : Dict[str, Any]
            Material properties
        test_params : Dict[str, Any]
            Test parameters (R ratio, frequency, control mode, etc.)
        results : FCGRResult
            Analysis results with Paris law coefficients
        validity_notes : List[str], optional
            List of validity check notes

        Returns
        -------
        Dict[str, Any]
            Flattened dictionary ready for template
        """
        data = {}

        # Test information
        data['test_project'] = test_info.get('test_project', '')
        data['customer'] = test_info.get('customer', '')
        data['specimen_id'] = test_info.get('specimen_id', '')
        data['material'] = test_info.get('material', '')
        data['certificate_number'] = test_info.get('certificate_number', '')
        data['test_date'] = test_info.get('test_date', '')
        data['test_standard'] = 'ASTM E647'
        data['test_equipment'] = 'MTS Landmark 500kN'

        # Specimen geometry
        data['specimen_type'] = specimen_data.get('specimen_type', 'C(T)')
        data['W'] = specimen_data.get('W', '')
        data['B'] = specimen_data.get('B', '')
        data['B_n'] = specimen_data.get('B_n', specimen_data.get('B', ''))
        data['a_0'] = specimen_data.get('a_0', '')
        data['notch_height'] = specimen_data.get('notch_height', '0')

        # Calculate a_0/W ratio
        try:
            W = float(specimen_data.get('W', 1))
            a_0 = float(specimen_data.get('a_0', 0))
            data['a_W_ratio'] = f"{a_0 / W:.3f}" if W > 0 else '-'
        except (ValueError, TypeError):
            data['a_W_ratio'] = '-'

        # Side grooves
        B = specimen_data.get('B')
        B_n = specimen_data.get('B_n')
        if B and B_n:
            try:
                data['side_grooves'] = 'Yes' if float(B_n) < float(B) else 'No'
            except (ValueError, TypeError):
                data['side_grooves'] = 'No'
        else:
            data['side_grooves'] = 'No'

        # Material properties
        data['yield_strength'] = material_data.get('yield_strength', '')
        data['ultimate_strength'] = material_data.get('ultimate_strength', '')
        data['youngs_modulus'] = material_data.get('youngs_modulus', '')
        data['poissons_ratio'] = material_data.get('poissons_ratio', '0.3')

        # Test parameters
        data['control_mode'] = test_params.get('control_mode', 'Load Control')
        data['load_ratio'] = test_params.get('load_ratio', '0.1')
        data['frequency'] = test_params.get('frequency', '10')
        data['temperature'] = test_params.get('temperature', '23')
        data['P_max'] = test_params.get('P_max', '-')
        data['K_max'] = test_params.get('K_max', '-')
        data['wave_shape'] = test_params.get('wave_shape', 'Sine')
        data['environment'] = test_params.get('environment', 'Laboratory Air')

        # da/dN calculation method
        data['dadn_method'] = test_params.get('dadn_method', 'Secant')
        data['outlier_threshold'] = test_params.get('outlier_threshold', '30')

        # Results from FCGRResult object
        if results:
            # Paris law results (final - after outlier removal)
            paris = results.paris_law
            if paris:
                data['paris_C'] = paris.C
                data['paris_m'] = paris.m
                data['paris_r_squared'] = paris.r_squared
                data['paris_n_points'] = paris.n_points
                data['paris_C_error'] = paris.std_error_C
                data['paris_m_error'] = paris.std_error_m
                data['delta_K_min'] = paris.delta_K_range[0] if paris.delta_K_range else '-'
                data['delta_K_max'] = paris.delta_K_range[1] if paris.delta_K_range else '-'
                data['da_dN_min'] = paris.da_dN_range[0] if paris.da_dN_range else '-'
                data['da_dN_max'] = paris.da_dN_range[1] if paris.da_dN_range else '-'

            # Paris law results (initial - all data)
            paris_init = results.paris_law_initial
            if paris_init:
                data['paris_C_initial'] = paris_init.C
                data['paris_m_initial'] = paris_init.m
                data['paris_r_squared_initial'] = paris_init.r_squared
                data['paris_n_points_initial'] = paris_init.n_points

            # Data summary
            data['n_valid_points'] = results.n_valid_points
            data['n_outliers'] = results.n_outliers
            data['total_cycles'] = results.total_cycles
            data['final_crack_length'] = results.final_crack_length
            data['threshold_delta_K'] = results.threshold_delta_K if results.threshold_delta_K > 0 else '-'

            # Validity
            data['is_valid'] = 'VALID' if results.is_valid else 'See Notes'
            data['validity_notes'] = '\n'.join(results.validity_notes) if results.validity_notes else 'All checks passed'
        else:
            # Default values if no results
            data['paris_C'] = '-'
            data['paris_m'] = '-'
            data['paris_r_squared'] = '-'
            data['paris_n_points'] = '-'
            data['n_valid_points'] = '-'
            data['n_outliers'] = '-'
            data['total_cycles'] = '-'
            data['final_crack_length'] = '-'
            data['is_valid'] = '-'
            data['validity_notes'] = '-'

        # Signatures (to be filled manually)
        data['tested_by'] = ''
        data['tested_date'] = ''
        data['reviewed_by'] = ''
        data['reviewed_date'] = ''
        data['approved_by'] = ''
        data['approved_date'] = ''

        return data
