"""
Vickers Hardness Word Report Generator per ASTM E92.

This module provides the VickersReportGenerator class for creating
Word documents with Vickers hardness test results, plots, and photos.
"""

from pathlib import Path
from typing import Optional, Dict, Any
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from datetime import datetime

import sys
sys.path.insert(0, str(Path(__file__).parent.parent.parent))

from utils.analysis.vickers_calculations import VickersResult


class VickersReportGenerator:
    """
    Word report generator for Vickers hardness tests per ASTM E92.

    Parameters
    ----------
    template_path : Path
        Path to Word template file
    """

    def __init__(self, template_path: Path):
        """Initialize with template path."""
        self.template_path = template_path

    def generate_report(self,
                       output_path: Path,
                       test_info: Dict[str, Any],
                       results: VickersResult,
                       uncertainty_budget: Dict[str, Any],
                       chart_path: Optional[Path] = None,
                       photo_path: Optional[Path] = None,
                       logo_path: Optional[Path] = None) -> Path:
        """
        Generate Vickers hardness Word report.

        Parameters
        ----------
        output_path : Path
            Output file path
        test_info : dict
            Test information (certificate, customer, etc.)
        results : VickersResult
            Analysis results
        uncertainty_budget : dict
            Uncertainty budget breakdown
        chart_path : Path, optional
            Path to hardness profile chart
        photo_path : Path, optional
            Path to indent photograph
        logo_path : Path, optional
            Path to company logo

        Returns
        -------
        Path
            Path to generated report
        """
        doc = Document(self.template_path)

        # Prepare replacement data
        data = self._prepare_report_data(test_info, results, uncertainty_budget)

        # Replace placeholders in paragraphs
        for paragraph in doc.paragraphs:
            self._replace_placeholders(paragraph, data)

        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_placeholders(paragraph, data)

        # Insert logo
        if logo_path and logo_path.exists():
            self._insert_logo(doc, logo_path)

        # Insert chart
        if chart_path and chart_path.exists():
            self._insert_image(doc, chart_path, "{{chart}}", width=Inches(6))

        # Insert photo
        if photo_path and Path(photo_path).exists():
            self._insert_image(doc, Path(photo_path), "{{photo}}", width=Inches(4))

        # Save document
        doc.save(output_path)

        return output_path

    def _prepare_report_data(self,
                            test_info: Dict[str, Any],
                            results: VickersResult,
                            uncertainty_budget: Dict[str, Any]) -> Dict[str, str]:
        """Prepare data dictionary for placeholder replacement."""
        # Generate report number
        report_number = f"VH-{datetime.now().strftime('%Y%m%d-%H%M%S')}"

        # Build readings table text
        readings_text = ""
        for r in results.readings:
            readings_text += f"{r.reading_number}\t{r.location}\t{r.hardness_value:.1f}\n"

        data = {
            # Report info
            '{{report_number}}': report_number,
            '{{report_date}}': datetime.now().strftime("%Y-%m-%d"),

            # Test info
            '{{certificate_number}}': test_info.get('certificate_number', ''),
            '{{test_project}}': test_info.get('test_project', ''),
            '{{customer}}': test_info.get('customer', ''),
            '{{specimen_id}}': test_info.get('specimen_id', ''),
            '{{material}}': test_info.get('material', ''),
            '{{test_date}}': test_info.get('test_date', ''),
            '{{operator}}': test_info.get('operator', ''),
            '{{load_level}}': test_info.get('load_level', ''),

            # Results
            '{{mean_hardness}}': f"{results.mean_hardness.value:.1f}",
            '{{uncertainty}}': f"{results.mean_hardness.uncertainty:.1f}",
            '{{std_dev}}': f"{results.std_dev:.1f}",
            '{{range}}': f"{results.range_value:.1f}",
            '{{min_value}}': f"{results.min_value:.1f}",
            '{{max_value}}': f"{results.max_value:.1f}",
            '{{n_readings}}': str(results.n_readings),
            '{{unit}}': results.load_level,

            # Uncertainty budget
            '{{u_A}}': f"{uncertainty_budget.get('u_A', 0):.2f}",
            '{{u_machine}}': f"{uncertainty_budget.get('u_machine', 0):.2f}",
            '{{u_diagonal}}': f"{uncertainty_budget.get('u_diagonal', 0):.2f}",
            '{{u_force}}': f"{uncertainty_budget.get('u_force', 0):.2f}",
            '{{u_combined}}': f"{uncertainty_budget.get('u_combined', 0):.2f}",
            '{{U_expanded}}': f"{uncertainty_budget.get('U_expanded', 0):.2f}",
            '{{k}}': str(uncertainty_budget.get('k', 2)),

            # Readings
            '{{readings}}': readings_text,
        }

        return data

    def _replace_placeholders(self, paragraph, data: Dict[str, str]):
        """Replace placeholders in a paragraph."""
        for placeholder, value in data.items():
            if placeholder in paragraph.text:
                # Handle inline replacement
                for run in paragraph.runs:
                    if placeholder in run.text:
                        run.text = run.text.replace(placeholder, str(value))

                # Handle split across runs
                if placeholder in paragraph.text:
                    paragraph.text = paragraph.text.replace(placeholder, str(value))

    def _insert_logo(self, doc: Document, logo_path: Path):
        """Insert logo at {{logo}} placeholder."""
        # Check paragraphs first
        for paragraph in doc.paragraphs:
            if '{{logo}}' in paragraph.text:
                paragraph.clear()
                run = paragraph.add_run()
                run.add_picture(str(logo_path), width=Inches(1.5))
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                return

        # Check tables (logo is in header table)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if '{{logo}}' in paragraph.text:
                            paragraph.clear()
                            run = paragraph.add_run()
                            run.add_picture(str(logo_path), width=Inches(1.5))
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                            return

    def _insert_image(self, doc: Document, image_path: Path, placeholder: str, width=Inches(5)):
        """Insert image at specified placeholder."""
        for paragraph in doc.paragraphs:
            if placeholder in paragraph.text:
                paragraph.clear()
                run = paragraph.add_run()
                run.add_picture(str(image_path), width=width)
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                return

        # Check tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        if placeholder in paragraph.text:
                            paragraph.clear()
                            run = paragraph.add_run()
                            run.add_picture(str(image_path), width=width)
                            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            return
