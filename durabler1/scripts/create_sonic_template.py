"""
Script to create the Sonic Resonance (E1875) report Word template.
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def set_table_column_widths(table, widths_inches):
    """Set column widths for a table."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)

    for row in table.rows:
        for idx, width in enumerate(widths_inches):
            if idx < len(row.cells):
                cell = row.cells[idx]
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(int(width * 1440)))
                tcW.set(qn('w:type'), 'dxa')
                existing = tcPr.find(qn('w:tcW'))
                if existing is not None:
                    tcPr.remove(existing)
                tcPr.insert(0, tcW)


def add_page_number_field(paragraph):
    """Add a PAGE field to the paragraph."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def add_page_header(doc):
    """Add page header with Certificate No and page number on right side."""
    section = doc.sections[0]
    header = section.header

    # Create a table for header layout (logo left, certificate right)
    header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    header_table.autofit = False
    header_table.columns[0].width = Inches(2)
    header_table.columns[1].width = Inches(4.5)

    # Logo cell (left)
    logo_cell = header_table.rows[0].cells[0]
    logo_para = logo_cell.paragraphs[0]
    logo_para.add_run("{{logo}}")
    logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Certificate and page number cell (right)
    cert_cell = header_table.rows[0].cells[1]
    cert_para = cert_cell.paragraphs[0]
    cert_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    cert_run = cert_para.add_run("Certificate No: {{certificate_number}}")
    cert_run.font.size = Pt(10)
    cert_run.bold = True

    # Add page number on new line
    page_para = cert_cell.add_paragraph()
    page_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    page_run = page_para.add_run("Page ")
    page_run.font.size = Pt(9)
    add_page_number_field(page_para)


def add_page_footer(doc):
    """Add page footer with legal text."""
    section = doc.sections[0]
    footer = section.footer

    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    footer_text = (
        "All work and services carried out by Durabler are subject to, and conducted in accordance with, "
        "Durabler standard terms and conditions, which are available at durabler.se. This document shall "
        "not be reproduced other than in full, except with prior written approval of the issuer. The results "
        "pertain only to the item(s) as sampled by the client unless otherwise indicated. "
        "Durabler a part of Subseatec S AB, Address: Durabler C/O Subseatec, Dalavägen 23, 68130 Kristinehamn, SWEDEN"
    )

    footer_run = footer_para.add_run(footer_text)
    footer_run.font.size = Pt(7)
    footer_run.italic = True


def create_template():
    doc = Document()

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.header_distance = Cm(0.5)
        section.footer_distance = Cm(0.5)

    # Add page header and footer
    add_page_header(doc)
    add_page_footer(doc)

    # Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("SONIC RESONANCE TEST REPORT")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    std_para = doc.add_paragraph()
    std_run = std_para.add_run("Modified ASTM E1875 - Ultrasonic Method")
    std_run.font.size = Pt(11)
    std_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Test Information Table (no certificate number)
    info_heading = doc.add_paragraph()
    info_heading.add_run("Test Information").bold = True

    info_table = doc.add_table(rows=6, cols=4)
    info_table.style = 'Table Grid'

    info_data = [
        ("Test Project:", "{{test_project}}", "Customer:", "{{customer}}"),
        ("Customer Order:", "{{customer_order}}", "Test Date:", "{{test_date}}"),
        ("Product/S/N:", "{{product_sn}}", "Test Standard:", "{{test_standard}}"),
        ("Material:", "{{material}}", "Specimen ID:", "{{specimen_id}}"),
        ("Location/Orient.:", "{{location_orientation}}", "Test Temperature:", "{{temperature}} °C"),
        ("Test Equipment:", "{{test_equipment}}", "", ""),
    ]

    for i, row_data in enumerate(info_data):
        row = info_table.rows[i]
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = text
            if j % 2 == 0 and text:
                cell.paragraphs[0].runs[0].bold = True
                set_cell_shading(cell, 'E8E8E8')

    doc.add_paragraph()

    # Specimen Geometry Table
    dim_heading = doc.add_paragraph()
    dim_heading.add_run("Specimen Geometry").bold = True

    dim_table = doc.add_table(rows=2, cols=6)
    dim_table.style = 'Table Grid'

    dim_data = [
        ("Type:", "{{specimen_type}}", "Diameter (mm):", "{{diameter}}", "Side (mm):", "{{side_length}}"),
        ("Length (mm):", "{{length}}", "Mass (g):", "{{mass}}", "Density (kg/m³):", "{{density}}"),
    ]

    for i, row_data in enumerate(dim_data):
        row = dim_table.rows[i]
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = text
            if j % 2 == 0 and text:
                cell.paragraphs[0].runs[0].bold = True
                set_cell_shading(cell, 'E8E8E8')

    doc.add_paragraph()

    # Velocity Measurements Table
    vel_heading = doc.add_paragraph()
    vel_heading.add_run("Ultrasonic Velocity Measurements").bold = True

    vel_table = doc.add_table(rows=3, cols=5)
    vel_table.style = 'Table Grid'

    # Header row
    vel_headers = ["Wave Type", "V₁ (m/s)", "V₂ (m/s)", "V₃ (m/s)", "Average (m/s)"]
    for i, header in enumerate(vel_headers):
        cell = vel_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D0D0D0')
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Longitudinal row
    long_data = ["Longitudinal (Vl)", "{{vl1}}", "{{vl2}}", "{{vl3}}", "{{vl_avg}}"]
    for i, text in enumerate(long_data):
        cell = vel_table.rows[1].cells[i]
        cell.text = text
        if i == 0:
            set_cell_shading(cell, 'F0F0F0')

    # Shear row
    shear_data = ["Shear (Vs)", "{{vs1}}", "{{vs2}}", "{{vs3}}", "{{vs_avg}}"]
    for i, text in enumerate(shear_data):
        cell = vel_table.rows[2].cells[i]
        cell.text = text
        if i == 0:
            set_cell_shading(cell, 'F0F0F0')

    doc.add_paragraph()

    # Results Table
    results_heading = doc.add_paragraph()
    results_heading.add_run("Test Results").bold = True

    results_table = doc.add_table(rows=7, cols=5)
    results_table.style = 'Table Grid'

    # Header row
    res_headers = ["Parameter", "Value", "U (k=2)", "Requirement", "Unit"]
    for i, header in enumerate(res_headers):
        cell = results_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D0D0D0')
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Results rows
    results_data = [
        ("Poisson's Ratio (ν)", "{{poissons_ratio}}", "{{poissons_ratio_unc}}", "{{poissons_ratio_req}}", "-"),
        ("Shear Modulus (G)", "{{shear_modulus}}", "{{shear_modulus_unc}}", "{{shear_modulus_req}}", "GPa"),
        ("Young's Modulus (E)", "{{youngs_modulus}}", "{{youngs_modulus_unc}}", "{{youngs_modulus_req}}", "GPa"),
        ("Flexural Frequency (ff)", "{{flexural_frequency}}", "{{flexural_frequency_unc}}", "-", "Hz"),
        ("Torsional Frequency (ft)", "{{torsional_frequency}}", "{{torsional_frequency_unc}}", "-", "Hz"),
        ("Validity", "{{validity_status}}", "-", "-", "-"),
    ]

    for i, row_data in enumerate(results_data):
        row = results_table.rows[i + 1]
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = text
            if j == 0:
                set_cell_shading(cell, 'F0F0F0')

    doc.add_paragraph()

    # Chart placeholder
    chart_heading = doc.add_paragraph()
    chart_heading.add_run("Velocity Measurements Chart").bold = True

    chart_para = doc.add_paragraph()
    chart_para.add_run("{{chart}}")
    chart_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Validity Notes
    validity_heading = doc.add_paragraph()
    validity_heading.add_run("Notes").bold = True

    validity_para = doc.add_paragraph()
    validity_para.add_run("{{validity_notes}}")

    doc.add_paragraph()

    # Signatures table
    sig_heading = doc.add_paragraph()
    sig_heading.add_run("Approvals").bold = True

    sig_table = doc.add_table(rows=2, cols=6)
    sig_table.style = 'Table Grid'

    sig_headers = ["Tested By:", "{{tested_by}}", "Reviewed By:", "{{reviewed_by}}",
                   "Approved By:", "{{approved_by}}"]
    sig_dates = ["Date:", "{{tested_date}}", "Date:", "{{reviewed_date}}",
                 "Date:", "{{approved_date}}"]

    for i, text in enumerate(sig_headers):
        cell = sig_table.rows[0].cells[i]
        cell.text = text
        if i % 2 == 0:
            cell.paragraphs[0].runs[0].bold = True
            set_cell_shading(cell, 'E8E8E8')

    for i, text in enumerate(sig_dates):
        cell = sig_table.rows[1].cells[i]
        cell.text = text
        if i % 2 == 0:
            cell.paragraphs[0].runs[0].bold = True
            set_cell_shading(cell, 'E8E8E8')

    # Save
    output_path = Path(__file__).parent.parent / "templates" / "sonic_e1875_report_template.docx"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"Template saved to: {output_path}")


if __name__ == "__main__":
    create_template()
