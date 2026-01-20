"""
Script to create the CTOD E1290 report Word template.
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Cm, Pt, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def set_table_column_widths(table, widths_inches):
    """Set column widths for a table using XML for reliable results."""
    # Disable auto-fit
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)

    # Set widths for each cell in each row
    for row in table.rows:
        for idx, width in enumerate(widths_inches):
            if idx < len(row.cells):
                cell = row.cells[idx]
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(int(width * 1440)))  # 1440 twips per inch
                tcW.set(qn('w:type'), 'dxa')
                existing = tcPr.find(qn('w:tcW'))
                if existing is not None:
                    tcPr.remove(existing)
                tcPr.insert(0, tcW)


def add_page_number_field(paragraph):
    """Add a PAGE field to the paragraph."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def add_page_header(doc):
    """Add page header with Certificate No and page number on right side."""
    section = doc.sections[0]
    header = section.header

    # Create a table for header layout (logo left, certificate right)
    header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    header_table.autofit = False
    header_table.columns[0].width = Inches(2)
    header_table.columns[1].width = Inches(4.5)

    # Logo cell (left)
    logo_cell = header_table.rows[0].cells[0]
    logo_para = logo_cell.paragraphs[0]
    logo_para.add_run("{{logo}}")
    logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Certificate and page number cell (right)
    cert_cell = header_table.rows[0].cells[1]
    cert_para = cert_cell.paragraphs[0]
    cert_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    cert_run = cert_para.add_run("Certificate No: {{certificate_number}}")
    cert_run.font.size = Pt(10)
    cert_run.bold = True

    # Add page number on new line
    page_para = cert_cell.add_paragraph()
    page_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    page_run = page_para.add_run("Page ")
    page_run.font.size = Pt(9)
    add_page_number_field(page_para)


def add_page_footer(doc):
    """Add page footer with legal text."""
    section = doc.sections[0]
    footer = section.footer

    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    footer_text = (
        "All work and services carried out by Durabler are subject to, and conducted in accordance with, "
        "Durabler standard terms and conditions, which are available at durabler.se. This document shall "
        "not be reproduced other than in full, except with prior written approval of the issuer. The results "
        "pertain only to the item(s) as sampled by the client unless otherwise indicated. "
        "Durabler a part of Subseatec S AB, Address: Durabler C/O Subseatec, Dalavägen 23, 68130 Kristinehamn, SWEDEN"
    )

    footer_run = footer_para.add_run(footer_text)
    footer_run.font.size = Pt(7)
    footer_run.italic = True


def create_template():
    doc = Document()

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.header_distance = Cm(0.5)
        section.footer_distance = Cm(0.5)

    # Add page header and footer
    add_page_header(doc)
    add_page_footer(doc)

    # Title
    title_para = doc.add_paragraph()
    title_run = title_para.add_run("CTOD TEST REPORT")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add standard reference below title
    std_para = doc.add_paragraph()
    std_run = std_para.add_run("ASTM E1290 / BS 7448")
    std_run.font.size = Pt(11)
    std_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Test Information Table (no certificate number)
    info_heading = doc.add_paragraph()
    info_heading.add_run("Test Information").bold = True

    info_table = doc.add_table(rows=6, cols=4)
    info_table.style = 'Table Grid'

    info_data = [
        ("Test Project:", "{{test_project}}", "Customer:", "{{customer}}"),
        ("Customer Order:", "{{customer_order}}", "Test Date:", "{{test_date}}"),
        ("Product/S/N:", "{{product_sn}}", "Test Standard:", "{{test_standard}}"),
        ("Material/HT:", "{{material}}", "Specimen ID:", "{{specimen_id}}"),
        ("Location/Orient.:", "{{location_orientation}}", "Test Temperature:", "{{test_temperature}} °C"),
        ("Test Equipment:", "{{test_equipment}}", "", ""),
    ]

    for i, row_data in enumerate(info_data):
        row = info_table.rows[i]
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = text
            if j % 2 == 0 and text:  # Labels
                cell.paragraphs[0].runs[0].bold = True
                set_cell_shading(cell, 'E8E8E8')

    doc.add_paragraph()

    # Specimen Dimensions Table
    dim_heading = doc.add_paragraph()
    dim_heading.add_run("Specimen Geometry").bold = True

    dim_table = doc.add_table(rows=4, cols=6)
    dim_table.style = 'Table Grid'

    dim_data = [
        ("Type:", "{{specimen_type}}", "W (mm):", "{{W}}", "B (mm):", "{{B}}"),
        ("Bₙ (mm):", "{{B_n}}", "a₀ (mm):", "{{a_0}}", "S (mm):", "{{S}}"),
        ("a₀/W:", "{{a_W_ratio}}", "aᶠ (mm):", "{{a_f}}", "Δa (mm):", "{{delta_a}}"),
        ("Notch Type:", "{{notch_type}}", "Side Grooves:", "{{side_grooves}}", "", ""),
    ]

    for i, row_data in enumerate(dim_data):
        row = dim_table.rows[i]
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = text
            if j % 2 == 0 and text:  # Labels
                cell.paragraphs[0].runs[0].bold = True
                set_cell_shading(cell, 'E8E8E8')

    doc.add_paragraph()

    # Material Properties Table
    mat_heading = doc.add_paragraph()
    mat_heading.add_run("Material Properties").bold = True

    mat_table = doc.add_table(rows=2, cols=6)
    mat_table.style = 'Table Grid'

    mat_data = [
        ("σᵧₛ (MPa):", "{{yield_strength}}", "σᵤₜₛ (MPa):", "{{ultimate_strength}}", "E (GPa):", "{{youngs_modulus}}"),
        ("ν:", "{{poissons_ratio}}", "", "", "", ""),
    ]

    for i, row_data in enumerate(mat_data):
        row = mat_table.rows[i]
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = text
            if j % 2 == 0 and text:  # Labels
                cell.paragraphs[0].runs[0].bold = True
                set_cell_shading(cell, 'E8E8E8')

    doc.add_paragraph()

    # Results Table
    results_heading = doc.add_paragraph()
    results_heading.add_run("Test Results").bold = True

    results_table = doc.add_table(rows=9, cols=6)
    results_table.style = 'Table Grid'

    # Set column widths (added Requirement column)
    widths = [2.0, 0.65, 0.65, 0.6, 1.2, 0.8]
    set_table_column_widths(results_table, widths)

    # Header row
    header_row = results_table.rows[0]
    headers = ["Parameter", "Value", "U (k=2)", "Unit", "Requirement", "Validity"]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D0D0D0')
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Results rows (added requirement column placeholder)
    results_data = [
        ("Pₘₐₓ (Maximum Force)", "{{P_max_value}}", "{{P_max_uncertainty}}", "kN", "{{P_max_req}}", "-"),
        ("CMOD at Pₘₐₓ", "{{CMOD_max_value}}", "{{CMOD_max_uncertainty}}", "mm", "{{CMOD_max_req}}", "-"),
        ("K at Pₘₐₓ", "{{K_max_value}}", "{{K_max_uncertainty}}", "MPa√m", "{{K_max_req}}", "-"),
        ("δc (CTOD at cleavage)", "{{delta_c_value}}", "{{delta_c_uncertainty}}", "mm", "{{delta_c_req}}", "{{delta_c_valid}}"),
        ("δu (CTOD with growth)", "{{delta_u_value}}", "{{delta_u_uncertainty}}", "mm", "{{delta_u_req}}", "{{delta_u_valid}}"),
        ("δm (CTOD at Pₘₐₓ)", "{{delta_m_value}}", "{{delta_m_uncertainty}}", "mm", "{{delta_m_req}}", "{{delta_m_valid}}"),
        ("CTOD Type Reported", "{{ctod_type}}", "-", "-", "-", "-"),
        ("Elastic Compliance", "{{compliance}}", "-", "mm/kN", "-", "-"),
    ]

    for i, row_data in enumerate(results_data):
        row = results_table.rows[i + 1]
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = text
            if j == 0:  # Parameter column
                set_cell_shading(cell, 'F0F0F0')

    doc.add_paragraph()

    # Chart placeholder
    chart_heading = doc.add_paragraph()
    chart_heading.add_run("Force vs CMOD Curve").bold = True

    chart_para = doc.add_paragraph()
    chart_para.add_run("{{chart}}")
    chart_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Crack Surface Photos section
    photos_heading = doc.add_paragraph()
    photos_heading.add_run("Crack Surface Documentation").bold = True

    photo_para = doc.add_paragraph()
    photo_para.add_run("{{photos}}")
    photo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # 9-Point Crack Measurements Table
    crack_heading = doc.add_paragraph()
    crack_heading.add_run("Crack Length Measurements (9-Point Average per E1290)").bold = True

    crack_table = doc.add_table(rows=2, cols=10)
    crack_table.style = 'Table Grid'

    # Header row
    crack_headers = ["Pos.", "a₁", "a₂", "a₃", "a₄", "a₅", "a₆", "a₇", "a₈", "a₉"]
    for i, header in enumerate(crack_headers):
        cell = crack_table.rows[0].cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D0D0D0')
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Data row
    crack_data = ["mm", "{{a1}}", "{{a2}}", "{{a3}}", "{{a4}}", "{{a5}}",
                  "{{a6}}", "{{a7}}", "{{a8}}", "{{a9}}"]
    for i, text in enumerate(crack_data):
        cell = crack_table.rows[1].cells[i]
        cell.text = text
        if i == 0:
            set_cell_shading(cell, 'E8E8E8')
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Validity statement
    validity = doc.add_paragraph()
    validity.add_run("Validity Status: ").bold = True
    validity.add_run("{{validity_status}}")

    validity_stmt = doc.add_paragraph()
    validity_stmt.add_run("{{validity_statement}}")

    doc.add_paragraph()

    # Signatures table
    sig_heading = doc.add_paragraph()
    sig_heading.add_run("Approvals").bold = True

    sig_table = doc.add_table(rows=2, cols=6)
    sig_table.style = 'Table Grid'

    sig_headers = ["Tested By:", "{{tested_by}}", "Reviewed By:", "{{reviewed_by}}",
                   "Approved By:", "{{approved_by}}"]
    sig_dates = ["Date:", "{{tested_date}}", "Date:", "{{reviewed_date}}",
                 "Date:", "{{approved_date}}"]

    for i, text in enumerate(sig_headers):
        cell = sig_table.rows[0].cells[i]
        cell.text = text
        if i % 2 == 0:
            cell.paragraphs[0].runs[0].bold = True
            set_cell_shading(cell, 'E8E8E8')

    for i, text in enumerate(sig_dates):
        cell = sig_table.rows[1].cells[i]
        cell.text = text
        if i % 2 == 0:
            cell.paragraphs[0].runs[0].bold = True
            set_cell_shading(cell, 'E8E8E8')

    # Save
    output_path = Path(__file__).parent.parent / "templates" / "ctod_e1290_report_template.docx"
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(output_path)
    print(f"Template saved to: {output_path}")


if __name__ == "__main__":
    create_template()
