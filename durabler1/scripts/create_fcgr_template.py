"""
Create FCGR E647 Report Template.

Generates a Word document template with placeholders for FCGR test reports.
Based on CTOD E1290 template layout.
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color):
    """Set cell background shading."""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def add_page_number_field(paragraph):
    """Add a PAGE field to the paragraph."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def add_page_header(doc):
    """Add page header with Certificate No and page number on right side."""
    section = doc.sections[0]
    header = section.header

    # Create a table for header layout (logo left, certificate right)
    header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    header_table.autofit = False
    header_table.columns[0].width = Inches(2)
    header_table.columns[1].width = Inches(4.5)

    # Logo cell (left)
    logo_cell = header_table.rows[0].cells[0]
    logo_para = logo_cell.paragraphs[0]
    logo_para.add_run("{{logo}}")
    logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Certificate and page number cell (right)
    cert_cell = header_table.rows[0].cells[1]
    cert_para = cert_cell.paragraphs[0]
    cert_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    cert_run = cert_para.add_run("Certificate No: {{certificate_number}}")
    cert_run.font.size = Pt(10)
    cert_run.bold = True

    # Add page number on new line
    page_para = cert_cell.add_paragraph()
    page_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    page_run = page_para.add_run("Page ")
    page_run.font.size = Pt(9)
    add_page_number_field(page_para)


def add_page_footer(doc):
    """Add page footer with legal text."""
    section = doc.sections[0]
    footer = section.footer

    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    footer_text = (
        "All work and services carried out by Durabler are subject to, and conducted in accordance with, "
        "Durabler standard terms and conditions, which are available at durabler.se. This document shall "
        "not be reproduced other than in full, except with prior written approval of the issuer. The results "
        "pertain only to the item(s) as sampled by the client unless otherwise indicated. "
        "Durabler a part of Subseatec S AB, Address: Durabler C/O Subseatec, Dalavägen 23, 68130 Kristinehamn, SWEDEN"
    )

    footer_run = footer_para.add_run(footer_text)
    footer_run.font.size = Pt(7)
    footer_run.italic = True


def create_fcgr_template():
    """Create the FCGR E647 report template."""
    doc = Document()

    # Set narrow margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.header_distance = Cm(0.5)
        section.footer_distance = Cm(0.5)

    # Add page header and footer
    add_page_header(doc)
    add_page_footer(doc)

    # === TITLE ===
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("FATIGUE CRACK GROWTH RATE TEST REPORT")
    title_run.bold = True
    title_run.font.size = Pt(14)

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_para.add_run("ASTM E647")

    doc.add_paragraph()

    # === TEST INFORMATION TABLE (no certificate number) ===
    doc.add_paragraph("TEST INFORMATION").runs[0].bold = True

    info_table = doc.add_table(rows=4, cols=4)
    info_table.style = 'Table Grid'

    info_data = [
        ("Test Project:", "{{test_project}}", "Customer:", "{{customer}}"),
        ("Specimen ID:", "{{specimen_id}}", "Material:", "{{material}}"),
        ("Test Date:", "{{test_date}}", "Temperature:", "{{temperature}} °C"),
        ("Test Standard:", "{{test_standard}}", "Equipment:", "{{test_equipment}}"),
    ]

    for i, row_data in enumerate(info_data):
        row = info_table.rows[i]
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = text
            if j % 2 == 0:  # Labels
                set_cell_shading(cell, "D9D9D9")
                cell.paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # === SPECIMEN GEOMETRY TABLE ===
    doc.add_paragraph("SPECIMEN GEOMETRY").runs[0].bold = True

    geom_table = doc.add_table(rows=5, cols=4)
    geom_table.style = 'Table Grid'

    geom_data = [
        ("Specimen Type:", "{{specimen_type}}", "Side Grooves:", "{{side_grooves}}"),
        ("Width W (mm):", "{{W}}", "Thickness B (mm):", "{{B}}"),
        ("Net Thickness Bn (mm):", "{{B_n}}", "Initial Notch a0 (mm):", "{{a_0}}"),
        ("Notch Height h (mm):", "{{notch_height}}", "a0/W Ratio:", "{{a_W_ratio}}"),
        ("Geometry Function:", "ASTM E647", "", ""),
    ]

    for i, row_data in enumerate(geom_data):
        row = geom_table.rows[i]
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = text
            if j % 2 == 0 and text:  # Labels
                set_cell_shading(cell, "D9D9D9")
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # === MATERIAL PROPERTIES TABLE ===
    doc.add_paragraph("MATERIAL PROPERTIES").runs[0].bold = True

    mat_table = doc.add_table(rows=2, cols=4)
    mat_table.style = 'Table Grid'

    mat_data = [
        ("Yield Strength (MPa):", "{{yield_strength}}", "Ultimate Strength (MPa):", "{{ultimate_strength}}"),
        ("Young's Modulus (GPa):", "{{youngs_modulus}}", "Poisson's Ratio:", "{{poissons_ratio}}"),
    ]

    for i, row_data in enumerate(mat_data):
        row = mat_table.rows[i]
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = text
            if j % 2 == 0:  # Labels
                set_cell_shading(cell, "D9D9D9")
                cell.paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # === TEST PARAMETERS TABLE ===
    doc.add_paragraph("TEST PARAMETERS").runs[0].bold = True

    param_table = doc.add_table(rows=4, cols=4)
    param_table.style = 'Table Grid'

    param_data = [
        ("Control Mode:", "{{control_mode}}", "Wave Shape:", "{{wave_shape}}"),
        ("Load Ratio R:", "{{load_ratio}}", "Frequency (Hz):", "{{frequency}}"),
        ("Maximum Load Pmax (kN):", "{{P_max}}", "Maximum Kmax (MPa√m):", "{{K_max}}"),
        ("Environment:", "{{environment}}", "da/dN Method:", "{{dadn_method}}"),
    ]

    for i, row_data in enumerate(param_data):
        row = param_table.rows[i]
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = text
            if j % 2 == 0:  # Labels
                set_cell_shading(cell, "D9D9D9")
                cell.paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # === PARIS LAW RESULTS TABLE ===
    doc.add_paragraph("PARIS LAW RESULTS: da/dN = C × (ΔK)^m").runs[0].bold = True

    paris_table = doc.add_table(rows=6, cols=4)
    paris_table.style = 'Table Grid'

    paris_data = [
        ("Coefficient C:", "{{paris_C}}", "Exponent m:", "{{paris_m}}"),
        ("C Std Error:", "{{paris_C_error}}", "m Std Error:", "{{paris_m_error}}"),
        ("R-squared:", "{{paris_r_squared}}", "Data Points:", "{{paris_n_points}}"),
        ("ΔK Min (MPa√m):", "{{delta_K_min}}", "ΔK Max (MPa√m):", "{{delta_K_max}}"),
        ("da/dN Min (mm/cycle):", "{{da_dN_min}}", "da/dN Max (mm/cycle):", "{{da_dN_max}}"),
        ("Outliers Removed:", "{{n_outliers}}", "Outlier Threshold:", "{{outlier_threshold}}%"),
    ]

    for i, row_data in enumerate(paris_data):
        row = paris_table.rows[i]
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = text
            if j % 2 == 0:  # Labels
                set_cell_shading(cell, "D9D9D9")
                cell.paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # === TEST SUMMARY TABLE ===
    doc.add_paragraph("TEST SUMMARY").runs[0].bold = True

    summary_table = doc.add_table(rows=3, cols=4)
    summary_table.style = 'Table Grid'

    summary_data = [
        ("Total Cycles:", "{{total_cycles}}", "Valid Data Points:", "{{n_valid_points}}"),
        ("Final Crack Length (mm):", "{{final_crack_length}}", "Threshold ΔK:", "{{threshold_delta_K}}"),
        ("Validity Status:", "{{is_valid}}", "", ""),
    ]

    for i, row_data in enumerate(summary_data):
        row = summary_table.rows[i]
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = text
            if j % 2 == 0 and text:  # Labels
                set_cell_shading(cell, "D9D9D9")
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # === PLOTS ===
    doc.add_paragraph("TEST PLOTS").runs[0].bold = True

    # Create table for side-by-side plots
    plot_table = doc.add_table(rows=2, cols=2)
    plot_table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Plot 1 title
    plot_table.rows[0].cells[0].text = "Crack Length vs Cycles"
    plot_table.rows[0].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    plot_table.rows[0].cells[0].paragraphs[0].runs[0].bold = True

    # Plot 2 title
    plot_table.rows[0].cells[1].text = "da/dN vs ΔK (Paris Law)"
    plot_table.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    plot_table.rows[0].cells[1].paragraphs[0].runs[0].bold = True

    # Plot placeholders
    plot_table.rows[1].cells[0].text = "{{plot1}}"
    plot_table.rows[1].cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    plot_table.rows[1].cells[1].text = "{{plot2}}"
    plot_table.rows[1].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # === CRACK SURFACE PHOTOS ===
    doc.add_paragraph("CRACK SURFACE PHOTOS").runs[0].bold = True
    doc.add_paragraph("{{photos}}")

    doc.add_paragraph()

    # === VALIDITY NOTES ===
    doc.add_paragraph("VALIDITY NOTES (ASTM E647)").runs[0].bold = True
    doc.add_paragraph("{{validity_notes}}")

    doc.add_paragraph()

    # === SIGNATURES TABLE ===
    doc.add_paragraph("SIGNATURES").runs[0].bold = True

    sig_table = doc.add_table(rows=4, cols=3)
    sig_table.style = 'Table Grid'

    sig_headers = ["", "Name", "Date"]
    for j, header in enumerate(sig_headers):
        cell = sig_table.rows[0].cells[j]
        cell.text = header
        set_cell_shading(cell, "D9D9D9")
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].bold = True

    sig_rows = [
        ("Tested by:", "{{tested_by}}", "{{tested_date}}"),
        ("Reviewed by:", "{{reviewed_by}}", "{{reviewed_date}}"),
        ("Approved by:", "{{approved_by}}", "{{approved_date}}"),
    ]

    for i, row_data in enumerate(sig_rows, start=1):
        row = sig_table.rows[i]
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = text
            if j == 0:  # Label column
                set_cell_shading(cell, "D9D9D9")
                cell.paragraphs[0].runs[0].bold = True

    # Save template
    output_path = Path(__file__).parent.parent / "templates" / "fcgr_e647_report_template.docx"
    output_path.parent.mkdir(exist_ok=True)
    doc.save(output_path)
    print(f"Template created: {output_path}")
    return output_path


if __name__ == "__main__":
    create_fcgr_template()
