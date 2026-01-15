"""
Script to create the tensile report Word template.
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Cm, Pt, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)

def set_column_width(column, width_inches):
    """Set column width."""
    for cell in column.cells:
        cell.width = Inches(width_inches)

def set_table_column_widths(table, widths_inches):
    """Set column widths for a table using XML for reliable results."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
    from docx.shared import Inches

    # Disable auto-fit
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    tblLayout = OxmlElement('w:tblLayout')
    tblLayout.set(qn('w:type'), 'fixed')
    tblPr.append(tblLayout)
    if tbl.tblPr is None:
        tbl.insert(0, tblPr)

    # Set widths for each cell in each row
    for row in table.rows:
        for idx, width in enumerate(widths_inches):
            if idx < len(row.cells):
                cell = row.cells[idx]
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = OxmlElement('w:tcW')
                tcW.set(qn('w:w'), str(int(width * 1440)))  # 1440 twips per inch
                tcW.set(qn('w:type'), 'dxa')
                # Remove existing tcW if present
                existing = tcPr.find(qn('w:tcW'))
                if existing is not None:
                    tcPr.remove(existing)
                tcPr.insert(0, tcW)

def create_template():
    doc = Document()

    # Set narrow margins
    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # Logo placeholder
    p = doc.add_paragraph()
    p.add_run("{{logo}}")

    # Title
    title = doc.add_paragraph()
    title_run = title.add_run("TENSILE TEST REPORT")
    title_run.bold = True
    title_run.font.size = Pt(16)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Standard reference
    std = doc.add_paragraph()
    std_run = std.add_run("ASTM E8/E8M-22")
    std_run.font.size = Pt(11)
    std.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Test Information Table
    info_heading = doc.add_paragraph()
    info_heading.add_run("Test Information").bold = True

    info_table = doc.add_table(rows=7, cols=4)
    info_table.style = 'Table Grid'

    info_data = [
        ("Test Project:", "{{test_project}}", "Certificate No:", "{{certificate_number}}"),
        ("Customer:", "{{customer}}", "Test Date:", "{{test_date}}"),
        ("Customer Order:", "{{customer_order}}", "Test Standard:", "{{test_standard}}"),
        ("Product/S/N:", "{{product_sn}}", "Yield Method:", "{{yield_method}}"),
        ("Specimen ID:", "{{specimen_id}}", "Strain Source:", "{{strain_source}}"),
        ("Location/Orient.:", "{{location_orientation}}", "Test Equipment:", "{{test_equipment}}"),
        ("Material/HT:", "{{material}}", "Temperature:", "{{test_temperature}} °C"),
    ]

    for i, row_data in enumerate(info_data):
        row = info_table.rows[i]
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = text
            if j % 2 == 0:  # Labels
                cell.paragraphs[0].runs[0].bold = True
                set_cell_shading(cell, 'E8E8E8')

    doc.add_paragraph()

    # Specimen Dimensions Table
    dim_heading = doc.add_paragraph()
    dim_heading.add_run("Specimen Dimensions").bold = True

    dim_table = doc.add_table(rows=4, cols=6)
    dim_table.style = 'Table Grid'

    dim_data = [
        ("Geometry:", "{{geometry_type}}", "d₀ (mm):", "{{d0}}", "w₀ (mm):", "{{w0}}"),
        ("L₀ (mm):", "{{L0}}", "Lc (mm):", "{{Lc}}", "t₀ (mm):", "{{t0}}"),
        ("L₁ (mm):", "{{L1}}", "df (mm):", "{{df}}", "A₀ (mm²):", "{{initial_area}}"),
        ("", "", "", "", "", ""),
    ]

    for i, row_data in enumerate(dim_data):
        if i >= len(dim_data) - 1:
            break
        row = dim_table.rows[i]
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = text
            if j % 2 == 0 and text:  # Labels
                cell.paragraphs[0].runs[0].bold = True
                set_cell_shading(cell, 'E8E8E8')

    doc.add_paragraph()

    # Results Table - adjusted column widths
    # Parameter wider (+8 chars ~0.5"), others narrower (-2 chars each ~0.125")
    results_heading = doc.add_paragraph()
    results_heading.add_run("Test Results").bold = True

    results_table = doc.add_table(rows=9, cols=5)
    results_table.style = 'Table Grid'

    # Set column widths: Parameter=2.4", Value=0.65", Uncertainty=0.65", Unit=0.4", Requirement=1.1"
    # Total ~5.2" (fits in page)
    widths = [2.4, 0.65, 0.65, 0.4, 1.1]
    set_table_column_widths(results_table, widths)

    # Header row
    header_row = results_table.rows[0]
    headers = ["Parameter", "Value", "U (k=2)", "Unit", "Requirement"]
    for i, header in enumerate(headers):
        cell = header_row.cells[i]
        cell.text = header
        cell.paragraphs[0].runs[0].bold = True
        set_cell_shading(cell, 'D0D0D0')
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Results rows
    results_data = [
        ("Rp0.2 (Yield 0.2%)", "{{Rp02_value}}", "{{Rp02_uncertainty}}", "MPa", "{{Rp02_req}}"),
        ("Rp0.5 (Yield 0.5%)", "{{Rp05_value}}", "{{Rp05_uncertainty}}", "MPa", "{{Rp05_req}}"),
        ("ReH (Upper Yield)", "{{ReH_value}}", "{{ReH_uncertainty}}", "MPa", "{{ReH_req}}"),
        ("ReL (Lower Yield)", "{{ReL_value}}", "{{ReL_uncertainty}}", "MPa", "{{ReL_req}}"),
        ("Rm (Ultimate)", "{{Rm_value}}", "{{Rm_uncertainty}}", "MPa", "{{Rm_req}}"),
        ("{{ratio_label}}", "{{yield_tensile_ratio}}", "-", "-", "-"),
        ("A5 (Elongation L1-L0)", "{{A5_value}}", "{{A5_uncertainty}}", "%", "{{A5_req}}"),
        ("Z (Reduction of Area)", "{{Z_value}}", "{{Z_uncertainty}}", "%", "{{Z_req}}"),
    ]

    for i, row_data in enumerate(results_data):
        row = results_table.rows[i + 1]
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = text
            if j == 0:  # Parameter column
                set_cell_shading(cell, 'F0F0F0')

    doc.add_paragraph()

    # Chart placeholder
    chart_para = doc.add_paragraph()
    chart_para.add_run("{{chart}}")
    chart_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Validity statement
    validity = doc.add_paragraph()
    validity.add_run("Status: ").bold = True
    validity.add_run("{{validity_status}}")

    validity_stmt = doc.add_paragraph()
    validity_stmt.add_run("{{validity_statement}}")

    doc.add_paragraph()

    # Signatures table
    sig_heading = doc.add_paragraph()
    sig_heading.add_run("Approvals").bold = True

    sig_table = doc.add_table(rows=2, cols=6)
    sig_table.style = 'Table Grid'

    sig_headers = ["Tested By:", "{{tested_by}}", "Reviewed By:", "{{reviewed_by}}", "Approved By:", "{{approved_by}}"]
    sig_dates = ["Date:", "{{tested_date}}", "Date:", "{{reviewed_date}}", "Date:", "{{approved_date}}"]

    for i, text in enumerate(sig_headers):
        cell = sig_table.rows[0].cells[i]
        cell.text = text
        if i % 2 == 0:
            cell.paragraphs[0].runs[0].bold = True
            set_cell_shading(cell, 'E8E8E8')

    for i, text in enumerate(sig_dates):
        cell = sig_table.rows[1].cells[i]
        cell.text = text
        if i % 2 == 0:
            cell.paragraphs[0].runs[0].bold = True
            set_cell_shading(cell, 'E8E8E8')

    doc.add_paragraph()
    doc.add_paragraph()

    # Footer
    footer = doc.add_paragraph()
    footer.add_run("Durabler a part of Subseatec S AB, Address: Durabler C/O Subseatec, Dalavägen 23, 68130 Kristinehamn, SWEDEN")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer.runs[0].font.size = Pt(8)
    footer.runs[0].italic = True

    # Save
    output_path = Path(__file__).parent.parent / "templates" / "tensile_report_template.docx"
    doc.save(output_path)
    print(f"Template saved to: {output_path}")

if __name__ == "__main__":
    create_template()
