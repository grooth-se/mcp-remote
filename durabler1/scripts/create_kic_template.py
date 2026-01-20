#!/usr/bin/env python3
"""
Create KIC E399 Report Template.

Generates a Word document template for KIC fracture toughness test reports
with placeholders for test data.
"""

from pathlib import Path

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
except ImportError:
    print("Error: python-docx is required. Install with: pip install python-docx")
    exit(1)


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_page_number_field(paragraph):
    """Add a PAGE field to the paragraph."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def add_page_header(doc):
    """Add page header with Certificate No and page number on right side."""
    section = doc.sections[0]
    header = section.header

    # Create a table for header layout (logo left, certificate right)
    header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    header_table.autofit = False
    header_table.columns[0].width = Inches(2)
    header_table.columns[1].width = Inches(4.5)

    # Logo cell (left)
    logo_cell = header_table.rows[0].cells[0]
    logo_para = logo_cell.paragraphs[0]
    logo_para.add_run("{{logo}}")
    logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Certificate and page number cell (right)
    cert_cell = header_table.rows[0].cells[1]
    cert_para = cert_cell.paragraphs[0]
    cert_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    cert_run = cert_para.add_run("Certificate No: {{certificate_number}}")
    cert_run.font.size = Pt(10)
    cert_run.bold = True

    # Add page number on new line
    page_para = cert_cell.add_paragraph()
    page_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    page_run = page_para.add_run("Page ")
    page_run.font.size = Pt(9)
    add_page_number_field(page_para)


def add_page_footer(doc):
    """Add page footer with legal text."""
    section = doc.sections[0]
    footer = section.footer

    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    footer_text = (
        "All work and services carried out by Durabler are subject to, and conducted in accordance with, "
        "Durabler standard terms and conditions, which are available at durabler.se. This document shall "
        "not be reproduced other than in full, except with prior written approval of the issuer. The results "
        "pertain only to the item(s) as sampled by the client unless otherwise indicated. "
        "Durabler a part of Subseatec S AB, Address: Durabler C/O Subseatec, Dalavägen 23, 68130 Kristinehamn, SWEDEN"
    )

    footer_run = footer_para.add_run(footer_text)
    footer_run.font.size = Pt(7)
    footer_run.italic = True


def create_kic_template():
    """Create KIC report template with placeholders."""
    doc = Document()

    # Set up styles and margins
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.header_distance = Cm(0.5)
        section.footer_distance = Cm(0.5)

    # Add page header and footer
    add_page_header(doc)
    add_page_footer(doc)

    # Title
    title = doc.add_heading('KIC Fracture Toughness Test Report', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Standard reference
    subtitle = doc.add_paragraph('ASTM E399 - Standard Test Method for Linear-Elastic')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle2 = doc.add_paragraph('Plane-Strain Fracture Toughness of Metallic Materials')
    subtitle2.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Report number
    report_para = doc.add_paragraph()
    report_para.add_run('Report Number: ').bold = True
    report_para.add_run('{{report_number}}')

    doc.add_paragraph()

    # Test Information section (no certificate number in table)
    doc.add_heading('1. Test Information', level=1)
    table = doc.add_table(rows=7, cols=2)
    table.style = 'Table Grid'

    info_fields = [
        ('Test Project:', '{{test_project}}'),
        ('Customer:', '{{customer}}'),
        ('Specimen ID:', '{{specimen_id}}'),
        ('Material:', '{{material}}'),
        ('Test Date:', '{{test_date}}'),
        ('Test Temperature:', '{{temperature}} °C'),
        ('Test Standard:', 'ASTM E399'),
    ]

    for i, (label, placeholder) in enumerate(info_fields):
        table.rows[i].cells[0].text = label
        table.rows[i].cells[1].text = placeholder
        table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[i].cells[0], 'D9D9D9')

    doc.add_paragraph()

    # Specimen Dimensions section
    doc.add_heading('2. Specimen Dimensions', level=1)
    table = doc.add_table(rows=7, cols=3)
    table.style = 'Table Grid'

    # Header
    headers = ['Parameter', 'Value', 'Unit']
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], 'D9D9D9')

    dim_fields = [
        ('Specimen Type', '{{specimen_type}}', '-'),
        ('Width W', '{{W}}', 'mm'),
        ('Thickness B', '{{B}}', 'mm'),
        ('Net Thickness B_n', '{{B_n}}', 'mm'),
        ('Crack Length a_0', '{{a_0}}', 'mm'),
        ('Span S (SE(B) only)', '{{S}}', 'mm'),
    ]

    for i, (param, placeholder, unit) in enumerate(dim_fields):
        table.rows[i+1].cells[0].text = param
        table.rows[i+1].cells[1].text = placeholder
        table.rows[i+1].cells[2].text = unit

    doc.add_paragraph()

    # Material Properties section
    doc.add_heading('3. Material Properties', level=1)
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'

    # Header
    for i, header in enumerate(headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], 'D9D9D9')

    mat_fields = [
        ('Yield Strength σ_ys', '{{yield_strength}}', 'MPa'),
        ("Young's Modulus E", '{{youngs_modulus}}', 'GPa'),
        ("Poisson's Ratio ν", '{{poissons_ratio}}', '-'),
    ]

    for i, (param, placeholder, unit) in enumerate(mat_fields):
        table.rows[i+1].cells[0].text = param
        table.rows[i+1].cells[1].text = placeholder
        table.rows[i+1].cells[2].text = unit

    doc.add_paragraph()

    # Results section
    doc.add_heading('4. Test Results', level=1)
    table = doc.add_table(rows=8, cols=5)
    table.style = 'Table Grid'

    # Header
    result_headers = ['Parameter', 'Value', 'U (k=2)', 'Requirement', 'Unit']
    for i, header in enumerate(result_headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], 'D9D9D9')

    result_fields = [
        ('Maximum Force P_max', '{{P_max}}', '±{{P_max_uncertainty}}', '{{P_max_req}}', 'kN'),
        ('Conditional Force P_Q', '{{P_Q}}', '±{{P_Q_uncertainty}}', '{{P_Q_req}}', 'kN'),
        ('P_max/P_Q Ratio', '{{P_ratio}}', '-', '{{P_ratio_req}}', '-'),
        ('Conditional K_Q', '{{K_Q}}', '±{{K_Q_uncertainty}}', '{{K_Q_req}}', 'MPa√m'),
        ('Fracture Toughness K_IC', '{{K_IC}}', '±{{K_IC_uncertainty}}', '{{K_IC_req}}', 'MPa√m'),
        ('Initial Compliance', '{{compliance}}', '-', '-', 'mm/kN'),
        ('Validity Status', '{{is_valid}}', '-', '-', '-'),
    ]

    for i, (param, value, unc, req, unit) in enumerate(result_fields):
        table.rows[i+1].cells[0].text = param
        table.rows[i+1].cells[1].text = value
        table.rows[i+1].cells[2].text = unc
        table.rows[i+1].cells[3].text = req
        table.rows[i+1].cells[4].text = unit

    doc.add_paragraph()

    # Force-Displacement Plot section
    doc.add_heading('5. Force vs Displacement', level=1)
    chart_para = doc.add_paragraph()
    chart_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    chart_para.add_run('{{chart}}')

    doc.add_paragraph()

    # Validity Assessment section
    doc.add_heading('6. Validity Assessment per ASTM E399', level=1)

    validity_para = doc.add_paragraph()
    validity_para.add_run('Overall Validity: ').bold = True
    validity_para.add_run('{{is_valid}}')

    doc.add_paragraph()
    doc.add_paragraph('Validity Checks:')
    doc.add_paragraph('{{validity_notes}}')

    doc.add_paragraph()

    # Notes section
    doc.add_heading('7. Notes', level=1)
    doc.add_paragraph('The 5% secant offset method was used to determine P_Q per ASTM E399.')
    doc.add_paragraph()
    doc.add_paragraph('Validity requirements per ASTM E399:')
    notes = [
        '- a/W ratio must be between 0.45 and 0.55',
        '- B ≥ 2.5(K_Q/σ_ys)² for plane-strain conditions',
        '- a ≥ 2.5(K_Q/σ_ys)²',
        '- (W-a) ≥ 2.5(K_Q/σ_ys)²',
        '- P_max/P_Q ≤ 1.10',
    ]
    for note in notes:
        doc.add_paragraph(note)

    doc.add_paragraph()

    # Approval section
    doc.add_heading('8. Approval', level=1)
    table = doc.add_table(rows=4, cols=3)
    table.style = 'Table Grid'

    approval_headers = ['Role', 'Name', 'Date/Signature']
    for i, header in enumerate(approval_headers):
        table.rows[0].cells[i].text = header
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
        set_cell_shading(table.rows[0].cells[i], 'D9D9D9')

    approval_roles = ['Tested by:', 'Reviewed by:', 'Approved by:']
    for i, role in enumerate(approval_roles):
        table.rows[i+1].cells[0].text = role
        set_cell_shading(table.rows[i+1].cells[0], 'D9D9D9')

    return doc


def main():
    """Create and save the KIC report template."""
    # Determine output path
    script_dir = Path(__file__).parent
    project_root = script_dir.parent
    template_dir = project_root / 'templates'
    template_dir.mkdir(exist_ok=True)

    output_path = template_dir / 'kic_e399_report_template.docx'

    print(f"Creating KIC E399 report template...")
    doc = create_kic_template()
    doc.save(output_path)
    print(f"Template saved to: {output_path}")


if __name__ == '__main__':
    main()
