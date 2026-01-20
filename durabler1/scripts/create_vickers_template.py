"""
Create Vickers Hardness Report Template for ASTM E92.

This script creates a Word document template with placeholders
for the Vickers hardness test report, matching the Tensile report style.
"""

from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


def set_cell_shading(cell, color):
    """Set cell background color."""
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    cell._tc.get_or_add_tcPr().append(shading)


def add_page_header(doc):
    """Add page header with Certificate No and page number on right side."""
    section = doc.sections[0]
    header = section.header

    # Create a table for header layout (logo left, certificate right)
    header_table = header.add_table(rows=1, cols=2, width=Inches(6.5))
    header_table.autofit = False
    header_table.columns[0].width = Inches(2)
    header_table.columns[1].width = Inches(4.5)

    # Logo cell (left)
    logo_cell = header_table.rows[0].cells[0]
    logo_para = logo_cell.paragraphs[0]
    logo_para.add_run("{{logo}}")
    logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # Certificate and page number cell (right)
    cert_cell = header_table.rows[0].cells[1]
    cert_para = cert_cell.paragraphs[0]
    cert_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    cert_run = cert_para.add_run("Certificate No: {{certificate_number}}")
    cert_run.font.size = Pt(10)
    cert_run.bold = True

    # Add page number on new line
    page_para = cert_cell.add_paragraph()
    page_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    page_run = page_para.add_run("Page ")
    page_run.font.size = Pt(9)

    # Add page number field
    add_page_number_field(page_para)


def add_page_number_field(paragraph):
    """Add a PAGE field to the paragraph."""
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = "PAGE"

    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')

    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def add_page_footer(doc):
    """Add page footer with legal text."""
    section = doc.sections[0]
    footer = section.footer

    footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    footer_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    footer_text = (
        "All work and services carried out by Durabler are subject to, and conducted in accordance with, "
        "Durabler standard terms and conditions, which are available at durabler.se. This document shall "
        "not be reproduced other than in full, except with prior written approval of the issuer. The results "
        "pertain only to the item(s) as sampled by the client unless otherwise indicated. "
        "Durabler a part of Subseatec S AB, Address: Durabler C/O Subseatec, Dalavägen 23, 68130 Kristinehamn, SWEDEN"
    )

    footer_run = footer_para.add_run(footer_text)
    footer_run.font.size = Pt(7)
    footer_run.italic = True


def create_vickers_template():
    """Create Vickers hardness report template."""
    doc = Document()

    # Set narrow margins
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.5)  # Extra space for header
        section.bottom_margin = Cm(2.5)  # Extra space for footer
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)
        section.header_distance = Cm(0.5)
        section.footer_distance = Cm(0.5)

    # Add page header with certificate and page number
    add_page_header(doc)

    # Add page footer with legal text
    add_page_footer(doc)

    # === TITLE ===
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title_para.add_run("VICKERS HARDNESS TEST REPORT")
    title_run.bold = True
    title_run.font.size = Pt(14)

    subtitle_para = doc.add_paragraph()
    subtitle_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle_para.add_run("ASTM E92 / ISO 6507-1")

    doc.add_paragraph()

    # === REPORT INFORMATION TABLE ===
    doc.add_paragraph("REPORT INFORMATION").runs[0].bold = True

    report_table = doc.add_table(rows=2, cols=4)
    report_table.style = 'Table Grid'

    report_data = [
        ("Report Number:", "{{report_number}}", "Report Date:", "{{report_date}}"),
        ("Load Level:", "{{load_level}}", "", ""),
    ]

    for i, row_data in enumerate(report_data):
        row = report_table.rows[i]
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = text
            if j % 2 == 0 and text:  # Labels
                set_cell_shading(cell, "D9D9D9")
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # === TEST INFORMATION TABLE (no certificate number) ===
    doc.add_paragraph("TEST INFORMATION").runs[0].bold = True

    info_table = doc.add_table(rows=4, cols=4)
    info_table.style = 'Table Grid'

    info_data = [
        ("Test Project:", "{{test_project}}", "Customer:", "{{customer}}"),
        ("Specimen ID:", "{{specimen_id}}", "Material:", "{{material}}"),
        ("Test Date:", "{{test_date}}", "Operator:", "{{operator}}"),
        ("Test Standard:", "ASTM E92", "Equipment:", "Vickers Hardness Tester"),
    ]

    for i, row_data in enumerate(info_data):
        row = info_table.rows[i]
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = text
            if j % 2 == 0:  # Labels
                set_cell_shading(cell, "D9D9D9")
                cell.paragraphs[0].runs[0].bold = True

    doc.add_paragraph()

    # === TEST RESULTS TABLE ===
    doc.add_paragraph("TEST RESULTS").runs[0].bold = True

    results_table = doc.add_table(rows=7, cols=5)
    results_table.style = 'Table Grid'

    # Header row
    header_row = results_table.rows[0]
    headers = ["Parameter", "Value", "U (k=2)", "Requirement", "Unit"]
    for j, header in enumerate(headers):
        cell = header_row.cells[j]
        cell.text = header
        set_cell_shading(cell, "D9D9D9")
        cell.paragraphs[0].runs[0].bold = True

    results_data = [
        ("Mean Hardness", "{{mean_hardness}}", "± {{uncertainty}}", "{{hardness_req}}", "{{unit}}"),
        ("Standard Deviation", "{{std_dev}}", "-", "-", "{{unit}}"),
        ("Range (Max - Min)", "{{range}}", "-", "-", "{{unit}}"),
        ("Minimum Value", "{{min_value}}", "-", "-", "{{unit}}"),
        ("Maximum Value", "{{max_value}}", "-", "-", "{{unit}}"),
        ("Number of Readings", "{{n_readings}}", "-", "-", "-"),
    ]

    for i, row_data in enumerate(results_data):
        row = results_table.rows[i + 1]
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = text

    doc.add_paragraph()

    # === UNCERTAINTY BUDGET TABLE ===
    doc.add_paragraph("UNCERTAINTY BUDGET (ISO 17025 / GUM)").runs[0].bold = True

    unc_table = doc.add_table(rows=7, cols=3)
    unc_table.style = 'Table Grid'

    # Header row
    unc_headers = ["Source", "Type", "Value (HV)"]
    for j, header in enumerate(unc_headers):
        cell = unc_table.rows[0].cells[j]
        cell.text = header
        set_cell_shading(cell, "D9D9D9")
        cell.paragraphs[0].runs[0].bold = True

    unc_data = [
        ("Repeatability (std of mean)", "A", "{{u_A}}"),
        ("Machine calibration", "B", "{{u_machine}}"),
        ("Diagonal measurement", "B", "{{u_diagonal}}"),
        ("Force application", "B", "{{u_force}}"),
        ("Combined standard uncertainty", "-", "{{u_combined}}"),
        ("Expanded uncertainty (k={{k}})", "-", "{{U_expanded}}"),
    ]

    for i, row_data in enumerate(unc_data):
        row = unc_table.rows[i + 1]
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = text
            if i >= 4:  # Highlight combined/expanded rows
                set_cell_shading(cell, "FFF2CC")

    doc.add_paragraph()

    # === HARDNESS PROFILE CHART ===
    doc.add_paragraph("HARDNESS PROFILE").runs[0].bold = True

    chart_para = doc.add_paragraph("{{chart}}")
    chart_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # === INDENT PHOTOGRAPH ===
    doc.add_paragraph("INDENT PHOTOGRAPH").runs[0].bold = True

    photo_para = doc.add_paragraph("{{photo}}")
    photo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # === SIGNATURES TABLE ===
    doc.add_paragraph("SIGNATURES").runs[0].bold = True

    sig_table = doc.add_table(rows=4, cols=3)
    sig_table.style = 'Table Grid'

    sig_headers = ["", "Name", "Date"]
    for j, header in enumerate(sig_headers):
        cell = sig_table.rows[0].cells[j]
        cell.text = header
        set_cell_shading(cell, "D9D9D9")
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].bold = True

    sig_rows = [
        ("Tested by:", "", ""),
        ("Reviewed by:", "", ""),
        ("Approved by:", "", ""),
    ]

    for i, row_data in enumerate(sig_rows, start=1):
        row = sig_table.rows[i]
        for j, text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = text
            if j == 0:  # Label column
                set_cell_shading(cell, "D9D9D9")
                cell.paragraphs[0].runs[0].bold = True

    # Save template
    template_dir = Path(__file__).parent.parent / "templates"
    template_dir.mkdir(exist_ok=True)
    template_path = template_dir / "vickers_e92_report_template.docx"
    doc.save(template_path)

    print(f"Template created: {template_path}")
    return template_path


if __name__ == "__main__":
    create_vickers_template()
