"""
Word report generator for CTOD E1290 test results.

Populates a Word template with test data, results, and crack surface photos.
"""

import re
from pathlib import Path
from typing import Dict, Any, Optional, List
from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


class CTODReportGenerator:
    """
    Generate CTOD test reports from Word template or from scratch.

    Uses placeholder syntax {{placeholder_name}} in template.
    """

    def __init__(self, template_path: Optional[Path] = None):
        """
        Initialize report generator.

        Parameters
        ----------
        template_path : Path, optional
            Path to Word template file. If None, creates report from scratch.
        """
        self.template_path = template_path
        if template_path and not template_path.exists():
            raise FileNotFoundError(f"Template not found: {template_path}")

    def generate_report(
        self,
        output_path: Path,
        data: Dict[str, Any],
        chart_path: Optional[Path] = None,
        logo_path: Optional[Path] = None,
        photo_paths: Optional[List[Path]] = None
    ) -> Path:
        """
        Generate report by populating template with data.

        Parameters
        ----------
        output_path : Path
            Path for output Word document
        data : Dict[str, Any]
            Dictionary of placeholder values
        chart_path : Path, optional
            Path to Force vs CMOD chart image
        logo_path : Path, optional
            Path to logo image to insert
        photo_paths : List[Path], optional
            List of paths to crack surface photos

        Returns
        -------
        Path
            Path to generated report
        """
        # Create from scratch if no template
        if not self.template_path or not self.template_path.exists():
            doc = self._create_report_from_scratch(data, chart_path, logo_path, photo_paths)
            doc.save(output_path)
            return output_path

        doc = Document(self.template_path)

        # Replace placeholders in page headers
        for section in doc.sections:
            header = section.header
            for paragraph in header.paragraphs:
                self._replace_in_paragraph(paragraph, data, chart_path, logo_path, photo_paths)
            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_in_paragraph(paragraph, data, chart_path, logo_path, photo_paths)

        # Replace placeholders in paragraphs
        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph, data, chart_path, logo_path, photo_paths)

        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, data, chart_path, logo_path, photo_paths)

        # Add disclaimer to page footer (visible on all pages)
        from docx.shared import Pt
        disclaimer_text = (
            "All work and services carried out by Durabler are subject to, and conducted in accordance with, "
            "Durabler standard terms and conditions, which are available at durabler.se. This document shall not "
            "be reproduced other than in full, except with prior written approval of the issuer. The results pertain "
            "only to the item(s) as sampled by the client unless otherwise indicated. Durabler a part of Subseatec S AB, "
            "Address: Durabler C/O Subseatec, Dalavägen 23, 68130 Kristinehamn, SWEDEN"
        )
        for section in doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.clear()
            footer_run = footer_para.add_run(disclaimer_text)
            footer_run.font.size = Pt(7)
            footer_run.italic = True

        doc.save(output_path)
        return output_path

    def _create_report_from_scratch(
        self,
        data: Dict[str, Any],
        chart_path: Optional[Path],
        logo_path: Optional[Path],
        photo_paths: Optional[List[Path]]
    ) -> Document:
        """Create report without template - matches Vickers layout."""
        from docx.shared import RGBColor

        doc = Document()

        # Set compact paragraph spacing for entire document
        style = doc.styles['Normal']
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.0
        style.font.name = 'Calibri'
        style.font.size = Pt(10)

        # Set compact heading styles with dark green color
        dark_green = RGBColor(0x00, 0x64, 0x00)
        for i in range(1, 4):
            heading_style = doc.styles[f'Heading {i}']
            heading_style.paragraph_format.space_before = Pt(8)
            heading_style.paragraph_format.space_after = Pt(4)
            heading_style.font.color.rgb = dark_green

        # Add header with logo on left, certificate info on right (5-line layout)
        for section in doc.sections:
            # Set narrower margins for compact layout
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(2.0)
            section.right_margin = Cm(2.0)
            header = section.header
            header.is_linked_to_previous = False

            # Row 1: Logo - left aligned
            logo_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            logo_para.paragraph_format.space_after = Pt(0)
            if logo_path and logo_path.exists():
                logo_run = logo_para.add_run()
                logo_run.add_picture(str(logo_path), width=Cm(5.0))  # 50mm width

            # Row 2: Title - centered, font size 12
            title_para = header.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_para.paragraph_format.space_before = Pt(0)
            title_para.paragraph_format.space_after = Pt(0)
            title_run = title_para.add_run('CTOD Fracture Toughness Test Report')
            title_run.bold = True
            title_run.font.size = Pt(12)

            # Row 3: Standard - centered, font size 8
            std_para = header.add_paragraph()
            std_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            std_para.paragraph_format.space_before = Pt(0)
            std_para.paragraph_format.space_after = Pt(0)
            std_run = std_para.add_run('ASTM E1290 / E1820')
            std_run.font.size = Pt(8)

            # Row 4: Certificate - right aligned, font size 8
            cert_para = header.add_paragraph()
            cert_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cert_para.paragraph_format.space_before = Pt(0)
            cert_para.paragraph_format.space_after = Pt(0)
            cert_run = cert_para.add_run(f"Certificate: {data.get('certificate_number', '')}")
            cert_run.font.size = Pt(8)

            # Row 5: Date - right aligned, font size 8
            date_para = header.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            date_para.paragraph_format.space_before = Pt(0)
            date_para.paragraph_format.space_after = Pt(0)
            date_run = date_para.add_run(f"Date: {data.get('test_date', '')}")
            date_run.font.size = Pt(8)

        # Test Information table (exclude certificate and date - now in header)
        heading = doc.add_heading('Test Information', level=1)
        heading.paragraph_format.space_before = Pt(0)
        heading.paragraph_format.space_after = Pt(6)

        table = doc.add_table(rows=8, cols=4)
        table.style = 'Table Grid'

        # Two-column layout: Label | Value | Label | Value (exclude cert/date)
        info_data = [
            ('Test Project:', data.get('test_project', ''), 'Temperature:', f"{data.get('test_temperature', '23')} °C"),
            ('Customer:', data.get('customer', ''), 'Test Standard:', 'ASTM E1290'),
            ('Customer Order:', data.get('customer_order', ''), 'Test Equipment:', 'MTS Landmark 500kN'),
            ('Product S/N:', data.get('product_sn', ''), 'Specimen ID:', data.get('specimen_id', '')),
            ('Material:', data.get('material', ''), 'Location/Orientation:', data.get('location_orientation', '')),
            ('Notch Type:', data.get('notch_type', 'Fatigue pre-crack'), 'Side Grooves:', data.get('side_grooves', 'No')),
            ('Specimen Type:', data.get('specimen_type', 'SE(B)'), 'a₀/W Ratio:', data.get('a_W_ratio', '-')),
            ('Ligament (W-a₀):', data.get('ligament', '-'), '', ''),
        ]

        for i, (label1, value1, label2, value2) in enumerate(info_data):
            table.rows[i].cells[0].text = label1
            table.rows[i].cells[1].text = str(value1)
            table.rows[i].cells[2].text = label2
            table.rows[i].cells[3].text = str(value2)
            # Bold the labels
            if table.rows[i].cells[0].paragraphs[0].runs:
                table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            if table.rows[i].cells[2].paragraphs[0].runs:
                table.rows[i].cells[2].paragraphs[0].runs[0].bold = True

        # Compact table rows
        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].paragraph_format.space_before = Pt(1)
                cell.paragraphs[0].paragraph_format.space_after = Pt(1)

        # Specimen Geometry table
        heading = doc.add_heading('Specimen Geometry', level=1)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)

        table = doc.add_table(rows=7, cols=3)
        table.style = 'Table Grid'

        # Header row
        headers = ['Parameter', 'Value', 'Unit']
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

        dim_data = [
            ('Specimen Type', data.get('specimen_type', 'SE(B)'), '-'),
            ('Width W', data.get('W', ''), 'mm'),
            ('Thickness B', data.get('B', ''), 'mm'),
            ('Net Thickness Bₙ', data.get('B_n', ''), 'mm'),
            ('Crack length a₀', data.get('a_0', ''), 'mm'),
            ('Span S', data.get('S', '-'), 'mm'),
        ]

        for i, (param, value, unit) in enumerate(dim_data):
            table.rows[i+1].cells[0].text = param
            table.rows[i+1].cells[1].text = str(value)
            table.rows[i+1].cells[2].text = unit

        # Compact table rows
        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].paragraph_format.space_before = Pt(1)
                cell.paragraphs[0].paragraph_format.space_after = Pt(1)

        # Material Properties table
        heading = doc.add_heading('Material Properties', level=1)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)

        table = doc.add_table(rows=5, cols=3)
        table.style = 'Table Grid'

        # Header row
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

        mat_data = [
            ('Yield Strength σys', data.get('yield_strength', ''), 'MPa'),
            ('Ultimate Strength σu', data.get('ultimate_strength', '-'), 'MPa'),
            ("Young's Modulus E", data.get('youngs_modulus', ''), 'GPa'),
            ("Poisson's Ratio ν", data.get('poissons_ratio', '0.3'), '-'),
        ]

        for i, (param, value, unit) in enumerate(mat_data):
            table.rows[i+1].cells[0].text = param
            table.rows[i+1].cells[1].text = str(value)
            table.rows[i+1].cells[2].text = unit

        # Compact table rows
        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].paragraph_format.space_before = Pt(1)
                cell.paragraphs[0].paragraph_format.space_after = Pt(1)

        # Test Results table
        heading = doc.add_heading('Test Results', level=1)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)

        table = doc.add_table(rows=8, cols=5)
        table.style = 'Table Grid'

        # Header row - new order: Parameter, Unit, Value, Requirement, U (k=2)
        result_headers = ['Parameter', 'Unit', 'Value', 'Requirement', 'U (k=2)']
        for i, header in enumerate(result_headers):
            table.rows[0].cells[i].text = header
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

        # Get CTOD requirement from data
        ctod_req = data.get('ctod_req', data.get('delta_c_req', data.get('delta_u_req', data.get('delta_m_req', '-'))))

        # Results data: (Parameter, Unit, Value, Requirement, Uncertainty)
        results_data = [
            ('Pmax', 'kN', data.get('P_max_value', '-'), data.get('P_max_req', '-'), data.get('P_max_uncertainty', '-')),
            ('CMOD at Pmax', 'mm', data.get('CMOD_max_value', '-'), data.get('CMOD_max_req', '-'), data.get('CMOD_max_uncertainty', '-')),
            ('Kmax', 'MPa√m', data.get('K_max_value', '-'), data.get('K_max_req', '-'), data.get('K_max_uncertainty', '-')),
            (f"CTOD ({data.get('ctod_type', 'δ')})", 'mm',
             data.get('delta_c_value', data.get('delta_u_value', data.get('delta_m_value', '-'))),
             ctod_req,
             data.get('delta_c_uncertainty', data.get('delta_u_uncertainty', data.get('delta_m_uncertainty', '-')))),
            ('Compliance', 'mm/kN', data.get('compliance', '-'), '-', '-'),
            ('Crack Growth Δa', 'mm', data.get('delta_a', '-'), '-', '-'),
            ('Validity', '-', data.get('validity_status', '-'), '-', '-'),
        ]

        for i, (param, unit, value, req, unc) in enumerate(results_data):
            table.rows[i+1].cells[0].text = param
            table.rows[i+1].cells[1].text = unit
            table.rows[i+1].cells[2].text = str(value)
            table.rows[i+1].cells[3].text = str(req)
            table.rows[i+1].cells[4].text = str(unc)

        # Compact table rows
        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].paragraph_format.space_before = Pt(1)
                cell.paragraphs[0].paragraph_format.space_after = Pt(1)

        # Force-CMOD Plot
        if chart_path and chart_path.exists():
            heading = doc.add_heading('Force vs CMOD', level=1)
            heading.paragraph_format.space_before = Pt(12)
            heading.paragraph_format.space_after = Pt(6)
            doc.add_picture(str(chart_path), width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add caption
            caption = doc.add_paragraph('Figure 1: Force-CMOD curve per ASTM E1290')
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.runs[0].font.size = Pt(10)
            caption.runs[0].font.italic = True

        # Crack Surface Photos (if available)
        if photo_paths:
            heading = doc.add_heading('Crack Surface', level=1)
            heading.paragraph_format.space_before = Pt(12)
            heading.paragraph_format.space_after = Pt(6)
            for i, photo_path in enumerate(photo_paths):
                if photo_path and photo_path.exists():
                    doc.add_picture(str(photo_path), width=Inches(4.0))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add caption
            caption = doc.add_paragraph('Figure 2: Crack surface photographs')
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.runs[0].font.size = Pt(10)
            caption.runs[0].font.italic = True

        # Validity Assessment
        heading = doc.add_heading('Validity Assessment per ASTM E1290', level=1)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)

        validity_status = data.get('validity_status', '-')
        validity_statement = data.get('validity_statement', '')

        status_para = doc.add_paragraph()
        status_run = status_para.add_run(f"Status: {validity_status}")
        status_run.bold = True

        # Add validity checks table
        a_W_ratio = data.get('a_W_ratio', '-')
        validity_table = doc.add_table(rows=2, cols=4)
        validity_table.style = 'Table Grid'

        # Header row
        validity_headers = ['Check', 'Requirement', 'Actual', 'Result']
        for i, header in enumerate(validity_headers):
            validity_table.rows[0].cells[i].text = header
            validity_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

        # a/W ratio check
        try:
            a_W_val = float(a_W_ratio)
            a_W_pass = 0.45 <= a_W_val <= 0.70
            a_W_result = 'PASS' if a_W_pass else 'FAIL'
        except (ValueError, TypeError):
            a_W_result = 'N/A'

        validity_table.rows[1].cells[0].text = 'a₀/W Ratio'
        validity_table.rows[1].cells[1].text = '0.45 - 0.70'
        validity_table.rows[1].cells[2].text = str(a_W_ratio)
        validity_table.rows[1].cells[3].text = a_W_result

        # Compact table rows
        for row in validity_table.rows:
            for cell in row.cells:
                cell.paragraphs[0].paragraph_format.space_before = Pt(1)
                cell.paragraphs[0].paragraph_format.space_after = Pt(1)

        if validity_statement:
            doc.add_paragraph()  # Add spacing
            doc.add_paragraph(validity_statement)

        # Approval Signatures
        heading = doc.add_heading('Approval', level=1)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)

        sig_table = doc.add_table(rows=4, cols=4)
        sig_table.style = 'Table Grid'

        sig_headers = ['Role', 'Name', 'Signature', 'Date']
        for i, header in enumerate(sig_headers):
            sig_table.rows[0].cells[i].text = header
            sig_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

        sig_table.rows[1].cells[0].text = 'Tested by:'
        sig_table.rows[2].cells[0].text = 'Reviewed by:'
        sig_table.rows[3].cells[0].text = 'Approved by:'

        # Compact signature table rows
        for row in sig_table.rows:
            for cell in row.cells:
                cell.paragraphs[0].paragraph_format.space_before = Pt(1)
                cell.paragraphs[0].paragraph_format.space_after = Pt(1)

        # Add disclaimer to page footer (visible on all pages)
        disclaimer_text = (
            "All work and services carried out by Durabler are subject to, and conducted in accordance with, "
            "Durabler standard terms and conditions, which are available at durabler.se. This document shall not "
            "be reproduced other than in full, except with prior written approval of the issuer. The results pertain "
            "only to the item(s) as sampled by the client unless otherwise indicated. Durabler a part of Subseatec S AB, "
            "Address: Durabler C/O Subseatec, Dalavägen 23, 68130 Kristinehamn, SWEDEN"
        )
        for section in doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.clear()
            footer_run = footer_para.add_run(disclaimer_text)
            footer_run.font.size = Pt(7)
            footer_run.italic = True

        return doc

    def _replace_in_paragraph(
        self,
        paragraph,
        data: Dict[str, Any],
        chart_path: Optional[Path],
        logo_path: Optional[Path],
        photo_paths: Optional[List[Path]]
    ):
        """Replace placeholders in a paragraph."""
        full_text = paragraph.text

        # Handle logo placeholder
        if '{{logo}}' in full_text and logo_path and logo_path.exists():
            paragraph.clear()
            run = paragraph.add_run()
            run.add_picture(str(logo_path), height=Cm(1.5))
            return

        # Handle chart placeholder
        if '{{chart}}' in full_text and chart_path and chart_path.exists():
            paragraph.clear()
            run = paragraph.add_run()
            run.add_picture(str(chart_path), width=Inches(5.5))
            return

        # Handle photos placeholder
        if '{{photos}}' in full_text:
            paragraph.clear()
            if photo_paths:
                for i, photo_path in enumerate(photo_paths):
                    if photo_path and photo_path.exists():
                        run = paragraph.add_run()
                        # Add photo with half page width (3.5 inches)
                        run.add_picture(str(photo_path), width=Inches(3.5))
                        if i < len(photo_paths) - 1:
                            paragraph.add_run("\n")  # Newline between photos
            else:
                paragraph.add_run("No crack surface photos attached.")
            return

        # Find all placeholders
        placeholders = re.findall(r'\{\{([^}]+)\}\}', full_text)

        if not placeholders:
            return

        # Replace each placeholder
        for placeholder in placeholders:
            key = placeholder
            value = data.get(key, '')

            # Handle None values
            if value is None:
                value = ''

            # Format numeric values
            if isinstance(value, float):
                if abs(value) < 0.001:
                    value = f"{value:.4g}"
                elif abs(value) < 1:
                    value = f"{value:.4f}"
                else:
                    value = f"{value:.2f}"
            elif isinstance(value, int):
                value = str(value)

            full_text = full_text.replace(f'{{{{{placeholder}}}}}', str(value))

        # Update paragraph text while preserving formatting
        if paragraph.runs:
            first_run = paragraph.runs[0]
            font_name = first_run.font.name
            font_size = first_run.font.size
            bold = first_run.bold
            italic = first_run.italic

            paragraph.clear()
            new_run = paragraph.add_run(full_text)
            new_run.font.name = font_name
            new_run.font.size = font_size
            new_run.bold = bold
            new_run.italic = italic
        else:
            paragraph.text = full_text

    @staticmethod
    def prepare_report_data(
        test_info: Dict[str, str],
        specimen_data: Dict[str, Any],
        material_data: Dict[str, Any],
        results: Dict[str, Any],
        crack_measurements: Optional[List[float]] = None
    ) -> Dict[str, Any]:
        """
        Prepare data dictionary for report generation.

        Parameters
        ----------
        test_info : Dict[str, str]
            Test information (project, customer, etc.)
        specimen_data : Dict[str, Any]
            Specimen geometry data
        material_data : Dict[str, Any]
            Material properties
        results : Dict[str, Any]
            Analysis results with CTODResult and MeasuredValue objects
        crack_measurements : List[float], optional
            9-point crack length measurements

        Returns
        -------
        Dict[str, Any]
            Flattened dictionary ready for template
        """
        data = {}

        # Test information
        data['test_project'] = test_info.get('test_project', '')
        data['customer'] = test_info.get('customer', '')
        data['customer_order'] = test_info.get('customer_order', '')
        data['product_sn'] = test_info.get('product_sn', '')
        data['specimen_id'] = test_info.get('specimen_id', '')
        data['location_orientation'] = test_info.get('location_orientation', '')
        data['material'] = test_info.get('material', '')
        data['certificate_number'] = test_info.get('certificate_number', '')
        data['test_date'] = test_info.get('test_date', '')
        data['test_standard'] = 'ASTM E1290'
        data['test_equipment'] = 'MTS Landmark 500kN'
        data['test_temperature'] = test_info.get('temperature', '23')

        # Specimen geometry
        data['specimen_type'] = specimen_data.get('specimen_type', 'SE(B)')
        data['W'] = specimen_data.get('W', '')
        data['B'] = specimen_data.get('B', '')
        data['B_n'] = specimen_data.get('B_n', specimen_data.get('B', ''))
        data['a_0'] = specimen_data.get('a_0', '')
        data['S'] = specimen_data.get('S', '')
        data['a_f'] = specimen_data.get('a_f', '-')
        data['notch_type'] = specimen_data.get('notch_type', 'Fatigue pre-crack')
        data['side_grooves'] = 'Yes' if specimen_data.get('B_n') and specimen_data.get('B_n') < specimen_data.get('B', 0) else 'No'

        # Calculate a_0/W ratio and delta_a
        W = float(specimen_data.get('W', 1))
        a_0 = float(specimen_data.get('a_0', 0))
        data['a_W_ratio'] = f"{a_0 / W:.3f}" if W > 0 else '-'

        a_f = specimen_data.get('a_f')
        if a_f and a_f != '-':
            try:
                delta_a = float(a_f) - a_0
                data['delta_a'] = f"{delta_a:.3f}"
            except (ValueError, TypeError):
                data['delta_a'] = '-'
        else:
            data['delta_a'] = '-'

        # Material properties
        data['yield_strength'] = material_data.get('yield_strength', '')
        data['ultimate_strength'] = material_data.get('ultimate_strength', '')
        data['youngs_modulus'] = material_data.get('youngs_modulus', '')
        data['poissons_ratio'] = material_data.get('poissons_ratio', '0.3')

        # Helper to extract value and uncertainty from MeasuredValue objects
        def get_measured(obj, default_value='-', default_unc='-'):
            if obj and hasattr(obj, 'value'):
                return f"{obj.value:.4g}", f"±{obj.uncertainty:.4g}"
            return default_value, default_unc

        # Results from MeasuredValue objects
        P_max = results.get('P_max')
        if P_max:
            data['P_max_value'], data['P_max_uncertainty'] = get_measured(P_max)
        else:
            data['P_max_value'] = '-'
            data['P_max_uncertainty'] = '-'
        data['P_max_req'] = '-'  # Requirement placeholder

        CMOD_max = results.get('CMOD_max')
        if CMOD_max:
            data['CMOD_max_value'], data['CMOD_max_uncertainty'] = get_measured(CMOD_max)
        else:
            data['CMOD_max_value'] = '-'
            data['CMOD_max_uncertainty'] = '-'
        data['CMOD_max_req'] = '-'  # Requirement placeholder

        K_max = results.get('K_max')
        if K_max:
            data['K_max_value'], data['K_max_uncertainty'] = get_measured(K_max)
        else:
            data['K_max_value'] = '-'
            data['K_max_uncertainty'] = '-'
        data['K_max_req'] = '-'  # Requirement placeholder

        # Parse CTOD requirement from test_info
        requirement_str = test_info.get('requirement', '')
        ctod_req = '-'
        if requirement_str:
            # Try to parse CTOD requirement from various formats
            import re
            # Match patterns like "CTOD: >0.15", "δ: ≥0.15mm", "CTOD >0.15 mm", etc.
            patterns = [
                r'CTOD[:\s]*([>≥<≤]?\s*[\d.]+)\s*(?:mm)?',
                r'δ[:\s]*([>≥<≤]?\s*[\d.]+)\s*(?:mm)?',
                r'delta[:\s]*([>≥<≤]?\s*[\d.]+)\s*(?:mm)?',
            ]
            for pattern in patterns:
                match = re.search(pattern, requirement_str, re.IGNORECASE)
                if match:
                    ctod_req = match.group(1).strip()
                    if not ctod_req.startswith(('>','<','≥','≤')):
                        ctod_req = '>' + ctod_req  # Default to minimum requirement
                    break

        # CTOD results (CTODResult objects)
        ctod_type_reported = None

        for ctod_key, ctod_label in [('delta_c', 'δc'), ('delta_u', 'δu'), ('delta_m', 'δm')]:
            ctod_result = results.get(ctod_key)
            if ctod_result:
                val, unc = get_measured(ctod_result.ctod_value)
                data[f'{ctod_key}_value'] = val
                data[f'{ctod_key}_uncertainty'] = unc
                data[f'{ctod_key}_valid'] = 'VALID' if ctod_result.is_valid else 'INVALID'

                # Set reported type (prefer δc > δu > δm)
                if ctod_type_reported is None:
                    ctod_type_reported = ctod_label
                    data['ctod_req'] = ctod_req  # Set requirement for the reported CTOD type
            else:
                data[f'{ctod_key}_value'] = '-'
                data[f'{ctod_key}_uncertainty'] = '-'
                data[f'{ctod_key}_valid'] = '-'
            data[f'{ctod_key}_req'] = ctod_req if ctod_type_reported == ctod_label else '-'

        data['ctod_type'] = ctod_type_reported or '-'
        # Ensure ctod_req is set even if no CTOD result found
        if 'ctod_req' not in data:
            data['ctod_req'] = ctod_req

        # Compliance
        compliance = results.get('compliance')
        if compliance:
            data['compliance'] = f"{compliance:.4f}"
        else:
            data['compliance'] = '-'

        # 9-point crack measurements
        if crack_measurements and len(crack_measurements) >= 9:
            for i, val in enumerate(crack_measurements[:9], 1):
                data[f'a{i}'] = f"{val:.2f}"
        else:
            for i in range(1, 10):
                data[f'a{i}'] = '-'

        # Validity
        is_valid = results.get('is_valid', False)
        data['validity_status'] = 'VALID' if is_valid else 'INVALID'
        validity_summary = results.get('validity_summary', '')
        # Convert Unicode symbols to plain text for Word compatibility
        if validity_summary:
            validity_summary = validity_summary.replace('✓', '[PASS]').replace('✗', '[FAIL]').replace('⚠', '[NOTE]')
        if is_valid:
            data['validity_statement'] = 'The test meets all validity requirements of ASTM E1290.'
        else:
            data['validity_statement'] = f'The test does not meet all validity requirements.\n{validity_summary}'

        # Signatures (to be filled manually)
        data['tested_by'] = ''
        data['tested_date'] = ''
        data['reviewed_by'] = ''
        data['reviewed_date'] = ''
        data['approved_by'] = ''
        data['approved_date'] = ''

        return data
