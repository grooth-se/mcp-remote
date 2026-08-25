"""Shared report header/footer/page-size layout for all test-method reports.

Every Durabler test report uses the same page furniture:

* A4 paper (210 x 297 mm) with compact 15/20 mm margins.
* A header with two logos side by side — the Durabler logo on the left and
  the Subseatec logo on the right, with the certificate number and date beneath
  the Subseatec logo — followed by a centred report title and test standard.
* A footer whose text is the standard Durabler disclaimer, preceded by the
  measurement-uncertainty statement and followed by the company org. number.

Keeping this in one place guarantees every report method renders identically.
"""
from pathlib import Path
from typing import Optional, Union

from docx.shared import Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


# Footer text pieces
FOOTER_PREFIX = (
    "When assessing results, the measurement of uncertainty has not been "
    "taken under consideration. "
)
FOOTER_SUFFIX = " Org. No SE556782-8255"
DISCLAIMER_BODY = (
    "All work and services carried out by Durabler are subject to, and conducted in accordance with, "
    "Durabler standard terms and conditions, which are available at durabler.se. This document shall not "
    "be reproduced other than in full, except with prior written approval of the issuer. The results pertain "
    "only to the item(s) as sampled by the client unless otherwise indicated. Durabler a part of Subseatec S AB, "
    "Address: Durabler C/O Subseatec, Dalavägen 23, 68130 Kristinehamn, SWEDEN"
)

# A4 dimensions
A4_WIDTH = Cm(21.0)
A4_HEIGHT = Cm(29.7)

PathLike = Union[str, Path, None]


def _exists(p: PathLike) -> bool:
    return bool(p) and Path(p).exists()


def resolve_logo_right(logo_left: PathLike, logo_right: PathLike) -> PathLike:
    """Derive the Subseatec (right) logo from the Durabler (left) logo folder
    when it is not supplied explicitly."""
    if logo_right:
        return logo_right
    if logo_left:
        candidate = Path(logo_left).parent / 'subseatec_logo.png'
        if candidate.exists():
            return candidate
    return None


def apply_a4(doc) -> None:
    """Set every section of the document to A4."""
    for section in doc.sections:
        section.page_width = A4_WIDTH
        section.page_height = A4_HEIGHT


def build_report_header(
    doc,
    title: str,
    standard: str,
    cert_number: str,
    report_date: str,
    logo_left: PathLike = None,
    logo_right: PathLike = None,
) -> None:
    """Build the standard two-logo header (and A4 page/margins) on all sections.

    Durabler logo on the left; Subseatec logo on the right with the certificate
    number and date beneath it; then a centred title and standard.
    """
    logo_right = resolve_logo_right(logo_left, logo_right)
    for section in doc.sections:
        # A4 paper + compact margins
        section.page_width = A4_WIDTH
        section.page_height = A4_HEIGHT
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)

        header = section.header
        header.is_linked_to_previous = False

        # Borderless 2-column table: Durabler logo | Subseatec logo + cert + date
        htable = header.add_table(rows=1, cols=2, width=Cm(17.0))
        htable.autofit = False
        left_cell, right_cell = htable.rows[0].cells
        left_cell.width = Cm(8.5)
        right_cell.width = Cm(8.5)

        # Left: Durabler logo
        lcell_p = left_cell.paragraphs[0]
        lcell_p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        lcell_p.paragraph_format.space_after = Pt(0)
        if _exists(logo_left):
            lcell_p.add_run().add_picture(str(logo_left), width=Cm(4.5))

        # Right: Subseatec logo, then Certificate, then Date (right-aligned)
        rcell_p = right_cell.paragraphs[0]
        rcell_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rcell_p.paragraph_format.space_after = Pt(0)
        if _exists(logo_right):
            rcell_p.add_run().add_picture(str(logo_right), width=Cm(4.5))
        rcert_p = right_cell.add_paragraph()
        rcert_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rcert_p.paragraph_format.space_before = Pt(0)
        rcert_p.paragraph_format.space_after = Pt(0)
        rcert_run = rcert_p.add_run(f"Certificate: {cert_number or ''}")
        rcert_run.font.size = Pt(8)
        rdate_p = right_cell.add_paragraph()
        rdate_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rdate_p.paragraph_format.space_before = Pt(0)
        rdate_p.paragraph_format.space_after = Pt(0)
        rdate_run = rdate_p.add_run(f"Date: {report_date or ''}")
        rdate_run.font.size = Pt(8)

        # Drop the empty leading paragraph so the table sits at the very top
        first_p = header.paragraphs[0]
        if not first_p.text and not first_p.runs:
            first_p._element.getparent().remove(first_p._element)

        # Title - centered
        title_para = header.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_before = Pt(2)
        title_para.paragraph_format.space_after = Pt(0)
        title_run = title_para.add_run(title)
        title_run.bold = True
        title_run.font.size = Pt(12)

        # Standard - centered
        std_para = header.add_paragraph()
        std_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        std_para.paragraph_format.space_before = Pt(0)
        std_para.paragraph_format.space_after = Pt(0)
        std_run = std_para.add_run(standard or '')
        std_run.font.size = Pt(8)


def apply_standard_footer(doc, body: Optional[str] = None) -> None:
    """Set the standard footer (uncertainty prefix + disclaimer + org. no.) on
    every section."""
    text = FOOTER_PREFIX + (body if body is not None else DISCLAIMER_BODY) + FOOTER_SUFFIX
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.clear()
        footer_run = footer_para.add_run(text)
        footer_run.font.size = Pt(7)
        footer_run.italic = True
