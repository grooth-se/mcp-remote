"""
Word report generator for FCGR E647 test results.

Populates a Word template with test data, Paris law results, plots, and photos.
Uses same layout style as CTOD E1290 and KIC E399 reports for consistency.
"""

import re
from pathlib import Path
from typing import Dict, Any, Optional, List
from datetime import datetime

from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK


class FCGRReportGenerator:
    """
    Generate FCGR test reports from Word template or from scratch.

    Uses placeholder syntax {{placeholder_name}} in template.
    """

    def __init__(self, template_path: Optional[Path] = None):
        """
        Initialize report generator.

        Parameters
        ----------
        template_path : Path, optional
            Path to Word template file. If None, creates report from scratch.
        """
        self.template_path = template_path
        if template_path and not template_path.exists():
            raise FileNotFoundError(f"Template not found: {template_path}")

    def generate_report(
        self,
        output_path: Path,
        data: Dict[str, Any],
        plot1_path: Optional[Path] = None,
        plot2_path: Optional[Path] = None,
        logo_path: Optional[Path] = None,
        photo_paths: Optional[List[Path]] = None
    ) -> Path:
        """
        Generate report by populating template with data or creating from scratch.

        Parameters
        ----------
        output_path : Path
            Path for output Word document
        data : Dict[str, Any]
            Dictionary of placeholder values
        plot1_path : Path, optional
            Path to Crack Length vs Cycles plot image
        plot2_path : Path, optional
            Path to da/dN vs Delta-K (Paris law) plot image
        logo_path : Path, optional
            Path to logo image to insert
        photo_paths : List[Path], optional
            List of paths to crack surface photos

        Returns
        -------
        Path
            Path to generated report
        """
        # Create from scratch if no template
        if not self.template_path or not self.template_path.exists():
            doc = self._create_report_from_scratch(data, plot1_path, plot2_path, logo_path, photo_paths)
            doc.save(output_path)
            return output_path

        doc = Document(self.template_path)

        # Replace placeholders in page headers
        for section in doc.sections:
            header = section.header
            for paragraph in header.paragraphs:
                self._replace_in_paragraph(paragraph, data, plot1_path, plot2_path, logo_path, photo_paths)
            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_in_paragraph(paragraph, data, plot1_path, plot2_path, logo_path, photo_paths)

        # Replace placeholders in paragraphs
        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph, data, plot1_path, plot2_path, logo_path, photo_paths)

        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, data, plot1_path, plot2_path, logo_path, photo_paths)

        doc.save(output_path)
        return output_path

    def _create_report_from_scratch(
        self,
        data: Dict[str, Any],
        plot1_path: Optional[Path],
        plot2_path: Optional[Path],
        logo_path: Optional[Path],
        photo_paths: Optional[List[Path]]
    ) -> Document:
        """Create report without template - matches Vickers/CTOD layout."""
        from docx.shared import RGBColor

        doc = Document()

        # Set compact paragraph spacing for entire document
        style = doc.styles['Normal']
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.0
        style.font.name = 'Calibri'
        style.font.size = Pt(10)

        # Set compact heading styles with dark green color
        dark_green = RGBColor(0x00, 0x64, 0x00)
        for i in range(1, 4):
            heading_style = doc.styles[f'Heading {i}']
            heading_style.paragraph_format.space_before = Pt(8)
            heading_style.paragraph_format.space_after = Pt(4)
            heading_style.font.color.rgb = dark_green

        # Add header with logo on left, certificate info on right (5-line layout)
        for section in doc.sections:
            # Set narrower margins for compact layout
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(2.0)
            section.right_margin = Cm(2.0)
            header = section.header
            header.is_linked_to_previous = False

            # Row 1: Logo - left aligned
            logo_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            logo_para.paragraph_format.space_after = Pt(0)
            if logo_path and logo_path.exists():
                logo_run = logo_para.add_run()
                logo_run.add_picture(str(logo_path), width=Cm(5.0))  # 50mm width

            # Row 2: Title - centered, font size 12
            title_para = header.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_para.paragraph_format.space_before = Pt(0)
            title_para.paragraph_format.space_after = Pt(0)
            title_run = title_para.add_run('Fatigue Crack Growth Rate Test Report')
            title_run.bold = True
            title_run.font.size = Pt(12)

            # Row 3: Standard - centered, font size 8
            std_para = header.add_paragraph()
            std_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            std_para.paragraph_format.space_before = Pt(0)
            std_para.paragraph_format.space_after = Pt(0)
            std_run = std_para.add_run('ASTM E647')
            std_run.font.size = Pt(8)

            # Row 4: Certificate - right aligned, font size 8
            cert_para = header.add_paragraph()
            cert_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cert_para.paragraph_format.space_before = Pt(0)
            cert_para.paragraph_format.space_after = Pt(0)
            cert_run = cert_para.add_run(f"Certificate: {data.get('certificate_number', '')}")
            cert_run.font.size = Pt(8)

            # Row 5: Date - right aligned, font size 8
            date_para = header.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            date_para.paragraph_format.space_before = Pt(0)
            date_para.paragraph_format.space_after = Pt(0)
            date_run = date_para.add_run(f"Date: {data.get('test_date', '')}")
            date_run.font.size = Pt(8)

        # Test Information table (exclude certificate and date - now in header)
        heading = doc.add_heading('Test Information', level=1)
        heading.paragraph_format.space_before = Pt(0)
        heading.paragraph_format.space_after = Pt(6)

        table = doc.add_table(rows=7, cols=4)
        table.style = 'Table Grid'

        # Two-column layout: Label | Value | Label | Value (exclude cert/date)
        info_data = [
            ('Test Project:', data.get('test_project', ''), 'Temperature:', f"{data.get('temperature', '23')} °C"),
            ('Customer:', data.get('customer', ''), 'Test Standard:', 'ASTM E647'),
            ('Customer Order:', data.get('customer_order', ''), 'Test Equipment:', 'MTS Landmark 500kN'),
            ('Product S/N:', data.get('product_sn', ''), 'Specimen ID:', data.get('specimen_id', '')),
            ('Material:', data.get('material', ''), 'Location/Orientation:', data.get('location_orientation', '')),
            ('Specimen Type:', data.get('specimen_type', 'C(T)'), 'Side Grooves:', data.get('side_grooves', 'No')),
            ('Environment:', data.get('environment', 'Laboratory Air'), 'a₀/W Ratio:', data.get('a_W_ratio', '-')),
        ]

        for i, (label1, value1, label2, value2) in enumerate(info_data):
            table.rows[i].cells[0].text = label1
            table.rows[i].cells[1].text = str(value1)
            table.rows[i].cells[2].text = label2
            table.rows[i].cells[3].text = str(value2)
            # Bold the labels
            if table.rows[i].cells[0].paragraphs[0].runs:
                table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            if table.rows[i].cells[2].paragraphs[0].runs:
                table.rows[i].cells[2].paragraphs[0].runs[0].bold = True

        # Compact table rows
        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].paragraph_format.space_before = Pt(1)
                cell.paragraphs[0].paragraph_format.space_after = Pt(1)

        # Specimen Geometry table
        heading = doc.add_heading('Specimen Geometry', level=1)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)

        table = doc.add_table(rows=6, cols=3)
        table.style = 'Table Grid'

        # Header row
        headers = ['Parameter', 'Value', 'Unit']
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

        dim_data = [
            ('Specimen Type', data.get('specimen_type', 'C(T)'), '-'),
            ('Width W', data.get('W', ''), 'mm'),
            ('Thickness B', data.get('B', ''), 'mm'),
            ('Net Thickness Bₙ', data.get('B_n', data.get('B', '')), 'mm'),
            ('Initial Crack a₀', data.get('a_0', ''), 'mm'),
        ]

        for i, (param, value, unit) in enumerate(dim_data):
            table.rows[i+1].cells[0].text = param
            table.rows[i+1].cells[1].text = str(value)
            table.rows[i+1].cells[2].text = unit

        # Compact table rows
        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].paragraph_format.space_before = Pt(1)
                cell.paragraphs[0].paragraph_format.space_after = Pt(1)

        # Material Properties table
        heading = doc.add_heading('Material Properties', level=1)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)

        table = doc.add_table(rows=5, cols=3)
        table.style = 'Table Grid'

        # Header row
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

        mat_data = [
            ('Yield Strength σys', data.get('yield_strength', ''), 'MPa'),
            ('Ultimate Strength σu', data.get('ultimate_strength', '-'), 'MPa'),
            ("Young's Modulus E", data.get('youngs_modulus', ''), 'GPa'),
            ("Poisson's Ratio ν", data.get('poissons_ratio', '0.3'), '-'),
        ]

        for i, (param, value, unit) in enumerate(mat_data):
            table.rows[i+1].cells[0].text = param
            table.rows[i+1].cells[1].text = str(value)
            table.rows[i+1].cells[2].text = unit

        # Compact table rows
        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].paragraph_format.space_before = Pt(1)
                cell.paragraphs[0].paragraph_format.space_after = Pt(1)

        # Test Parameters table (FCGR-specific)
        heading = doc.add_heading('Test Parameters', level=1)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)

        table = doc.add_table(rows=6, cols=3)
        table.style = 'Table Grid'

        # Header row
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

        param_data = [
            ('Load Ratio R', data.get('load_ratio', '0.1'), '-'),
            ('Frequency f', data.get('frequency', '10'), 'Hz'),
            ('Control Mode', data.get('control_mode', 'Load Control'), '-'),
            ('Wave Shape', data.get('wave_shape', 'Sine'), '-'),
            ('da/dN Method', data.get('dadn_method', 'Secant'), '-'),
        ]

        for i, (param, value, unit) in enumerate(param_data):
            table.rows[i+1].cells[0].text = param
            table.rows[i+1].cells[1].text = str(value)
            table.rows[i+1].cells[2].text = unit

        # Compact table rows
        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].paragraph_format.space_before = Pt(1)
                cell.paragraphs[0].paragraph_format.space_after = Pt(1)

        # Paris Law Results table
        heading = doc.add_heading('Paris Law Results', level=1)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)

        table = doc.add_table(rows=9, cols=5)
        table.style = 'Table Grid'

        # Header row - new order: Parameter, Unit, Value, Requirement, U (k=2)
        result_headers = ['Parameter', 'Unit', 'Value', 'Requirement', 'U (k=2)']
        for i, header in enumerate(result_headers):
            table.rows[0].cells[i].text = header
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

        # Format Paris law coefficient C in scientific notation
        paris_C = data.get('paris_C', '-')
        paris_C_error = data.get('paris_C_error', '-')
        if isinstance(paris_C, (int, float)) and paris_C != '-':
            paris_C = f"{paris_C:.4e}"
        if isinstance(paris_C_error, (int, float)) and paris_C_error != '-':
            paris_C_error = f"±{paris_C_error:.4e}"

        paris_m = data.get('paris_m', '-')
        paris_m_error = data.get('paris_m_error', '-')
        if isinstance(paris_m, (int, float)) and paris_m != '-':
            paris_m = f"{paris_m:.4f}"
        if isinstance(paris_m_error, (int, float)) and paris_m_error != '-':
            paris_m_error = f"±{paris_m_error:.4f}"

        # Get requirement from data (parsed from certificate)
        requirement = data.get('requirement', '-') or '-'

        # Format: (Parameter, Unit, Value, Requirement, Uncertainty)
        results_data = [
            ('Paris Coefficient C', 'm/cycle/(MPa√m)^m', paris_C, requirement, paris_C_error),
            ('Paris Exponent m', '-', paris_m, '-', paris_m_error),
            ('R² (fit quality)', '-', self._format_value(data.get('paris_r_squared', '-')), '-', '-'),
            ('Valid Data Points', '-', str(data.get('paris_n_points', '-')), '-', '-'),
            ('ΔK Range (min)', 'MPa√m', self._format_value(data.get('delta_K_min', '-')), '-', '-'),
            ('ΔK Range (max)', 'MPa√m', self._format_value(data.get('delta_K_max', '-')), '-', '-'),
            ('Total Cycles', '-', str(data.get('total_cycles', '-')), '-', '-'),
            ('Final Crack Length', 'mm', self._format_value(data.get('final_crack_length', '-')), '-', '-'),
        ]

        for i, (param, unit, value, req, unc) in enumerate(results_data):
            table.rows[i+1].cells[0].text = param
            table.rows[i+1].cells[1].text = unit
            table.rows[i+1].cells[2].text = str(value)
            table.rows[i+1].cells[3].text = str(req)
            table.rows[i+1].cells[4].text = str(unc)

        # Compact table rows
        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].paragraph_format.space_before = Pt(1)
                cell.paragraphs[0].paragraph_format.space_after = Pt(1)

        # Data Quality table
        heading = doc.add_heading('Data Quality', level=1)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)

        table = doc.add_table(rows=4, cols=2)
        table.style = 'Table Grid'

        quality_data = [
            ('Valid Data Points:', str(data.get('n_valid_points', '-'))),
            ('Outliers Removed:', str(data.get('n_outliers', '-'))),
            ('Outlier Threshold:', f"{data.get('outlier_threshold', '2.5')}σ"),
        ]

        for i, (label, value) in enumerate(quality_data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value
            if table.rows[i].cells[0].paragraphs[0].runs:
                table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

        # Add totals row
        table.rows[3].cells[0].text = 'Total Data Points:'
        n_valid = data.get('n_valid_points', 0)
        n_outliers = data.get('n_outliers', 0)
        try:
            total = int(n_valid) + int(n_outliers)
            table.rows[3].cells[1].text = str(total)
        except (ValueError, TypeError):
            table.rows[3].cells[1].text = '-'
        if table.rows[3].cells[0].paragraphs[0].runs:
            table.rows[3].cells[0].paragraphs[0].runs[0].bold = True

        # Compact table rows
        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].paragraph_format.space_before = Pt(1)
                cell.paragraphs[0].paragraph_format.space_after = Pt(1)

        # Plot 1: Crack Length vs Cycles
        if plot1_path and plot1_path.exists():
            heading = doc.add_heading('Crack Length vs Cycles', level=1)
            heading.paragraph_format.space_before = Pt(12)
            heading.paragraph_format.space_after = Pt(6)
            doc.add_picture(str(plot1_path), width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add caption
            caption = doc.add_paragraph('Figure 1: Crack length as a function of fatigue cycles')
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.runs[0].font.size = Pt(10)
            caption.runs[0].font.italic = True

        # Plot 2: da/dN vs Delta-K (Paris Law)
        if plot2_path and plot2_path.exists():
            heading = doc.add_heading('da/dN vs ΔK (Paris Law)', level=1)
            heading.paragraph_format.space_before = Pt(12)
            heading.paragraph_format.space_after = Pt(6)
            doc.add_picture(str(plot2_path), width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add caption
            caption = doc.add_paragraph('Figure 2: Fatigue crack growth rate (da/dN) vs stress intensity factor range (ΔK)')
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.runs[0].font.size = Pt(10)
            caption.runs[0].font.italic = True

        # Crack Surface Photos (if available)
        if photo_paths:
            heading = doc.add_heading('Crack Surface', level=1)
            heading.paragraph_format.space_before = Pt(12)
            heading.paragraph_format.space_after = Pt(6)
            for i, photo_path in enumerate(photo_paths):
                if photo_path and photo_path.exists():
                    doc.add_picture(str(photo_path), width=Inches(4.0))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add caption
            caption = doc.add_paragraph('Figure 3: Crack surface photographs')
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.runs[0].font.size = Pt(10)
            caption.runs[0].font.italic = True

        # Validity Assessment
        heading = doc.add_heading('Validity Assessment per ASTM E647', level=1)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)

        validity_status = data.get('is_valid', '-')

        # Validity statement paragraph
        if validity_status == 'VALID':
            validity_text = "The test meets all validity requirements of ASTM E647. The Paris law coefficients are valid for the tested ΔK range."
        else:
            validity_text = "The test data should be reviewed. See validity notes below."

        status_para = doc.add_paragraph()
        status_run = status_para.add_run(f"Status: {validity_status}")
        status_run.bold = True

        doc.add_paragraph(validity_text)

        # List validity notes if any
        validity_notes = data.get('validity_notes', '')
        if validity_notes and validity_notes != 'All checks passed':
            notes_para = doc.add_paragraph()
            notes_para.add_run("Validity Notes:").bold = True
            for note in validity_notes.split('\n'):
                if note.strip():
                    doc.add_paragraph(f"• {note.strip()}")

        # Approval Signatures
        heading = doc.add_heading('Approval', level=1)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)

        sig_table = doc.add_table(rows=4, cols=4)
        sig_table.style = 'Table Grid'

        sig_headers = ['Role', 'Name', 'Signature', 'Date']
        for i, header in enumerate(sig_headers):
            sig_table.rows[0].cells[i].text = header
            sig_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

        sig_table.rows[1].cells[0].text = 'Tested by:'
        sig_table.rows[2].cells[0].text = 'Reviewed by:'
        sig_table.rows[3].cells[0].text = 'Approved by:'

        # Compact signature table rows
        for row in sig_table.rows:
            for cell in row.cells:
                cell.paragraphs[0].paragraph_format.space_before = Pt(1)
                cell.paragraphs[0].paragraph_format.space_after = Pt(1)

        # Add disclaimer to page footer (visible on all pages)
        disclaimer_text = (
            "All work and services carried out by Durabler are subject to, and conducted in accordance with, "
            "Durabler standard terms and conditions, which are available at durabler.se. This document shall not "
            "be reproduced other than in full, except with prior written approval of the issuer. The results pertain "
            "only to the item(s) as sampled by the client unless otherwise indicated. Durabler a part of Subseatec S AB, "
            "Address: Durabler C/O Subseatec, Dalavägen 23, 68130 Kristinehamn, SWEDEN"
        )
        for section in doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.clear()
            footer_run = footer_para.add_run(disclaimer_text)
            footer_run.font.size = Pt(7)
            footer_run.italic = True

        return doc

    def _format_value(self, value) -> str:
        """Format a value for display."""
        if value is None or value == '-':
            return '-'
        if isinstance(value, float):
            if abs(value) < 0.0001 or abs(value) > 10000:
                return f"{value:.4e}"
            elif abs(value) < 1:
                return f"{value:.4f}"
            else:
                return f"{value:.2f}"
        return str(value)

    def _replace_in_paragraph(
        self,
        paragraph,
        data: Dict[str, Any],
        plot1_path: Optional[Path],
        plot2_path: Optional[Path],
        logo_path: Optional[Path],
        photo_paths: Optional[List[Path]]
    ):
        """Replace placeholders in a paragraph."""
        full_text = paragraph.text

        # Handle logo placeholder
        if '{{logo}}' in full_text and logo_path and logo_path.exists():
            paragraph.clear()
            run = paragraph.add_run()
            run.add_picture(str(logo_path), height=Cm(1.5))
            return

        # Handle plot1 placeholder (Crack Length vs Cycles)
        if '{{plot1}}' in full_text and plot1_path and plot1_path.exists():
            paragraph.clear()
            run = paragraph.add_run()
            run.add_picture(str(plot1_path), width=Inches(6.5))
            return

        # Handle plot2 placeholder (da/dN vs Delta-K)
        if '{{plot2}}' in full_text and plot2_path and plot2_path.exists():
            paragraph.clear()
            run = paragraph.add_run()
            run.add_picture(str(plot2_path), width=Inches(6.5))
            return

        # Handle combined plots placeholder (each plot on its own row, full width)
        if '{{plots}}' in full_text:
            paragraph.clear()
            if plot1_path and plot1_path.exists():
                run = paragraph.add_run()
                run.add_picture(str(plot1_path), width=Inches(6.5))
            if plot2_path and plot2_path.exists():
                # Add line breaks to separate plots vertically
                run = paragraph.add_run()
                run.add_break(WD_BREAK.LINE)
                run.add_break(WD_BREAK.LINE)
                run = paragraph.add_run()
                run.add_picture(str(plot2_path), width=Inches(6.5))
            return

        # Handle photos placeholder
        if '{{photos}}' in full_text:
            paragraph.clear()
            if photo_paths:
                for i, photo_path in enumerate(photo_paths):
                    if photo_path and photo_path.exists():
                        run = paragraph.add_run()
                        run.add_picture(str(photo_path), width=Inches(5.5))
                        paragraph.add_run()  # Line break between photos
            else:
                paragraph.add_run("No crack surface photos attached.")
            return

        # Find all placeholders
        placeholders = re.findall(r'\{\{([^}]+)\}\}', full_text)

        if not placeholders:
            return

        # Replace each placeholder
        for placeholder in placeholders:
            key = placeholder
            value = data.get(key, '')

            # Handle None values
            if value is None:
                value = ''

            # Format numeric values
            if isinstance(value, float):
                if abs(value) < 1e-10:
                    value = "0"
                elif abs(value) < 0.0001 or abs(value) > 10000:
                    value = f"{value:.4e}"
                elif abs(value) < 1:
                    value = f"{value:.4f}"
                else:
                    value = f"{value:.2f}"
            elif isinstance(value, int):
                value = str(value)

            full_text = full_text.replace(f'{{{{{placeholder}}}}}', str(value))

        # Update paragraph text while preserving formatting
        if paragraph.runs:
            first_run = paragraph.runs[0]
            font_name = first_run.font.name
            font_size = first_run.font.size
            bold = first_run.bold
            italic = first_run.italic

            paragraph.clear()
            new_run = paragraph.add_run(full_text)
            new_run.font.name = font_name
            new_run.font.size = font_size
            new_run.bold = bold
            new_run.italic = italic
        else:
            paragraph.text = full_text

    @staticmethod
    def prepare_report_data(
        test_info: Dict[str, str],
        specimen_data: Dict[str, Any],
        material_data: Dict[str, Any],
        test_params: Dict[str, Any],
        results: Any,
        validity_notes: Optional[List[str]] = None
    ) -> Dict[str, Any]:
        """
        Prepare data dictionary for report generation.

        Parameters
        ----------
        test_info : Dict[str, str]
            Test information (project, customer, certificate, etc.)
        specimen_data : Dict[str, Any]
            Specimen geometry data
        material_data : Dict[str, Any]
            Material properties
        test_params : Dict[str, Any]
            Test parameters (R ratio, frequency, control mode, etc.)
        results : FCGRResult
            Analysis results with Paris law coefficients
        validity_notes : List[str], optional
            List of validity check notes

        Returns
        -------
        Dict[str, Any]
            Flattened dictionary ready for template
        """
        data = {}

        # Test information
        data['test_project'] = test_info.get('test_project', '')
        data['customer'] = test_info.get('customer', '')
        data['specimen_id'] = test_info.get('specimen_id', '')
        data['material'] = test_info.get('material', '')
        data['certificate_number'] = test_info.get('certificate_number', '')
        data['test_date'] = test_info.get('test_date', '')
        data['test_standard'] = 'ASTM E647'
        data['test_equipment'] = 'MTS Landmark 500kN'
        data['requirement'] = test_info.get('requirement', '')

        # Specimen geometry
        data['specimen_type'] = specimen_data.get('specimen_type', 'C(T)')
        data['W'] = specimen_data.get('W', '')
        data['B'] = specimen_data.get('B', '')
        data['B_n'] = specimen_data.get('B_n', specimen_data.get('B', ''))
        data['a_0'] = specimen_data.get('a_0', '')
        data['notch_height'] = specimen_data.get('notch_height', '0')

        # Calculate a_0/W ratio
        try:
            W = float(specimen_data.get('W', 1))
            a_0 = float(specimen_data.get('a_0', 0))
            data['a_W_ratio'] = f"{a_0 / W:.3f}" if W > 0 else '-'
        except (ValueError, TypeError):
            data['a_W_ratio'] = '-'

        # Side grooves
        B = specimen_data.get('B')
        B_n = specimen_data.get('B_n')
        if B and B_n:
            try:
                data['side_grooves'] = 'Yes' if float(B_n) < float(B) else 'No'
            except (ValueError, TypeError):
                data['side_grooves'] = 'No'
        else:
            data['side_grooves'] = 'No'

        # Material properties
        data['yield_strength'] = material_data.get('yield_strength', '')
        data['ultimate_strength'] = material_data.get('ultimate_strength', '')
        data['youngs_modulus'] = material_data.get('youngs_modulus', '')
        data['poissons_ratio'] = material_data.get('poissons_ratio', '0.3')

        # Test parameters
        data['control_mode'] = test_params.get('control_mode', 'Load Control')
        data['load_ratio'] = test_params.get('load_ratio', '0.1')
        data['frequency'] = test_params.get('frequency', '10')
        data['temperature'] = test_params.get('temperature', '23')
        data['P_max'] = test_params.get('P_max', '-')
        data['K_max'] = test_params.get('K_max', '-')
        data['wave_shape'] = test_params.get('wave_shape', 'Sine')
        data['environment'] = test_params.get('environment', 'Laboratory Air')

        # da/dN calculation method
        data['dadn_method'] = test_params.get('dadn_method', 'Secant')
        data['outlier_threshold'] = test_params.get('outlier_threshold', '2.5')

        # Results from FCGRResult object
        if results:
            # Paris law results (final - after outlier removal)
            paris = results.paris_law
            if paris:
                data['paris_C'] = paris.C
                data['paris_m'] = paris.m
                data['paris_r_squared'] = paris.r_squared
                data['paris_n_points'] = paris.n_points
                data['paris_C_error'] = paris.std_error_C
                data['paris_m_error'] = paris.std_error_m
                data['delta_K_min'] = paris.delta_K_range[0] if paris.delta_K_range else '-'
                data['delta_K_max'] = paris.delta_K_range[1] if paris.delta_K_range else '-'
                data['da_dN_min'] = paris.da_dN_range[0] if paris.da_dN_range else '-'
                data['da_dN_max'] = paris.da_dN_range[1] if paris.da_dN_range else '-'

            # Paris law results (initial - all data)
            paris_init = results.paris_law_initial
            if paris_init:
                data['paris_C_initial'] = paris_init.C
                data['paris_m_initial'] = paris_init.m
                data['paris_r_squared_initial'] = paris_init.r_squared
                data['paris_n_points_initial'] = paris_init.n_points

            # Data summary
            data['n_valid_points'] = results.n_valid_points
            data['n_outliers'] = results.n_outliers
            data['total_cycles'] = results.total_cycles
            data['final_crack_length'] = results.final_crack_length
            data['threshold_delta_K'] = results.threshold_delta_K if results.threshold_delta_K > 0 else '-'

            # Validity
            data['is_valid'] = 'VALID' if results.is_valid else 'See Notes'
            data['validity_notes'] = '\n'.join(results.validity_notes) if results.validity_notes else 'All checks passed'
        else:
            # Default values if no results
            data['paris_C'] = '-'
            data['paris_m'] = '-'
            data['paris_r_squared'] = '-'
            data['paris_n_points'] = '-'
            data['n_valid_points'] = '-'
            data['n_outliers'] = '-'
            data['total_cycles'] = '-'
            data['final_crack_length'] = '-'
            data['is_valid'] = '-'
            data['validity_notes'] = '-'

        # Signatures (to be filled manually)
        data['tested_by'] = ''
        data['tested_date'] = ''
        data['reviewed_by'] = ''
        data['reviewed_date'] = ''
        data['approved_by'] = ''
        data['approved_date'] = ''

        return data
