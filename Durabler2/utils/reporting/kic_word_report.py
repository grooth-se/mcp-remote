"""
KIC Fracture Toughness Report Generator.

Generates Word document reports for KIC (ASTM E399) fracture toughness tests.
Reports include test information, specimen dimensions, material properties,
results with uncertainties, validity assessment, and force-displacement plot.

Uses same layout style as CTOD E1290 reports for consistency.
"""

import re
from pathlib import Path
from typing import Optional, Dict, Any, List
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Inches, Pt, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("Warning: python-docx not installed. Report generation will be limited.")


class KICReportGenerator:
    """
    Generator for KIC fracture toughness test reports.

    Creates Word documents with all test information, results,
    and validity assessment per ASTM E399.

    Parameters
    ----------
    template_path : Path, optional
        Path to Word template file. If None, creates report from scratch.
    """

    def __init__(self, template_path: Optional[Path] = None):
        self.template_path = template_path

        if not DOCX_AVAILABLE:
            print("Warning: python-docx not available. Install with: pip install python-docx")

    def generate_report(self,
                        output_path: Path,
                        test_info: Dict[str, str],
                        dimensions: Dict[str, str],
                        material_props: Dict[str, str],
                        results: Any,
                        chart_path: Optional[Path] = None,
                        logo_path: Optional[Path] = None,
                        precrack_measurements: Optional[List[float]] = None,
                        crack_photo_path: Optional[Path] = None) -> Path:
        """
        Generate KIC test report.

        Parameters
        ----------
        output_path : Path
            Output file path for the Word document
        test_info : Dict[str, str]
            Test information (certificate, project, customer, etc.)
        dimensions : Dict[str, str]
            Specimen dimensions (W, B, a_0, S, etc.)
        material_props : Dict[str, str]
            Material properties (yield strength, E, nu)
        results : KICResult
            Analysis results from KICAnalyzer
        chart_path : Path, optional
            Path to chart image to include
        logo_path : Path, optional
            Path to company logo
        precrack_measurements : List[float], optional
            Precrack crack length measurements (mm)
        crack_photo_path : Path, optional
            Path to crack surface photo

        Returns
        -------
        Path
            Path to generated report file
        """
        if not DOCX_AVAILABLE:
            raise ImportError("python-docx is required for report generation")

        # Create document from template or new
        if self.template_path and self.template_path.exists():
            doc = Document(self.template_path)
            self._fill_template(doc, test_info, dimensions, material_props, results,
                               chart_path, logo_path, precrack_measurements, crack_photo_path)
        else:
            doc = self._create_report_from_scratch(
                test_info, dimensions, material_props, results, chart_path, logo_path,
                precrack_measurements, crack_photo_path
            )

        # Save document
        doc.save(output_path)
        return output_path

    def _create_report_from_scratch(self,
                                     test_info: Dict[str, str],
                                     dimensions: Dict[str, str],
                                     material_props: Dict[str, str],
                                     results: Any,
                                     chart_path: Optional[Path],
                                     logo_path: Optional[Path],
                                     precrack_measurements: Optional[List[float]] = None,
                                     crack_photo_path: Optional[Path] = None) -> Document:
        """Create report without template - matches Vickers layout."""
        from docx.shared import RGBColor

        doc = Document()

        # Set compact paragraph spacing for entire document
        style = doc.styles['Normal']
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.0
        style.font.name = 'Calibri'
        style.font.size = Pt(10)

        # Set compact heading styles with dark green color
        dark_green = RGBColor(0x00, 0x64, 0x00)
        for i in range(1, 4):
            heading_style = doc.styles[f'Heading {i}']
            heading_style.paragraph_format.space_before = Pt(8)
            heading_style.paragraph_format.space_after = Pt(4)
            heading_style.font.color.rgb = dark_green

        # Add header with logo on left, certificate info on right (5-line layout)
        for section in doc.sections:
            # Set narrower margins for compact layout
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(2.0)
            section.right_margin = Cm(2.0)
            header = section.header
            header.is_linked_to_previous = False

            # Row 1: Logo - left aligned
            logo_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            logo_para.paragraph_format.space_after = Pt(0)
            if logo_path and logo_path.exists():
                logo_run = logo_para.add_run()
                logo_run.add_picture(str(logo_path), width=Cm(5.0))  # 50mm width

            # Row 2: Title - centered, font size 12
            title_para = header.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_para.paragraph_format.space_before = Pt(0)
            title_para.paragraph_format.space_after = Pt(0)
            title_run = title_para.add_run('KIC Fracture Toughness Test Report')
            title_run.bold = True
            title_run.font.size = Pt(12)

            # Row 3: Standard - centered, font size 8
            std_para = header.add_paragraph()
            std_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            std_para.paragraph_format.space_before = Pt(0)
            std_para.paragraph_format.space_after = Pt(0)
            std_run = std_para.add_run('ASTM E399')
            std_run.font.size = Pt(8)

            # Row 4: Certificate - right aligned, font size 8
            cert_para = header.add_paragraph()
            cert_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cert_para.paragraph_format.space_before = Pt(0)
            cert_para.paragraph_format.space_after = Pt(0)
            cert_run = cert_para.add_run(f"Certificate: {test_info.get('certificate_number', '')}")
            cert_run.font.size = Pt(8)

            # Row 5: Date - right aligned, font size 8
            date_para = header.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            date_para.paragraph_format.space_before = Pt(0)
            date_para.paragraph_format.space_after = Pt(0)
            date_run = date_para.add_run(f"Date: {test_info.get('test_date', '')}")
            date_run.font.size = Pt(8)

        # Test Information table (exclude certificate and date - now in header)
        heading = doc.add_heading('Test Information', level=1)
        heading.paragraph_format.space_before = Pt(0)
        heading.paragraph_format.space_after = Pt(6)

        table = doc.add_table(rows=8, cols=4)
        table.style = 'Table Grid'

        # Two-column layout: Label | Value | Label | Value (exclude cert/date)
        info_data = [
            ('Test Project:', test_info.get('test_project', ''), 'Temperature:', f"{test_info.get('temperature', '23')} °C"),
            ('Customer:', test_info.get('customer', ''), 'Test Standard:', 'ASTM E399'),
            ('Customer Order:', test_info.get('customer_order', ''), 'Test Equipment:', 'MTS Landmark 500kN'),
            ('Product S/N:', test_info.get('product_sn', ''), 'Specimen ID:', test_info.get('specimen_id', '')),
            ('Material:', test_info.get('material', ''), 'Location/Orientation:', test_info.get('location_orientation', '')),
            ('Notch Type:', dimensions.get('notch_type', 'Fatigue pre-crack'), 'Side Grooves:', dimensions.get('side_grooves', 'No')),
            ('Specimen Type:', dimensions.get('specimen_type', 'SE(B)'), 'a₀/W Ratio:', self._calc_a_w_ratio(dimensions)),
            ('Ligament (W-a₀):', self._calc_ligament(dimensions), '', ''),
        ]

        for i, (label1, value1, label2, value2) in enumerate(info_data):
            table.rows[i].cells[0].text = label1
            table.rows[i].cells[1].text = str(value1)
            table.rows[i].cells[2].text = label2
            table.rows[i].cells[3].text = str(value2)
            # Bold the labels
            if table.rows[i].cells[0].paragraphs[0].runs:
                table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            if table.rows[i].cells[2].paragraphs[0].runs:
                table.rows[i].cells[2].paragraphs[0].runs[0].bold = True

        # Compact table rows
        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].paragraph_format.space_before = Pt(1)
                cell.paragraphs[0].paragraph_format.space_after = Pt(1)

        # Specimen Geometry table
        heading = doc.add_heading('Specimen Geometry', level=1)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)

        table = doc.add_table(rows=7, cols=3)
        table.style = 'Table Grid'

        # Header row
        headers = ['Parameter', 'Value', 'Unit']
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

        dim_data = [
            ('Specimen Type', dimensions.get('specimen_type', 'SE(B)'), '-'),
            ('Width W', dimensions.get('W', ''), 'mm'),
            ('Thickness B', dimensions.get('B', ''), 'mm'),
            ('Net Thickness Bₙ', dimensions.get('B_n', dimensions.get('B', '')), 'mm'),
            ('Crack length a₀', dimensions.get('a_0', ''), 'mm'),
            ('Span S', dimensions.get('S', '-'), 'mm'),
        ]

        for i, (param, value, unit) in enumerate(dim_data):
            table.rows[i+1].cells[0].text = param
            table.rows[i+1].cells[1].text = str(value)
            table.rows[i+1].cells[2].text = unit

        # Compact table rows
        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].paragraph_format.space_before = Pt(1)
                cell.paragraphs[0].paragraph_format.space_after = Pt(1)

        # Material Properties table
        heading = doc.add_heading('Material Properties', level=1)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)

        table = doc.add_table(rows=5, cols=3)
        table.style = 'Table Grid'

        # Header row
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

        mat_data = [
            ('Yield Strength σys', material_props.get('yield_strength', ''), 'MPa'),
            ('Ultimate Strength σu', material_props.get('ultimate_strength', '-'), 'MPa'),
            ("Young's Modulus E", material_props.get('youngs_modulus', ''), 'GPa'),
            ("Poisson's Ratio ν", material_props.get('poissons_ratio', '0.3'), '-'),
        ]

        for i, (param, value, unit) in enumerate(mat_data):
            table.rows[i+1].cells[0].text = param
            table.rows[i+1].cells[1].text = str(value)
            table.rows[i+1].cells[2].text = unit

        # Compact table rows
        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].paragraph_format.space_before = Pt(1)
                cell.paragraphs[0].paragraph_format.space_after = Pt(1)

        # Precrack Measurements table (if available)
        if precrack_measurements and len(precrack_measurements) > 0:
            heading = doc.add_heading('Crack Length Measurements', level=1)
            heading.paragraph_format.space_before = Pt(12)
            heading.paragraph_format.space_after = Pt(6)

            num_rows = len(precrack_measurements) + 2  # +1 for header, +1 for average
            table = doc.add_table(rows=num_rows, cols=3)
            table.style = 'Table Grid'

            # Header row
            precrack_headers = ['Position', 'Value', 'Unit']
            for i, header in enumerate(precrack_headers):
                table.rows[0].cells[i].text = header
                table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

            # Individual measurements with position labels
            position_labels = ['Surface (1/4B)', 'Quarter (1/4)', 'Center', 'Quarter (3/4)', 'Surface (3/4B)']
            for i, meas in enumerate(precrack_measurements, 1):
                label = position_labels[i-1] if i <= len(position_labels) else f'a{i}'
                table.rows[i].cells[0].text = label
                table.rows[i].cells[1].text = f'{meas:.3f}'
                table.rows[i].cells[2].text = 'mm'

            # Calculate and add average (E399 formula for 5 measurements)
            if len(precrack_measurements) == 5:
                a = precrack_measurements
                avg_crack = (0.5 * a[0] + a[1] + a[2] + a[3] + 0.5 * a[4]) / 4
            else:
                avg_crack = sum(precrack_measurements) / len(precrack_measurements)

            avg_row = len(precrack_measurements) + 1
            table.rows[avg_row].cells[0].text = 'Average (E399)'
            table.rows[avg_row].cells[1].text = f'{avg_crack:.3f}'
            table.rows[avg_row].cells[2].text = 'mm'
            # Bold the average row
            for cell in table.rows[avg_row].cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        run.bold = True

            # Compact table rows
            for row in table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].paragraph_format.space_before = Pt(1)
                    cell.paragraphs[0].paragraph_format.space_after = Pt(1)

        # Test Results table
        heading = doc.add_heading('Test Results', level=1)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)

        table = doc.add_table(rows=8, cols=4)
        table.style = 'Table Grid'

        # Header row
        result_headers = ['Parameter', 'Value', 'U (k=2)', 'Unit']
        for i, header in enumerate(result_headers):
            table.rows[0].cells[i].text = header
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

        if results:
            # Format results table
            results_data = [
                ('Pmax', f'{results.P_max.value:.2f}', f'±{results.P_max.uncertainty:.2f}', 'kN'),
                ('PQ (5% secant)', f'{results.P_Q.value:.2f}', f'±{results.P_Q.uncertainty:.2f}', 'kN'),
                ('Pmax/PQ ratio', f'{results.P_ratio:.3f}', '-', '≤1.10'),
                ('KQ (conditional)', f'{results.K_Q.value:.2f}', f'±{results.K_Q.uncertainty:.2f}', 'MPa√m'),
                ('KIC', f'{results.K_IC.value:.2f}' if results.K_IC else 'CONDITIONAL',
                 f'±{results.K_IC.uncertainty:.2f}' if results.K_IC else '-', 'MPa√m'),
                ('Compliance', f'{results.compliance:.6f}', '-', 'mm/kN'),
                ('Validity', 'VALID' if results.is_valid else 'CONDITIONAL', '-', '-'),
            ]

            for i, (param, value, unc, unit) in enumerate(results_data):
                table.rows[i+1].cells[0].text = param
                table.rows[i+1].cells[1].text = value
                table.rows[i+1].cells[2].text = unc
                table.rows[i+1].cells[3].text = unit

        # Compact table rows
        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].paragraph_format.space_before = Pt(1)
                cell.paragraphs[0].paragraph_format.space_after = Pt(1)

        # Force-Displacement Plot
        if chart_path and chart_path.exists():
            heading = doc.add_heading('Force vs Displacement', level=1)
            heading.paragraph_format.space_before = Pt(12)
            heading.paragraph_format.space_after = Pt(6)
            doc.add_picture(str(chart_path), width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add caption
            caption = doc.add_paragraph('Figure 1: Force-displacement curve with PQ determination per ASTM E399')
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.runs[0].font.size = Pt(10)
            caption.runs[0].font.italic = True

        # Crack Surface Photo (if available)
        if crack_photo_path and crack_photo_path.exists():
            heading = doc.add_heading('Crack Surface', level=1)
            heading.paragraph_format.space_before = Pt(12)
            heading.paragraph_format.space_after = Pt(6)
            doc.add_picture(str(crack_photo_path), width=Inches(4.0))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add caption
            caption = doc.add_paragraph('Figure 2: Crack surface showing fatigue precrack and final fracture')
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.runs[0].font.size = Pt(10)
            caption.runs[0].font.italic = True

        # Validity Assessment
        heading = doc.add_heading('Validity Assessment per ASTM E399', level=1)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)

        validity_status = "VALID" if results and results.is_valid else "CONDITIONAL"

        # Validity statement paragraph
        if results and results.is_valid:
            validity_text = "The test meets all validity requirements of ASTM E399. The result is a valid plane-strain fracture toughness (KIC)."
        else:
            validity_text = "The test does not meet all validity requirements of ASTM E399. The result is reported as conditional (KQ)."

        status_para = doc.add_paragraph()
        status_run = status_para.add_run(f"Status: {validity_status}")
        status_run.bold = True

        doc.add_paragraph(validity_text)

        # List validity notes if any
        if results and results.validity_notes:
            notes_para = doc.add_paragraph()
            notes_para.add_run("Validity Notes:").bold = True
            for note in results.validity_notes:
                doc.add_paragraph(f"• {note}")

        # Approval Signatures
        heading = doc.add_heading('Approval', level=1)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)

        sig_table = doc.add_table(rows=4, cols=4)
        sig_table.style = 'Table Grid'

        sig_headers = ['Role', 'Name', 'Signature', 'Date']
        for i, header in enumerate(sig_headers):
            sig_table.rows[0].cells[i].text = header
            sig_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

        sig_table.rows[1].cells[0].text = 'Tested by:'
        sig_table.rows[2].cells[0].text = 'Reviewed by:'
        sig_table.rows[3].cells[0].text = 'Approved by:'

        # Compact signature table rows
        for row in sig_table.rows:
            for cell in row.cells:
                cell.paragraphs[0].paragraph_format.space_before = Pt(1)
                cell.paragraphs[0].paragraph_format.space_after = Pt(1)

        # Add disclaimer to page footer (visible on all pages)
        disclaimer_text = (
            "All work and services carried out by Durabler are subject to, and conducted in accordance with, "
            "Durabler standard terms and conditions, which are available at durabler.se. This document shall not "
            "be reproduced other than in full, except with prior written approval of the issuer. The results pertain "
            "only to the item(s) as sampled by the client unless otherwise indicated. Durabler a part of Subseatec S AB, "
            "Address: Durabler C/O Subseatec, Dalavägen 23, 68130 Kristinehamn, SWEDEN"
        )
        for section in doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.clear()
            footer_run = footer_para.add_run(disclaimer_text)
            footer_run.font.size = Pt(7)
            footer_run.italic = True

        return doc

    def _calc_a_w_ratio(self, dimensions: Dict[str, str]) -> str:
        """Calculate a₀/W ratio."""
        try:
            W = float(dimensions.get('W', 0))
            a_0 = float(dimensions.get('a_0', 0))
            if W > 0:
                return f"{a_0 / W:.3f}"
        except (ValueError, TypeError):
            pass
        return '-'

    def _calc_ligament(self, dimensions: Dict[str, str]) -> str:
        """Calculate ligament (W - a₀)."""
        try:
            W = float(dimensions.get('W', 0))
            a_0 = float(dimensions.get('a_0', 0))
            if W > 0:
                return f"{W - a_0:.2f} mm"
        except (ValueError, TypeError):
            pass
        return '-'

    def _fill_template(self,
                       doc: Document,
                       test_info: Dict[str, str],
                       dimensions: Dict[str, str],
                       material_props: Dict[str, str],
                       results: Any,
                       chart_path: Optional[Path],
                       logo_path: Optional[Path],
                       precrack_measurements: Optional[List[float]] = None,
                       crack_photo_path: Optional[Path] = None):
        """Fill in template placeholders."""
        # Build replacement dictionary
        replacements = {
            '{{certificate_number}}': test_info.get('certificate_number', ''),
            '{{test_project}}': test_info.get('test_project', ''),
            '{{customer}}': test_info.get('customer', ''),
            '{{specimen_id}}': test_info.get('specimen_id', ''),
            '{{material}}': test_info.get('material', ''),
            '{{test_date}}': test_info.get('test_date', ''),
            '{{temperature}}': test_info.get('temperature', '23'),
            '{{specimen_type}}': dimensions.get('specimen_type', ''),
            '{{W}}': dimensions.get('W', ''),
            '{{B}}': dimensions.get('B', ''),
            '{{a_0}}': dimensions.get('a_0', ''),
            '{{S}}': dimensions.get('S', '-'),
            '{{B_n}}': dimensions.get('B_n', '-'),
            '{{yield_strength}}': material_props.get('yield_strength', ''),
            '{{youngs_modulus}}': material_props.get('youngs_modulus', ''),
            '{{poissons_ratio}}': material_props.get('poissons_ratio', ''),
        }

        if results:
            replacements.update({
                '{{P_max}}': f'{results.P_max.value:.2f}',
                '{{P_max_uncertainty}}': f'{results.P_max.uncertainty:.2f}',
                '{{P_Q}}': f'{results.P_Q.value:.2f}',
                '{{P_Q_uncertainty}}': f'{results.P_Q.uncertainty:.2f}',
                '{{P_ratio}}': f'{results.P_ratio:.3f}',
                '{{K_Q}}': f'{results.K_Q.value:.2f}',
                '{{K_Q_uncertainty}}': f'{results.K_Q.uncertainty:.2f}',
                '{{K_IC}}': f'{results.K_IC.value:.2f}' if results.K_IC else 'CONDITIONAL',
                '{{K_IC_uncertainty}}': f'{results.K_IC.uncertainty:.2f}' if results.K_IC else '-',
                '{{compliance}}': f'{results.compliance:.4f}',
                '{{is_valid}}': 'VALID' if results.is_valid else 'CONDITIONAL',
                '{{validity_notes}}': '\n'.join(results.validity_notes) if results.validity_notes else '',
            })

        # Replace in page headers
        for section in doc.sections:
            header = section.header
            for para in header.paragraphs:
                for key, value in replacements.items():
                    if key in para.text:
                        para.text = para.text.replace(key, str(value))
            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            for key, value in replacements.items():
                                if key in para.text:
                                    para.text = para.text.replace(key, str(value))

        # Replace in paragraphs
        for para in doc.paragraphs:
            for key, value in replacements.items():
                if key in para.text:
                    para.text = para.text.replace(key, str(value))

        # Replace in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        for key, value in replacements.items():
                            if key in para.text:
                                para.text = para.text.replace(key, str(value))

        # Insert chart if placeholder exists
        if chart_path and chart_path.exists():
            for para in doc.paragraphs:
                if '{{chart}}' in para.text:
                    para.text = ''
                    run = para.add_run()
                    run.add_picture(str(chart_path), width=Inches(5.5))
                    break

        # Insert logo if placeholder exists (check header first, then body)
        if logo_path and logo_path.exists():
            logo_inserted = False
            # Check page header first
            for section in doc.sections:
                if logo_inserted:
                    break
                header = section.header
                for para in header.paragraphs:
                    if '{{logo}}' in para.text:
                        para.text = ''
                        run = para.add_run()
                        run.add_picture(str(logo_path), height=Cm(1.5))
                        logo_inserted = True
                        break
                if not logo_inserted:
                    for table in header.tables:
                        for row in table.rows:
                            for cell in row.cells:
                                for para in cell.paragraphs:
                                    if '{{logo}}' in para.text:
                                        para.text = ''
                                        run = para.add_run()
                                        run.add_picture(str(logo_path), height=Cm(1.5))
                                        logo_inserted = True
                                        break
                                if logo_inserted:
                                    break
                            if logo_inserted:
                                break
                        if logo_inserted:
                            break
            # If not in header, check body paragraphs
            if not logo_inserted:
                for para in doc.paragraphs:
                    if '{{logo}}' in para.text:
                        para.text = ''
                        run = para.add_run()
                        run.add_picture(str(logo_path), width=Inches(2))
                        break

        # Insert crack photo if placeholder exists
        if crack_photo_path and crack_photo_path.exists():
            for para in doc.paragraphs:
                if '{{crack_photo}}' in para.text:
                    para.text = ''
                    run = para.add_run()
                    run.add_picture(str(crack_photo_path), width=Inches(4.5))
                    break

        # Add precrack measurements placeholders
        if precrack_measurements:
            # Calculate E399 average
            if len(precrack_measurements) == 5:
                a = precrack_measurements
                avg_crack = (0.5 * a[0] + a[1] + a[2] + a[3] + 0.5 * a[4]) / 4
            else:
                avg_crack = sum(precrack_measurements) / len(precrack_measurements)

            # Add precrack placeholders to replacements
            for i, meas in enumerate(precrack_measurements, 1):
                for para in doc.paragraphs:
                    placeholder = f'{{{{crack_{i}}}}}'
                    if placeholder in para.text:
                        para.text = para.text.replace(placeholder, f'{meas:.2f}')
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            for para in cell.paragraphs:
                                placeholder = f'{{{{crack_{i}}}}}'
                                if placeholder in para.text:
                                    para.text = para.text.replace(placeholder, f'{meas:.2f}')

            # Replace average crack placeholder
            for para in doc.paragraphs:
                if '{{crack_avg}}' in para.text:
                    para.text = para.text.replace('{{crack_avg}}', f'{avg_crack:.2f}')
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            if '{{crack_avg}}' in para.text:
                                para.text = para.text.replace('{{crack_avg}}', f'{avg_crack:.2f}')

    @staticmethod
    def prepare_report_data(
        test_info: Dict[str, str],
        specimen_data: Dict[str, Any],
        material_data: Dict[str, Any],
        results: Dict[str, Any],
        crack_measurements: Optional[List[float]] = None
    ) -> Dict[str, Any]:
        """
        Prepare data dictionary for report generation (CTOD E1290 style).

        Parameters
        ----------
        test_info : Dict[str, str]
            Test information (project, customer, etc.)
        specimen_data : Dict[str, Any]
            Specimen geometry data
        material_data : Dict[str, Any]
            Material properties
        results : Dict[str, Any]
            Analysis results with KICResult and MeasuredValue objects
        crack_measurements : List[float], optional
            5-point crack length measurements (E399)

        Returns
        -------
        Dict[str, Any]
            Flattened dictionary ready for template
        """
        data = {}

        # Test information
        data['test_project'] = test_info.get('test_project', '')
        data['customer'] = test_info.get('customer', '')
        data['customer_order'] = test_info.get('customer_order', '')
        data['product_sn'] = test_info.get('product_sn', '')
        data['specimen_id'] = test_info.get('specimen_id', '')
        data['location_orientation'] = test_info.get('location_orientation', '')
        data['material'] = test_info.get('material', '')
        data['certificate_number'] = test_info.get('certificate_number', '')
        data['test_date'] = test_info.get('test_date', '')
        data['test_standard'] = 'ASTM E399'
        data['test_equipment'] = 'MTS Landmark 500kN'
        data['test_temperature'] = test_info.get('temperature', '23')
        data['temperature'] = test_info.get('temperature', '23')
        data['operator'] = test_info.get('operator', '')

        # Specimen geometry
        data['specimen_type'] = specimen_data.get('specimen_type', 'C(T)')
        data['W'] = specimen_data.get('W', '')
        data['B'] = specimen_data.get('B', '')
        data['B_n'] = specimen_data.get('B_n', specimen_data.get('B', ''))
        data['a_0'] = specimen_data.get('a_0', '')
        data['S'] = specimen_data.get('S', '-')
        data['notch_type'] = specimen_data.get('notch_type', 'Fatigue pre-crack')
        data['side_grooves'] = 'Yes' if specimen_data.get('B_n') and float(specimen_data.get('B_n', 0)) < float(specimen_data.get('B', 0)) else 'No'

        # Calculate a_0/W ratio and ligament
        try:
            W = float(specimen_data.get('W', 1))
            a_0 = float(specimen_data.get('a_0', 0))
            data['a_W_ratio'] = f"{a_0 / W:.3f}" if W > 0 else '-'
            data['ligament'] = f"{W - a_0:.2f}" if W > 0 else '-'
        except (ValueError, TypeError):
            data['a_W_ratio'] = '-'
            data['ligament'] = '-'

        # Material properties
        data['yield_strength'] = material_data.get('yield_strength', '')
        data['ultimate_strength'] = material_data.get('ultimate_strength', '')
        data['youngs_modulus'] = material_data.get('youngs_modulus', '')
        data['poissons_ratio'] = material_data.get('poissons_ratio', '0.3')

        # Helper to extract value and uncertainty from MeasuredValue objects
        def get_measured(obj, default_value='-', default_unc='-'):
            if obj and hasattr(obj, 'value'):
                return f"{obj.value:.4g}", f"±{obj.uncertainty:.4g}"
            return default_value, default_unc

        # Results from MeasuredValue objects
        P_max = results.get('P_max')
        if P_max:
            data['P_max_value'], data['P_max_uncertainty'] = get_measured(P_max)
        else:
            data['P_max_value'] = '-'
            data['P_max_uncertainty'] = '-'
        data['P_max_req'] = '-'  # Requirement placeholder

        P_Q = results.get('P_Q')
        if P_Q:
            data['P_Q_value'], data['P_Q_uncertainty'] = get_measured(P_Q)
        else:
            data['P_Q_value'] = '-'
            data['P_Q_uncertainty'] = '-'
        data['P_Q_req'] = '-'  # Requirement placeholder

        K_Q = results.get('K_Q')
        if K_Q:
            data['K_Q_value'], data['K_Q_uncertainty'] = get_measured(K_Q)
        else:
            data['K_Q_value'] = '-'
            data['K_Q_uncertainty'] = '-'
        data['K_Q_req'] = '-'  # Requirement placeholder

        K_IC = results.get('K_IC')
        if K_IC:
            data['K_IC_value'], data['K_IC_uncertainty'] = get_measured(K_IC)
            data['K_IC_valid'] = 'VALID'
        else:
            data['K_IC_value'] = 'CONDITIONAL'
            data['K_IC_uncertainty'] = '-'
            data['K_IC_valid'] = 'CONDITIONAL'
        data['K_IC_req'] = '-'  # Requirement placeholder

        # P ratio
        P_ratio = results.get('P_ratio')
        if P_ratio:
            data['P_ratio'] = f"{P_ratio:.3f}"
            data['P_ratio_valid'] = 'PASS' if P_ratio <= 1.10 else 'FAIL'
        else:
            data['P_ratio'] = '-'
            data['P_ratio_valid'] = '-'
        data['P_ratio_req'] = '≤ 1.10'  # E399 requirement

        # Compliance
        compliance = results.get('compliance')
        if compliance:
            data['compliance'] = f"{compliance:.4f}"
        else:
            data['compliance'] = '-'

        # R-squared (fit quality)
        r_squared = results.get('r_squared')
        if r_squared:
            data['r_squared'] = f"{r_squared:.4f}"
        else:
            data['r_squared'] = '-'

        # 5-point crack measurements (E399)
        if crack_measurements and len(crack_measurements) >= 5:
            for i, val in enumerate(crack_measurements[:5], 1):
                data[f'a{i}'] = f"{val:.2f}"
            # Calculate E399 average
            a = crack_measurements[:5]
            avg_crack = (0.5 * a[0] + a[1] + a[2] + a[3] + 0.5 * a[4]) / 4
            data['a_avg'] = f"{avg_crack:.2f}"
        else:
            for i in range(1, 6):
                data[f'a{i}'] = '-'
            data['a_avg'] = '-'

        # Validity
        is_valid = results.get('is_valid', False)
        data['validity_status'] = 'VALID' if is_valid else 'CONDITIONAL'
        validity_notes = results.get('validity_notes', [])
        if is_valid:
            data['validity_statement'] = 'The test meets all validity requirements of ASTM E399.'
        else:
            data['validity_statement'] = 'The test does not meet all validity requirements of ASTM E399.'
        data['validity_notes'] = '\n'.join(validity_notes) if validity_notes else ''

        # Signatures (to be filled manually)
        data['tested_by'] = ''
        data['tested_date'] = ''
        data['reviewed_by'] = ''
        data['reviewed_date'] = ''
        data['approved_by'] = ''
        data['approved_date'] = ''

        return data

    def generate_from_template(
        self,
        output_path: Path,
        data: Dict[str, Any],
        chart_path: Optional[Path] = None,
        logo_path: Optional[Path] = None,
        photo_paths: Optional[List[Path]] = None
    ) -> Path:
        """
        Generate report by populating template with data (CTOD E1290 style).

        Parameters
        ----------
        output_path : Path
            Path for output Word document
        data : Dict[str, Any]
            Dictionary of placeholder values from prepare_report_data()
        chart_path : Path, optional
            Path to Force vs Displacement chart image
        logo_path : Path, optional
            Path to logo image to insert
        photo_paths : List[Path], optional
            List of paths to crack surface photos

        Returns
        -------
        Path
            Path to generated report
        """
        if not self.template_path or not self.template_path.exists():
            raise FileNotFoundError(f"Template not found: {self.template_path}")

        doc = Document(self.template_path)

        # Replace placeholders in page headers
        for section in doc.sections:
            header = section.header
            for paragraph in header.paragraphs:
                self._replace_in_paragraph(paragraph, data, chart_path, logo_path, photo_paths)
            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_in_paragraph(paragraph, data, chart_path, logo_path, photo_paths)

        # Replace placeholders in paragraphs
        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph, data, chart_path, logo_path, photo_paths)

        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, data, chart_path, logo_path, photo_paths)

        doc.save(output_path)
        return output_path

    def _replace_in_paragraph(
        self,
        paragraph,
        data: Dict[str, Any],
        chart_path: Optional[Path],
        logo_path: Optional[Path],
        photo_paths: Optional[List[Path]]
    ):
        """Replace placeholders in a paragraph (CTOD E1290 style)."""
        full_text = paragraph.text

        # Handle logo placeholder
        if '{{logo}}' in full_text and logo_path and logo_path.exists():
            paragraph.clear()
            run = paragraph.add_run()
            run.add_picture(str(logo_path), height=Cm(1.5))
            return

        # Handle chart placeholder
        if '{{chart}}' in full_text and chart_path and chart_path.exists():
            paragraph.clear()
            run = paragraph.add_run()
            run.add_picture(str(chart_path), width=Inches(5.5))
            return

        # Handle photos placeholder
        if '{{photos}}' in full_text:
            paragraph.clear()
            if photo_paths:
                for i, photo_path in enumerate(photo_paths):
                    if photo_path and photo_path.exists():
                        run = paragraph.add_run()
                        run.add_picture(str(photo_path), width=Inches(3.5))
                        if i < len(photo_paths) - 1:
                            paragraph.add_run("\n")
            else:
                paragraph.add_run("No crack surface photos attached.")
            return

        # Find all placeholders
        placeholders = re.findall(r'\{\{([^}]+)\}\}', full_text)

        if not placeholders:
            return

        # Replace each placeholder
        for placeholder in placeholders:
            key = placeholder
            value = data.get(key, '')

            # Handle None values
            if value is None:
                value = ''

            # Format numeric values
            if isinstance(value, float):
                if abs(value) < 0.001:
                    value = f"{value:.4g}"
                elif abs(value) < 1:
                    value = f"{value:.4f}"
                else:
                    value = f"{value:.2f}"
            elif isinstance(value, int):
                value = str(value)

            full_text = full_text.replace(f'{{{{{placeholder}}}}}', str(value))

        # Update paragraph text while preserving formatting
        if paragraph.runs:
            first_run = paragraph.runs[0]
            font_name = first_run.font.name
            font_size = first_run.font.size
            bold = first_run.bold
            italic = first_run.italic

            paragraph.clear()
            new_run = paragraph.add_run(full_text)
            new_run.font.name = font_name
            new_run.font.size = font_size
            new_run.bold = bold
            new_run.italic = italic
        else:
            paragraph.text = full_text
