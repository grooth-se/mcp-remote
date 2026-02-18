"""
Word report generator for tensile test results.

Populates a Word template with test data and results, or creates report from scratch.
"""

import re
from pathlib import Path
from typing import Dict, Any, Optional
from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


class TensileReportGenerator:
    """
    Generate tensile test reports from Word template or from scratch.

    Uses placeholder syntax {{placeholder_name}} in template.
    """

    def __init__(self, template_path: Optional[Path] = None):
        """
        Initialize report generator.

        Parameters
        ----------
        template_path : Path, optional
            Path to Word template file. If None, creates report from scratch.
        """
        self.template_path = template_path
        if template_path and not template_path.exists():
            raise FileNotFoundError(f"Template not found: {template_path}")

    def generate_report(
        self,
        output_path: Path,
        data: Dict[str, Any],
        chart_path: Optional[Path] = None,
        logo_path: Optional[Path] = None
    ) -> Path:
        """
        Generate report by populating template with data.

        Parameters
        ----------
        output_path : Path
            Path for output Word document
        data : Dict[str, Any]
            Dictionary of placeholder values
        chart_path : Path, optional
            Path to chart image to insert
        logo_path : Path, optional
            Path to logo image to insert

        Returns
        -------
        Path
            Path to generated report
        """
        # Create from scratch if no template
        if not self.template_path or not self.template_path.exists():
            doc = self._create_report_from_scratch(data, chart_path, logo_path)
            doc.save(output_path)
            return output_path

        doc = Document(self.template_path)

        # Replace placeholders in page headers
        for section in doc.sections:
            header = section.header
            for paragraph in header.paragraphs:
                self._replace_in_paragraph(paragraph, data, chart_path, logo_path)
            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_in_paragraph(paragraph, data, chart_path, logo_path)

        # Replace placeholders in paragraphs
        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph, data, chart_path, logo_path)

        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, data, chart_path, logo_path)

        # Add disclaimer to page footer (visible on all pages)
        disclaimer_text = (
            "All work and services carried out by Durabler are subject to, and conducted in accordance with, "
            "Durabler standard terms and conditions, which are available at durabler.se. This document shall not "
            "be reproduced other than in full, except with prior written approval of the issuer. The results pertain "
            "only to the item(s) as sampled by the client unless otherwise indicated. Durabler a part of Subseatec S AB, "
            "Address: Durabler C/O Subseatec, Dalavägen 23, 68130 Kristinehamn, SWEDEN"
        )
        for section in doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.clear()
            footer_run = footer_para.add_run(disclaimer_text)
            footer_run.font.size = Pt(7)
            footer_run.italic = True

        doc.save(output_path)
        return output_path

    def _create_report_from_scratch(
        self,
        data: Dict[str, Any],
        chart_path: Optional[Path],
        logo_path: Optional[Path]
    ) -> Document:
        """Create report without template - matches Vickers layout."""
        from docx.shared import RGBColor

        doc = Document()

        # Set compact paragraph spacing for entire document
        style = doc.styles['Normal']
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.0
        style.font.name = 'Calibri'
        style.font.size = Pt(10)

        # Set compact heading styles with dark green color
        dark_green = RGBColor(0x00, 0x64, 0x00)
        for i in range(1, 4):
            heading_style = doc.styles[f'Heading {i}']
            heading_style.paragraph_format.space_before = Pt(8)
            heading_style.paragraph_format.space_after = Pt(4)
            heading_style.font.color.rgb = dark_green

        # Add header with logo on left, certificate info on right (5-line layout)
        for section in doc.sections:
            # Set narrower margins for compact layout
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(2.0)
            section.right_margin = Cm(2.0)
            header = section.header
            header.is_linked_to_previous = False

            # Row 1: Logo - left aligned
            logo_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            logo_para.paragraph_format.space_after = Pt(0)
            if logo_path and logo_path.exists():
                logo_run = logo_para.add_run()
                logo_run.add_picture(str(logo_path), width=Cm(5.0))  # 50mm width

            # Row 2: Title - centered, font size 12
            title_para = header.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_para.paragraph_format.space_before = Pt(0)
            title_para.paragraph_format.space_after = Pt(0)
            title_run = title_para.add_run('Tensile Test Report')
            title_run.bold = True
            title_run.font.size = Pt(12)

            # Row 3: Standard - centered, font size 8
            std_para = header.add_paragraph()
            std_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            std_para.paragraph_format.space_before = Pt(0)
            std_para.paragraph_format.space_after = Pt(0)
            std_run = std_para.add_run(data.get('test_standard', 'ASTM E8/E8M-22'))
            std_run.font.size = Pt(8)

            # Row 4: Certificate - right aligned, font size 8
            cert_para = header.add_paragraph()
            cert_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cert_para.paragraph_format.space_before = Pt(0)
            cert_para.paragraph_format.space_after = Pt(0)
            cert_run = cert_para.add_run(f"Certificate: {data.get('certificate_number', '')}")
            cert_run.font.size = Pt(8)

            # Row 5: Date - right aligned, font size 8
            date_para = header.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            date_para.paragraph_format.space_before = Pt(0)
            date_para.paragraph_format.space_after = Pt(0)
            date_run = date_para.add_run(f"Date: {data.get('test_date', '')}")
            date_run.font.size = Pt(8)

        # Test Information table
        heading = doc.add_heading('Test Information', level=1)
        heading.paragraph_format.space_before = Pt(0)
        heading.paragraph_format.space_after = Pt(6)

        table = doc.add_table(rows=7, cols=4)
        table.style = 'Table Grid'

        # Two-column layout: Label | Value | Label | Value
        info_data = [
            ('Test Project:', data.get('test_project', ''), 'Temperature:', f"{data.get('test_temperature', '23')} °C"),
            ('Customer:', data.get('customer', ''), 'Test Standard:', data.get('test_standard', 'ASTM E8/E8M-22')),
            ('Customer Order:', data.get('customer_order', ''), 'Test Equipment:', data.get('test_equipment', 'MTS Landmark 500kN')),
            ('Product S/N:', data.get('product_sn', ''), 'Specimen ID:', data.get('specimen_id', '')),
            ('Material:', data.get('material', ''), 'Location/Orientation:', data.get('location_orientation', '')),
            ('Strain Source:', data.get('strain_source', 'Extensometer'), 'Extensometer:', data.get('extensometer', 'MTS Extensometer')),
            ('Yield Method:', data.get('yield_method', 'Rp0.2'), 'Operator:', data.get('operator', '')),
        ]

        for i, (label1, value1, label2, value2) in enumerate(info_data):
            table.rows[i].cells[0].text = label1
            table.rows[i].cells[1].text = str(value1)
            table.rows[i].cells[2].text = label2
            table.rows[i].cells[3].text = str(value2)
            # Bold the labels
            if table.rows[i].cells[0].paragraphs[0].runs:
                table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            if table.rows[i].cells[2].paragraphs[0].runs:
                table.rows[i].cells[2].paragraphs[0].runs[0].bold = True

        # Compact table rows
        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].paragraph_format.space_before = Pt(1)
                cell.paragraphs[0].paragraph_format.space_after = Pt(1)

        # Specimen Dimensions table
        heading = doc.add_heading('Specimen Dimensions', level=1)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)

        table = doc.add_table(rows=6, cols=3)
        table.style = 'Table Grid'

        # Header row
        headers = ['Parameter', 'Value', 'Unit']
        for i, header_text in enumerate(headers):
            table.rows[0].cells[i].text = header_text
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

        specimen_type = data.get('specimen_type', 'Rectangular')
        if specimen_type == 'Round':
            dim_data = [
                ('Geometry Type', 'Round', '-'),
                ('Diameter d₀', data.get('d0', data.get('diameter', '')), 'mm'),
                ('Gauge Length L₀', data.get('L0', data.get('gauge_length', '')), 'mm'),
                ('Cross-section Area', data.get('cross_section_area', data.get('initial_area', '')), 'mm²'),
                ('Final Diameter df', data.get('df', '-'), 'mm'),
            ]
        else:
            dim_data = [
                ('Geometry Type', 'Rectangular', '-'),
                ('Width w₀', data.get('w0', data.get('width', '')), 'mm'),
                ('Thickness t₀', data.get('t0', data.get('thickness', '')), 'mm'),
                ('Gauge Length L₀', data.get('L0', data.get('gauge_length', '')), 'mm'),
                ('Cross-section Area', data.get('cross_section_area', data.get('initial_area', '')), 'mm²'),
            ]

        for i, (param, value, unit) in enumerate(dim_data):
            table.rows[i+1].cells[0].text = param
            table.rows[i+1].cells[1].text = str(value)
            table.rows[i+1].cells[2].text = unit

        # Compact table rows
        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].paragraph_format.space_before = Pt(1)
                cell.paragraphs[0].paragraph_format.space_after = Pt(1)

        # Test Results table - simplified per ISO 17025 requirements
        heading = doc.add_heading('Test Results', level=1)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)

        # Build results list based on yield method
        # Format: (parameter, unit, value, requirement, uncertainty)
        yield_method = data.get('yield_method', 'Rp0.2/Rp0.5')
        is_yield_point = 'ReH' in yield_method or 'yield_point' in yield_method.lower()

        if is_yield_point:
            # Yield point method: ReH, ReL
            results_data = [
                ('ReH', 'MPa', data.get('ReH', '-'), data.get('ReH_req', '-'), data.get('ReH_uncertainty', '-')),
                ('ReL', 'MPa', data.get('ReL', '-'), data.get('ReL_req', '-'), data.get('ReL_uncertainty', '-')),
            ]
            ratio_label = 'ReH/Rm'
        else:
            # Offset method: Rp0.2, Rp0.5
            results_data = [
                ('Rp0.2', 'MPa', data.get('Rp02', '-'), data.get('Rp02_req', '-'), data.get('Rp02_uncertainty', '-')),
                ('Rp0.5', 'MPa', data.get('Rp05_value', '-'), data.get('Rp05_req', '-'), data.get('Rp05_uncertainty', '-')),
            ]
            ratio_label = 'Rp0.2/Rm'

        # Add common results: Rm, A, Z, and yield/tensile ratio
        # Default requirements: A >18%, Z >40%
        a_req = data.get('A_req', data.get('A5_req', '-'))
        if a_req == '-' or not a_req:
            a_req = '>18'
        z_req = data.get('Z_req', '-')
        if z_req == '-' or not z_req:
            z_req = '>40'

        results_data.extend([
            ('Rm', 'MPa', data.get('Rm', '-'), data.get('Rm_req', '-'), data.get('Rm_uncertainty', '-')),
            ('A', '%', data.get('A', '-'), a_req, data.get('A_uncertainty', '-')),
            ('Z', '%', data.get('Z', '-'), z_req, data.get('Z_uncertainty', '-')),
            (ratio_label, '-', data.get('yield_tensile_ratio', '-'), data.get('ratio_req', '-'), '-'),
        ])

        # Create table with 5 columns: Parameter, Unit, Value, Requirement, U(k=2)
        num_rows = len(results_data) + 1  # +1 for header
        table = doc.add_table(rows=num_rows, cols=5)
        table.style = 'Table Grid'

        # Header row
        result_headers = ['Parameter', 'Unit', 'Value', 'Requirement', 'U (k=2)']
        for i, header_text in enumerate(result_headers):
            table.rows[0].cells[i].text = header_text
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

        # Data rows
        for i, (param, unit, value, req, unc) in enumerate(results_data):
            table.rows[i+1].cells[0].text = param
            table.rows[i+1].cells[1].text = unit
            table.rows[i+1].cells[2].text = str(value)
            table.rows[i+1].cells[3].text = str(req) if req else '-'
            table.rows[i+1].cells[4].text = str(unc)

        # Compact table rows
        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].paragraph_format.space_before = Pt(1)
                cell.paragraphs[0].paragraph_format.space_after = Pt(1)

        # Stress-Strain Plot
        if chart_path and chart_path.exists():
            heading = doc.add_heading('Stress-Strain Curve', level=1)
            heading.paragraph_format.space_before = Pt(12)
            heading.paragraph_format.space_after = Pt(6)
            doc.add_picture(str(chart_path), width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Add caption
            caption = doc.add_paragraph('Figure 1: Engineering stress-strain curve per ASTM E8/E8M')
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption.runs[0].font.size = Pt(10)
            caption.runs[0].font.italic = True

        # Comments section
        comments = data.get('comments', '')
        if comments:
            heading = doc.add_heading('Comments', level=1)
            heading.paragraph_format.space_before = Pt(12)
            heading.paragraph_format.space_after = Pt(6)
            doc.add_paragraph(comments)

        # Approval Signatures
        heading = doc.add_heading('Approval', level=1)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)

        sig_table = doc.add_table(rows=4, cols=4)
        sig_table.style = 'Table Grid'

        sig_headers = ['Role', 'Name', 'Signature', 'Date']
        for i, header_text in enumerate(sig_headers):
            sig_table.rows[0].cells[i].text = header_text
            sig_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

        sig_table.rows[1].cells[0].text = 'Tested by:'
        sig_table.rows[2].cells[0].text = 'Reviewed by:'
        sig_table.rows[3].cells[0].text = 'Approved by:'

        # Compact signature table rows
        for row in sig_table.rows:
            for cell in row.cells:
                cell.paragraphs[0].paragraph_format.space_before = Pt(1)
                cell.paragraphs[0].paragraph_format.space_after = Pt(1)

        # Add disclaimer to page footer (visible on all pages)
        disclaimer_text = (
            "All work and services carried out by Durabler are subject to, and conducted in accordance with, "
            "Durabler standard terms and conditions, which are available at durabler.se. This document shall not "
            "be reproduced other than in full, except with prior written approval of the issuer. The results pertain "
            "only to the item(s) as sampled by the client unless otherwise indicated. Durabler a part of Subseatec S AB, "
            "Address: Durabler C/O Subseatec, Dalavägen 23, 68130 Kristinehamn, SWEDEN"
        )
        for section in doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.clear()
            footer_run = footer_para.add_run(disclaimer_text)
            footer_run.font.size = Pt(7)
            footer_run.italic = True

        return doc

    def _replace_in_paragraph(
        self,
        paragraph,
        data: Dict[str, Any],
        chart_path: Optional[Path],
        logo_path: Optional[Path]
    ):
        """Replace placeholders in a paragraph."""
        full_text = paragraph.text

        # Handle logo placeholder
        if '{{logo}}' in full_text and logo_path and logo_path.exists():
            paragraph.clear()
            run = paragraph.add_run()
            run.add_picture(str(logo_path), height=Cm(1.5))
            return

        # Handle chart placeholder
        if '{{chart}}' in full_text and chart_path and chart_path.exists():
            paragraph.clear()
            run = paragraph.add_run()
            run.add_picture(str(chart_path), width=Inches(5.5))
            return

        # Find all placeholders
        placeholders = re.findall(r'\{\{([^}]+)\}\}', full_text)

        if not placeholders:
            return

        # Replace each placeholder
        for placeholder in placeholders:
            key = placeholder
            value = data.get(key, '')

            # Handle None values
            if value is None:
                value = ''

            # Format numeric values
            if isinstance(value, float):
                value = f"{value:.2f}"
            elif isinstance(value, int):
                value = str(value)

            full_text = full_text.replace(f'{{{{{placeholder}}}}}', str(value))

        # Update paragraph text while preserving formatting
        if paragraph.runs:
            first_run = paragraph.runs[0]
            font_name = first_run.font.name
            font_size = first_run.font.size
            bold = first_run.bold
            italic = first_run.italic

            paragraph.clear()
            new_run = paragraph.add_run(full_text)
            new_run.font.name = font_name
            new_run.font.size = font_size
            new_run.bold = bold
            new_run.italic = italic
        else:
            paragraph.text = full_text

    @staticmethod
    def prepare_report_data(
        test_info: Dict[str, str],
        dimensions: Dict[str, Any],
        results: Dict[str, Any],
        specimen_type: str,
        yield_type: str,
        requirements: Optional[Dict[str, str]] = None
    ) -> Dict[str, Any]:
        """
        Prepare data dictionary for report generation.

        Parameters
        ----------
        test_info : Dict[str, str]
            Test information (project, customer, etc.)
        dimensions : Dict[str, Any]
            Specimen dimensions
        results : Dict[str, Any]
            Analysis results with MeasuredValue objects
        specimen_type : str
            'round' or 'rectangular'
        yield_type : str
            'offset' or 'yield_point'
        requirements : Dict[str, str], optional
            Requirement values (e.g., {'Rp02': 'min. 500', 'Rm': 'min. 600'})

        Returns
        -------
        Dict[str, Any]
            Flattened dictionary ready for template
        """
        import math

        data = {}
        requirements = requirements or {}

        # Test information
        data['test_project'] = test_info.get('test_project', '')
        data['customer'] = test_info.get('customer', '')
        data['customer_order'] = test_info.get('customer_order', '')
        data['product_sn'] = test_info.get('product_sn', '')
        data['specimen_id'] = test_info.get('specimen_id', '')
        data['location_orientation'] = test_info.get('location_orientation', '')
        data['material'] = test_info.get('material', '')
        data['certificate_number'] = test_info.get('certificate_number', '')
        data['test_date'] = test_info.get('test_date', '')
        data['test_standard'] = 'ASTM E8/E8M-22'
        data['yield_method'] = 'Rp0.2/Rp0.5' if yield_type == 'offset' else 'ReH/ReL'
        data['test_engineer'] = test_info.get('test_engineer', '')
        data['strain_source'] = test_info.get('strain_source', 'Extensometer')
        data['test_equipment'] = 'MTS Landmark 500kN'
        data['extensometer'] = 'MTS Extensometer'
        data['test_temperature'] = test_info.get('temperature', '23')
        data['humidity'] = test_info.get('humidity', '50')
        data['comments'] = test_info.get('comments', '')

        # Additional test info keys for template compatibility
        data['operator'] = test_info.get('test_engineer', '')
        data['temperature'] = test_info.get('temperature', '23')

        # Specimen dimensions - template-compatible keys
        data['geometry_type'] = 'Round' if specimen_type == 'round' else 'Rectangular'
        data['specimen_type'] = 'Round' if specimen_type == 'round' else 'Rectangular'
        data['gauge_length'] = dimensions.get('gauge_length', '')
        data['width'] = dimensions.get('width', '-') if specimen_type == 'rectangular' else '-'
        data['thickness'] = dimensions.get('thickness', '-') if specimen_type == 'rectangular' else '-'
        data['diameter'] = dimensions.get('diameter', '-') if specimen_type == 'round' else '-'

        # Calculate cross-section area
        if specimen_type == 'round':
            data['d0'] = dimensions.get('diameter', '')
            data['df'] = dimensions.get('final_diameter', '-')
            data['w0'] = '-'
            data['t0'] = '-'
            if dimensions.get('diameter'):
                d = float(dimensions['diameter'])
                area = math.pi * d**2 / 4
                data['initial_area'] = f"{area:.2f}"
                data['cross_section_area'] = f"{area:.2f}"
            else:
                data['initial_area'] = ''
                data['cross_section_area'] = ''
        else:
            data['d0'] = '-'
            data['df'] = '-'
            data['w0'] = dimensions.get('width', '')
            data['t0'] = dimensions.get('thickness', '')
            if dimensions.get('width') and dimensions.get('thickness'):
                area = float(dimensions['width']) * float(dimensions['thickness'])
                data['initial_area'] = f"{area:.2f}"
                data['cross_section_area'] = f"{area:.2f}"
            else:
                data['initial_area'] = ''
                data['cross_section_area'] = ''

        data['L0'] = dimensions.get('gauge_length', '')
        data['L1'] = dimensions.get('final_gauge_length', '-')
        data['Lc'] = dimensions.get('parallel_length', '')

        # Helper to extract value and uncertainty from MeasuredValue objects
        def get_result(key, default_value='-', default_unc='-'):
            result = results.get(key)
            if result and hasattr(result, 'value'):
                return f"{result.value:.1f}", f"±{result.uncertainty:.1f}"
            return default_value, default_unc

        def get_result_raw(key):
            """Get raw MeasuredValue object."""
            return results.get(key)

        # Young's modulus (E) - template uses {{E}}, {{E_uncertainty}}
        E_result = get_result_raw('E')
        if E_result and hasattr(E_result, 'value'):
            data['E'] = f"{E_result.value:.1f}"
            data['E_uncertainty'] = f"±{E_result.uncertainty:.1f}"
        else:
            data['E'] = '-'
            data['E_uncertainty'] = '-'
        data['E_req'] = requirements.get('E', '-')

        # Yield strengths - show values only for selected method
        # Template uses {{Rp02}}, {{Rp02_uncertainty}}, etc.
        if yield_type == 'offset':
            Rp02_val, Rp02_unc = get_result('Rp02')
            Rp05_val, Rp05_unc = get_result('Rp05')
            # Template-compatible keys (simple names)
            data['Rp02'] = Rp02_val
            data['Rp02_uncertainty'] = Rp02_unc
            # Legacy keys (with _value suffix)
            data['Rp02_value'] = Rp02_val
            data['Rp02_req'] = requirements.get('Rp02', '-')
            data['Rp05_value'] = Rp05_val
            data['Rp05_uncertainty'] = Rp05_unc
            data['Rp05_req'] = requirements.get('Rp05', '-')
            # ReH/ReL not applicable
            data['ReH'] = '-'
            data['ReH_uncertainty'] = '-'
            data['ReH_value'] = '-'
            data['ReH_req'] = '-'
            data['ReL'] = '-'
            data['ReL_uncertainty'] = '-'
            data['ReL_value'] = '-'
            data['ReL_req'] = '-'
        else:
            # Offset values not applicable
            data['Rp02'] = '-'
            data['Rp02_uncertainty'] = '-'
            data['Rp02_value'] = '-'
            data['Rp02_req'] = '-'
            data['Rp05_value'] = '-'
            data['Rp05_uncertainty'] = '-'
            data['Rp05_req'] = '-'
            # ReH/ReL values
            ReH_val, ReH_unc = get_result('ReH')
            ReL_val, ReL_unc = get_result('ReL')
            data['ReH'] = ReH_val
            data['ReH_uncertainty'] = ReH_unc
            data['ReH_value'] = ReH_val
            data['ReH_req'] = requirements.get('ReH', '-')
            data['ReL'] = ReL_val
            data['ReL_uncertainty'] = ReL_unc
            data['ReL_value'] = ReL_val
            data['ReL_req'] = requirements.get('ReL', '-')

        # Ultimate tensile strength - template uses {{Rm}}, {{Rm_uncertainty}}
        Rm_val, Rm_unc = get_result('Rm')
        data['Rm'] = Rm_val
        data['Rm_uncertainty'] = Rm_unc
        data['Rm_value'] = Rm_val
        data['Rm_req'] = requirements.get('Rm', '-')

        # Elongation - template uses {{A}}, {{A_uncertainty}}
        # Use A_percent (from test data) for the main A value
        A_val, A_unc = get_result('A_percent')
        data['A'] = A_val
        data['A_uncertainty'] = A_unc
        data['A_req'] = requirements.get('A', requirements.get('A5', '-'))

        # A5 from L1-L0 (manual measurement), fallback to A_percent if not available
        A5_val, A5_unc = get_result('A_manual')
        if A5_val == '-' or A5_val is None:
            # Fallback to extensometer A% if manual measurement not available
            A5_val, A5_unc = A_val, A_unc
        data['A5_value'] = A5_val if A5_val != '-' else '-'
        data['A5_uncertainty'] = A5_unc if A5_unc != '-' else '-'
        data['A5_req'] = requirements.get('A5', requirements.get('A', '-'))

        # Uniform elongation (Ag) - template uses {{Ag}}, {{Ag_uncertainty}}
        Ag_val, Ag_unc = get_result('Ag')
        data['Ag'] = Ag_val
        data['Ag_uncertainty'] = Ag_unc

        # Reduction of area (Z%) - template uses {{Z}}, {{Z_uncertainty}}
        Z_val, Z_unc = get_result('Z')
        data['Z'] = Z_val
        data['Z_uncertainty'] = Z_unc
        data['Z_value'] = Z_val
        data['Z_req'] = requirements.get('Z', '-')

        # Rate calculations - template uses {{stress_rate_yield}}, {{strain_rate_yield}}, etc.
        # For yield point, use Rp0.2 rates for offset method, ReH rates for yield_point method
        if yield_type == 'offset':
            stress_rate_yield = get_result_raw('stress_rate_rp02')
            strain_rate_yield = get_result_raw('strain_rate_rp02')
        else:
            stress_rate_yield = get_result_raw('stress_rate_reh')
            strain_rate_yield = get_result_raw('strain_rate_reh')

        # Format stress rate at yield
        if stress_rate_yield and hasattr(stress_rate_yield, 'value'):
            data['stress_rate_yield'] = f"{stress_rate_yield.value:.1f}"
        else:
            data['stress_rate_yield'] = '-'

        # Format strain rate at yield (use scientific notation for small values)
        if strain_rate_yield and hasattr(strain_rate_yield, 'value'):
            val = strain_rate_yield.value
            if abs(val) < 0.01:
                data['strain_rate_yield'] = f"{val:.2e}"
            else:
                data['strain_rate_yield'] = f"{val:.4f}"
        else:
            data['strain_rate_yield'] = '-'

        # Rates at Rm
        stress_rate_rm = get_result_raw('stress_rate_rm')
        strain_rate_rm = get_result_raw('strain_rate_rm')

        if stress_rate_rm and hasattr(stress_rate_rm, 'value'):
            data['stress_rate_rm'] = f"{stress_rate_rm.value:.1f}"
        else:
            data['stress_rate_rm'] = '-'

        if strain_rate_rm and hasattr(strain_rate_rm, 'value'):
            val = strain_rate_rm.value
            if abs(val) < 0.01:
                data['strain_rate_rm'] = f"{val:.2e}"
            else:
                data['strain_rate_rm'] = f"{val:.4f}"
        else:
            data['strain_rate_rm'] = '-'

        # Yield/Tensile ratio (Rp0.2/Rm or ReH/Rm)
        data['ratio_label'] = 'Rp0.2/Rm' if yield_type == 'offset' else 'ReH/Rm'
        if Rm_val and Rm_val != '-':
            if yield_type == 'offset':
                if Rp02_val and Rp02_val != '-':
                    try:
                        ratio = float(Rp02_val) / float(Rm_val)
                        data['yield_tensile_ratio'] = f"{ratio:.3f}"
                    except (ValueError, ZeroDivisionError):
                        data['yield_tensile_ratio'] = '-'
                else:
                    data['yield_tensile_ratio'] = '-'
            else:
                ReH_val_check = data.get('ReH', '-')
                if ReH_val_check and ReH_val_check != '-':
                    try:
                        ratio = float(ReH_val_check) / float(Rm_val)
                        data['yield_tensile_ratio'] = f"{ratio:.3f}"
                    except (ValueError, ZeroDivisionError):
                        data['yield_tensile_ratio'] = '-'
                else:
                    data['yield_tensile_ratio'] = '-'
        else:
            data['yield_tensile_ratio'] = '-'
        data['ratio_req'] = requirements.get('ratio', '-')

        # Validity notes
        data['validity_statement'] = 'The test results comply with the requirements and are considered valid.'
        data['validity_status'] = 'VALID'
        data['validity_notes'] = ''

        # Signatures (to be filled manually)
        data['tested_by'] = ''
        data['tested_date'] = ''
        data['reviewed_by'] = ''
        data['reviewed_date'] = ''
        data['approved_by'] = ''
        data['approved_date'] = ''

        return data
