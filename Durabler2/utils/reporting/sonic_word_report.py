"""
Word report generator for Sonic Resonance (E1875) test results.

Populates a Word template with test data and results.
"""

import re
from pathlib import Path
from typing import Dict, Any, Optional
from docx import Document
from docx.shared import Inches, Cm, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


class SonicReportGenerator:
    """
    Generate Sonic Resonance test reports from Word template or from scratch.

    Uses placeholder syntax {{placeholder_name}} in template.
    """

    def __init__(self, template_path: Optional[Path] = None):
        """
        Initialize report generator.

        Parameters
        ----------
        template_path : Path, optional
            Path to Word template file. If None, creates report from scratch.
        """
        self.template_path = template_path
        if template_path and not template_path.exists():
            raise FileNotFoundError(f"Template not found: {template_path}")

    def generate_report(
        self,
        output_path: Path,
        data: Dict[str, Any],
        chart_path: Optional[Path] = None,
        logo_path: Optional[Path] = None
    ) -> Path:
        """
        Generate report by populating template with data.

        Parameters
        ----------
        output_path : Path
            Path for output Word document
        data : Dict[str, Any]
            Dictionary of placeholder values
        chart_path : Path, optional
            Path to velocity chart image
        logo_path : Path, optional
            Path to logo image

        Returns
        -------
        Path
            Path to generated report
        """
        # Create from scratch if no template
        if not self.template_path or not self.template_path.exists():
            doc = self._create_report_from_scratch(data, chart_path, logo_path)
            doc.save(output_path)
            return output_path

        doc = Document(self.template_path)

        # Replace placeholders in page headers
        for section in doc.sections:
            header = section.header
            for paragraph in header.paragraphs:
                self._replace_in_paragraph(paragraph, data, chart_path, logo_path)
            for table in header.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            self._replace_in_paragraph(paragraph, data, chart_path, logo_path)

        # Replace placeholders in paragraphs
        for paragraph in doc.paragraphs:
            self._replace_in_paragraph(paragraph, data, chart_path, logo_path)

        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        self._replace_in_paragraph(paragraph, data, chart_path, logo_path)

        # Add disclaimer to page footer (visible on all pages)
        disclaimer_text = (
            "All work and services carried out by Durabler are subject to, and conducted in accordance with, "
            "Durabler standard terms and conditions, which are available at durabler.se. This document shall not "
            "be reproduced other than in full, except with prior written approval of the issuer. The results pertain "
            "only to the item(s) as sampled by the client unless otherwise indicated. Durabler a part of Subseatec S AB, "
            "Address: Durabler C/O Subseatec, Dalavägen 23, 68130 Kristinehamn, SWEDEN"
        )
        for section in doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.clear()
            footer_run = footer_para.add_run(disclaimer_text)
            footer_run.font.size = Pt(7)
            footer_run.italic = True

        doc.save(output_path)
        return output_path

    def _create_report_from_scratch(
        self,
        data: Dict[str, Any],
        chart_path: Optional[Path],
        logo_path: Optional[Path]
    ) -> Document:
        """Create report without template - matches Vickers layout."""
        from docx.shared import RGBColor

        doc = Document()

        # Set compact paragraph spacing for entire document
        style = doc.styles['Normal']
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(3)
        style.paragraph_format.line_spacing = 1.0
        style.font.name = 'Calibri'
        style.font.size = Pt(10)

        # Set compact heading styles with dark green color
        dark_green = RGBColor(0x00, 0x64, 0x00)
        for i in range(1, 4):
            heading_style = doc.styles[f'Heading {i}']
            heading_style.paragraph_format.space_before = Pt(8)
            heading_style.paragraph_format.space_after = Pt(4)
            heading_style.font.color.rgb = dark_green

        # Add header with logo on left, certificate info on right (5-line layout)
        for section in doc.sections:
            # Set narrower margins for compact layout
            section.top_margin = Cm(1.5)
            section.bottom_margin = Cm(1.5)
            section.left_margin = Cm(2.0)
            section.right_margin = Cm(2.0)
            header = section.header
            header.is_linked_to_previous = False

            # Row 1: Logo - left aligned
            logo_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
            logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            logo_para.paragraph_format.space_after = Pt(0)
            if logo_path and logo_path.exists():
                logo_run = logo_para.add_run()
                logo_run.add_picture(str(logo_path), width=Cm(5.0))  # 50mm width

            # Row 2: Title - centered, font size 12
            title_para = header.add_paragraph()
            title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            title_para.paragraph_format.space_before = Pt(0)
            title_para.paragraph_format.space_after = Pt(0)
            title_run = title_para.add_run('Sonic Resonance Test Report')
            title_run.bold = True
            title_run.font.size = Pt(12)

            # Row 3: Standard - centered, font size 8
            std_para = header.add_paragraph()
            std_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            std_para.paragraph_format.space_before = Pt(0)
            std_para.paragraph_format.space_after = Pt(0)
            std_run = std_para.add_run('Modified ASTM E1875')
            std_run.font.size = Pt(8)

            # Row 4: Certificate - right aligned, font size 8
            cert_para = header.add_paragraph()
            cert_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            cert_para.paragraph_format.space_before = Pt(0)
            cert_para.paragraph_format.space_after = Pt(0)
            cert_run = cert_para.add_run(f"Certificate: {data.get('certificate_number', '')}")
            cert_run.font.size = Pt(8)

            # Row 5: Date - right aligned, font size 8
            date_para = header.add_paragraph()
            date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            date_para.paragraph_format.space_before = Pt(0)
            date_para.paragraph_format.space_after = Pt(0)
            date_run = date_para.add_run(f"Date: {data.get('test_date', '')}")
            date_run.font.size = Pt(8)

        # Test Information table (exclude certificate and date - now in header)
        heading = doc.add_heading('Test Information', level=1)
        heading.paragraph_format.space_before = Pt(0)
        heading.paragraph_format.space_after = Pt(6)

        table = doc.add_table(rows=6, cols=4)
        table.style = 'Table Grid'

        # Two-column layout: Label | Value | Label | Value (exclude cert/date)
        info_data = [
            ('Test Project:', data.get('test_project', ''), 'Temperature:', f"{data.get('temperature', '23')} °C"),
            ('Customer:', data.get('customer', ''), 'Test Standard:', 'Modified ASTM E1875'),
            ('Customer Order:', data.get('customer_order', ''), 'Test Equipment:', 'Ultrasonic Tester'),
            ('Product S/N:', data.get('product_sn', ''), 'Specimen ID:', data.get('specimen_id', '')),
            ('Material:', data.get('material', ''), 'Location/Orientation:', data.get('location_orientation', '')),
            ('Specimen Type:', data.get('specimen_type', ''), '', ''),
        ]

        for i, (label1, value1, label2, value2) in enumerate(info_data):
            table.rows[i].cells[0].text = label1
            table.rows[i].cells[1].text = str(value1)
            table.rows[i].cells[2].text = label2
            table.rows[i].cells[3].text = str(value2)
            # Bold the labels
            if table.rows[i].cells[0].paragraphs[0].runs:
                table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
            if table.rows[i].cells[2].paragraphs[0].runs:
                table.rows[i].cells[2].paragraphs[0].runs[0].bold = True

        # Compact table rows
        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].paragraph_format.space_before = Pt(1)
                cell.paragraphs[0].paragraph_format.space_after = Pt(1)

        # Specimen Geometry table
        heading = doc.add_heading('Specimen Geometry', level=1)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)

        table = doc.add_table(rows=5, cols=3)
        table.style = 'Table Grid'

        # Header row
        headers = ['Parameter', 'Value', 'Unit']
        for i, header in enumerate(headers):
            table.rows[0].cells[i].text = header
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

        # Show diameter or side length based on specimen type
        dim_data = [
            ('Specimen Type', data.get('specimen_type', ''), '-'),
            ('Diameter' if data.get('specimen_type') == 'Cylinder' else 'Side Length',
             data.get('diameter', data.get('side_length', '')), 'mm'),
            ('Length', data.get('length', ''), 'mm'),
            ('Mass', data.get('mass', ''), 'g'),
        ]

        for i, (param, value, unit) in enumerate(dim_data):
            table.rows[i+1].cells[0].text = param
            table.rows[i+1].cells[1].text = str(value)
            table.rows[i+1].cells[2].text = unit

        # Compact table rows
        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].paragraph_format.space_before = Pt(1)
                cell.paragraphs[0].paragraph_format.space_after = Pt(1)

        # Velocity Measurements table
        heading = doc.add_heading('Velocity Measurements', level=1)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)

        table = doc.add_table(rows=4, cols=5)
        table.style = 'Table Grid'

        # Header row
        vel_headers = ['Type', 'V1', 'V2', 'V3', 'Average']
        for i, header in enumerate(vel_headers):
            table.rows[0].cells[i].text = header
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

        vel_data = [
            ('Longitudinal (m/s)', data.get('vl1', '-'), data.get('vl2', '-'), data.get('vl3', '-'), data.get('vl_avg', '-')),
            ('Shear (m/s)', data.get('vs1', '-'), data.get('vs2', '-'), data.get('vs3', '-'), data.get('vs_avg', '-')),
            ('Density (kg/m³)', data.get('density', '-'), '-', '-', '-'),
        ]

        for i, (type_name, v1, v2, v3, avg) in enumerate(vel_data):
            table.rows[i+1].cells[0].text = type_name
            table.rows[i+1].cells[1].text = str(v1)
            table.rows[i+1].cells[2].text = str(v2)
            table.rows[i+1].cells[3].text = str(v3)
            table.rows[i+1].cells[4].text = str(avg)

        # Compact table rows
        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].paragraph_format.space_before = Pt(1)
                cell.paragraphs[0].paragraph_format.space_after = Pt(1)

        # Test Results table
        heading = doc.add_heading('Elastic Moduli Results', level=1)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)

        table = doc.add_table(rows=6, cols=4)
        table.style = 'Table Grid'

        # Header row
        result_headers = ['Parameter', 'Value', 'U (k=2)', 'Unit']
        for i, header in enumerate(result_headers):
            table.rows[0].cells[i].text = header
            table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

        results_data = [
            ("Young's Modulus E", data.get('youngs_modulus', '-'), data.get('youngs_modulus_unc', '-'), 'GPa'),
            ('Shear Modulus G', data.get('shear_modulus', '-'), data.get('shear_modulus_unc', '-'), 'GPa'),
            ("Poisson's Ratio ν", data.get('poissons_ratio', '-'), data.get('poissons_ratio_unc', '-'), '-'),
            ('Flexural Frequency', data.get('flexural_frequency', '-'), data.get('flexural_frequency_unc', '-'), 'Hz'),
            ('Torsional Frequency', data.get('torsional_frequency', '-'), data.get('torsional_frequency_unc', '-'), 'Hz'),
        ]

        for i, (param, value, unc, unit) in enumerate(results_data):
            table.rows[i+1].cells[0].text = param
            table.rows[i+1].cells[1].text = str(value)
            table.rows[i+1].cells[2].text = str(unc)
            table.rows[i+1].cells[3].text = unit

        # Compact table rows
        for row in table.rows:
            for cell in row.cells:
                cell.paragraphs[0].paragraph_format.space_before = Pt(1)
                cell.paragraphs[0].paragraph_format.space_after = Pt(1)

        # Chart
        if chart_path and chart_path.exists():
            heading = doc.add_heading('Velocity Chart', level=1)
            heading.paragraph_format.space_before = Pt(12)
            heading.paragraph_format.space_after = Pt(6)
            doc.add_picture(str(chart_path), width=Inches(5.5))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Validity Assessment
        heading = doc.add_heading('Validity Assessment', level=1)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)

        validity_status = data.get('validity_status', '-')
        validity_notes = data.get('validity_notes', '')

        status_para = doc.add_paragraph()
        status_run = status_para.add_run(f"Status: {validity_status}")
        status_run.bold = True

        if validity_notes:
            doc.add_paragraph(validity_notes)

        # Approval Signatures
        heading = doc.add_heading('Approval', level=1)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)

        sig_table = doc.add_table(rows=4, cols=4)
        sig_table.style = 'Table Grid'

        sig_headers = ['Role', 'Name', 'Signature', 'Date']
        for i, header in enumerate(sig_headers):
            sig_table.rows[0].cells[i].text = header
            sig_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

        sig_table.rows[1].cells[0].text = 'Tested by:'
        sig_table.rows[2].cells[0].text = 'Reviewed by:'
        sig_table.rows[3].cells[0].text = 'Approved by:'

        # Compact signature table rows
        for row in sig_table.rows:
            for cell in row.cells:
                cell.paragraphs[0].paragraph_format.space_before = Pt(1)
                cell.paragraphs[0].paragraph_format.space_after = Pt(1)

        # Add disclaimer to page footer (visible on all pages)
        disclaimer_text = (
            "All work and services carried out by Durabler are subject to, and conducted in accordance with, "
            "Durabler standard terms and conditions, which are available at durabler.se. This document shall not "
            "be reproduced other than in full, except with prior written approval of the issuer. The results pertain "
            "only to the item(s) as sampled by the client unless otherwise indicated. Durabler a part of Subseatec S AB, "
            "Address: Durabler C/O Subseatec, Dalavägen 23, 68130 Kristinehamn, SWEDEN"
        )
        for section in doc.sections:
            footer = section.footer
            footer.is_linked_to_previous = False
            footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
            footer_para.clear()
            footer_run = footer_para.add_run(disclaimer_text)
            footer_run.font.size = Pt(7)
            footer_run.italic = True

        return doc

    def _replace_in_paragraph(
        self,
        paragraph,
        data: Dict[str, Any],
        chart_path: Optional[Path],
        logo_path: Optional[Path]
    ):
        """Replace placeholders in a paragraph."""
        full_text = paragraph.text

        # Handle logo placeholder
        if '{{logo}}' in full_text and logo_path and logo_path.exists():
            paragraph.clear()
            run = paragraph.add_run()
            run.add_picture(str(logo_path), height=Cm(1.5))
            return

        # Handle chart placeholder
        if '{{chart}}' in full_text and chart_path and chart_path.exists():
            paragraph.clear()
            run = paragraph.add_run()
            run.add_picture(str(chart_path), width=Inches(5.5))
            return

        # Find all placeholders
        placeholders = re.findall(r'\{\{([^}]+)\}\}', full_text)

        if not placeholders:
            return

        # Replace each placeholder
        for placeholder in placeholders:
            key = placeholder
            value = data.get(key, '')

            if value is None:
                value = ''

            # Format numeric values
            if isinstance(value, float):
                if abs(value) < 0.001:
                    value = f"{value:.4g}"
                elif abs(value) < 1:
                    value = f"{value:.4f}"
                else:
                    value = f"{value:.2f}"
            elif isinstance(value, int):
                value = str(value)

            full_text = full_text.replace(f'{{{{{placeholder}}}}}', str(value))

        # Update paragraph text
        if paragraph.runs:
            first_run = paragraph.runs[0]
            font_name = first_run.font.name
            font_size = first_run.font.size
            bold = first_run.bold
            italic = first_run.italic

            paragraph.clear()
            new_run = paragraph.add_run(full_text)
            new_run.font.name = font_name
            new_run.font.size = font_size
            new_run.bold = bold
            new_run.italic = italic
        else:
            paragraph.text = full_text

    @staticmethod
    def prepare_report_data(
        test_info: Dict[str, str],
        specimen_data: Dict[str, Any],
        velocity_data: Dict[str, Any],
        results: Any
    ) -> Dict[str, Any]:
        """
        Prepare data dictionary for report generation.

        Parameters
        ----------
        test_info : Dict[str, str]
            Test information (project, customer, etc.)
        specimen_data : Dict[str, Any]
            Specimen geometry data
        velocity_data : Dict[str, Any]
            Velocity measurements
        results : SonicResults
            Analysis results

        Returns
        -------
        Dict[str, Any]
            Flattened dictionary ready for template
        """
        data = {}

        # Test information
        data['test_project'] = test_info.get('test_project', '')
        data['customer'] = test_info.get('customer', '')
        data['customer_order'] = test_info.get('customer_order', '')
        data['product_sn'] = test_info.get('product_sn', '')
        data['specimen_id'] = test_info.get('specimen_id', '')
        data['location_orientation'] = test_info.get('location_orientation', '')
        data['material'] = test_info.get('material', '')
        data['certificate_number'] = test_info.get('certificate_number', '')
        data['test_date'] = test_info.get('test_date', '')
        data['temperature'] = test_info.get('temperature', '23')
        data['test_standard'] = 'Modified ASTM E1875'
        data['test_equipment'] = 'Ultrasonic Tester'

        # Specimen geometry
        data['specimen_type'] = specimen_data.get('specimen_type', '')
        data['diameter'] = specimen_data.get('diameter', '-')
        data['side_length'] = specimen_data.get('side_length', '-')
        data['length'] = specimen_data.get('length', '')
        data['mass'] = specimen_data.get('mass', '')

        # Velocity measurements
        data['vl1'] = velocity_data.get('vl1', '-')
        data['vl2'] = velocity_data.get('vl2', '-')
        data['vl3'] = velocity_data.get('vl3', '-')
        data['vs1'] = velocity_data.get('vs1', '-')
        data['vs2'] = velocity_data.get('vs2', '-')
        data['vs3'] = velocity_data.get('vs3', '-')

        # Results
        if results:
            data['density'] = f"{results.density.value:.1f}"
            data['vl_avg'] = f"{results.longitudinal_velocity.value:.1f}"
            data['vs_avg'] = f"{results.shear_velocity.value:.1f}"

            data['poissons_ratio'] = f"{results.poissons_ratio.value:.4f}"
            data['poissons_ratio_unc'] = f"±{results.poissons_ratio.uncertainty:.4f}"

            data['shear_modulus'] = f"{results.shear_modulus.value:.2f}"
            data['shear_modulus_unc'] = f"±{results.shear_modulus.uncertainty:.2f}"

            data['youngs_modulus'] = f"{results.youngs_modulus.value:.2f}"
            data['youngs_modulus_unc'] = f"±{results.youngs_modulus.uncertainty:.2f}"

            data['flexural_frequency'] = f"{results.flexural_frequency.value:.1f}"
            data['flexural_frequency_unc'] = f"±{results.flexural_frequency.uncertainty:.1f}"

            data['torsional_frequency'] = f"{results.torsional_frequency.value:.1f}"
            data['torsional_frequency_unc'] = f"±{results.torsional_frequency.uncertainty:.1f}"

            data['validity_status'] = 'VALID' if results.is_valid else 'CHECK'
            data['validity_notes'] = results.validity_notes
        else:
            data['density'] = '-'
            data['vl_avg'] = '-'
            data['vs_avg'] = '-'
            data['poissons_ratio'] = '-'
            data['poissons_ratio_unc'] = '-'
            data['shear_modulus'] = '-'
            data['shear_modulus_unc'] = '-'
            data['youngs_modulus'] = '-'
            data['youngs_modulus_unc'] = '-'
            data['flexural_frequency'] = '-'
            data['flexural_frequency_unc'] = '-'
            data['torsional_frequency'] = '-'
            data['torsional_frequency_unc'] = '-'
            data['validity_status'] = '-'
            data['validity_notes'] = ''

        # Signatures
        data['tested_by'] = ''
        data['tested_date'] = ''
        data['reviewed_by'] = ''
        data['reviewed_date'] = ''
        data['approved_by'] = ''
        data['approved_date'] = ''

        return data
