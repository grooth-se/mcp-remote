"""Routes for Charpy Impact test module - ASTM E23 / ISO 148-1."""
import os
import json
from datetime import datetime
from pathlib import Path

from flask import (
    render_template, redirect, url_for, flash, request,
    current_app, send_file, Response
)
from flask_login import login_required, current_user
from werkzeug.utils import secure_filename

from . import charpy_bp
from .forms import SpecimenForm, ReportForm
from app.extensions import db
from app.models import (TestRecord, AnalysisResult, AuditLog, Certificate, TestPhoto,
                        ReportApproval, STATUS_DRAFT, STATUS_REJECTED)

from utils.analysis.charpy_calculations import CharpyAnalyzer, CharpyResult
from utils.models.charpy_specimen import CharpyTestData, CharpyReading, CharpySpecimenConfig


def generate_charpy_test_id():
    """Generate unique test ID for Charpy test."""
    today = datetime.now()
    prefix = f"CV-{today.strftime('%y%m%d')}"
    count = TestRecord.query.filter(
        TestRecord.test_id.like(f'{prefix}%')
    ).count()
    return f"{prefix}-{count + 1:03d}"


def create_energy_bar_chart(readings, test_temperature=23.0):
    """Create interactive bar chart of absorbed energy per specimen.

    Parameters
    ----------
    readings : list
        List of dicts with 'specimen_id' and 'absorbed_energy'
    test_temperature : float
        Test temperature in degrees Celsius

    Returns
    -------
    str
        HTML string with Plotly figure
    """
    import plotly.graph_objects as go
    import numpy as np

    ids = []
    energies = []
    for i, r in enumerate(readings, 1):
        if isinstance(r, dict):
            ids.append(r.get('specimen_id', f'#{i}'))
            energies.append(r.get('absorbed_energy', 0))
        else:
            ids.append(r.specimen_id or f'#{i}')
            energies.append(r.absorbed_energy)

    mean_val = np.mean(energies)

    fig = go.Figure()

    fig.add_trace(go.Bar(
        x=ids,
        y=energies,
        marker_color='#0d6efd',
        text=[f'{e:.1f} J' for e in energies],
        textposition='outside',
        hovertemplate='%{x}<br>%{y:.1f} J<extra></extra>'
    ))

    fig.add_hline(
        y=mean_val,
        line_dash='dot',
        line_color='grey',
        line_width=2,
        annotation_text=f'Mean: {mean_val:.1f} J',
        annotation_position='right',
        annotation_font=dict(color='grey')
    )

    fig.update_layout(
        title=f'Absorbed Energy at {test_temperature}°C',
        xaxis_title='Specimen',
        yaxis_title='Absorbed Energy (J)',
        template='plotly_white',
        showlegend=False,
        height=400,
        yaxis=dict(range=[0, max(energies) * 1.25 if energies else 10])
    )

    return fig.to_html(full_html=False, include_plotlyjs='cdn')


@charpy_bp.route('/')
@login_required
def index():
    """List all Charpy tests."""
    tests = TestRecord.query.filter_by(test_method='CHARPY').order_by(
        TestRecord.created_at.desc()
    ).all()
    return render_template('charpy/index.html', tests=tests)


@charpy_bp.route('/new', methods=['GET', 'POST'])
@login_required
def new():
    """Create a new Charpy impact test."""
    form = SpecimenForm()

    # Populate certificate dropdown
    certificates = Certificate.query.order_by(
        Certificate.year.desc(),
        Certificate.cert_id.desc()
    ).all()
    form.certificate_id.choices = [(0, '-- Select Certificate --')] + [
        (c.id, f"{c.certificate_number_with_rev} - {c.customer or 'No customer'}")
        for c in certificates
    ]

    # Pre-fill from certificate if coming from cert page
    cert_id = request.args.get('certificate', type=int)
    if cert_id and request.method == 'GET':
        form.certificate_id.data = cert_id
        cert = Certificate.query.get(cert_id)
        if cert:
            form.material.data = cert.material
            form.specimen_id.data = cert.test_article_sn
            form.customer_specimen_info.data = cert.customer_specimen_info
            form.requirement.data = cert.requirement
            form.location_orientation.data = cert.location_orientation
            if cert.temperature:
                try:
                    temp_str = cert.temperature.replace('\u00b0C', '').replace('C', '').strip()
                    form.test_temperature.data = float(temp_str)
                except (ValueError, AttributeError):
                    form.test_temperature.data = -40.0
            else:
                form.test_temperature.data = -40.0

    if request.method == 'POST' and not form.validate():
        for field_name, errors in form.errors.items():
            for error in errors:
                flash(f'{field_name}: {error}', 'danger')

    if form.validate_on_submit():
        # Link to certificate
        certificate_id = None
        cert_number = None
        if form.certificate_id.data and form.certificate_id.data != 0:
            cert = Certificate.query.get(form.certificate_id.data)
            if cert:
                certificate_id = cert.id
                cert_number = cert.certificate_number_with_rev

        test_id = generate_charpy_test_id()

        # Parse specimen size
        size_parts = form.specimen_size.data.split('x')
        width = float(size_parts[0])
        height = float(size_parts[1])
        notch_depth = 2.0 if form.notch_type.data == 'V' else 5.0

        test_temperature = form.test_temperature.data if form.test_temperature.data is not None else -40.0

        # Build test parameters
        test_params = {
            'notch_type': form.notch_type.data,
            'specimen_size': form.specimen_size.data,
            'specimen_width': width,
            'specimen_height': height,
            'notch_depth': notch_depth,
            'test_temperature': test_temperature,
            'num_specimens': form.num_specimens.data,
            'location_orientation': form.location_orientation.data,
            'notes': form.notes.data,
        }

        test_params['uncertainty_inputs'] = {
            'machine_pct': form.machine_uncertainty.data or 1.0,
            'temperature_c': form.temperature_uncertainty.data or 1.0,
            'dimension_pct': form.dimension_uncertainty.data or 0.5,
        }

        # Create test record
        test = TestRecord(
            test_id=test_id,
            test_method='CHARPY',
            specimen_id=form.specimen_id.data,
            material=form.material.data,
            test_date=form.test_date.data or datetime.now(),
            temperature=test_temperature,
            geometry=test_params,
            status='DRAFT',
            certificate_id=certificate_id,
            certificate_number=cert_number,
            operator_id=current_user.id
        )

        # Collect readings from manual entry
        readings = []
        num_specimens = int(form.num_specimens.data or 3)
        for i in range(1, num_specimens + 1):
            energy_field = getattr(form, f'specimen_{i}_energy', None)
            if energy_field and energy_field.data is not None and energy_field.data >= 0:
                specimen_id_field = getattr(form, f'specimen_{i}_id', None)
                lat_exp_field = getattr(form, f'specimen_{i}_lateral_exp', None)
                shear_field = getattr(form, f'specimen_{i}_shear_area', None)
                readings.append({
                    'specimen_number': len(readings) + 1,
                    'specimen_id': specimen_id_field.data if specimen_id_field and specimen_id_field.data else f'{form.specimen_id.data}-{i}',
                    'absorbed_energy': energy_field.data,
                    'lateral_expansion': lat_exp_field.data if lat_exp_field else None,
                    'shear_fracture_area': shear_field.data if shear_field else None,
                })

        test_params['readings'] = readings

        # Handle photo
        if form.photo.data:
            photo = form.photo.data
            filename = secure_filename(photo.filename)
            filename = f"charpy_{test_id}_{filename}"
            filepath = os.path.join(current_app.config['UPLOAD_FOLDER'], filename)
            photo.save(filepath)
            test_params['photo_path'] = filename

        test.geometry = test_params

        db.session.add(test)
        db.session.flush()

        # Run analysis
        if readings:
            try:
                reading_objects = [
                    CharpyReading(
                        specimen_number=r['specimen_number'],
                        specimen_id=r['specimen_id'],
                        absorbed_energy=r['absorbed_energy'],
                        lateral_expansion=r.get('lateral_expansion'),
                        shear_fracture_area=r.get('shear_fracture_area'),
                    )
                    for r in readings
                ]

                specimen_config = CharpySpecimenConfig(
                    notch_type=form.notch_type.data,
                    width=width,
                    height=height,
                    notch_depth=notch_depth,
                )

                test_data = CharpyTestData(
                    readings=reading_objects,
                    specimen_config=specimen_config,
                    test_temperature=test_temperature,
                    specimen_id=form.specimen_id.data or '',
                    material=form.material.data or '',
                    specimen_orientation=form.location_orientation.data or '',
                )

                analyzer = CharpyAnalyzer(
                    machine_uncertainty=(form.machine_uncertainty.data or 1.0) / 100,
                    temperature_uncertainty=form.temperature_uncertainty.data or 1.0,
                    dimension_uncertainty=(form.dimension_uncertainty.data or 0.5) / 100,
                )
                result = analyzer.run_analysis(test_data)

                import numpy as np
                values = np.array([r['absorbed_energy'] for r in readings])
                uncertainty_budget = analyzer.get_uncertainty_budget(values, result.mean_energy.value)

                # Store results
                results_to_store = [
                    ('mean_energy', result.mean_energy.value, result.mean_energy.uncertainty, 'J'),
                    ('std_dev', result.std_dev, None, 'J'),
                    ('range', result.range_value, None, 'J'),
                    ('min_value', result.min_value, None, 'J'),
                    ('max_value', result.max_value, None, 'J'),
                    ('n_specimens', result.n_specimens, None, '-'),
                ]
                if result.mean_lateral_expansion is not None:
                    results_to_store.append(
                        ('mean_lateral_expansion', result.mean_lateral_expansion, None, 'mm'))
                if result.mean_shear_area is not None:
                    results_to_store.append(
                        ('mean_shear_area', result.mean_shear_area, None, '%'))

                for param_name, value, uncertainty, unit in results_to_store:
                    analysis = AnalysisResult(
                        test_record_id=test.id,
                        parameter_name=param_name,
                        value=value,
                        uncertainty=uncertainty,
                        unit=unit,
                        calculated_by_id=current_user.id
                    )
                    db.session.add(analysis)

                test_params['uncertainty_budget'] = uncertainty_budget
                test.geometry = test_params
                test.status = 'ANALYZED'
                flash(f'Analysis completed: Mean = {result.mean_energy.value:.1f} \u00b1 '
                      f'{result.mean_energy.uncertainty:.1f} J at {test_temperature}\u00b0C', 'success')
            except Exception as e:
                flash(f'Analysis error: {e}', 'warning')
                test.status = 'DRAFT'
        else:
            flash('No specimen readings entered.', 'warning')

        # Store photo in database
        if form.photo.data:
            photo_file = form.photo.data
            photo_file.seek(0)
            photo_data = photo_file.read()
            db_photo = TestPhoto(
                test_record_id=test.id,
                photo_number=1,
                description='Charpy fracture surface photo',
                uploaded_by_id=current_user.id
            )
            db_photo.set_image(photo_data, photo_file.filename)
            db.session.add(db_photo)

        # Audit log
        audit = AuditLog(
            user_id=current_user.id,
            action='CREATE',
            table_name='test_record',
            record_id=test.id,
            new_values=json.dumps({'test_id': test.test_id, 'test_method': 'CHARPY'})
        )
        db.session.add(audit)
        db.session.commit()

        flash(f'Charpy test {test.test_id} created.', 'success')
        return redirect(url_for('charpy.view', test_id=test.id))

    return render_template('charpy/new.html', form=form)


@charpy_bp.route('/<int:test_id>')
@login_required
def view(test_id):
    """View Charpy test details and results."""
    test = TestRecord.query.get_or_404(test_id)

    if test.test_method != 'CHARPY':
        flash('Invalid test type.', 'error')
        return redirect(url_for('charpy.index'))

    test_params = test.geometry if test.geometry else {}
    readings = test_params.get('readings', [])

    # Get analysis results
    results = {}
    analysis_records = AnalysisResult.query.filter_by(test_record_id=test.id).all()
    for ar in analysis_records:
        if ar.parameter_name == 'mean_energy':
            results['mean_energy'] = {'value': ar.value, 'uncertainty': ar.uncertainty}
        else:
            results[ar.parameter_name] = ar.value
    results['uncertainty_budget'] = test_params.get('uncertainty_budget', {})

    # Create bar chart
    energy_chart = None
    if readings:
        temp = test_params.get('test_temperature', test.temperature or 23)
        energy_chart = create_energy_bar_chart(readings, temp)

    # Get photo
    photo_url = None
    db_photo = test.photos.first()
    if db_photo:
        photo_url = url_for('charpy.photo', test_id=test_id, photo_id=db_photo.id)
    elif test_params.get('photo_path'):
        photo_url = url_for('static', filename=f'uploads/{test_params["photo_path"]}')

    return render_template('charpy/view.html',
                           test=test,
                           test_params=test_params,
                           readings=readings,
                           results=results,
                           energy_chart=energy_chart,
                           photo_url=photo_url)


@charpy_bp.route('/<int:test_id>/photo/<int:photo_id>')
@login_required
def photo(test_id, photo_id):
    """Serve photo from database."""
    p = TestPhoto.query.filter_by(id=photo_id, test_record_id=test_id).first_or_404()
    return Response(
        p.data,
        mimetype=p.mime_type or 'image/jpeg',
        headers={'Content-Disposition': f'inline; filename="{p.original_filename}"'}
    )


@charpy_bp.route('/<int:test_id>/report', methods=['GET', 'POST'])
@login_required
def report(test_id):
    """Generate Charpy impact test report."""
    test = TestRecord.query.get_or_404(test_id)

    if test.test_method != 'CHARPY':
        flash('Invalid test type.', 'error')
        return redirect(url_for('charpy.index'))

    form = ReportForm()

    if request.method == 'GET':
        form.certificate_number.data = test.certificate_number or test.test_id

    if form.validate_on_submit():
        try:
            test_params = test.geometry if test.geometry else {}
            readings = test_params.get('readings', [])

            # Get analysis results
            results = {}
            analysis_records = AnalysisResult.query.filter_by(test_record_id=test.id).all()
            for ar in analysis_records:
                if ar.parameter_name == 'mean_energy':
                    results['mean_energy'] = {'value': ar.value, 'uncertainty': ar.uncertainty}
                else:
                    results[ar.parameter_name] = ar.value
            results['uncertainty_budget'] = test_params.get('uncertainty_budget', {})

            # Generate chart image
            chart_path = None
            if readings:
                import matplotlib
                matplotlib.use('Agg')
                import matplotlib.pyplot as plt
                import numpy as np

                energies = [r['absorbed_energy'] for r in readings]
                specimen_ids = [r.get('specimen_id', f'#{i+1}') for i, r in enumerate(readings)]
                mean_val = np.mean(energies)

                fig, ax = plt.subplots(figsize=(8, 5))
                x = range(len(energies))
                bars = ax.bar(x, energies, color='#0d6efd', width=0.6)
                ax.axhline(y=mean_val, color='grey', linestyle=':', linewidth=2, label='Mean')

                for bar, e in zip(bars, energies):
                    ax.text(bar.get_x() + bar.get_width()/2., bar.get_height() + max(energies)*0.02,
                            f'{e:.1f}', ha='center', va='bottom', fontsize=9)

                temp = test_params.get('test_temperature', test.temperature or 23)
                ax.set_xlabel('Specimen')
                ax.set_ylabel('Absorbed Energy (J)')
                ax.set_title(f'Absorbed Energy at {temp}\u00b0C')
                ax.set_xticks(x)
                ax.set_xticklabels(specimen_ids, rotation=45 if len(specimen_ids) > 4 else 0, ha='right')
                ax.legend(loc='upper right')
                ax.set_ylim(0, max(energies) * 1.25)
                ax.grid(True, alpha=0.3, axis='y')

                chart_path = Path(current_app.config['REPORTS_FOLDER']) / f'charpy_chart_{test.id}.png'
                fig.savefig(chart_path, dpi=150, bbox_inches='tight')
                plt.close(fig)

            test_info = {
                'certificate_number': form.certificate_number.data or test.test_id,
                'test_project': test.certificate.test_project if test.certificate else '',
                'customer': test.certificate.customer if test.certificate else '',
                'customer_order': test.certificate.customer_order if test.certificate else '',
                'product_sn': test.certificate.product_sn if test.certificate else '',
                'specimen_id': test.specimen_id or '',
                'customer_specimen_info': test.certificate.customer_specimen_info if test.certificate else '',
                'material': test.material or '',
                'requirement': test.certificate.requirement if test.certificate else '',
                'location_orientation': test_params.get('location_orientation', ''),
                'test_date': test.test_date.strftime('%Y-%m-%d') if test.test_date else '',
                'test_temperature': test_params.get('test_temperature', test.temperature or 23),
                'notch_type': test_params.get('notch_type', 'V'),
                'specimen_size': test_params.get('specimen_size', '10x10'),
                'notes': test_params.get('notes', ''),
                'operator': current_user.full_name if current_user.full_name else current_user.username,
            }

            # Result proxy
            class ResultProxy:
                def __init__(self, data):
                    self._data = data
                @property
                def mean_energy(self):
                    d = self._data.get('mean_energy', {})
                    return type('MV', (), {'value': d.get('value', 0), 'uncertainty': d.get('uncertainty', 0)})()
                @property
                def std_dev(self):
                    return self._data.get('std_dev', 0)
                @property
                def range_value(self):
                    return self._data.get('range', 0)
                @property
                def min_value(self):
                    return self._data.get('min_value', 0)
                @property
                def max_value(self):
                    return self._data.get('max_value', 0)
                @property
                def n_specimens(self):
                    return self._data.get('n_specimens', 0)
                @property
                def mean_lateral_expansion(self):
                    return self._data.get('mean_lateral_expansion')
                @property
                def mean_shear_area(self):
                    return self._data.get('mean_shear_area')

            result_proxy = ResultProxy(results)
            uncertainty_budget = results.get('uncertainty_budget', {})
            requirement_value = test_info.get('requirement', '') or '-'

            # Build Word report
            from docx import Document
            from docx.shared import Inches, Pt, Cm, RGBColor
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            style = doc.styles['Normal']
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after = Pt(3)
            style.paragraph_format.line_spacing = 1.0
            style.font.size = Pt(10)

            dark_green = RGBColor(0x00, 0x64, 0x00)
            for i in range(1, 4):
                heading_style = doc.styles[f'Heading {i}']
                heading_style.paragraph_format.space_before = Pt(8)
                heading_style.paragraph_format.space_after = Pt(4)
                heading_style.font.color.rgb = dark_green

            # Header with logo
            logo_path = Path(current_app.root_path) / 'static' / 'images' / 'logo.png'

            for section in doc.sections:
                section.top_margin = Cm(1.5)
                section.bottom_margin = Cm(1.5)
                section.left_margin = Cm(2.0)
                section.right_margin = Cm(2.0)
                header = section.header
                header.is_linked_to_previous = False

                logo_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
                logo_para.paragraph_format.space_after = Pt(0)
                if logo_path.exists():
                    logo_run = logo_para.add_run()
                    logo_run.add_picture(str(logo_path), width=Cm(5.0))

                title_para = header.add_paragraph()
                title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                title_para.paragraph_format.space_before = Pt(0)
                title_para.paragraph_format.space_after = Pt(0)
                title_run = title_para.add_run('Charpy Impact Test Report')
                title_run.bold = True
                title_run.font.size = Pt(12)

                std_para = header.add_paragraph()
                std_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                std_para.paragraph_format.space_before = Pt(0)
                std_para.paragraph_format.space_after = Pt(0)
                std_run = std_para.add_run('ASTM E23 / ISO 148-1')
                std_run.font.size = Pt(8)

                cert_para = header.add_paragraph()
                cert_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                cert_para.paragraph_format.space_before = Pt(0)
                cert_para.paragraph_format.space_after = Pt(0)
                cert_run = cert_para.add_run(f"Certificate: {test_info.get('certificate_number', '')}")
                cert_run.font.size = Pt(8)

                date_para = header.add_paragraph()
                date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                date_para.paragraph_format.space_before = Pt(0)
                date_para.paragraph_format.space_after = Pt(0)
                date_run = date_para.add_run(f"Date: {test_info.get('test_date', '')}")
                date_run.font.size = Pt(8)

            # Test Information table
            heading = doc.add_heading('Test Information', level=1)
            heading.paragraph_format.space_before = Pt(0)
            heading.paragraph_format.space_after = Pt(6)

            notch_label = 'V-notch (45\u00b0, r=0.25mm)' if test_info['notch_type'] == 'V' else 'U-notch (r=1mm)'
            info_data = [
                ('Test Project:', test_info.get('test_project', ''), 'Test Temperature:', f"{test_info.get('test_temperature', '-40')} \u00b0C"),
                ('Customer:', test_info.get('customer', ''), 'Test Standard:', 'ASTM E23 / ISO 148-1'),
                ('Customer Order:', test_info.get('customer_order', ''), 'Test Equipment:', 'Analog Charpy impact tester'),
                ('Product S/N:', test_info.get('specimen_id', ''), 'Notch Type:', notch_label),
                ('Material:', test_info.get('material', ''), 'Specimen Size:', f"{test_info.get('specimen_size', '10x10')} mm"),
                ('Customer Specimen Info:', test_info.get('customer_specimen_info', ''), 'Notch Orientation:', test_info.get('location_orientation', '')),
                ('Requirement:', test_info.get('requirement', ''), 'Operator:', test_info.get('operator', '')),
            ]

            table = doc.add_table(rows=len(info_data), cols=4)
            table.style = 'Table Grid'
            for i, (label1, value1, label2, value2) in enumerate(info_data):
                table.rows[i].cells[0].text = label1
                table.rows[i].cells[1].text = str(value1) if value1 else ''
                table.rows[i].cells[2].text = label2
                table.rows[i].cells[3].text = str(value2) if value2 else ''
                if table.rows[i].cells[0].paragraphs[0].runs:
                    table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
                if table.rows[i].cells[2].paragraphs[0].runs:
                    table.rows[i].cells[2].paragraphs[0].runs[0].bold = True
                for cell in table.rows[i].cells:
                    cell.paragraphs[0].paragraph_format.space_before = Pt(1)
                    cell.paragraphs[0].paragraph_format.space_after = Pt(1)

            # Results Summary
            heading = doc.add_heading('Results Summary', level=1)
            heading.paragraph_format.space_before = Pt(12)
            heading.paragraph_format.space_after = Pt(6)

            result_rows = [
                ('Mean Absorbed Energy', f'{result_proxy.mean_energy.value:.1f} \u00b1 {result_proxy.mean_energy.uncertainty:.1f}', 'J'),
                ('Standard Deviation', f'{result_proxy.std_dev:.1f}', 'J'),
                ('Range', f'{result_proxy.range_value:.1f}', 'J'),
                ('Minimum', f'{result_proxy.min_value:.1f}', 'J'),
                ('Maximum', f'{result_proxy.max_value:.1f}', 'J'),
                ('Number of Specimens', str(result_proxy.n_specimens), '-'),
            ]
            if result_proxy.mean_lateral_expansion is not None:
                result_rows.append(('Mean Lateral Expansion', f'{result_proxy.mean_lateral_expansion:.2f}', 'mm'))
            if result_proxy.mean_shear_area is not None:
                result_rows.append(('Mean Shear Fracture Area', f'{result_proxy.mean_shear_area:.0f}', '%'))
            result_rows.append(('Requirement', requirement_value, '-'))

            table = doc.add_table(rows=len(result_rows) + 1, cols=3)
            table.style = 'Table Grid'
            for i, h in enumerate(['Parameter', 'Value', 'Unit']):
                table.rows[0].cells[i].text = h
                table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
            for i, (param, value, unit) in enumerate(result_rows):
                table.rows[i+1].cells[0].text = param
                table.rows[i+1].cells[1].text = value
                table.rows[i+1].cells[2].text = unit
            for row in table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].paragraph_format.space_before = Pt(1)
                    cell.paragraphs[0].paragraph_format.space_after = Pt(1)

            # Individual Specimens table
            heading = doc.add_heading('Individual Specimens', level=1)
            heading.paragraph_format.space_before = Pt(12)
            heading.paragraph_format.space_after = Pt(6)

            has_le = any(r.get('lateral_expansion') is not None for r in readings)
            has_sa = any(r.get('shear_fracture_area') is not None for r in readings)
            cols = ['#', 'Specimen ID', 'Energy (J)']
            if has_le:
                cols.append('Lat. Exp. (mm)')
            if has_sa:
                cols.append('Shear (%)')

            table = doc.add_table(rows=len(readings) + 1, cols=len(cols))
            table.style = 'Table Grid'
            for i, h in enumerate(cols):
                table.rows[0].cells[i].text = h
                table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
            for i, r in enumerate(readings):
                table.rows[i+1].cells[0].text = str(r.get('specimen_number', i+1))
                table.rows[i+1].cells[1].text = r.get('specimen_id', f'#{i+1}')
                table.rows[i+1].cells[2].text = f"{r.get('absorbed_energy', 0):.1f}"
                col_idx = 3
                if has_le:
                    le = r.get('lateral_expansion')
                    table.rows[i+1].cells[col_idx].text = f"{le:.2f}" if le is not None else '-'
                    col_idx += 1
                if has_sa:
                    sa = r.get('shear_fracture_area')
                    table.rows[i+1].cells[col_idx].text = f"{sa:.0f}" if sa is not None else '-'
            for row in table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].paragraph_format.space_before = Pt(1)
                    cell.paragraphs[0].paragraph_format.space_after = Pt(1)

            # Chart
            if chart_path and chart_path.exists():
                heading = doc.add_heading('Absorbed Energy Chart', level=1)
                heading.paragraph_format.space_before = Pt(12)
                heading.paragraph_format.space_after = Pt(6)
                doc.add_picture(str(chart_path), width=Inches(5.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Uncertainty Budget
            if form.include_uncertainty_budget.data == 'yes' and uncertainty_budget:
                heading = doc.add_heading('Uncertainty Budget (k=2)', level=1)
                heading.paragraph_format.space_before = Pt(12)
                heading.paragraph_format.space_after = Pt(6)

                table = doc.add_table(rows=5, cols=2)
                table.style = 'Table Grid'
                unc_data = [
                    ('Type A (Repeatability)', f"{uncertainty_budget.get('u_A', 0):.2f}"),
                    ('Machine Calibration', f"{uncertainty_budget.get('u_machine', 0):.2f}"),
                    ('Specimen Dimensions', f"{uncertainty_budget.get('u_dimension', 0):.2f}"),
                    ('Combined Standard (u_c)', f"{uncertainty_budget.get('u_combined', 0):.2f}"),
                    ('Expanded (U, k=2)', f"{uncertainty_budget.get('U_expanded', 0):.2f}"),
                ]
                for i, (comp, val) in enumerate(unc_data):
                    table.rows[i].cells[0].text = comp
                    table.rows[i].cells[1].text = val
                for row in table.rows:
                    for cell in row.cells:
                        cell.paragraphs[0].paragraph_format.space_before = Pt(1)
                        cell.paragraphs[0].paragraph_format.space_after = Pt(1)

            # Fracture Photo
            import tempfile
            temp_photo_files = []
            if form.include_photo.data == 'yes':
                photo_path = None
                db_photo = test.photos.first()
                if db_photo and db_photo.data:
                    ext = (db_photo.original_filename or 'photo.jpg').rsplit('.', 1)[-1]
                    tmp = tempfile.NamedTemporaryFile(suffix=f'.{ext}', delete=False)
                    tmp.write(db_photo.data)
                    tmp.close()
                    photo_path = Path(tmp.name)
                    temp_photo_files.append(photo_path)
                elif test_params.get('photo_path'):
                    p = Path(current_app.config['UPLOAD_FOLDER']) / test_params['photo_path']
                    if p.exists():
                        photo_path = p

                if photo_path and photo_path.exists():
                    heading = doc.add_heading('Fracture Surface', level=1)
                    heading.paragraph_format.space_before = Pt(12)
                    heading.paragraph_format.space_after = Pt(6)
                    doc.add_picture(str(photo_path), width=Inches(4))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Approval signatures
            heading = doc.add_heading('Approval', level=1)
            heading.paragraph_format.space_before = Pt(12)
            heading.paragraph_format.space_after = Pt(6)

            sig_table = doc.add_table(rows=4, cols=2)
            sig_table.style = 'Table Grid'
            sig_table.rows[0].cells[0].text = 'Role'
            sig_table.rows[0].cells[0].paragraphs[0].runs[0].bold = True
            sig_table.rows[0].cells[1].text = 'Name / Signature'
            sig_table.rows[0].cells[1].paragraphs[0].runs[0].bold = True
            sig_table.rows[1].cells[0].text = 'Test Engineer:'
            sig_table.rows[1].cells[0].paragraphs[0].runs[0].bold = True
            sig_table.rows[1].cells[1].text = test_info.get('operator', '')
            sig_table.rows[2].cells[0].text = 'Approved by:'
            sig_table.rows[2].cells[0].paragraphs[0].runs[0].bold = True
            sig_table.rows[2].cells[1].text = ''
            sig_table.rows[3].cells[0].text = 'Third Party Approval:'
            sig_table.rows[3].cells[0].paragraphs[0].runs[0].bold = True
            sig_table.rows[3].cells[1].text = ''

            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            sig_row_height = Cm(1.5)
            for row in sig_table.rows:
                for cell in row.cells:
                    cell.paragraphs[0].paragraph_format.space_before = Pt(1)
                    cell.paragraphs[0].paragraph_format.space_after = Pt(1)
            for row in [sig_table.rows[1], sig_table.rows[2], sig_table.rows[3]]:
                tr = row._tr
                trPr = tr.get_or_add_trPr()
                trHeight = trPr.find(qn('w:trHeight'))
                if trHeight is None:
                    trHeight = OxmlElement('w:trHeight')
                    trPr.append(trHeight)
                trHeight.set(qn('w:val'), str(int(sig_row_height.emu / 635)))
                trHeight.set(qn('w:hRule'), 'exact')

            # Disclaimer footer
            disclaimer_text = (
                "All work and services carried out by Durabler are subject to, and conducted in accordance with, "
                "Durabler standard terms and conditions, which are available at durabler.se. This document shall not "
                "be reproduced other than in full, except with prior written approval of the issuer. The results pertain "
                "only to the item(s) as sampled by the client unless otherwise indicated. Durabler a part of Subseatec S AB, "
                "Address: Durabler C/O Subseatec, Dalavägen 23, 68130 Kristinehamn, SWEDEN"
            )
            for section in doc.sections:
                footer = section.footer
                footer.is_linked_to_previous = False
                footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
                footer_para.clear()
                footer_run = footer_para.add_run(disclaimer_text)
                footer_run.font.size = Pt(7)
                footer_run.italic = True

            # Save
            reports_folder = Path(current_app.config['REPORTS_FOLDER'])
            drafts_folder = reports_folder / 'drafts'
            drafts_folder.mkdir(parents=True, exist_ok=True)

            safe_cert_num = (test.certificate.certificate_number_with_rev.replace(' ', '_').replace('/', '-')
                             if test.certificate else test.test_id.replace(' ', '_'))
            timestamp_str = datetime.now().strftime('%Y%m%d_%H%M%S')
            output_filename = f"{safe_cert_num}_{timestamp_str}.docx"
            output_path = drafts_folder / output_filename
            doc.save(output_path)

            # Clean up
            for f in temp_photo_files:
                try:
                    os.remove(f)
                except OSError:
                    pass
            if chart_path and chart_path.exists():
                os.remove(chart_path)

            # Update approval record
            if test.certificate:
                approval = test.certificate.approval
                if not approval:
                    approval = ReportApproval.get_or_create_for_certificate(
                        test.certificate, current_user)
                if approval.status in (STATUS_DRAFT, STATUS_REJECTED, None):
                    approval.word_report_path = str(
                        output_path.relative_to(reports_folder))
                    approval.status = STATUS_DRAFT

            audit = AuditLog(
                user_id=current_user.id,
                action='REPORT',
                table_name='test_record',
                record_id=test.id,
                new_values=json.dumps({'report': output_filename})
            )
            db.session.add(audit)
            db.session.commit()

            flash(f'Report generated: {output_filename}', 'success')

            if test.certificate:
                return redirect(url_for('certificates.view',
                                        cert_id=test.certificate.id))
            return send_file(
                output_path,
                as_attachment=True,
                download_name=output_filename,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

        except Exception as e:
            flash(f'Error generating report: {e}', 'error')
            return redirect(url_for('charpy.view', test_id=test.id))

    return render_template('charpy/report.html', test=test, form=form)


@charpy_bp.route('/<int:test_id>/delete', methods=['POST'])
@login_required
def delete(test_id):
    """Delete a Charpy test (admin only)."""
    if current_user.role != 'admin':
        flash('Admin access required.', 'error')
        return redirect(url_for('charpy.index'))

    test = TestRecord.query.get_or_404(test_id)

    if test.test_method != 'CHARPY':
        flash('Invalid test type.', 'error')
        return redirect(url_for('charpy.index'))

    test_id_str = test.test_id

    audit = AuditLog(
        user_id=current_user.id,
        action='DELETE',
        table_name='test_record',
        record_id=test.id,
        old_values=json.dumps({'test_id': test_id_str, 'test_method': 'CHARPY'})
    )
    db.session.add(audit)

    AnalysisResult.query.filter_by(test_record_id=test.id).delete()
    db.session.delete(test)
    db.session.commit()

    flash(f'Charpy test {test_id_str} deleted.', 'success')
    return redirect(url_for('charpy.index'))
