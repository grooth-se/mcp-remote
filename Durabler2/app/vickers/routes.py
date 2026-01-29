"""Routes for Vickers Hardness test module - ASTM E92 / ISO 6507."""
import os
import json
from datetime import datetime
from pathlib import Path

from flask import (
    render_template, redirect, url_for, flash, request,
    current_app, send_file
)
from flask_login import login_required, current_user
from werkzeug.utils import secure_filename

from . import vickers_bp
from .forms import SpecimenForm, ReportForm
from app.extensions import db
from app.models import TestRecord, AnalysisResult, AuditLog, Certificate

# Import calculation utilities
from utils.analysis.vickers_calculations import VickersAnalyzer, VickersResult
from utils.models.vickers_specimen import VickersTestData, VickersReading, VickersLoadLevel


def generate_vickers_test_id():
    """Generate unique test ID for Vickers test."""
    today = datetime.now()
    prefix = f"VH-{today.strftime('%y%m%d')}"
    count = TestRecord.query.filter(
        TestRecord.test_id.like(f'{prefix}%')
    ).count()
    return f"{prefix}-{count + 1:03d}"


def create_hardness_profile_plot(readings, load_level):
    """Create interactive Hardness Profile plot using Plotly.

    Parameters
    ----------
    readings : list
        List of VickersReading objects or dicts
    load_level : str
        Load level designation (e.g., "HV 10")

    Returns
    -------
    str
        HTML string with Plotly figure
    """
    import plotly.graph_objects as go
    import numpy as np

    # Extract data
    locations = []
    values = []
    for i, r in enumerate(readings, 1):
        if isinstance(r, dict):
            locations.append(r.get('location', f'Point {i}'))
            values.append(r.get('hardness_value', 0))
        else:
            locations.append(r.location or f'Point {i}')
            values.append(r.hardness_value)

    # Calculate statistics
    mean_val = np.mean(values)
    std_val = np.std(values, ddof=1) if len(values) > 1 else 0

    fig = go.Figure()

    # Line plot with markers (darkred)
    fig.add_trace(go.Scatter(
        x=list(range(1, len(values) + 1)),
        y=values,
        mode='lines+markers+text',
        text=[f'{v:.1f}' for v in values],
        textposition='top center',
        textfont=dict(size=10),
        name='Readings',
        line=dict(color='darkred', width=2),
        marker=dict(color='darkred', size=10, symbol='circle'),
        hovertemplate='%{customdata}<br>%{y:.1f} ' + load_level + '<extra></extra>',
        customdata=locations
    ))

    # Mean line (grey dotted)
    fig.add_hline(
        y=mean_val,
        line_dash='dot',
        line_color='grey',
        line_width=2,
        annotation_text=f'Mean: {mean_val:.1f}',
        annotation_position='right',
        annotation_font=dict(color='grey')
    )

    fig.update_layout(
        title=f'Hardness Profile ({load_level})',
        xaxis_title='Reading Number',
        yaxis_title=f'Hardness ({load_level})',
        template='plotly_white',
        showlegend=False,
        height=400,
        xaxis=dict(tickmode='linear', tick0=1, dtick=1),
        yaxis=dict(range=[0, max(values) * 1.15])
    )

    return fig.to_html(full_html=False, include_plotlyjs='cdn')


@vickers_bp.route('/')
@login_required
def index():
    """List all Vickers tests."""
    tests = TestRecord.query.filter_by(test_method='VICKERS').order_by(
        TestRecord.created_at.desc()
    ).all()
    return render_template('vickers/index.html', tests=tests)


@vickers_bp.route('/new', methods=['GET', 'POST'])
@login_required
def new():
    """Create a new Vickers test."""
    form = SpecimenForm()

    # Populate certificate dropdown
    certificates = Certificate.query.order_by(
        Certificate.year.desc(),
        Certificate.cert_id.desc()
    ).all()
    form.certificate_id.choices = [(0, '-- Select Certificate --')] + [
        (c.id, f"{c.certificate_number_with_rev} - {c.customer or 'No customer'}")
        for c in certificates
    ]

    # Check if coming from certificate page
    cert_id = request.args.get('certificate', type=int)
    if cert_id and request.method == 'GET':
        form.certificate_id.data = cert_id
        cert = Certificate.query.get(cert_id)
        if cert:
            # Pre-fill form from certificate data (certificate register is the master)
            form.material.data = cert.material
            form.specimen_id.data = cert.test_article_sn  # Specimen SN
            form.customer_specimen_info.data = cert.customer_specimen_info
            form.requirement.data = cert.requirement
            form.location_orientation.data = cert.location_orientation
            # Parse temperature - handle string format
            if cert.temperature:
                try:
                    temp_str = cert.temperature.replace('°C', '').replace('C', '').strip()
                    form.temperature.data = float(temp_str)
                except (ValueError, AttributeError):
                    form.temperature.data = 23.0
            else:
                form.temperature.data = 23.0

    if form.validate_on_submit():
        # Link to certificate if selected
        certificate_id = None
        cert_number = None
        if form.certificate_id.data and form.certificate_id.data != 0:
            cert = Certificate.query.get(form.certificate_id.data)
            if cert:
                certificate_id = cert.id
                cert_number = cert.certificate_number_with_rev

        # Generate test ID automatically
        test_id = generate_vickers_test_id()

        # Build test parameters/geometry
        test_params = {
            'load_level': form.load_level.data,
            'dwell_time': form.dwell_time.data,
            'num_readings': form.num_readings.data,
            'location_orientation': form.location_orientation.data,
            'notes': form.notes.data,
        }

        # Create test record
        test = TestRecord(
            test_id=test_id,
            test_method='VICKERS',
            specimen_id=form.specimen_id.data,
            material=form.material.data,
            test_date=form.test_date.data or datetime.now(),
            temperature=form.temperature.data if form.temperature.data else 23.0,
            geometry=test_params,
            status='DRAFT',
            certificate_id=certificate_id,
            certificate_number=cert_number,
            operator_id=current_user.id
        )

        # Collect readings (up to 20)
        readings = []
        num_readings = int(form.num_readings.data or 5)
        for i in range(1, num_readings + 1):
            location = getattr(form, f'reading_{i}_location').data
            value = getattr(form, f'reading_{i}_value').data
            if value is not None and value > 0:
                readings.append({
                    'reading_number': len(readings) + 1,
                    'location': location or f'Point {len(readings) + 1}',
                    'hardness_value': value,
                })

        # Store readings in geometry
        test_params['readings'] = readings

        # Handle photo upload
        if form.photo.data:
            photo = form.photo.data
            filename = secure_filename(photo.filename)
            # Add test_id prefix to avoid collisions
            filename = f"vickers_{test_id}_{filename}"
            filepath = os.path.join(current_app.config['UPLOAD_FOLDER'], filename)
            photo.save(filepath)
            test_params['photo_path'] = filename

        # Update geometry with all data
        test.geometry = test_params

        # Add test to session and flush to get ID before creating analysis records
        db.session.add(test)
        db.session.flush()

        # Run analysis if we have readings
        if readings:
            try:
                # Create VickersReading objects
                reading_objects = [
                    VickersReading(
                        reading_number=r['reading_number'],
                        location=r['location'],
                        hardness_value=r['hardness_value']
                    )
                    for r in readings
                ]

                # Create load level
                load_value = float(form.load_level.data.replace('HV ', ''))
                load_level = VickersLoadLevel(form.load_level.data, load_value)

                # Create test data
                test_data = VickersTestData(
                    readings=reading_objects,
                    load_level=load_level,
                    specimen_id=form.specimen_id.data or '',
                    material=form.material.data or ''
                )

                # Run analysis
                analyzer = VickersAnalyzer()
                result = analyzer.run_analysis(test_data)

                # Get uncertainty budget
                import numpy as np
                values = np.array([r['hardness_value'] for r in readings])
                uncertainty_budget = analyzer.get_uncertainty_budget(values, result.mean_hardness.value)

                # Store results as individual AnalysisResult records
                load_lvl = result.load_level
                results_to_store = [
                    ('mean_hardness', result.mean_hardness.value, result.mean_hardness.uncertainty, load_lvl),
                    ('std_dev', result.std_dev, None, load_lvl),
                    ('range', result.range_value, None, load_lvl),
                    ('min_value', result.min_value, None, load_lvl),
                    ('max_value', result.max_value, None, load_lvl),
                    ('n_readings', result.n_readings, None, '-'),
                ]

                for param_name, value, uncertainty, unit in results_to_store:
                    analysis = AnalysisResult(
                        test_record_id=test.id,
                        parameter_name=param_name,
                        value=value,
                        uncertainty=uncertainty,
                        unit=unit,
                        calculated_by_id=current_user.id
                    )
                    db.session.add(analysis)

                # Store uncertainty budget in geometry for report generation
                test_params['uncertainty_budget'] = uncertainty_budget
                test.geometry = test_params
                test.status = 'ANALYZED'
                flash(f'Analysis completed: Mean = {result.mean_hardness.value:.1f} +/- {result.mean_hardness.uncertainty:.1f} {result.load_level}', 'success')
            except Exception as e:
                flash(f'Analysis error: {e}', 'warning')
                test.status = 'DRAFT'
        else:
            flash('No hardness readings entered.', 'warning')

        # Audit log
        audit = AuditLog(
            user_id=current_user.id,
            action='CREATE',
            table_name='test_record',
            record_id=test.id,
            new_values=json.dumps({'test_id': test.test_id, 'test_method': 'VICKERS'})
        )
        db.session.add(audit)
        db.session.commit()

        flash(f'Vickers test {test.test_id} created.', 'success')
        return redirect(url_for('vickers.view', test_id=test.id))

    return render_template('vickers/new.html', form=form)


@vickers_bp.route('/<int:test_id>')
@login_required
def view(test_id):
    """View Vickers test details and results."""
    test = TestRecord.query.get_or_404(test_id)

    if test.test_method != 'VICKERS':
        flash('Invalid test type.', 'error')
        return redirect(url_for('vickers.index'))

    # Parse stored data (geometry is JSON dict, not string)
    test_params = test.geometry if test.geometry else {}
    readings = test_params.get('readings', [])

    # Get analysis results (individual records per parameter)
    results = {}
    analysis_records = AnalysisResult.query.filter_by(test_record_id=test.id).all()
    for ar in analysis_records:
        if ar.parameter_name == 'mean_hardness':
            results['mean_hardness'] = {'value': ar.value, 'uncertainty': ar.uncertainty}
        else:
            results[ar.parameter_name] = ar.value
    # Get load level and uncertainty budget from geometry
    results['load_level'] = test_params.get('load_level', 'HV')
    results['uncertainty_budget'] = test_params.get('uncertainty_budget', {})

    # Create plot if we have readings
    hardness_plot = None
    if readings:
        load_level = test_params.get('load_level', 'HV')
        hardness_plot = create_hardness_profile_plot(readings, load_level)

    # Get photo path
    photo_url = None
    if test_params.get('photo_path'):
        photo_url = url_for('static', filename=f'uploads/{test_params["photo_path"]}')

    return render_template('vickers/view.html',
                           test=test,
                           test_params=test_params,
                           readings=readings,
                           results=results,
                           hardness_plot=hardness_plot,
                           photo_url=photo_url)


@vickers_bp.route('/<int:test_id>/report', methods=['GET', 'POST'])
@login_required
def report(test_id):
    """Generate Vickers test report."""
    test = TestRecord.query.get_or_404(test_id)

    if test.test_method != 'VICKERS':
        flash('Invalid test type.', 'error')
        return redirect(url_for('vickers.index'))

    form = ReportForm()

    # Pre-fill certificate number
    if request.method == 'GET':
        form.certificate_number.data = test.certificate_number or test.test_id

    if form.validate_on_submit():
        try:
            # Parse stored data (geometry is JSON dict, not string)
            test_params = test.geometry if test.geometry else {}
            readings = test_params.get('readings', [])

            # Get analysis results (individual records per parameter)
            results = {}
            analysis_records = AnalysisResult.query.filter_by(test_record_id=test.id).all()
            for ar in analysis_records:
                if ar.parameter_name == 'mean_hardness':
                    results['mean_hardness'] = {'value': ar.value, 'uncertainty': ar.uncertainty}
                else:
                    results[ar.parameter_name] = ar.value
            results['load_level'] = test_params.get('load_level', 'HV')
            results['uncertainty_budget'] = test_params.get('uncertainty_budget', {})

            # Generate chart image for report
            chart_path = None
            if readings:
                import matplotlib
                matplotlib.use('Agg')
                import matplotlib.pyplot as plt
                import numpy as np

                values = [r['hardness_value'] for r in readings]
                mean_val = np.mean(values)

                fig, ax = plt.subplots(figsize=(8, 5))

                # Line plot with markers (darkred)
                x = list(range(1, len(values) + 1))
                ax.plot(x, values, color='darkred', linewidth=2, marker='o',
                        markersize=10, markerfacecolor='darkred', markeredgecolor='darkred')

                # Add value labels above points
                for xi, val in zip(x, values):
                    ax.text(xi, val + max(values) * 0.02, f'{val:.1f}',
                            ha='center', va='bottom', fontsize=9)

                # Mean line (grey dotted)
                ax.axhline(y=mean_val, color='grey', linestyle=':', linewidth=2,
                           label=f'Mean: {mean_val:.1f}')

                ax.set_xlabel('Reading Number')
                ax.set_ylabel(f'Hardness ({test_params.get("load_level", "HV")})')
                ax.set_title('Hardness Profile')
                ax.legend(loc='upper right')
                ax.set_xlim(0.5, len(values) + 0.5)
                ax.set_ylim(0, max(values) * 1.15)
                ax.set_xticks(x)
                ax.grid(True, alpha=0.3, axis='y')

                chart_path = Path(current_app.config['REPORTS_FOLDER']) / f'vickers_chart_{test.id}.png'
                fig.savefig(chart_path, dpi=150, bbox_inches='tight')
                plt.close(fig)

            # Prepare report data - include all test information fields
            test_info = {
                'certificate_number': form.certificate_number.data or test.test_id,
                'test_project': test.certificate.test_project if test.certificate else '',
                'customer': test.certificate.customer if test.certificate else '',
                'specimen_id': test.specimen_id or '',
                'customer_specimen_info': test.certificate.customer_specimen_info if test.certificate else '',
                'material': test.material or '',
                'requirement': test.certificate.requirement if test.certificate else '',
                'location_orientation': test_params.get('location_orientation', ''),
                'test_date': test.test_date.strftime('%Y-%m-%d') if test.test_date else '',
                'temperature': test.temperature or 23,
                'load_level': test_params.get('load_level', 'HV 10'),
                'dwell_time': test_params.get('dwell_time', '15'),
                'notes': test_params.get('notes', ''),
            }

            # Create VickersResult for report
            class ResultProxy:
                def __init__(self, data, readings_list):
                    self._data = data
                    self._readings = readings_list

                @property
                def mean_hardness(self):
                    d = self._data.get('mean_hardness', {})
                    return type('MV', (), {
                        'value': d.get('value', 0),
                        'uncertainty': d.get('uncertainty', 0)
                    })()

                @property
                def std_dev(self):
                    return self._data.get('std_dev', 0)

                @property
                def range_value(self):
                    return self._data.get('range_value', 0)

                @property
                def min_value(self):
                    return self._data.get('min_value', 0)

                @property
                def max_value(self):
                    return self._data.get('max_value', 0)

                @property
                def n_readings(self):
                    return self._data.get('n_readings', 0)

                @property
                def load_level(self):
                    return self._data.get('load_level', 'HV')

                @property
                def readings(self):
                    return [
                        type('R', (), {
                            'reading_number': r.get('reading_number', i),
                            'location': r.get('location', f'Point {i}'),
                            'hardness_value': r.get('hardness_value', 0)
                        })()
                        for i, r in enumerate(self._readings, 1)
                    ]

            result_proxy = ResultProxy(results, readings)
            uncertainty_budget = results.get('uncertainty_budget', {})

            # Create report from scratch (no template dependency)
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH

            doc = Document()

            # Title
            title = doc.add_heading('Vickers Hardness Test Report', level=0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            subtitle = doc.add_paragraph('ASTM E92 / ISO 6507')
            subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()

            # Test Information - all fields except status columns
            doc.add_heading('Test Information', level=1)

            # Build info data - include all relevant fields
            info_data = [
                ('Certificate Number:', test_info.get('certificate_number', '')),
                ('Customer:', test_info.get('customer', '')),
                ('Test Project:', test_info.get('test_project', '')),
                ('Specimen ID:', test_info.get('specimen_id', '')),
                ('Customer Specimen Info:', test_info.get('customer_specimen_info', '')),
                ('Material:', test_info.get('material', '')),
                ('Requirement:', test_info.get('requirement', '')),
                ('Location/Orientation:', test_info.get('location_orientation', '')),
                ('Test Date:', test_info.get('test_date', '')),
                ('Temperature:', f"{test_info.get('temperature', '23')} °C"),
                ('Load Level:', test_info.get('load_level', '')),
                ('Dwell Time:', f"{test_info.get('dwell_time', '15')} s"),
                ('Test Equipment:', 'q-ness ATM test machine with automatic indenter'),
            ]

            # Filter out empty values for cleaner report
            info_data = [(label, value) for label, value in info_data if value and str(value).strip()]

            table = doc.add_table(rows=len(info_data), cols=2)
            table.style = 'Table Grid'

            for i, (label, value) in enumerate(info_data):
                table.rows[i].cells[0].text = label
                table.rows[i].cells[1].text = str(value)
                table.rows[i].cells[0].paragraphs[0].runs[0].bold = True

            doc.add_paragraph()

            # Results Summary
            doc.add_heading('Results Summary', level=1)
            table = doc.add_table(rows=7, cols=3)
            table.style = 'Table Grid'

            headers = ['Parameter', 'Value', 'Unit']
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h
                table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

            results_data = [
                ('Mean Hardness', f'{result_proxy.mean_hardness.value:.1f} +/- {result_proxy.mean_hardness.uncertainty:.1f}', result_proxy.load_level),
                ('Standard Deviation', f'{result_proxy.std_dev:.1f}', result_proxy.load_level),
                ('Range', f'{result_proxy.range_value:.1f}', result_proxy.load_level),
                ('Minimum', f'{result_proxy.min_value:.1f}', result_proxy.load_level),
                ('Maximum', f'{result_proxy.max_value:.1f}', result_proxy.load_level),
                ('Number of Readings', str(result_proxy.n_readings), '-'),
            ]

            for i, (param, value, unit) in enumerate(results_data):
                table.rows[i+1].cells[0].text = param
                table.rows[i+1].cells[1].text = value
                table.rows[i+1].cells[2].text = unit

            doc.add_paragraph()

            # Individual Readings
            doc.add_heading('Individual Readings', level=1)
            table = doc.add_table(rows=len(readings) + 1, cols=3)
            table.style = 'Table Grid'

            headers = ['#', 'Location', 'Hardness']
            for i, h in enumerate(headers):
                table.rows[0].cells[i].text = h
                table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

            for i, r in enumerate(readings):
                table.rows[i+1].cells[0].text = str(r.get('reading_number', i+1))
                table.rows[i+1].cells[1].text = r.get('location', f'Point {i+1}')
                table.rows[i+1].cells[2].text = f"{r.get('hardness_value', 0):.1f}"

            doc.add_paragraph()

            # Notes (if any)
            if test_info.get('notes'):
                doc.add_heading('Notes', level=1)
                doc.add_paragraph(test_info['notes'])
                doc.add_paragraph()

            # Uncertainty Budget (if requested)
            if form.include_uncertainty_budget.data == 'yes' and uncertainty_budget:
                doc.add_heading('Uncertainty Budget (k=2)', level=1)
                table = doc.add_table(rows=6, cols=2)
                table.style = 'Table Grid'

                unc_data = [
                    ('Type A (Repeatability)', f"{uncertainty_budget.get('u_A', 0):.2f}"),
                    ('Machine Calibration', f"{uncertainty_budget.get('u_machine', 0):.2f}"),
                    ('Diagonal Measurement', f"{uncertainty_budget.get('u_diagonal', 0):.2f}"),
                    ('Force Application', f"{uncertainty_budget.get('u_force', 0):.2f}"),
                    ('Combined Standard (u_c)', f"{uncertainty_budget.get('u_combined', 0):.2f}"),
                    ('Expanded (U, k=2)', f"{uncertainty_budget.get('U_expanded', 0):.2f}"),
                ]

                for i, (comp, val) in enumerate(unc_data):
                    table.rows[i].cells[0].text = comp
                    table.rows[i].cells[1].text = val

                doc.add_paragraph()

            # Chart
            if chart_path and chart_path.exists():
                doc.add_heading('Hardness Profile', level=1)
                doc.add_picture(str(chart_path), width=Inches(5.5))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

            doc.add_paragraph()

            # Indent Photo (if requested and available)
            if form.include_photo.data == 'yes' and test_params.get('photo_path'):
                photo_path = Path(current_app.config['UPLOAD_FOLDER']) / test_params['photo_path']
                if photo_path.exists():
                    doc.add_heading('Indent Photo', level=1)
                    doc.add_picture(str(photo_path), width=Inches(4))
                    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()

            # Approval Signatures (4 rows: header + tested + reviewed + approved)
            doc.add_heading('Approval', level=1)
            sig_table = doc.add_table(rows=4, cols=4)
            sig_table.style = 'Table Grid'

            sig_headers = ['Role', 'Name', 'Signature', 'Date']
            for i, header in enumerate(sig_headers):
                sig_table.rows[0].cells[i].text = header
                sig_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

            sig_table.rows[1].cells[0].text = 'Tested by:'
            sig_table.rows[2].cells[0].text = 'Reviewed by:'
            sig_table.rows[3].cells[0].text = 'Approved by:'

            doc.add_paragraph()

            # Footer
            footer = doc.add_paragraph()
            footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
            footer.add_run(f"Report generated: {datetime.now().strftime('%Y-%m-%d %H:%M')} | ASTM E92 / ISO 6507 Vickers Hardness Test Report").font.size = Pt(9)

            # Save document
            output_filename = f"Vickers_Report_{test.test_id.replace(' ', '_')}.docx"
            output_path = Path(current_app.config['REPORTS_FOLDER']) / output_filename
            doc.save(output_path)

            # Audit log
            audit = AuditLog(
                user_id=current_user.id,
                action='REPORT',
                table_name='test_record',
                record_id=test.id,
                new_values=json.dumps({'report': output_filename})
            )
            db.session.add(audit)
            db.session.commit()

            return send_file(
                output_path,
                as_attachment=True,
                download_name=output_filename,
                mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )

        except Exception as e:
            flash(f'Error generating report: {e}', 'error')
            return redirect(url_for('vickers.view', test_id=test.id))

    return render_template('vickers/report.html', test=test, form=form)


@vickers_bp.route('/<int:test_id>/delete', methods=['POST'])
@login_required
def delete(test_id):
    """Delete a Vickers test (admin only)."""
    if current_user.role != 'admin':
        flash('Admin access required.', 'error')
        return redirect(url_for('vickers.index'))

    test = TestRecord.query.get_or_404(test_id)

    if test.test_method != 'VICKERS':
        flash('Invalid test type.', 'error')
        return redirect(url_for('vickers.index'))

    test_id_str = test.test_id

    # Audit log before deletion
    audit = AuditLog(
        user_id=current_user.id,
        action='DELETE',
        table_name='test_record',
        record_id=test.id,
        old_values=json.dumps({'test_id': test_id_str, 'test_method': 'VICKERS'})
    )
    db.session.add(audit)

    # Delete analysis results first
    AnalysisResult.query.filter_by(test_record_id=test.id).delete()
    db.session.delete(test)
    db.session.commit()

    flash(f'Vickers test {test_id_str} deleted.', 'success')
    return redirect(url_for('vickers.index'))
