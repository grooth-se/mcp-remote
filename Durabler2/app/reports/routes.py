"""Reports routes for certificate-centric approval workflow.

Workflow:
1. Certificate view → [Generate Report] → Creates Word doc, saves to server, status=DRAFT
2. Certificate view → [Submit for Review] → status=PENDING_REVIEW
3. Approver → [Review] → See certificate data + Word doc link
4. Approver → [Approve] → Generates PDF, status=PUBLISHED
5. Approver → [Reject] → status=REJECTED with comments, back to engineer
"""
from datetime import datetime
from pathlib import Path
from flask import (
    render_template, redirect, url_for, flash, request,
    current_app, send_file, abort
)
from flask_login import login_required, current_user

from . import reports_bp
from app.extensions import db
from app.models import (
    ReportApproval, TestRecord, AuditLog, Certificate,
    STATUS_DRAFT, STATUS_PENDING, STATUS_APPROVED, STATUS_REJECTED, STATUS_PUBLISHED,
    APPROVAL_STATUS_LABELS, STATUS_COLORS,
    approver_required, engineer_required
)


@reports_bp.route('/')
@login_required
def index():
    """List all reports with approval status."""
    status_filter = request.args.get('status', '')

    query = ReportApproval.query

    if status_filter:
        query = query.filter(ReportApproval.status == status_filter)

    reports = query.order_by(ReportApproval.created_at.desc()).all()

    return render_template('reports/index.html',
                           reports=reports,
                           status_filter=status_filter,
                           status_labels=APPROVAL_STATUS_LABELS,
                           status_colors=STATUS_COLORS)


@reports_bp.route('/pending')
@login_required
@approver_required
def pending():
    """List reports pending approval (approvers only)."""
    reports = ReportApproval.query.filter_by(status=STATUS_PENDING)\
        .order_by(ReportApproval.submitted_at.asc())\
        .all()

    return render_template('reports/pending.html',
                           reports=reports,
                           status_labels=APPROVAL_STATUS_LABELS,
                           status_colors=STATUS_COLORS)


@reports_bp.route('/certificate/<int:cert_id>/start', methods=['POST'])
@login_required
@engineer_required
def start_approval(cert_id):
    """Start approval workflow for a certificate (create approval record)."""
    certificate = Certificate.query.get_or_404(cert_id)

    # Check if approval already exists
    if certificate.approval:
        flash(f'Approval workflow already exists for {certificate.certificate_number_with_rev}.', 'info')
        return redirect(url_for('certificates.view', cert_id=cert_id))

    # Create new approval record
    approval = ReportApproval.get_or_create_for_certificate(certificate, current_user)

    # Audit log
    audit = AuditLog(
        user_id=current_user.id,
        action='START_APPROVAL',
        table_name='report_approvals',
        record_id=approval.id,
        new_values={
            'certificate_number': certificate.certificate_number_with_rev,
            'status': STATUS_DRAFT
        },
        ip_address=request.remote_addr
    )
    db.session.add(audit)
    db.session.commit()

    flash(f'Approval workflow started for {certificate.certificate_number_with_rev}.', 'success')
    return redirect(url_for('certificates.view', cert_id=cert_id))


@reports_bp.route('/certificate/<int:cert_id>/generate-report', methods=['POST'])
@login_required
@engineer_required
def generate_report(cert_id):
    """Generate Word report for certificate and save to server."""
    certificate = Certificate.query.get_or_404(cert_id)

    # Get or create approval record
    if not certificate.approval:
        approval = ReportApproval.get_or_create_for_certificate(certificate, current_user)
        db.session.commit()
    else:
        approval = certificate.approval

    # Check if we can edit (DRAFT or REJECTED status)
    if approval.status not in [STATUS_DRAFT, STATUS_REJECTED]:
        flash('Cannot generate report - approval is already in progress.', 'warning')
        return redirect(url_for('certificates.view', cert_id=cert_id))

    # Get test records for this certificate
    test_records = certificate.test_records.all()
    if not test_records:
        flash('No test records linked to this certificate.', 'warning')
        return redirect(url_for('certificates.view', cert_id=cert_id))

    try:
        # Import report generators
        from utils.reporting.word_report import TensileReportGenerator
        from utils.reporting.ctod_word_report import CTODReportGenerator
        from utils.reporting.sonic_word_report import SonicReportGenerator
        from utils.reporting.fcgr_word_report import FCGRReportGenerator
        from utils.reporting.kic_word_report import KICReportGenerator

        # Create reports folder
        reports_folder = Path(current_app.config['REPORTS_FOLDER'])
        drafts_folder = reports_folder / 'drafts'
        drafts_folder.mkdir(parents=True, exist_ok=True)

        # Generate filename
        safe_cert_num = certificate.certificate_number_with_rev.replace(' ', '_').replace('/', '-')
        timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
        report_filename = f"{safe_cert_num}_{timestamp}.docx"
        output_path = drafts_folder / report_filename

        # For now, generate report based on first test record's type
        # TODO: Support multi-test combined reports
        test_record = test_records[0]

        if test_record.test_method == 'TENSILE':
            # Generate tensile report
            results = {r.parameter_name: r for r in test_record.results.all()}
            geometry = test_record.geometry or {}

            test_info = {
                'test_project': certificate.test_order or '',
                'customer': certificate.customer or '',
                'customer_order': certificate.customer_order or '',
                'product_sn': certificate.product_sn or '',
                'specimen_id': test_record.specimen_id or '',
                'location_orientation': certificate.location_orientation or '',
                'material': certificate.material or '',
                'certificate_number': certificate.certificate_number_with_rev,
                'test_date': test_record.test_date.strftime('%Y-%m-%d') if test_record.test_date else '',
                'test_engineer': current_user.username,
                'temperature': str(test_record.temperature) if test_record.temperature else '23',
                'strain_source': 'Displacement Only' if geometry.get('use_displacement_only') else 'Extensometer',
                'comments': ''
            }

            specimen_type = geometry.get('type', 'round')
            if specimen_type == 'round':
                dimensions = {
                    'diameter': geometry.get('D0'),
                    'final_diameter': geometry.get('D1'),
                    'gauge_length': geometry.get('L0'),
                    'final_gauge_length': geometry.get('L1'),
                    'parallel_length': geometry.get('Lp')
                }
            else:
                dimensions = {
                    'width': geometry.get('a0'),
                    'thickness': geometry.get('b0'),
                    'gauge_length': geometry.get('L0'),
                    'final_gauge_length': geometry.get('L1'),
                    'parallel_length': geometry.get('Lp')
                }

            # Convert results
            results_for_report = {}
            for name, result in results.items():
                class ResultValue:
                    def __init__(self, v, u):
                        self.value = v
                        self.uncertainty = u
                results_for_report[name] = ResultValue(result.value, result.uncertainty)

            # Map names
            result_mapping = {
                'Rp0.2': 'Rp02', 'Rp0.5': 'Rp05', 'A%': 'A_percent', 'Z%': 'Z',
                'Stress_rate_Rp02': 'stress_rate_rp02', 'Strain_rate_Rp02': 'strain_rate_rp02',
                'Stress_rate_Rm': 'stress_rate_rm', 'Strain_rate_Rm': 'strain_rate_rm'
            }
            for db_name, report_name in result_mapping.items():
                if db_name in results_for_report:
                    results_for_report[report_name] = results_for_report[db_name]

            yield_type = geometry.get('yield_method', 'offset')

            # Parse requirements from certificate
            # Format expected: "Rp0.2 min 500, Rm 600-800, A min 15%, Z min 40%"
            requirements = {}
            if certificate.requirement:
                req_text = certificate.requirement
                # Try to parse individual requirements
                import re
                # Match patterns like "Rp0.2 min 500" or "Rm: 600-800" or "A >= 15%"
                patterns = [
                    (r'Rp0\.?2[:\s]*(.*?)(?:,|;|$)', 'Rp02'),
                    (r'Rp0\.?5[:\s]*(.*?)(?:,|;|$)', 'Rp05'),
                    (r'ReH[:\s]*(.*?)(?:,|;|$)', 'ReH'),
                    (r'ReL[:\s]*(.*?)(?:,|;|$)', 'ReL'),
                    (r'Rm[:\s]*(.*?)(?:,|;|$)', 'Rm'),
                    (r'\bA[:\s]*(.*?)(?:,|;|$)', 'A'),
                    (r'\bZ[:\s]*(.*?)(?:,|;|$)', 'Z'),
                ]
                for pattern, key in patterns:
                    match = re.search(pattern, req_text, re.IGNORECASE)
                    if match:
                        requirements[key] = match.group(1).strip()

                # If no structured format, use entire text as general requirement
                if not requirements:
                    requirements['general'] = req_text

            report_data = TensileReportGenerator.prepare_report_data(
                test_info=test_info,
                dimensions=dimensions,
                results=results_for_report,
                specimen_type=specimen_type,
                yield_type=yield_type,
                requirements=requirements
            )

            logo_path = Path(current_app.root_path).parent / 'templates' / 'logo.png'
            generator = TensileReportGenerator(None)
            generator.generate_report(
                output_path=output_path,
                data=report_data,
                chart_path=None,
                logo_path=logo_path if logo_path.exists() else None
            )
        elif test_record.test_method == 'CTOD':
            # Generate CTOD report
            _generate_ctod_report(certificate, test_record, output_path)

        elif test_record.test_method == 'SONIC':
            # Generate Sonic Resonance report
            _generate_sonic_report(certificate, test_record, output_path)

        elif test_record.test_method == 'FCGR':
            # Generate FCGR report
            _generate_fcgr_report(certificate, test_record, output_path)

        elif test_record.test_method == 'KIC':
            # Generate KIC report
            _generate_kic_report(certificate, test_record, output_path)

        elif test_record.test_method == 'VICKERS':
            # Generate Vickers report
            _generate_vickers_report(certificate, test_record, output_path)

        else:
            flash(f'Report generation for {test_record.test_method} not yet implemented.', 'warning')
            return redirect(url_for('certificates.view', cert_id=cert_id))

        # Update approval record with Word report path
        approval.word_report_path = str(output_path.relative_to(reports_folder))
        approval.status = STATUS_DRAFT

        # Audit log
        audit = AuditLog(
            user_id=current_user.id,
            action='GENERATE_REPORT',
            table_name='report_approvals',
            record_id=approval.id,
            new_values={
                'certificate_number': certificate.certificate_number_with_rev,
                'word_report_path': approval.word_report_path
            },
            ip_address=request.remote_addr
        )
        db.session.add(audit)
        db.session.commit()

        flash(f'Report generated and saved: {report_filename}', 'success')

    except Exception as e:
        import traceback
        traceback.print_exc()
        flash(f'Report generation failed: {str(e)}', 'danger')

    return redirect(url_for('certificates.view', cert_id=cert_id))


@reports_bp.route('/certificate/<int:cert_id>/download-word')
@login_required
def download_word(cert_id):
    """Download the Word report for editing."""
    certificate = Certificate.query.get_or_404(cert_id)

    if not certificate.approval or not certificate.approval.word_report_path:
        flash('No Word report available.', 'warning')
        return redirect(url_for('certificates.view', cert_id=cert_id))

    reports_folder = Path(current_app.config['REPORTS_FOLDER'])
    word_path = reports_folder / certificate.approval.word_report_path

    if not word_path.exists():
        flash('Word report file not found.', 'danger')
        return redirect(url_for('certificates.view', cert_id=cert_id))

    return send_file(
        word_path,
        as_attachment=True,
        download_name=f"{certificate.certificate_number_with_rev.replace(' ', '_')}.docx",
        mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
    )


@reports_bp.route('/certificate/<int:cert_id>/upload-word-new', methods=['POST'])
@login_required
@engineer_required
def upload_word_new(cert_id):
    """Upload Word report to start approval workflow (without auto-generating)."""
    from werkzeug.utils import secure_filename

    certificate = Certificate.query.get_or_404(cert_id)

    if 'word_file' not in request.files:
        flash('No file uploaded.', 'warning')
        return redirect(url_for('certificates.view', cert_id=cert_id))

    file = request.files['word_file']
    if file.filename == '':
        flash('No file selected.', 'warning')
        return redirect(url_for('certificates.view', cert_id=cert_id))

    if not file.filename.endswith('.docx'):
        flash('Please upload a Word document (.docx).', 'warning')
        return redirect(url_for('certificates.view', cert_id=cert_id))

    # Create or get approval record
    if not certificate.approval:
        approval = ReportApproval.get_or_create_for_certificate(certificate, current_user)
        db.session.commit()
    else:
        approval = certificate.approval
        if approval.status not in [STATUS_DRAFT, STATUS_REJECTED]:
            flash('Cannot upload - approval is already in progress.', 'warning')
            return redirect(url_for('certificates.view', cert_id=cert_id))

    # Save uploaded file
    reports_folder = Path(current_app.config['REPORTS_FOLDER'])
    drafts_folder = reports_folder / 'drafts'
    drafts_folder.mkdir(parents=True, exist_ok=True)

    safe_cert_num = certificate.certificate_number_with_rev.replace(' ', '_').replace('/', '-')
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f"{safe_cert_num}_{timestamp}.docx"
    filepath = drafts_folder / filename

    file.save(filepath)

    # Update approval record
    approval.word_report_path = str(filepath.relative_to(reports_folder))
    approval.status = STATUS_DRAFT

    # Audit log
    audit = AuditLog(
        user_id=current_user.id,
        action='UPLOAD_REPORT_NEW',
        table_name='report_approvals',
        record_id=approval.id,
        new_values={
            'word_report_path': approval.word_report_path,
            'certificate_number': certificate.certificate_number_with_rev
        },
        ip_address=request.remote_addr
    )
    db.session.add(audit)
    db.session.commit()

    flash(f'Word report uploaded: {filename}. You can now submit for approval.', 'success')
    return redirect(url_for('certificates.view', cert_id=cert_id))


@reports_bp.route('/certificate/<int:cert_id>/upload-word', methods=['POST'])
@login_required
@engineer_required
def upload_word(cert_id):
    """Upload edited Word report back to server."""
    from werkzeug.utils import secure_filename

    certificate = Certificate.query.get_or_404(cert_id)

    if not certificate.approval:
        flash('No approval workflow started.', 'warning')
        return redirect(url_for('certificates.view', cert_id=cert_id))

    if certificate.approval.status not in [STATUS_DRAFT, STATUS_REJECTED]:
        flash('Cannot upload - approval is already in progress.', 'warning')
        return redirect(url_for('certificates.view', cert_id=cert_id))

    if 'word_file' not in request.files:
        flash('No file uploaded.', 'warning')
        return redirect(url_for('certificates.view', cert_id=cert_id))

    file = request.files['word_file']
    if file.filename == '':
        flash('No file selected.', 'warning')
        return redirect(url_for('certificates.view', cert_id=cert_id))

    if not file.filename.endswith('.docx'):
        flash('Please upload a Word document (.docx).', 'warning')
        return redirect(url_for('certificates.view', cert_id=cert_id))

    # Save uploaded file
    reports_folder = Path(current_app.config['REPORTS_FOLDER'])
    drafts_folder = reports_folder / 'drafts'
    drafts_folder.mkdir(parents=True, exist_ok=True)

    safe_cert_num = certificate.certificate_number_with_rev.replace(' ', '_').replace('/', '-')
    timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
    filename = f"{safe_cert_num}_{timestamp}.docx"
    filepath = drafts_folder / filename

    file.save(filepath)

    # Update approval record
    old_path = certificate.approval.word_report_path
    certificate.approval.word_report_path = str(filepath.relative_to(reports_folder))

    # Audit log
    audit = AuditLog(
        user_id=current_user.id,
        action='UPLOAD_REPORT',
        table_name='report_approvals',
        record_id=certificate.approval.id,
        old_values={'word_report_path': old_path},
        new_values={'word_report_path': certificate.approval.word_report_path},
        ip_address=request.remote_addr
    )
    db.session.add(audit)
    db.session.commit()

    flash(f'Word report uploaded: {filename}', 'success')
    return redirect(url_for('certificates.view', cert_id=cert_id))


@reports_bp.route('/certificate/<int:cert_id>/submit', methods=['POST'])
@login_required
@engineer_required
def submit(cert_id):
    """Submit certificate report for approval."""
    certificate = Certificate.query.get_or_404(cert_id)

    if not certificate.approval:
        flash('No approval workflow started. Generate a report first.', 'warning')
        return redirect(url_for('certificates.view', cert_id=cert_id))

    if not certificate.approval.can_submit:
        flash('This report cannot be submitted for approval.', 'danger')
        return redirect(url_for('certificates.view', cert_id=cert_id))

    if not certificate.approval.word_report_path:
        flash('Please generate a Word report first.', 'warning')
        return redirect(url_for('certificates.view', cert_id=cert_id))

    certificate.approval.submit_for_approval(current_user)

    # Audit log
    audit = AuditLog(
        user_id=current_user.id,
        action='SUBMIT_FOR_APPROVAL',
        table_name='report_approvals',
        record_id=certificate.approval.id,
        new_values={
            'status': STATUS_PENDING,
            'certificate_number': certificate.certificate_number_with_rev
        },
        ip_address=request.remote_addr
    )
    db.session.add(audit)
    db.session.commit()

    flash(f'Report submitted for approval: {certificate.certificate_number_with_rev}', 'success')
    return redirect(url_for('certificates.view', cert_id=cert_id))


@reports_bp.route('/certificate/<int:cert_id>/review')
@login_required
@approver_required
def review(cert_id):
    """Review page for approving/rejecting certificate report."""
    certificate = Certificate.query.get_or_404(cert_id)

    if not certificate.approval:
        flash('No approval workflow for this certificate.', 'warning')
        return redirect(url_for('certificates.view', cert_id=cert_id))

    if not certificate.approval.can_review:
        flash('This report is not pending review.', 'warning')
        return redirect(url_for('certificates.view', cert_id=cert_id))

    # Get all test records for this certificate
    test_records = certificate.test_records.all()

    return render_template('reports/review_certificate.html',
                           certificate=certificate,
                           approval=certificate.approval,
                           test_records=test_records,
                           status_labels=APPROVAL_STATUS_LABELS,
                           status_colors=STATUS_COLORS)


@reports_bp.route('/certificate/<int:cert_id>/approve', methods=['POST'])
@login_required
@approver_required
def approve(cert_id):
    """Approve the certificate report with uploaded signed PDF.

    Simplified workflow for test phase:
    1. Approver downloads Word doc, reviews externally
    2. Approver saves as PDF and signs with external PDF software
    3. Approver uploads signed PDF here to approve and publish
    """
    import hashlib
    from werkzeug.utils import secure_filename

    certificate = Certificate.query.get_or_404(cert_id)

    if not certificate.approval or not certificate.approval.can_review:
        flash('This report cannot be approved.', 'danger')
        return redirect(url_for('certificates.view', cert_id=cert_id))

    # Check for uploaded signed PDF
    if 'signed_pdf' not in request.files:
        flash('Please upload the signed PDF file.', 'warning')
        return redirect(url_for('reports.review', cert_id=cert_id))

    pdf_file = request.files['signed_pdf']
    if pdf_file.filename == '':
        flash('No PDF file selected.', 'warning')
        return redirect(url_for('reports.review', cert_id=cert_id))

    if not pdf_file.filename.lower().endswith('.pdf'):
        flash('Please upload a PDF file.', 'warning')
        return redirect(url_for('reports.review', cert_id=cert_id))

    # Save the uploaded signed PDF
    reports_folder = Path(current_app.config['REPORTS_FOLDER'])
    year = datetime.now().year
    signed_folder = reports_folder / 'signed' / str(year)
    signed_folder.mkdir(parents=True, exist_ok=True)

    safe_cert_num = certificate.certificate_number_with_rev.replace(' ', '_').replace('/', '-')
    pdf_filename = f"{safe_cert_num}_signed.pdf"
    output_pdf = signed_folder / pdf_filename

    pdf_file.save(output_pdf)

    # Calculate hash for integrity verification
    with open(output_pdf, 'rb') as f:
        pdf_hash = hashlib.sha256(f.read()).hexdigest()

    pdf_path = str(output_pdf.relative_to(reports_folder))
    timestamp = datetime.utcnow()
    approver_name = current_user.full_name or current_user.username

    # Mark as approved and published
    certificate.approval.approve(current_user)
    certificate.approval.publish(pdf_path=pdf_path, pdf_hash=pdf_hash)
    certificate.approval.signature_timestamp = timestamp
    certificate.reported = True

    # Audit log
    audit = AuditLog(
        user_id=current_user.id,
        action='APPROVE_AND_PUBLISH',
        table_name='report_approvals',
        record_id=certificate.approval.id,
        new_values={
            'status': STATUS_PUBLISHED,
            'certificate_number': certificate.certificate_number_with_rev,
            'approved_by': approver_name,
            'signed_pdf_uploaded': True,
            'pdf_hash': pdf_hash[:16] + '...',
            'timestamp': timestamp.isoformat()
        },
        ip_address=request.remote_addr
    )
    db.session.add(audit)
    db.session.commit()

    flash(f'Report approved and published: {certificate.certificate_number_with_rev}', 'success')
    return redirect(url_for('certificates.view', cert_id=cert_id))


@reports_bp.route('/certificate/<int:cert_id>/reject', methods=['POST'])
@login_required
@approver_required
def reject(cert_id):
    """Reject the certificate report with comments."""
    certificate = Certificate.query.get_or_404(cert_id)

    if not certificate.approval or not certificate.approval.can_review:
        flash('This report cannot be rejected.', 'danger')
        return redirect(url_for('certificates.view', cert_id=cert_id))

    comments = request.form.get('comments', '').strip()
    if not comments:
        flash('Please provide a reason for rejection.', 'warning')
        return redirect(url_for('reports.review', cert_id=cert_id))

    certificate.approval.reject(current_user, comments)

    # Audit log
    audit = AuditLog(
        user_id=current_user.id,
        action='REJECT_REPORT',
        table_name='report_approvals',
        record_id=certificate.approval.id,
        new_values={
            'status': STATUS_REJECTED,
            'certificate_number': certificate.certificate_number_with_rev,
            'rejected_by': current_user.full_name or current_user.username,
            'comments': comments
        },
        ip_address=request.remote_addr
    )
    db.session.add(audit)
    db.session.commit()

    flash(f'Report rejected: {certificate.certificate_number_with_rev}', 'warning')
    return redirect(url_for('reports.pending'))


@reports_bp.route('/certificate/<int:cert_id>/download-pdf')
@login_required
def download_pdf(cert_id):
    """Download the published PDF."""
    certificate = Certificate.query.get_or_404(cert_id)

    if not certificate.approval or not certificate.approval.can_download_signed:
        flash('PDF is not available.', 'warning')
        return redirect(url_for('certificates.view', cert_id=cert_id))

    reports_folder = Path(current_app.config['REPORTS_FOLDER'])
    pdf_path = reports_folder / certificate.approval.signed_pdf_path

    if not pdf_path.exists():
        flash('PDF file not found.', 'danger')
        return redirect(url_for('certificates.view', cert_id=cert_id))

    return send_file(
        pdf_path,
        as_attachment=True,
        download_name=f"{certificate.certificate_number_with_rev.replace(' ', '_')}.pdf",
        mimetype='application/pdf'
    )


@reports_bp.route('/certificate/<int:cert_id>/replace-pdf', methods=['POST'])
@login_required
@approver_required
def replace_pdf(cert_id):
    """Replace signed PDF for an approved/published certificate."""
    import hashlib

    certificate = Certificate.query.get_or_404(cert_id)

    if not certificate.approval:
        flash('No approval record found.', 'danger')
        return redirect(url_for('certificates.view', cert_id=cert_id))

    if certificate.approval.status not in [STATUS_APPROVED, STATUS_PUBLISHED]:
        flash('Can only replace PDF for approved reports.', 'danger')
        return redirect(url_for('certificates.view', cert_id=cert_id))

    # Check for uploaded signed PDF
    if 'signed_pdf' not in request.files:
        flash('Please upload the signed PDF file.', 'warning')
        return redirect(url_for('certificates.view', cert_id=cert_id))

    pdf_file = request.files['signed_pdf']
    if pdf_file.filename == '':
        flash('No PDF file selected.', 'warning')
        return redirect(url_for('certificates.view', cert_id=cert_id))

    if not pdf_file.filename.lower().endswith('.pdf'):
        flash('Please upload a PDF file.', 'warning')
        return redirect(url_for('certificates.view', cert_id=cert_id))

    # Save the uploaded signed PDF
    reports_folder = Path(current_app.config['REPORTS_FOLDER'])
    year = datetime.now().year
    signed_folder = reports_folder / 'signed' / str(year)
    signed_folder.mkdir(parents=True, exist_ok=True)

    safe_cert_num = certificate.certificate_number_with_rev.replace(' ', '_').replace('/', '-')
    pdf_filename = f"{safe_cert_num}_signed.pdf"
    output_pdf = signed_folder / pdf_filename

    # Archive old PDF if exists
    old_pdf_path = certificate.approval.signed_pdf_path
    if old_pdf_path:
        old_pdf_full = reports_folder / old_pdf_path
        if old_pdf_full.exists():
            archive_folder = reports_folder / 'archive' / str(year)
            archive_folder.mkdir(parents=True, exist_ok=True)
            timestamp = datetime.now().strftime('%Y%m%d_%H%M%S')
            archive_name = f"{safe_cert_num}_{timestamp}.pdf"
            old_pdf_full.rename(archive_folder / archive_name)

    pdf_file.save(output_pdf)

    # Calculate hash for integrity verification
    with open(output_pdf, 'rb') as f:
        pdf_hash = hashlib.sha256(f.read()).hexdigest()

    pdf_path = str(output_pdf.relative_to(reports_folder))
    timestamp = datetime.utcnow()

    certificate.approval.publish(pdf_path=pdf_path, pdf_hash=pdf_hash)
    certificate.approval.signature_timestamp = timestamp
    certificate.reported = True

    # Audit log
    audit = AuditLog(
        user_id=current_user.id,
        action='REPLACE_PDF',
        table_name='report_approvals',
        record_id=certificate.approval.id,
        old_values={'signed_pdf_path': old_pdf_path},
        new_values={
            'certificate_number': certificate.certificate_number_with_rev,
            'signed_pdf_path': pdf_path,
            'pdf_hash': pdf_hash[:16] + '...',
            'timestamp': timestamp.isoformat()
        },
        ip_address=request.remote_addr
    )
    db.session.add(audit)
    db.session.commit()

    flash(f'Signed PDF replaced: {certificate.certificate_number_with_rev}', 'success')
    return redirect(url_for('certificates.view', cert_id=cert_id))


# Legacy routes for backwards compatibility
@reports_bp.route('/<int:id>')
@login_required
def view(id):
    """View report approval details (legacy - redirects to certificate)."""
    approval = ReportApproval.query.get_or_404(id)
    if approval.certificate_id:
        return redirect(url_for('certificates.view', cert_id=approval.certificate_id))
    # Old test-record based approval
    return render_template('reports/view.html',
                           approval=approval,
                           status_labels=APPROVAL_STATUS_LABELS,
                           status_colors=STATUS_COLORS)


@reports_bp.route('/<int:id>/download')
@login_required
def download(id):
    """Download signed PDF (legacy route)."""
    approval = ReportApproval.query.get_or_404(id)

    if not approval.can_download_signed:
        flash('PDF is not available.', 'warning')
        if approval.certificate_id:
            return redirect(url_for('certificates.view', cert_id=approval.certificate_id))
        return redirect(url_for('reports.index'))

    pdf_path = Path(current_app.config['REPORTS_FOLDER']) / approval.signed_pdf_path

    if not pdf_path.exists():
        flash('PDF file not found.', 'danger')
        return redirect(url_for('reports.index'))

    return send_file(
        pdf_path,
        as_attachment=True,
        download_name=f"{approval.certificate_number}.pdf",
        mimetype='application/pdf'
    )


# ---------------------------------------------------------------------------
# Helper functions for generating reports per test type
# ---------------------------------------------------------------------------

def _get_logo_path():
    """Get path to company logo."""
    logo_path = Path(current_app.root_path).parent / 'templates' / 'logo.png'
    if not logo_path.exists():
        logo_path = Path(current_app.root_path) / 'static' / 'images' / 'logo.png'
    return logo_path if logo_path.exists() else None


def _generate_ctod_report(certificate, test_record, output_path):
    """Generate CTOD Word report for certificate approval workflow."""
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import numpy as np
    from utils.reporting.ctod_word_report import CTODReportGenerator
    from utils.models.ctod_specimen import CTODSpecimen

    results = {r.parameter_name: r for r in test_record.results.all()}
    geometry = test_record.geometry or {}

    test_info = {
        'test_project': certificate.test_order or '',
        'customer': certificate.customer or '',
        'customer_order': certificate.customer_order or '',
        'product_sn': certificate.product_sn or '',
        'specimen_id': test_record.specimen_id or '',
        'location_orientation': certificate.location_orientation or '',
        'material': certificate.material or '',
        'certificate_number': certificate.certificate_number_with_rev,
        'test_date': test_record.test_date.strftime('%Y-%m-%d') if test_record.test_date else '',
        'temperature': str(test_record.temperature) if test_record.temperature else '23',
        'requirement': certificate.requirement or '',
    }

    specimen_data = {
        'specimen_type': geometry.get('type', 'SE(B)'),
        'W': geometry.get('W', ''),
        'B': geometry.get('B', ''),
        'B_n': geometry.get('B_n', ''),
        'a_0': geometry.get('a_0', ''),
        'S': geometry.get('S', ''),
        'notch_type': geometry.get('notch_type', 'fatigue'),
    }

    material_data = {
        'yield_strength': geometry.get('yield_strength', ''),
        'ultimate_strength': geometry.get('ultimate_strength', ''),
        'youngs_modulus': geometry.get('youngs_modulus', ''),
        'poissons_ratio': geometry.get('poissons_ratio', 0.3),
    }

    # Build results dict with mock MeasuredValue/CTODResult objects
    class MockMeasured:
        def __init__(self, value, uncertainty):
            self.value = value
            self.uncertainty = uncertainty

    class MockCTODResult:
        def __init__(self, ctod_val, force_val, cmod_val, is_valid):
            self.ctod_value = MockMeasured(ctod_val[0], ctod_val[1])
            self.force = MockMeasured(force_val[0], force_val[1])
            self.cmod = MockMeasured(cmod_val[0], cmod_val[1])
            self.is_valid = is_valid

    report_results = {}

    for key in ['P_max', 'CMOD_max', 'K_max']:
        r = results.get(key)
        if r:
            report_results[key] = MockMeasured(r.value, r.uncertainty)

    report_results['compliance'] = results.get('compliance').value if results.get('compliance') else None

    # Recalculate validity
    W = geometry.get('W', 0)
    a_0 = geometry.get('a_0', 0)
    S = geometry.get('S', 0)
    specimen_type = geometry.get('type', 'SE(B)')

    if W > 0 and a_0 > 0:
        try:
            specimen_obj = CTODSpecimen(
                specimen_id=test_record.specimen_id or '',
                specimen_type=specimen_type,
                W=W, B=geometry.get('B', 1), a_0=a_0,
                S=S if S > 0 else W * 4,
                B_n=geometry.get('B_n')
            )
            report_results['is_valid'] = specimen_obj.is_valid_geometry
            report_results['validity_summary'] = specimen_obj.validity_summary()
        except Exception:
            report_results['is_valid'] = geometry.get('is_valid', False)
            report_results['validity_summary'] = geometry.get('validity_summary', '')
    else:
        report_results['is_valid'] = geometry.get('is_valid', False)
        report_results['validity_summary'] = geometry.get('validity_summary', '')

    # CTOD results
    ctod_points = geometry.get('ctod_points', {})
    for ctod_type in ['delta_c', 'delta_u', 'delta_m']:
        r = results.get(ctod_type)
        pt = ctod_points.get(ctod_type, {})
        if r and pt:
            report_results[ctod_type] = MockCTODResult(
                (r.value, r.uncertainty),
                (pt.get('force', 0), 0.1),
                (pt.get('cmod', 0), 0.01),
                report_results['is_valid']
            )

    report_data = CTODReportGenerator.prepare_report_data(
        test_info=test_info,
        specimen_data=specimen_data,
        material_data=material_data,
        results=report_results,
        crack_measurements=geometry.get('crack_measurements', [])
    )

    # Generate chart
    chart_path = None
    force = np.array(geometry.get('force', []))
    cmod = np.array(geometry.get('cmod', []))

    if len(force) > 0 and len(cmod) > 0:
        from app.ctod.routes import truncate_at_break

        cmod_plot, force_plot = truncate_at_break(cmod, force, break_threshold=0.5)
        fig, ax = plt.subplots(figsize=(6, 4))
        ax.plot(cmod_plot, force_plot, color='darkred', linewidth=1.5, label='Test Data')

        idx_max = np.argmax(force)
        P_max_val = force[idx_max]
        V_max = cmod[idx_max]

        elastic_coeffs = geometry.get('elastic_coeffs')
        if elastic_coeffs:
            slope, intercept = elastic_coeffs
            cmod_elastic = np.linspace(0, max(cmod) * 0.6, 100)
            force_elastic = (cmod_elastic - intercept) / slope
            force_elastic = np.maximum(force_elastic, 0)
            ax.plot(cmod_elastic, force_elastic, '--', color='grey', linewidth=1, label='Elastic Line')

            Vp = V_max - P_max_val * slope
            if Vp < 0:
                Vp = 0
            ax.plot([Vp, V_max], [0, P_max_val], ':', color='grey', linewidth=1,
                    label=f'Plastic Line (Vp={Vp:.4f} mm)')
            ax.plot(Vp, 0, '^', color='grey', markersize=8, label=f'Vp = {Vp:.4f} mm')

        for ctod_type, marker in [('delta_m', 'o'), ('delta_c', 'D'), ('delta_u', 's')]:
            pt = ctod_points.get(ctod_type)
            if pt:
                ax.plot(pt['cmod'], pt['force'], marker, color='grey', markersize=10,
                        markerfacecolor='none', markeredgewidth=2,
                        label=f'{ctod_type}: \u03b4={pt["ctod"]:.4f} mm')

        ax.set_xlabel('CMOD (mm)')
        ax.set_ylabel('Force (kN)')
        ax.set_title(f'Force vs CMOD - {test_record.specimen_id}')
        y_max = P_max_val * 1.1
        ax.set_ylim(0, y_max)
        ax.yaxis.set_major_locator(plt.MultipleLocator(y_max / 12))
        ax.xaxis.set_major_locator(plt.MultipleLocator(1))
        ax.legend(fontsize=7, loc='upper right')
        ax.grid(True, alpha=0.3)

        chart_path = Path(current_app.config['UPLOAD_FOLDER']) / f'ctod_chart_{test_record.id}.png'
        fig.savefig(chart_path, dpi=150, bbox_inches='tight')
        plt.close(fig)

    logo_path = _get_logo_path()
    generator = CTODReportGenerator(None)
    generator.generate_report(
        output_path=output_path,
        data=report_data,
        chart_path=chart_path,
        logo_path=logo_path
    )

    if chart_path and chart_path.exists():
        import os
        os.remove(chart_path)


def _generate_sonic_report(certificate, test_record, output_path):
    """Generate Sonic Resonance Word report for certificate approval workflow."""
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    from utils.reporting.sonic_word_report import SonicReportGenerator

    results = {r.parameter_name: r for r in test_record.results.all()}
    geometry = test_record.geometry or {}

    test_info = {
        'test_project': certificate.test_order or '',
        'customer': certificate.customer or '',
        'customer_order': certificate.customer_order or '',
        'product_sn': certificate.product_sn or '',
        'specimen_id': test_record.specimen_id or '',
        'location_orientation': certificate.location_orientation or '',
        'material': certificate.material or '',
        'certificate_number': certificate.certificate_number_with_rev,
        'test_date': test_record.test_date.strftime('%Y-%m-%d') if test_record.test_date else '',
        'temperature': str(test_record.temperature) if test_record.temperature else '23',
    }

    specimen_data = {
        'specimen_type': 'Round' if geometry.get('type') == 'round' else 'Square',
        'diameter': geometry.get('diameter', '-'),
        'side_length': geometry.get('side_length', '-'),
        'length': geometry.get('length', ''),
        'mass': geometry.get('mass', ''),
    }

    vl1 = geometry.get('vl1', 0)
    vl2 = geometry.get('vl2', 0)
    vl3 = geometry.get('vl3', 0)
    vs1 = geometry.get('vs1', 0)
    vs2 = geometry.get('vs2', 0)
    vs3 = geometry.get('vs3', 0)

    velocity_data = {
        'vl1': vl1 if vl1 else '-', 'vl2': vl2 if vl2 else '-', 'vl3': vl3 if vl3 else '-',
        'vs1': vs1 if vs1 else '-', 'vs2': vs2 if vs2 else '-', 'vs3': vs3 if vs3 else '-',
    }

    def get_result(key, default_val=0, default_unc=0):
        r = results.get(key)
        return (r.value, r.uncertainty) if r else (default_val, default_unc)

    class MockValue:
        def __init__(self, value, uncertainty):
            self.value = value
            self.uncertainty = uncertainty

    class MockResults:
        pass

    mock_results = MockResults()
    for attr, key in [('density', 'Density'), ('longitudinal_velocity', 'Vl'),
                      ('shear_velocity', 'Vs'), ('poissons_ratio', 'Poisson'),
                      ('shear_modulus', 'G'), ('youngs_modulus', 'E'),
                      ('flexural_frequency', 'ff'), ('torsional_frequency', 'ft')]:
        v, u = get_result(key)
        setattr(mock_results, attr, MockValue(v, u))

    notes = geometry.get('notes', '')
    mock_results.is_valid = 'Valid: True' in (notes or '')
    mock_results.validity_notes = notes or ''

    report_data = SonicReportGenerator.prepare_report_data(
        test_info=test_info,
        specimen_data=specimen_data,
        velocity_data=velocity_data,
        results=mock_results
    )

    # Generate velocity chart
    chart_path = None
    if vl1 > 0 and vs1 > 0:
        try:
            fig, ax = plt.subplots(figsize=(7, 4))
            x_vl = [0.8, 1.0, 1.2]
            x_vs = [1.8, 2.0, 2.2]
            bars_vl = ax.bar(x_vl, [vl1, vl2, vl3], width=0.15, label='Longitudinal (Vl)',
                             color='darkred', edgecolor='black')
            bars_vs = ax.bar(x_vs, [vs1, vs2, vs3], width=0.15, label='Shear (Vs)',
                             color='black', edgecolor='black')
            for bar, val in zip(bars_vl, [vl1, vl2, vl3]):
                ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 50,
                        f'{val:.0f}', ha='center', va='bottom', fontsize=8)
            for bar, val in zip(bars_vs, [vs1, vs2, vs3]):
                ax.text(bar.get_x() + bar.get_width()/2, bar.get_height() + 50,
                        f'{val:.0f}', ha='center', va='bottom', fontsize=8, color='white')
            vl_avg = (vl1 + vl2 + vl3) / 3
            vs_avg = (vs1 + vs2 + vs3) / 3
            ax.axhline(y=vl_avg, xmin=0.1, xmax=0.45, color='grey', linestyle='--', linewidth=1)
            ax.axhline(y=vs_avg, xmin=0.55, xmax=0.9, color='grey', linestyle=':', linewidth=1)
            ax.set_ylabel('Velocity (m/s)')
            ax.set_title(f'Sound Velocity Measurements - {test_record.specimen_id}')
            ax.set_xticks([1.0, 2.0])
            ax.set_xticklabels(['Longitudinal', 'Shear'])
            ax.legend(loc='upper right')
            ax.grid(True, alpha=0.3, axis='y')
            min_val = min(vl1, vl2, vl3, vs1, vs2, vs3) * 0.9
            max_val = max(vl1, vl2, vl3) * 1.1
            ax.set_ylim(min_val, max_val)
            chart_path = Path(current_app.config['UPLOAD_FOLDER']) / f'sonic_chart_{test_record.id}.png'
            fig.savefig(chart_path, dpi=150, bbox_inches='tight')
            plt.close(fig)
        except Exception:
            chart_path = None

    logo_path = _get_logo_path()
    generator = SonicReportGenerator(None)
    generator.generate_report(
        output_path=output_path,
        data=report_data,
        chart_path=chart_path,
        logo_path=logo_path
    )

    if chart_path and chart_path.exists():
        import os
        os.remove(chart_path)


def _generate_fcgr_report(certificate, test_record, output_path):
    """Generate FCGR Word report for certificate approval workflow."""
    import matplotlib
    matplotlib.use('Agg')
    import matplotlib.pyplot as plt
    import numpy as np
    from utils.reporting.fcgr_word_report import FCGRReportGenerator

    results = {r.parameter_name: r for r in test_record.results.all()}
    geometry = test_record.geometry or {}

    test_info = {
        'test_project': certificate.test_order or '',
        'customer': certificate.customer or '',
        'specimen_id': test_record.specimen_id or '',
        'material': certificate.material or '',
        'certificate_number': certificate.certificate_number_with_rev,
        'test_date': test_record.test_date.strftime('%Y-%m-%d') if test_record.test_date else '',
        'requirement': certificate.requirement or '',
    }

    specimen_data = {
        'specimen_type': geometry.get('type', 'C(T)'),
        'W': geometry.get('W', ''),
        'B': geometry.get('B', ''),
        'B_n': geometry.get('B_n', ''),
        'a_0': geometry.get('a_0', ''),
        'notch_height': geometry.get('notch_height', 0),
    }

    material_data = {
        'yield_strength': geometry.get('yield_strength', ''),
        'ultimate_strength': geometry.get('ultimate_strength', ''),
        'youngs_modulus': geometry.get('youngs_modulus', ''),
        'poissons_ratio': geometry.get('poissons_ratio', 0.3),
    }

    test_params = {
        'control_mode': geometry.get('control_mode', 'Load Control'),
        'load_ratio': geometry.get('load_ratio', 0.1),
        'frequency': geometry.get('frequency', 10),
        'temperature': test_record.temperature or 23,
        'wave_shape': geometry.get('wave_shape', 'Sine'),
        'environment': 'Laboratory Air',
        'dadn_method': geometry.get('dadn_method', 'Secant'),
        'outlier_threshold': geometry.get('outlier_threshold', 2.5),
    }

    # Build mock results
    class MockParis:
        def __init__(self, C, m, r2, n, dK_range, dadN_range, C_err, m_err):
            self.C = C
            self.m = m
            self.r_squared = r2
            self.n_points = n
            self.delta_K_range = dK_range
            self.da_dN_range = dadN_range
            self.std_error_C = C_err
            self.std_error_m = m_err

    class MockResults:
        def __init__(self):
            self.paris_law = None
            self.paris_law_initial = None
            self.n_valid_points = 0
            self.n_outliers = 0
            self.total_cycles = 0
            self.final_crack_length = 0
            self.threshold_delta_K = 0
            self.is_valid = True
            self.validity_notes = []

    mock_results = MockResults()

    paris_C = results.get('paris_C')
    paris_m = results.get('paris_m')
    r_squared = results.get('r_squared')
    n_points = results.get('n_points')
    dK_min = results.get('delta_K_min')
    dK_max = results.get('delta_K_max')
    final_crack = results.get('final_crack')
    total_cycles = results.get('total_cycles')

    if paris_C and paris_m:
        mock_results.paris_law = MockParis(
            paris_C.value, paris_m.value,
            r_squared.value if r_squared else 0,
            int(n_points.value) if n_points else 0,
            (dK_min.value if dK_min else 0, dK_max.value if dK_max else 0),
            (1e-7, 1e-3),
            paris_C.uncertainty, paris_m.uncertainty
        )
    mock_results.n_valid_points = int(n_points.value) if n_points else 0
    mock_results.n_outliers = sum(geometry.get('outlier_mask', []))
    mock_results.total_cycles = int(total_cycles.value) if total_cycles else 0
    mock_results.final_crack_length = final_crack.value if final_crack else 0
    mock_results.is_valid = geometry.get('is_valid', False)
    mock_results.validity_notes = geometry.get('validity_notes', [])

    report_data = FCGRReportGenerator.prepare_report_data(
        test_info=test_info,
        specimen_data=specimen_data,
        material_data=material_data,
        test_params=test_params,
        results=mock_results
    )

    # Generate plots
    plot1_path = None
    plot2_path = None

    raw_cycles = geometry.get('raw_cycles', geometry.get('cycles', []))
    raw_crack_lengths = geometry.get('raw_crack_lengths', geometry.get('crack_lengths', []))
    delta_K = np.array(geometry.get('delta_K', []))
    da_dN = np.array(geometry.get('da_dN', []))
    outlier_mask = np.array(geometry.get('outlier_mask', []))

    if len(raw_cycles) > 0:
        fig1, ax1 = plt.subplots(figsize=(8, 5))
        ax1.plot(raw_cycles, raw_crack_lengths, color='darkred', linewidth=1.5, marker='o', markersize=3,
                 label='Crack Length, a (mm)')
        ax1.set_xlabel('Cycles (N)')
        ax1.set_ylabel('Crack Length, a (mm)')
        ax1.set_title(f'Crack Growth - {test_record.specimen_id}')
        ax1.set_xlim(left=0)
        ax1.set_ylim(bottom=0)
        ax1.yaxis.set_major_locator(plt.MultipleLocator(1))
        ax1.legend(fontsize=7, loc='upper left')
        ax1.grid(True, alpha=0.3)
        plot1_path = Path(current_app.config['UPLOAD_FOLDER']) / f'fcgr_plot1_{test_record.id}.png'
        fig1.savefig(plot1_path, dpi=150, bbox_inches='tight')
        plt.close(fig1)

    if len(delta_K) > 0 and paris_C:
        fig2, ax2 = plt.subplots(figsize=(8, 5))
        valid_mask = ~outlier_mask if len(outlier_mask) == len(delta_K) else np.ones(len(delta_K), dtype=bool)
        ax2.loglog(delta_K[valid_mask], da_dN[valid_mask], 'o', color='darkred', markersize=4, label='Valid Data')
        if np.any(~valid_mask):
            ax2.loglog(delta_K[~valid_mask], da_dN[~valid_mask], 'x', color='grey', markersize=5, label='Outliers')

        dK_fit = np.logspace(np.log10(dK_min.value * 0.9), np.log10(dK_max.value * 1.1), 100)
        dadN_fit = paris_C.value * dK_fit ** paris_m.value
        ax2.loglog(dK_fit, dadN_fit, '-', color='black', linewidth=2,
                   label=f'Paris: C={paris_C.value:.2e}, m={paris_m.value:.2f}')
        ax2.set_xlabel('\u0394K (MPa\u221am)')
        ax2.set_ylabel('da/dN (mm/cycle)')
        ax2.set_title(f'Paris Law - {test_record.specimen_id}')
        ax2.legend(fontsize=6, loc='lower right')
        ax2.grid(True, alpha=0.3, which='both')
        plot2_path = Path(current_app.config['UPLOAD_FOLDER']) / f'fcgr_plot2_{test_record.id}.png'
        fig2.savefig(plot2_path, dpi=150, bbox_inches='tight')
        plt.close(fig2)

    # Photo paths
    photo_paths = []
    for photo in geometry.get('photos', []):
        photo_path = Path(current_app.root_path) / 'static' / 'uploads' / photo['filename']
        if photo_path.exists():
            photo_paths.append(photo_path)

    logo_path = _get_logo_path()
    generator = FCGRReportGenerator(None)
    generator.generate_report(
        output_path=output_path,
        data=report_data,
        plot1_path=plot1_path,
        plot2_path=plot2_path,
        logo_path=logo_path,
        photo_paths=photo_paths if photo_paths else None
    )

    import os
    if plot1_path and plot1_path.exists():
        os.remove(plot1_path)
    if plot2_path and plot2_path.exists():
        os.remove(plot2_path)


def _generate_kic_report(certificate, test_record, output_path):
    """Generate KIC Word report for certificate approval workflow."""
    import re
    import numpy as np
    from utils.reporting.kic_word_report import KICReportGenerator
    from app.models import AnalysisResult

    geometry_data = test_record.geometry if test_record.geometry else {}
    geometry = geometry_data.get('specimen_geometry', {})
    material_props = geometry_data.get('material_properties', {})
    raw_data = geometry_data.get('raw_data', {})

    # Get analysis results
    results = {}
    analysis_records = AnalysisResult.query.filter_by(test_record_id=test_record.id).all()
    validity_notes = []
    is_valid = True
    for ar in analysis_records:
        if ar.parameter_name in ('P_max', 'P_Q', 'K_Q', 'K_IC'):
            results[ar.parameter_name] = {'value': ar.value, 'uncertainty': ar.uncertainty}
        else:
            results[ar.parameter_name] = ar.value
        if ar.validity_notes:
            validity_notes.append(ar.validity_notes)
        if not ar.is_valid:
            is_valid = False
    results['is_valid'] = is_valid
    results['validity_notes'] = validity_notes

    # Generate chart
    chart_path = None
    if 'force' in raw_data and 'displacement' in raw_data:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt
        from app.kic.routes import truncate_at_break

        fig, ax = plt.subplots(figsize=(8, 6))
        force = np.array(raw_data['force'])
        displacement = np.array(raw_data['displacement'])
        disp_plot, force_plot = truncate_at_break(displacement, force, break_threshold=0.5)
        ax.plot(disp_plot, force_plot, color='darkred', linewidth=1.5, label='Force vs Displacement')

        P_Q_data = results.get('P_Q')
        if P_Q_data and isinstance(P_Q_data, dict):
            P_Q = P_Q_data.get('value')
            if P_Q:
                idx = np.argmin(np.abs(force - P_Q))
                ax.plot(displacement[idx], P_Q, 'D', color='grey', markersize=10,
                        markerfacecolor='none', markeredgewidth=2, label=f'PQ = {P_Q:.2f} kN')

        P_max_data = results.get('P_max')
        if P_max_data and isinstance(P_max_data, dict):
            P_max = P_max_data.get('value')
            if P_max:
                idx = np.argmax(force)
                ax.plot(displacement[idx], P_max, 's', color='grey', markersize=12,
                        markerfacecolor='none', markeredgewidth=2, label=f'Pmax = {P_max:.2f} kN')

        ax.set_xlabel('Displacement (mm)')
        ax.set_ylabel('Force (kN)')
        ax.set_title('Force vs Displacement (ASTM E399)')
        ax.legend()
        ax.grid(True, alpha=0.3)

        chart_path = Path(current_app.config['REPORTS_FOLDER']) / f'kic_chart_{test_record.id}.png'
        fig.savefig(chart_path, dpi=150, bbox_inches='tight')
        plt.close(fig)

    # Parse KIC requirement
    kic_req = '-'
    if certificate.requirement:
        req_str = certificate.requirement
        patterns = [
            r'K_?IC[:\s]*([>≥<≤]?\s*[\d.]+)\s*(?:MPa)?',
            r'fracture\s*toughness[:\s]*([>≥<≤]?\s*[\d.]+)',
        ]
        for pattern in patterns:
            match = re.search(pattern, req_str, re.IGNORECASE)
            if match:
                kic_req = match.group(1).strip()
                if not kic_req.startswith(('>', '<', '\u2265', '\u2264')):
                    kic_req = '>' + kic_req
                break

    test_info = {
        'certificate_number': certificate.certificate_number_with_rev,
        'test_project': certificate.test_order or '',
        'customer': certificate.customer or '',
        'specimen_id': test_record.specimen_id or '',
        'material': certificate.material or '',
        'test_date': test_record.test_date.strftime('%Y-%m-%d') if test_record.test_date else '',
        'temperature': test_record.temperature or '23',
        'kic_req': kic_req,
    }

    dimensions = {
        'specimen_type': geometry.get('type', 'SE(B)'),
        'W': str(geometry.get('W', '')),
        'B': str(geometry.get('B', '')),
        'B_n': str(geometry.get('B_n', '')),
        'a_0': str(geometry.get('a_0', '')),
        'S': str(geometry.get('S', '-')),
    }

    mat_props = {
        'yield_strength': str(material_props.get('yield_strength', '')),
        'ultimate_strength': str(material_props.get('ultimate_strength', '')),
        'youngs_modulus': str(material_props.get('youngs_modulus', '')),
        'poissons_ratio': str(material_props.get('poissons_ratio', '0.3')),
    }

    # Create result proxy
    class ResultProxy:
        def __init__(self, data):
            self._data = data

        @property
        def P_max(self):
            d = self._data.get('P_max', {})
            if isinstance(d, dict):
                return type('MV', (), {'value': d.get('value', 0), 'uncertainty': d.get('uncertainty', 0)})()
            return type('MV', (), {'value': 0, 'uncertainty': 0})()

        @property
        def P_Q(self):
            d = self._data.get('P_Q', {})
            if isinstance(d, dict):
                return type('MV', (), {'value': d.get('value', 0), 'uncertainty': d.get('uncertainty', 0)})()
            return type('MV', (), {'value': 0, 'uncertainty': 0})()

        @property
        def K_Q(self):
            d = self._data.get('K_Q', {})
            if isinstance(d, dict):
                return type('MV', (), {'value': d.get('value', 0), 'uncertainty': d.get('uncertainty', 0)})()
            return type('MV', (), {'value': 0, 'uncertainty': 0})()

        @property
        def K_IC(self):
            d = self._data.get('K_IC')
            if d and isinstance(d, dict):
                return type('MV', (), {'value': d.get('value', 0), 'uncertainty': d.get('uncertainty', 0)})()
            return None

        @property
        def P_ratio(self):
            return self._data.get('P_ratio', 0)

        @property
        def compliance(self):
            return self._data.get('compliance', 0)

        @property
        def is_valid(self):
            return self._data.get('is_valid', False)

        @property
        def validity_notes(self):
            return self._data.get('validity_notes', [])

    result_proxy = ResultProxy(results)
    crack_measurements = geometry.get('crack_measurements', [])

    logo_path = _get_logo_path()
    generator = KICReportGenerator()
    generator.generate_report(
        output_path=output_path,
        test_info=test_info,
        dimensions=dimensions,
        material_props=mat_props,
        results=result_proxy,
        chart_path=chart_path if chart_path and chart_path.exists() else None,
        logo_path=logo_path,
        precrack_measurements=crack_measurements if len(crack_measurements) == 5 else None
    )

    if chart_path and chart_path.exists():
        import os
        os.remove(chart_path)


def _generate_vickers_report(certificate, test_record, output_path):
    """Generate Vickers Word report for certificate approval workflow."""
    import numpy as np
    from docx import Document
    from docx.shared import Inches, Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from app.models import AnalysisResult

    test_params = test_record.geometry if test_record.geometry else {}
    readings = test_params.get('readings', [])

    # Get analysis results
    results = {}
    analysis_records = AnalysisResult.query.filter_by(test_record_id=test_record.id).all()
    for ar in analysis_records:
        if ar.parameter_name == 'mean_hardness':
            results['mean_hardness'] = {'value': ar.value, 'uncertainty': ar.uncertainty}
        else:
            results[ar.parameter_name] = ar.value
    results['load_level'] = test_params.get('load_level', 'HV')
    results['uncertainty_budget'] = test_params.get('uncertainty_budget', {})

    # Generate chart
    chart_path = None
    if readings:
        import matplotlib
        matplotlib.use('Agg')
        import matplotlib.pyplot as plt

        values = [r['hardness_value'] for r in readings]
        mean_val = np.mean(values)

        fig, ax = plt.subplots(figsize=(8, 5))
        x = list(range(1, len(values) + 1))
        ax.plot(x, values, color='darkred', linewidth=2, marker='o',
                markersize=10, markerfacecolor='darkred', markeredgecolor='darkred')
        for xi, val in zip(x, values):
            ax.text(xi, val + max(values) * 0.02, f'{val:.1f}',
                    ha='center', va='bottom', fontsize=9)
        ax.axhline(y=mean_val, color='grey', linestyle=':', linewidth=2,
                   label=f'Mean: {mean_val:.1f}')
        ax.set_xlabel('Reading Number')
        ax.set_ylabel(f'Hardness ({test_params.get("load_level", "HV")})')
        ax.set_title('Hardness Profile')
        ax.legend(loc='upper right')
        ax.set_xlim(0.5, len(values) + 0.5)
        ax.set_ylim(0, max(values) * 1.15)
        ax.set_xticks(x)
        ax.grid(True, alpha=0.3, axis='y')

        chart_path = Path(current_app.config['REPORTS_FOLDER']) / f'vickers_chart_{test_record.id}.png'
        fig.savefig(chart_path, dpi=150, bbox_inches='tight')
        plt.close(fig)

    test_info = {
        'certificate_number': certificate.certificate_number_with_rev,
        'test_project': certificate.test_order or '',
        'customer': certificate.customer or '',
        'customer_order': certificate.customer_order or '',
        'specimen_id': test_record.specimen_id or '',
        'customer_specimen_info': certificate.customer_specimen_info or '',
        'material': certificate.material or '',
        'requirement': certificate.requirement or '',
        'location_orientation': test_params.get('location_orientation', certificate.location_orientation or ''),
        'test_date': test_record.test_date.strftime('%Y-%m-%d') if test_record.test_date else '',
        'temperature': test_record.temperature or 23,
        'load_level': test_params.get('load_level', 'HV 10'),
        'dwell_time': test_params.get('dwell_time', '15'),
        'notes': test_params.get('notes', ''),
        'operator': current_user.full_name if current_user.full_name else current_user.username,
    }

    # Build result proxy
    class ResultProxy:
        def __init__(self, data, readings_list):
            self._data = data
            self._readings = readings_list

        @property
        def mean_hardness(self):
            d = self._data.get('mean_hardness', {})
            return type('MV', (), {'value': d.get('value', 0), 'uncertainty': d.get('uncertainty', 0)})()

        @property
        def std_dev(self):
            return self._data.get('std_dev', 0)

        @property
        def range_value(self):
            return self._data.get('range_value', 0)

        @property
        def min_value(self):
            return self._data.get('min_value', 0)

        @property
        def max_value(self):
            return self._data.get('max_value', 0)

        @property
        def n_readings(self):
            return self._data.get('n_readings', 0)

        @property
        def load_level(self):
            return self._data.get('load_level', 'HV')

    result_proxy = ResultProxy(results, readings)
    uncertainty_budget = results.get('uncertainty_budget', {})
    requirement_value = test_info.get('requirement', '') or '-'

    # Build report from scratch (same as vickers/routes.py)
    doc = Document()

    style = doc.styles['Normal']
    style.paragraph_format.space_before = Pt(0)
    style.paragraph_format.space_after = Pt(3)
    style.paragraph_format.line_spacing = 1.0
    style.font.size = Pt(10)

    dark_green = RGBColor(0x00, 0x64, 0x00)
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_style.paragraph_format.space_before = Pt(8)
        heading_style.paragraph_format.space_after = Pt(4)
        heading_style.font.color.rgb = dark_green

    logo_path = _get_logo_path()

    for section in doc.sections:
        section.top_margin = Cm(1.5)
        section.bottom_margin = Cm(1.5)
        section.left_margin = Cm(2.0)
        section.right_margin = Cm(2.0)
        header = section.header
        header.is_linked_to_previous = False

        logo_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        logo_para.alignment = WD_ALIGN_PARAGRAPH.LEFT
        logo_para.paragraph_format.space_after = Pt(0)
        if logo_path:
            logo_run = logo_para.add_run()
            logo_run.add_picture(str(logo_path), width=Cm(5.0))

        title_para = header.add_paragraph()
        title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_para.paragraph_format.space_before = Pt(0)
        title_para.paragraph_format.space_after = Pt(0)
        title_run = title_para.add_run('Vickers Hardness Test Report')
        title_run.bold = True
        title_run.font.size = Pt(12)

        std_para = header.add_paragraph()
        std_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        std_para.paragraph_format.space_before = Pt(0)
        std_para.paragraph_format.space_after = Pt(0)
        std_run = std_para.add_run('ASTM E92 / ISO 6507')
        std_run.font.size = Pt(8)

        cert_para = header.add_paragraph()
        cert_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        cert_para.paragraph_format.space_before = Pt(0)
        cert_para.paragraph_format.space_after = Pt(0)
        cert_run = cert_para.add_run(f"Certificate: {test_info.get('certificate_number', '')}")
        cert_run.font.size = Pt(8)

        date_para = header.add_paragraph()
        date_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_para.paragraph_format.space_before = Pt(0)
        date_para.paragraph_format.space_after = Pt(0)
        date_run = date_para.add_run(f"Date: {test_info.get('test_date', '')}")
        date_run.font.size = Pt(8)

    # Test Information
    heading = doc.add_heading('Test Information', level=1)
    heading.paragraph_format.space_before = Pt(0)
    heading.paragraph_format.space_after = Pt(6)

    info_data = [
        ('Test Project:', test_info.get('test_project', ''), 'Temperature:', f"{test_info.get('temperature', '23')} \u00b0C"),
        ('Customer:', test_info.get('customer', ''), 'Test Standard:', 'ASTM E92 / ISO 6507'),
        ('Customer Order:', test_info.get('customer_order', ''), 'Test Equipment:', 'q-ness ATM test machine'),
        ('Product S/N:', test_info.get('specimen_id', ''), 'Load Level:', test_info.get('load_level', '')),
        ('Material:', test_info.get('material', ''), 'Dwell Time:', f"{test_info.get('dwell_time', '15')} s"),
        ('Customer Specimen Info:', test_info.get('customer_specimen_info', ''), 'Location/Orientation:', test_info.get('location_orientation', '')),
        ('Requirement:', test_info.get('requirement', ''), 'Operator:', test_info.get('operator', '')),
    ]

    table = doc.add_table(rows=len(info_data), cols=4)
    table.style = 'Table Grid'
    for i, (label1, value1, label2, value2) in enumerate(info_data):
        table.rows[i].cells[0].text = label1
        table.rows[i].cells[1].text = str(value1) if value1 else ''
        table.rows[i].cells[2].text = label2
        table.rows[i].cells[3].text = str(value2) if value2 else ''
        if table.rows[i].cells[0].paragraphs[0].runs:
            table.rows[i].cells[0].paragraphs[0].runs[0].bold = True
        if table.rows[i].cells[2].paragraphs[0].runs:
            table.rows[i].cells[2].paragraphs[0].runs[0].bold = True
        for cell in table.rows[i].cells:
            cell.paragraphs[0].paragraph_format.space_before = Pt(1)
            cell.paragraphs[0].paragraph_format.space_after = Pt(1)

    # Results Summary
    heading = doc.add_heading('Results Summary', level=1)
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)

    table = doc.add_table(rows=8, cols=3)
    table.style = 'Table Grid'
    for i, h in enumerate(['Parameter', 'Value', 'Unit']):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True

    results_data = [
        ('Mean Hardness', f'{result_proxy.mean_hardness.value:.1f} \u00b1 {result_proxy.mean_hardness.uncertainty:.1f}', result_proxy.load_level),
        ('Standard Deviation', f'{result_proxy.std_dev:.1f}', result_proxy.load_level),
        ('Range', f'{result_proxy.range_value:.1f}', result_proxy.load_level),
        ('Minimum', f'{result_proxy.min_value:.1f}', result_proxy.load_level),
        ('Maximum', f'{result_proxy.max_value:.1f}', result_proxy.load_level),
        ('Number of Readings', str(result_proxy.n_readings), '-'),
        ('Requirement', requirement_value, '-'),
    ]
    for i, (param, value, unit) in enumerate(results_data):
        table.rows[i+1].cells[0].text = param
        table.rows[i+1].cells[1].text = value
        table.rows[i+1].cells[2].text = unit
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_before = Pt(1)
            cell.paragraphs[0].paragraph_format.space_after = Pt(1)

    # Individual Readings
    heading = doc.add_heading('Individual Readings', level=1)
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)

    table = doc.add_table(rows=len(readings) + 1, cols=3)
    table.style = 'Table Grid'
    for i, h in enumerate(['#', 'Location', 'Hardness']):
        table.rows[0].cells[i].text = h
        table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    for i, r in enumerate(readings):
        table.rows[i+1].cells[0].text = str(r.get('reading_number', i+1))
        table.rows[i+1].cells[1].text = r.get('location', f'Point {i+1}')
        table.rows[i+1].cells[2].text = f"{r.get('hardness_value', 0):.1f}"
    for row in table.rows:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_before = Pt(1)
            cell.paragraphs[0].paragraph_format.space_after = Pt(1)

    # Chart
    if chart_path and chart_path.exists():
        heading = doc.add_heading('Hardness Profile', level=1)
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(6)
        doc.add_picture(str(chart_path), width=Inches(5.5))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Approval Signatures
    heading = doc.add_heading('Approval', level=1)
    heading.paragraph_format.space_before = Pt(12)
    heading.paragraph_format.space_after = Pt(6)

    sig_table = doc.add_table(rows=4, cols=4)
    sig_table.style = 'Table Grid'
    for i, h in enumerate(['Role', 'Name', 'Signature', 'Date']):
        sig_table.rows[0].cells[i].text = h
        sig_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    sig_table.rows[1].cells[0].text = 'Tested by:'
    sig_table.rows[2].cells[0].text = 'Reviewed by:'
    sig_table.rows[3].cells[0].text = 'Approved by:'
    for row in sig_table.rows:
        for cell in row.cells:
            cell.paragraphs[0].paragraph_format.space_before = Pt(1)
            cell.paragraphs[0].paragraph_format.space_after = Pt(1)

    # Disclaimer footer
    disclaimer_text = (
        "All work and services carried out by Durabler are subject to, and conducted in accordance with, "
        "Durabler standard terms and conditions, which are available at durabler.se. This document shall not "
        "be reproduced other than in full, except with prior written approval of the issuer. The results pertain "
        "only to the item(s) as sampled by the client unless otherwise indicated. Durabler a part of Subseatec S AB, "
        "Address: Durabler C/O Subseatec, Dalavägen 23, 68130 Kristinehamn, SWEDEN"
    )
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        footer_para = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        footer_para.clear()
        footer_run = footer_para.add_run(disclaimer_text)
        footer_run.font.size = Pt(7)
        footer_run.italic = True

    doc.save(output_path)

    if chart_path and chart_path.exists():
        import os
        os.remove(chart_path)


# Context processor for pending count in navbar
@reports_bp.app_context_processor
def inject_pending_count():
    """Inject pending reports count into all templates."""
    if current_user.is_authenticated and current_user.can_approve:
        return {'pending_reports_count': ReportApproval.get_pending_count()}
    return {'pending_reports_count': 0}
